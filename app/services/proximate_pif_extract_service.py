"""ProximatePifExtractService — Blue Nile round intake (July 2026).

Partner Information Forms arrive as Word/PDF files (Arabic originals +
English translations). This service turns one uploaded PIF into the
structured dict stored on ProximatePartner.intake_form under the 'pif'
key, so a folder of 31 forms can be ingested instead of retyped.

Same shape as proximate_grant_extract_service (Phase 721b): extract
text locally, one AI call with a tool schema, return a PROPOSAL — the
caller decides what to persist. The importing endpoint also accepts a
pre-parsed fields_json so bulk ingestion never depends on AI uptime.
"""

import logging

logger = logging.getLogger('kuja')

EXTRACT_MODEL = 'claude-sonnet-4-6'
MAX_DOC_TEXT_CHARS = 24000

# The canonical PIF shape (sections A-F of the paper form). Keys stay
# stable — the partner detail UI renders straight from this dict.
PIF_FIELDS = (
    'org_name', 'org_name_ar', 'abbreviated_name', 'legal_status',
    'registration_number', 'country_of_registration', 'year_established',
    'headquarters_address', 'operational_offices', 'website_social',
    'contact_person', 'contact_title', 'contact_phone', 'contact_email',
    'alternate_contact', 'thematic_areas', 'primary_target_group',
    'geographic_areas', 'key_activities_12m', 'beneficiaries_12m',
    'current_donor_programs', 'annual_budget', 'funding_sources_12m',
    'financial_system', 'bank_account_holder', 'bank_name_branch',
    'bank_account_number', 'currency', 'mobile_money_number',
    'transfer_capabilities', 'ever_blacklisted', 'external_audits_2y',
    'has_coi_policy', 'senior_staff_peps', 'declaration_name',
    'declaration_title', 'declaration_date',
)


def extract_docx_text(filepath: str) -> str:
    """Text from a .docx PIF — paragraphs plus table cells (the form IS
    a table, so tables carry the payload)."""
    try:
        import docx as _docx
    except ImportError:
        logger.error('python-docx not installed — cannot extract PIF text')
        return ''
    try:
        d = _docx.Document(filepath)
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(' | '.join(cells))
        return '\n'.join(parts).strip()
    except Exception as e:
        logger.warning(f'PIF docx text extraction failed: {e}')
        return ''


def extract_pif_text(filepath: str) -> str:
    if filepath.lower().endswith(('.docx', '.doc')):
        return extract_docx_text(filepath)
    # PDF path — reuse the agreement extractor's reader.
    from app.services.proximate_grant_extract_service import extract_pdf_text
    return extract_pdf_text(filepath)


def extract_pif_fields(*, doc_text: str, filename: str = '') -> dict | None:
    """One Claude call: PIF text -> {field: value} for PIF_FIELDS.
    Returns None when AI is unavailable (caller falls back to a stub
    partner + manual completion)."""
    try:
        from app.services.ai_service import AIService
    except Exception:
        return None

    system_prompt = (
        'You are a grants officer at Adeso digitising Partner Information '
        'Forms (PIFs) submitted by grassroots Sudanese organisations. The '
        'form has sections A-F: organisational identity, contact details, '
        'programmatic information, financial information, risk & '
        'compliance, and a signed declaration. Forms may be in English or '
        'Arabic, and English versions may contain machine-translation '
        'artifacts — prefer the plainly intended meaning. Extract ONLY '
        'what the form states; use null for anything absent. Keep numbers '
        'as written (do not convert currencies).'
    )
    user_message = (
        f'PIF FILE: {filename or "(uploaded)"}\n'
        f'=== FORM TEXT ({len(doc_text)} chars) ===\n'
        f'{doc_text[:MAX_DOC_TEXT_CHARS]}\n'
        '=== END ===\n\n'
        'Return the extraction via the record_pif tool.'
    )
    tool_schema = {
        'type': 'object',
        'properties': {f: {'type': ['string', 'null']} for f in PIF_FIELDS},
        'required': [],
    }
    try:
        parsed = AIService._call_claude_tool(
            system_prompt,
            user_message,
            tool_name='record_pif',
            tool_description=(
                'Structured extraction of one Partner Information Form '
                'for the Proximate partner register.'
            ),
            tool_schema=tool_schema,
            max_tokens=3000,
            endpoint='proximate_pif_extract',
        )
        if not isinstance(parsed, dict):
            return None
        return {k: v for k, v in parsed.items() if k in PIF_FIELDS and v}
    except Exception as e:
        logger.warning(f'PIF AI extraction failed: {e}')
        return None
