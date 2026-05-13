"""
Generate Kuja Grant Management System — Business Requirements Document
Version 5.0 — May 2026

Scope: Business requirements only. User-facing guidance has moved to a
separate User Guide. This document is the authoritative requirements
artefact for sponsors, donors, partner integrators, and the engineering
team.

Voice: "the system shall…" forward-looking requirement form. Content
reflects the system actually built; this document is derived from that
system but written as a requirements specification so it remains
useful for downstream integrations, audits, and partner conversations.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
import sys

# Regenerate the diagrams every time the BRD is rebuilt so they
# stay in lockstep with content edits.
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import brd_diagrams as _diagrams
DIAGRAM_PATHS = _diagrams.generate_all()
DIAGRAM = {os.path.basename(p).split("_", 1)[1].rsplit(".", 1)[0]: p
           for p in DIAGRAM_PATHS}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_table(doc, headers, rows, col_widths=None, header_color="1B3A5C"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(hdr_cells[i], header_color)
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        stripe = "F2F6FA" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            set_cell_shading(row_cells[c_idx], stripe)
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)
    return table


def add_bullet_list(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
        p.paragraph_format.space_after = Pt(2)
        if isinstance(item, tuple):
            run_b = p.add_run(item[0])
            run_b.bold = True
            run_b.font.size = Pt(10)
            run_b.font.name = "Calibri"
            run_n = p.add_run(item[1])
            run_n.font.size = Pt(10)
            run_n.font.name = "Calibri"
        else:
            run = p.add_run(str(item))
            run.font.size = Pt(10)
            run.font.name = "Calibri"


def add_numbered_list(doc, items, prefix="The system shall "):
    """Number the items as requirements. Each item is the body text;
    the helper prepends `prefix` so the requirement form is consistent."""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        p.paragraph_format.space_after = Pt(3)
        run_num = p.add_run(f"{i}.  ")
        run_num.bold = True
        run_num.font.size = Pt(10)
        run_num.font.name = "Calibri"
        run_b = p.add_run(prefix)
        run_b.font.size = Pt(10)
        run_b.font.name = "Calibri"
        run_n = p.add_run(item)
        run_n.font.size = Pt(10)
        run_n.font.name = "Calibri"


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(6)
    return p


def add_diagram(doc, key, caption=None, width_inches=6.5):
    """Embed a generated diagram with an optional caption beneath."""
    path = DIAGRAM.get(key)
    if not path or not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c.add_run(caption)
        run.italic = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        run.font.name = "Calibri"


def add_callout(doc, label, text, color="EEF2F7"):
    """A subtle shaded paragraph used for narrative use cases at the
    top of each functional section."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, color)
    cell.text = ""
    p = cell.paragraphs[0]
    run_b = p.add_run(label + "  ")
    run_b.bold = True
    run_b.font.size = Pt(9)
    run_b.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run_b.font.name = "Calibri"
    run_n = p.add_run(text)
    run_n.font.size = Pt(10)
    run_n.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(4)
    return table


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
            run.font.size = Pt(13)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)
            run.font.size = Pt(11)
    return h


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(10)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for lvl in range(1, 4):
    hs = doc.styles[f"Heading {lvl}"]
    hs.font.name = "Calibri"
    if lvl == 1:
        hs.font.size = Pt(18)
        hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    elif lvl == 2:
        hs.font.size = Pt(13)
        hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        hs.paragraph_format.space_before = Pt(14)
        hs.paragraph_format.space_after = Pt(6)
    elif lvl == 3:
        hs.font.size = Pt(11)
        hs.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)
        hs.paragraph_format.space_before = Pt(10)
        hs.paragraph_format.space_after = Pt(4)

for section in doc.sections:
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)


# ===========================================================================
# COVER PAGE
# ===========================================================================

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
run = p.add_run("Kuja")
run.bold = True
run.font.size = Pt(56)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Grant Management System")
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
run = p.add_run("Business Requirements Document")
run.font.size = Pt(18)
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
run = p.add_run("Findable. Fundable. Trusted.")
run.italic = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(240)
run = p.add_run("Version 5.0  |  May 2026")
run.font.size = Pt(12)
run.font.name = "Calibri"

doc.add_page_break()


# ===========================================================================
# DOCUMENT CONTROL
# ===========================================================================

heading(doc, "Document Control", 1)

heading(doc, "Revision History", 2)
add_formatted_table(doc,
    ["Version", "Date", "Author", "Description"],
    [
        ["1.0", "January 2026", "Kuja Team", "Initial release — core grant lifecycle"],
        ["2.0", "February 2026", "Kuja Team", "AI services, capacity assessment, sanctions screening"],
        ["3.0", "March 2026", "Kuja Team", "Live sanctions, registry verification, AI reporting, donor requirements"],
        ["4.0", "April 2026", "Kuja Team", "Match engine, grant Q&A, diligence room, two-phase intake lifecycle, NGO This-Week home, decision audit, donor portfolio diagnostics"],
        ["5.0", "May 2026", "Kuja Team", "Capacity assessment passporting, deep AI co-pilot integration, compliance health & trajectory, risk register, organisational memory & provenance, web push, async AI dispatcher, admin self-service & system health"],
    ],
    col_widths=[0.7, 1.1, 1.1, 3.6]
)

doc.add_paragraph()
heading(doc, "Approval", 2)
add_formatted_table(doc,
    ["Role", "Name", "Signature", "Date"],
    [
        ["Project Sponsor", "", "", ""],
        ["Product Owner", "", "", ""],
        ["Technical Lead", "", "", ""],
        ["QA Lead", "", "", ""],
        ["Partner Reviewer", "", "", ""],
    ],
    col_widths=[1.5, 2.0, 1.8, 1.2]
)

doc.add_page_break()


# ===========================================================================
# TABLE OF CONTENTS
# ===========================================================================

heading(doc, "Table of Contents", 1)

toc_entries = [
    ("1. Executive Summary", 0),
    ("2. The Problem We Address", 0),
    ("    2.1 The NGO experience today", 1),
    ("    2.2 The donor experience today", 1),
    ("    2.3 The reviewer experience today", 1),
    ("3. Vision and Objectives", 0),
    ("4. Target Users", 0),
    ("5. Solution Overview", 0),
    ("    5.1 System Architecture", 1),
    ("    5.2 Action-Driven Workspaces", 1),
    ("6. Functional Requirements", 0),
    ("    6.1 Identity, Access and Authentication", 1),
    ("    6.2 Grant Lifecycle Management (two creation paths)", 1),
    ("    6.3 Application Lifecycle Management", 1),
    ("    6.4 Capacity Assessment and Passporting", 1),
    ("    6.5 Document Management and Real-Time Analysis", 1),
    ("    6.6 AI-Powered Assistance Across the Workflow", 1),
    ("    6.7 Compliance and Sanctions Screening", 1),
    ("    6.8 Government Registry Verification", 1),
    ("    6.9 Reporting and Outcome Tracking", 1),
    ("    6.10 Compliance and Risk — Simplified for Both Sides", 1),
    ("        6.10.1 For NGOs — AI as compliance co-pilot", 2),
    ("        6.10.2 For donors — pre-assessed, scored, monitored", 2),
    ("        6.10.3 Shared mechanics", 2),
    ("    6.11 Collaboration and Notifications", 1),
    ("    6.12 Organisational Memory and Provenance", 1),
    ("    6.13 Saved Searches and Personalisation", 1),
    ("    6.14 Administration and System Health", 1),
    ("    6.15 Internationalisation and Localisation", 1),
    ("    6.16 Field Operations: Mobile and Offline Work", 1),
    ("7. Non-Functional Requirements", 0),
    ("    7.1 Performance and Scalability", 1),
    ("    7.2 Reliability and Availability", 1),
    ("    7.3 Security and Privacy", 1),
    ("    7.4 Observability and Operability", 1),
    ("    7.5 Usability, Accessibility, and Inclusivity", 1),
    ("8. Use Cases (UC-001 through UC-017)", 0),
    ("9. Feature Flags and Phased Rollout", 0),
    ("APPENDICES", 0),
    ("    A. Status Definitions", 1),
    ("    B. Glossary", 1),
]
for entry, level in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(entry)
    run.font.name = "Calibri"
    if level == 0:
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    elif level == 1:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    else:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        run.italic = True

doc.add_page_break()


# ===========================================================================
# 1. EXECUTIVE SUMMARY
# ===========================================================================

heading(doc, "1. Executive Summary", 1)

add_body(doc,
    "Kuja is an end-to-end grant management system purpose-built for the "
    "Global South. It exists to make Global South non-profits "
    "findable by international donors and fundable when they apply — "
    "and to make those grants trustworthy, accountable, and easier to "
    "deliver for everyone involved. Kuja covers the full grant lifecycle: "
    "from organisation onboarding and capacity assessment, through grant "
    "discovery and matching, application drafting and submission, donor "
    "review and decisioning, compliance and sanctions screening, contract "
    "issuance, progress reporting, and ongoing risk and compliance health "
    "monitoring."
)

add_body(doc,
    "The system serves four primary user groups: non-governmental "
    "organisations (NGOs) seeking funding; donors and foundations "
    "deploying capital; independent reviewers evaluating proposals; and "
    "platform administrators operating the service. Each group has a "
    "dedicated workspace, but the value of Kuja lies in how the "
    "workspaces interlock — an NGO's capacity assessment becomes a "
    "donor's confidence signal; an AI-assisted application becomes a "
    "reviewer's structured evidence summary; a compliance flag becomes a "
    "tracked risk with an owner and a due date."
)

add_body(doc,
    "Kuja's strategic differentiation rests on five commitments. First, "
    "the platform treats Global South NGOs as first-class participants "
    "rather than translated afterthoughts: every workflow is available "
    "in six languages (English, Arabic, French, Swahili, Somali, and "
    "Spanish) with role-appropriate tone, and the system gracefully "
    "handles scanned documents, low-bandwidth connections, and inconsistent "
    "data quality typical of grassroots organisations. Second, AI is "
    "embedded as a working co-pilot rather than a chatbot — the system "
    "drafts applications, predicts donor concerns, surfaces compliance "
    "issues before submission, evaluates evidence per criterion, and "
    "generates auditable rationales, all grounded in the NGO's own "
    "history and the donor's stated criteria. Third, capacity assessments "
    "are completed once and passported across applications, "
    "regardless of which framework the donor prefers. Fourth, compliance "
    "is built in rather than bolted on: live sanctions screening, "
    "government registry verification, and document-by-document evidence "
    "analysis run continuously. Fifth, every AI claim is traceable to its "
    "source, every decision leaves an audit trail, and the donor-NGO "
    "relationship is supported by structured collaboration tools rather "
    "than email threads."
)

add_body(doc,
    "This document specifies the requirements the system shall meet. It "
    "is written in requirement form (\"the system shall…\") so it can "
    "serve as the basis for partner integrations, accreditation reviews, "
    "and continued product evolution."
)


# ===========================================================================
# 2. THE PROBLEM WE ADDRESS
# ===========================================================================

heading(doc, "2. The Problem We Address", 1)

add_body(doc,
    "Grant funding for Global South non-profits is structurally inefficient. "
    "Capital exists; capacity exists; but the matching, evaluation, and "
    "compliance machinery sits on infrastructure designed for a different "
    "set of users. The result is a market that systematically undervalues "
    "smaller, grassroots, non-English-first organisations — even when their "
    "programmes are objectively well-designed and well-run. Kuja addresses "
    "specific, well-known frictions on each side of the market."
)

heading(doc, "2.1 The NGO experience today", 2)
add_bullet_list(doc, [
    ("Capacity assessment fatigue. ",
     "Every donor wants their own version of the same baseline information. "
     "Strong NGOs spend disproportionate time answering the same 80–120 "
     "questions across CHS, UN-HACT, NUPAS, STEP, and donor-specific "
     "variants — time that could go to programming."),
    ("Application overhead with no help. ",
     "Donor portals collect detailed responses, but applicants have no way "
     "to know whether their answer is strong, whether they have addressed "
     "the criterion, or what evidence the reviewer expects."),
    ("Compliance evidence scattered. ",
     "Audits, board minutes, beneficiary data, and registration certificates "
     "live across email, drives, and physical binders. Producing them on "
     "demand under a tight reporting deadline is expensive."),
    ("Reporting templates that vary by donor. ",
     "Each donor has different indicators, periods, and narrative "
     "expectations. Reporters rewrite the same programme story five "
     "different ways with no help interpreting what each donor cares about."),
    ("Decision opacity. ",
     "Most applications are rejected with little or no feedback. There is "
     "no learning loop and no way for an NGO to improve over time."),
    ("Language exclusion. ",
     "Most donor portals are English-only. Francophone, Arabic, Swahili, "
     "and Somali NGOs are forced to write in their second or third "
     "language, often at a measurable disadvantage."),
])

heading(doc, "2.2 The donor experience today", 2)
add_bullet_list(doc, [
    ("Sourcing is slow and intuition-driven. ",
     "Finding strong NGOs in the Global South requires personal networks, "
     "regional intermediaries, or trial-and-error open calls."),
    ("Capacity signal is non-standard. ",
     "Each donor maintains their own capacity framework; cross-donor "
     "comparison is impossible without manual reconciliation."),
    ("Application quality varies widely. ",
     "Variance comes from language proficiency and writing experience, not "
     "from programme quality. Reviewers anchor on prose rather than "
     "evidence."),
    ("Compliance is reactive. ",
     "Sanctions screening, government registry verification, and document "
     "validation happen at the last minute — or after award, when problems "
     "are expensive to unwind."),
    ("Reports do not roll up. ",
     "Each report is written for the donor's template; aggregating "
     "outcomes across a portfolio requires manual extraction."),
    ("No early-warning signal. ",
     "Grants drift off track without visibility until a missed deadline or "
     "a compliance flag forces attention."),
])

heading(doc, "2.3 The reviewer experience today", 2)
add_bullet_list(doc, [
    ("Each application takes hours to read carefully. ",
     "Long narratives, embedded attachments, and inconsistent structure "
     "make systematic comparison difficult."),
    ("Rationales drift. ",
     "Reviewers anchor on each other rather than on the rubric, especially "
     "when reviewing in cohorts."),
    ("Patterns across the portfolio are invisible. ",
     "Reviewers cannot easily see which criteria are discriminative, which "
     "are answered generically, or where applicants systematically "
     "underclaim or overclaim."),
])

doc.add_page_break()


# ===========================================================================
# 3. VISION AND OBJECTIVES
# ===========================================================================

heading(doc, "3. Vision and Objectives", 1)

add_body(doc,
    "Kuja's vision is a grant economy in which Global South non-profits "
    "are discoverable, fundable, and accountable on equal footing with "
    "their global peers. The system pursues that vision through five "
    "concrete objectives."
)

add_formatted_table(doc,
    ["Objective", "How the system delivers it"],
    [
        ["O1. Make Global South NGOs findable.",
         "Match engine ranks NGOs against grants on capacity, sector, "
         "geography, eligibility, and track record. Donors discover "
         "qualified applicants without depending on intermediaries."],
        ["O2. Make Global South NGOs fundable.",
         "Multilingual application drafting, AI co-authoring grounded in "
         "the NGO's own evidence, pre-submission gap analysis, and "
         "coaching tone for first-time applicants raise the floor on "
         "application quality."],
        ["O3. Reduce duplicated work.",
         "Capacity assessment passporting, organisational memory, and "
         "evidence reuse across applications eliminate the rewrite tax "
         "NGOs pay today."],
        ["O4. Make trust legible.",
         "Live sanctions screening, registry verification, hash-chained "
         "audit trail, provenance on every AI claim, and a 4-pillar "
         "compliance health score give donors a defensible basis for "
         "decisions and continued funding."],
        ["O5. Make the relationship continuous.",
         "Structured comments with @mentions, web push notifications, "
         "tracked risks with owners and due dates, and trajectory-based "
         "compliance forecasting keep the donor-NGO relationship active "
         "between submission and final report."],
    ],
    col_widths=[2.0, 4.4]
)

doc.add_page_break()


# ===========================================================================
# 4. TARGET USERS
# ===========================================================================

heading(doc, "4. Target Users", 1)

add_body(doc,
    "The system serves four user roles, each with distinct goals, "
    "workflows, and access scopes."
)

add_formatted_table(doc,
    ["Role", "Primary goal", "Key actions"],
    [
        ["NGO",
         "Secure funding aligned to mission and capacity.",
         "Complete capacity passport · discover matching grants · draft "
         "and submit applications · respond to clarifications · upload "
         "evidence · submit progress reports · monitor compliance health"],
        ["Donor",
         "Deploy capital to highest-impact, lowest-risk programmes.",
         "Publish grants · define eligibility and rubric · review "
         "applications · score and decide · monitor compliance and "
         "trajectory · review reports · flag risks · request clarification"],
        ["Reviewer",
         "Evaluate applications consistently against the rubric.",
         "Score per criterion · cite evidence · raise red flags · compare "
         "applications · submit decision package"],
        ["Administrator",
         "Operate the platform safely and observably.",
         "Manage users and organisations · monitor system health · review "
         "AI surface health · manage feature flags · review failed logins "
         "· trigger sanctions rescreening · review audit chain"],
    ],
    col_widths=[1.0, 1.8, 3.6]
)

doc.add_page_break()


# ===========================================================================
# 5. SOLUTION OVERVIEW
# ===========================================================================

heading(doc, "5. Solution Overview", 1)

add_body(doc,
    "Kuja is a single web-based platform with role-aware workspaces. "
    "NGOs, donors, reviewers, and administrators sign in through the "
    "same authentication surface and land on a dashboard tailored to "
    "their goals for the day. The platform is built around an explicit "
    "principle: every workspace is action-driven — users see what they "
    "need to do next, not just what data exists."
)

heading(doc, "5.1 System Architecture", 2)

add_diagram(doc, "system_overview",
    "Figure 1 — User roles, the Kuja platform's capability pillars, "
    "and the external services the platform integrates with.")

add_body(doc,
    "The system shall be delivered as a single web platform with twelve "
    "logically separable capability pillars. Each pillar is reachable "
    "from any user's workspace where their role permits, and pillars "
    "interlock: an NGO's capacity passport informs an AI-co-authored "
    "application; that application's score becomes part of the donor's "
    "portfolio insights; the resulting grant's compliance health flows "
    "from the document analyses underneath the application — all from "
    "the same source of record."
)

add_bullet_list(doc, [
    ("Identity and access. ",
     "Multi-tenant organisations; role-based access control; two-factor "
     "authentication for administrators; per-organisation data isolation."),
    ("Lifecycle management. ",
     "Grants, applications, capacity assessments, and reports follow "
     "explicit state machines with audit trails and lifecycle telemetry."),
    ("Document intelligence. ",
     "Uploaded documents are extracted, analysed against the grant's "
     "evidence requirements, and scored for completeness with clear "
     "fallbacks for scanned or low-quality files."),
    ("AI co-pilot. ",
     "AI assistance is woven into every workflow as a working partner — "
     "drafting, evaluating, surfacing gaps, generating rationales, "
     "translating findings — never as a generic chatbot."),
    ("Compliance and trust. ",
     "Live sanctions screening, government registry verification, "
     "compliance health scoring, risk register, provenance tracking, "
     "and hash-chained audit log."),
    ("Operations and observability. ",
     "Admin self-service for system health, AI cost forecasting, demo "
     "readiness scanning, flagship AI surface health monitoring, audit "
     "retention configuration, and failed-login analytics."),
])

heading(doc, "5.2 Action-Driven Workspaces", 2)

add_body(doc,
    "Every role's primary surface shall be a dashboard organised around "
    "what the user must do, not just what records the system holds. The "
    "system shall surface to-dos, deadlines, blockers, and opportunities "
    "the moment the user signs in — the dashboard answers the question "
    "\"what should I do today?\" before the user has to ask."
)

heading(doc, "NGO workspace — \"This Week\"", 3)
add_body(doc,
    "An NGO's home shall present a single \"This Week\" surface that "
    "consolidates every obligation, opportunity, and decision the user "
    "owes attention to. The system shall not require the NGO to navigate "
    "across separate screens to discover what is due."
)
add_bullet_list(doc, [
    ("Upcoming reports due. ",
     "Every progress report due within the next four weeks, in date "
     "order, with the report draft state and an indication of which "
     "ones have AI-prepared draft narratives waiting."),
    ("Applications in progress. ",
     "Draft applications grouped by deadline urgency, each with the "
     "next action the system recommends (complete section X, upload "
     "document Y, address pre-flight finding Z)."),
    ("Grants matched but not started. ",
     "Newly published grants where the NGO's match score exceeds the "
     "threshold, surfaced with top strength and top blocker so the NGO "
     "can decide quickly whether to invest the effort."),
    ("Compliance to-dos. ",
     "Open clarifications, donor revision requests, risks awaiting "
     "response, and capacity passport gaps that, if closed, would "
     "improve the NGO's match scoring."),
    ("Recent feedback. ",
     "Donor feedback on accepted or revised reports, with AI-translated "
     "summaries of what each donor would like to see differently."),
    ("Capacity passport status. ",
     "Completion percentage across the embedded frameworks so the NGO "
     "knows how application-ready the organisation is."),
])

heading(doc, "Donor workspace", 3)
add_body(doc,
    "A donor's home shall present a portfolio-level command centre: "
    "applications awaiting decision, grants with overdue reports, "
    "compliance health bands across the portfolio, slips-in-N-days "
    "warnings, risks awaiting donor response, and recent applicant "
    "activity. The system shall also surface a single AI-generated "
    "portfolio insights panel with a headline, anomalies, and the "
    "next decisions the donor owes."
)

heading(doc, "Reviewer workspace", 3)
add_body(doc,
    "A reviewer's home shall present an action queue of applications "
    "assigned, scored, awaiting their second pass, and (where the "
    "donor has enabled it) comparable applications they have not yet "
    "rated. Every row shall surface whether an AI summary has been "
    "produced and whether the reviewer has cited evidence per criterion."
)

heading(doc, "Administrator workspace", 3)
add_body(doc,
    "An administrator's home shall present an operations command "
    "centre: system health, AI cost projection vs. budget, AI surface "
    "health (the daily probe of every flagship AI surface), demo "
    "readiness findings, stuck document extractions, organisations "
    "pending verification, and failed-login analytics."
)

doc.add_page_break()


# ===========================================================================
# 6. FUNCTIONAL REQUIREMENTS
# ===========================================================================

heading(doc, "6. Functional Requirements", 1)
add_body(doc,
    "This chapter specifies what the system shall do, organised by "
    "capability area. Each section opens with a narrative use case "
    "establishing the user context, followed by enumerated requirements "
    "in the form \"the system shall…\". Where AI plays a material role, "
    "a dedicated subsection details how AI is leveraged."
)


# --- 6.1 Identity --------------------------------------------------------
heading(doc, "6.1 Identity, Access and Authentication", 2)

add_callout(doc, "USE CASE",
    "Fatima, a programme officer at an NGO in Kakamega, signs in to apply "
    "for a maternal-health grant. She belongs to the Amani Foundation, "
    "which her colleague registered last month. She is presented only "
    "with applications, grants, and reports that belong to her "
    "organisation; an administrator visiting later sees the same "
    "organisation in their fuller administrative view."
)

add_numbered_list(doc, [
    "support user registration and authentication via email and password, with bcrypt-hashed credentials stored at rest.",
    "support four distinct roles (NGO, donor, reviewer, administrator), each with documented permissions and a dedicated dashboard.",
    "enforce per-organisation data isolation such that users see only records belonging to their organisation, with administrators as the sole cross-organisation exception.",
    "support time-based one-time password (TOTP) two-factor authentication, with recovery codes issued at enrolment and verified at sign-in.",
    "support administrator-level enforcement of two-factor authentication on write actions, configurable as a hard gate when the organisation chooses to enforce it.",
    "throttle authentication attempts per source address and per email account, locking out abusive sources without locking out legitimate users on shared networks.",
    "log every authentication event into a tamper-evident audit chain so that the sequence of authentication activity is independently verifiable.",
    "allow users to change their interface language at any time, persisting the choice on their profile.",
])


# --- 6.2 Grant Lifecycle -----------------------------------------------
heading(doc, "6.2 Grant Lifecycle Management", 2)

add_callout(doc, "USE CASE",
    "David, a programme director at a donor foundation, wants to publish "
    "a new $2 million health-systems grant. He has two equally fast paths: "
    "type a two-line brief and let the system draft the full grant, or "
    "upload the existing brief document his team has drafted in Word and "
    "let the system extract it into the wizard. Either way he lands in the "
    "same editable six-step wizard with everything prefilled. He refines a "
    "few criteria, runs the burden check (which warns him that one document "
    "requirement is unrealistic for small NGOs and proposes an alternative), "
    "and publishes."
)

add_body(doc,
    "Grant creation shall be a wizard-driven experience with two equally "
    "supported entry points so donors are never blocked on \"start from "
    "scratch\" friction. Both entry points land the donor in the same "
    "six-step wizard, prefilled by AI; the donor refines and publishes."
)

add_diagram(doc, "two_grant_paths",
    "Figure 2 — Two paths to start a grant. Both deposit the donor into "
    "the same editable wizard with AI prefill.")

add_numbered_list(doc, [
    "represent a grant as a structured entity with title, description, sector tags, geographic scope, total funding, currency, deadline, eligibility requirements, evaluation criteria with weights, required document types with specific requirements, reporting requirements, and status.",
    "support a six-step grant publication wizard covering basic information, eligibility, criteria, document requirements, reporting configuration, and review-and-publish.",
    "offer two entry points into the wizard: (a) provide a short donor prompt and let the system draft a complete grant scaffold, or (b) upload an existing grant brief in PDF, DOCX, or plain text and let the system extract its content into the wizard.",
    "allow the donor to modify, add to, or remove anything the AI prefilled at any step of the wizard before publishing, treating AI output as a starting point rather than a final answer.",
    "track each grant through a documented lifecycle of draft, open, review, closed, and awarded, with state transitions audited.",
    "allow donors to update grants prior to publication and to amend reporting requirements post-award where contractually permitted.",
    "support live drafting visibility, so donors can see how many NGOs are actively preparing applications without identifying the NGOs themselves.",
    "support a grant-level question-and-answer thread allowing NGOs to ask clarifying questions and donors to publish responses visible to all applicants.",
])

heading(doc, "AI integration", 3)
add_bullet_list(doc, [
    ("Path A — Grant brief generation from prompt. ",
     "The system shall draft a complete grant scaffold — title, "
     "description, criteria with weights, eligibility, document "
     "requirements, reporting cadence, recommended deadline — from a "
     "short donor prompt, producing an editable starting point in "
     "seconds. Every field shall be modifiable in the wizard."),
    ("Path B — Grant import from uploaded document. ",
     "The system shall extract the structured content of an existing "
     "grant brief uploaded as PDF, DOCX, or plain text, mapping the "
     "extracted content into the wizard's six steps. The donor "
     "reviews each step, refines, adds, or removes content, and "
     "publishes. Where the source document is scanned or image-only, "
     "the system shall fall back to vision-based extraction rather "
     "than rejecting the upload."),
    ("Applicant burden critique. ",
     "Before publishing, the system shall analyse the grant for "
     "applicant burden, identifying vague criteria any answer could "
     "satisfy, criteria asking for evidence small NGOs cannot "
     "realistically produce, and document requirements likely to "
     "depress submission quality. It shall propose specific, less "
     "burdensome alternatives."),
    ("Median-NGO preview. ",
     "The system shall predict what the median qualifying NGO will "
     "submit against each criterion and rate how well the criterion "
     "will discriminate strong from weak applicants, allowing donors "
     "to tighten the rubric before opening submissions."),
    ("Reporting requirements auto-population. ",
     "Where the donor uploads a draft grant agreement (rather than a "
     "shorter brief), the system shall extract the reporting cadence, "
     "indicator definitions, narrative section template, and budget "
     "reporting format from the document and use them to populate the "
     "grant's reporting configuration."),
])

add_body(doc,
    "After the grant is published, the system shall surface it through "
    "the match engine to NGOs whose capacity passport, sector, "
    "geography, and track record align — the same end-to-end lifecycle "
    "illustrated below."
)

add_diagram(doc, "grant_lifecycle",
    "Figure 3 — The complete grant lifecycle from publication through "
    "compliance monitoring, with AI assistance (clay) and compliance "
    "checkpoints (sky) at every stage.")


# --- 6.3 Application Lifecycle -----------------------------------------
heading(doc, "6.3 Application Lifecycle Management", 2)

add_callout(doc, "USE CASE",
    "Fatima at Amani Foundation finds the maternal-health grant. The "
    "system has pre-scored Amani as an 87% match. She opens the "
    "application; the system has prefilled the eligibility section "
    "from her organisation profile, drafted initial responses to each "
    "criterion drawing on her capacity passport, prior applications, "
    "and her uploaded annual report, and flagged a single piece of "
    "missing evidence (last year's audited financials). She refines two "
    "responses, uploads the missing audit, and submits. Before "
    "submission the system runs a pre-flight check, surfaces no blockers, "
    "and confirms her application is ready."
)

add_numbered_list(doc, [
    "represent an application as a structured entity with grant reference, applicant organisation, per-criterion responses, eligibility responses, attached documents, total funding requested, status, and lifecycle timestamps.",
    "support a four-step application wizard covering eligibility check, proposal responses, document upload, and review-and-submit.",
    "track each application through a documented lifecycle of draft, submitted, under-review, scored, accepted, and rejected, with state transitions audited.",
    "allow NGOs to save draft applications and return to them later, preserving partial work and the AI assistance state.",
    "auto-save applicant input at regular intervals so that connection drops do not lose work.",
    "support reassignment of applications to alternate reviewers by administrators.",
    "allow donors to issue clarification requests to applicants, with structured responses returned in-band rather than via email.",
])

heading(doc, "AI integration", 3)
add_bullet_list(doc, [
    ("Application co-authoring. ",
     "When an NGO begins an application, the system shall generate a "
     "first draft of each response by drawing on the NGO's "
     "organisation profile, capacity passport, prior applications, "
     "uploaded documents, and organisational memory. Every draft "
     "claim shall be linked to its source so the applicant can verify "
     "and refine."),
    ("Pre-submission readiness check. ",
     "Before submission the system shall run a structured gap "
     "analysis identifying missing evidence per criterion, overclaims "
     "lacking supporting data, generic answers that any applicant "
     "could give, and weak passages where a one-click rewrite is "
     "available. Each finding shall be classified blocker, weak, or "
     "polish."),
    ("Compliance pre-emption. ",
     "Before submission the system shall scan the application for "
     "compliance risks — eligibility gaps, missing required documents, "
     "financial inconsistencies, narrative-versus-data contradictions — "
     "and propose specific fixes."),
    ("Match scoring. ",
     "The system shall compute an alignment score between each NGO and "
     "each open grant on eligibility, sector, geography, capacity, and "
     "track record, surfacing top strengths and top blockers so the NGO "
     "knows where to focus."),
    ("Pattern intelligence across applications. ",
     "The system shall identify cross-application patterns the donor "
     "would otherwise miss, including criteria that fail to discriminate, "
     "common applicant misunderstandings, and systematic gaps."),
])


# --- 6.4 Capacity Assessment & Passporting -----------------------------
heading(doc, "6.4 Capacity Assessment and Passporting", 2)

add_callout(doc, "USE CASE",
    "Aisha runs the Sahel Women's Network. Two years ago she completed "
    "the Kuja capacity framework. This month she applies to three "
    "grants: one accepting Kuja, one requiring UN-HACT, and one with a "
    "donor-defined custom framework. The system passports her existing "
    "assessment to all three: the Kuja grant accepts it as-is; the "
    "UN-HACT grant is prefilled from the mapped equivalents in her "
    "Kuja passport with provenance shown on each prefilled answer; the "
    "custom framework is prefilled where mappings exist and clearly "
    "marks the donor-specific questions she still needs to answer. "
    "Aisha finishes all three capacity sections in under an hour rather "
    "than re-doing twelve hours of work."
)

add_numbered_list(doc, [
    "support five embedded capacity assessment frameworks: Kuja Capacity Framework, STEP (Standard Tool for the Evaluation of Performance), UN-HACT (Harmonised Approach to Cash Transfers), CHS (Core Humanitarian Standard), and NUPAS (Non-U.S. Organisation Pre-Award Survey).",
    "represent each framework as a configurable collection of question groups, individual questions, weighting, scoring rubric, and required supporting evidence.",
    "treat each NGO's completed assessments as a capacity passport — a versioned, persistent record attached to the organisation profile and reusable across applications.",
    "allow donors to specify which framework or frameworks they accept for a given grant, including the option to require a single specific framework.",
    "allow donors to define a custom capacity framework tailored to their portfolio, composed from a question bank with configurable weights, evidence requirements, and scoring rubric.",
    "passport an NGO's existing capacity assessment to any new application by mapping equivalent questions across frameworks via maintained translation tables, prefilling matching responses, and clearly marking provenance on each prefilled answer (source assessment, framework, completion date).",
    "allow NGOs to override or augment prefilled responses before submission, with the override clearly distinguished from the prefilled baseline.",
    "score completed assessments automatically against the framework rubric and surface gaps where evidence is missing or weak.",
    "allow NGOs to update their passport at any time; subsequent prefills shall reflect the latest answers without disturbing previously submitted applications.",
    "surface a passport completion percentage on the NGO dashboard, indicating how application-ready the organisation is across the major frameworks.",
])

heading(doc, "AI integration", 3)
add_bullet_list(doc, [
    ("Cross-framework mapping. ",
     "Where no curated translation exists between a passported answer "
     "and a target framework question, the system shall use AI to "
     "propose a mapping and shall flag the proposal for human "
     "confirmation. Confirmed mappings shall enrich the translation "
     "table over time."),
    ("Drafted responses. ",
     "The system shall use AI to draft initial responses to "
     "assessment questions by drawing on prior assessments, uploaded "
     "supporting documents, and organisational memory, presenting "
     "each draft with traceable source attribution."),
    ("Gap prioritisation. ",
     "The system shall use AI to predict which capacity gaps are "
     "most likely to disqualify the NGO from a specific grant, "
     "helping them prioritise where to invest improvement effort."),
    ("Custom framework drafting. ",
     "When a donor defines a custom framework, the system shall "
     "propose a complete question bank, weights, evidence "
     "requirements, and scoring rubric derived from the grant's focus "
     "area, which the donor then refines."),
])

doc.add_paragraph()
add_diagram(doc, "capacity_passporting",
    "Figure 5 — Capacity assessment passporting. The NGO completes the "
    "work once; the system carries it forward across every donor's "
    "framework with traceable provenance.")


# --- 6.5 Document Management & Analysis -------------------------------
heading(doc, "6.5 Document Management and Real-Time Analysis", 2)

add_callout(doc, "USE CASE",
    "Peter is applying to a maternal-health grant and uploads his "
    "organisation's audited financial statements. Before the upload "
    "even completes, the system tells him what it's looking for: two "
    "years of audited data with auditor signature and a breakdown of "
    "grants received. As soon as extraction finishes, the dashboard "
    "shows a score of 78 percent — the auditor signature is on page 4 "
    "and the prior-year comparable is included, but the donor asked "
    "for three years and Peter only has two. The system suggests "
    "either uploading the missing year or, if unavailable, attaching "
    "a clarification note explaining why. Peter chooses the clarification "
    "path, types two sentences, and his score rises to 92 percent — all "
    "without leaving the page."
)

heading(doc, "Supported file types", 3)
add_body(doc,
    "The system shall accept the document formats real NGOs actually "
    "produce in the field — including the messy ones. Every supported "
    "format shall receive equivalent AI analysis; no format is "
    "second-class."
)

add_formatted_table(doc,
    ["Format", "Use", "Extraction"],
    [
        ["PDF (text)",
         "Audited financials, formal reports, signed agreements",
         "Direct text extraction with structure preservation"],
        ["PDF (scanned / image)",
         "Photographed registration certificates, hand-scanned audits, MoUs photographed in the field",
         "Vision-based extraction fallback — no silent failure"],
        ["DOCX",
         "Narrative reports, M&E plans, organisation profiles",
         "Native Word parser with table and list preservation"],
        ["XLSX",
         "Budget actuals, indicator data, beneficiary lists",
         "Sheet-by-sheet parsing; numeric values preserved"],
        ["Plain text (TXT)",
         "Quick notes, exported logs, simple narratives",
         "Direct read with encoding detection"],
        ["Images (JPG, PNG)",
         "Photographs of beneficiaries, before/after evidence, signed forms",
         "Vision-based content description and OCR where text is present"],
    ],
    col_widths=[1.4, 2.5, 2.5]
)

heading(doc, "Functional requirements", 3)
add_numbered_list(doc, [
    "accept document uploads in PDF (text and scanned), DOCX, XLSX, plain text, JPG, and PNG formats, up to a configured per-file size limit, with clear error responses on oversized or unsupported files.",
    "treat scanned PDFs and image-only documents as first-class inputs by falling back to vision-based extraction rather than rejecting low-quality field-collected evidence.",
    "extract text and structure from each uploaded document using a format-appropriate parser, preserving tables, lists, signatures, and numeric values where present.",
    "auto-detect the document type (audited financials, MoU, M&E plan, organisation profile, registration certificate, budget actuals, beneficiary list) from the document's content rather than relying solely on the user's classification.",
    "analyse each document against the donor's specific requirements for that document type, producing a real-time score, classified findings, and specific actionable recommendations.",
    "version uploaded documents per application or report so that re-uploading the same document type produces a new version with prior versions retained for audit.",
    "perform analysis asynchronously where extraction is non-trivial, returning the document immediately and surfacing the analysis result to the user as soon as it completes.",
    "track the extraction lifecycle of each document (pending, running, completed, failed) with timing and error code so administrators can diagnose stuck extractions.",
    "allow NGO uploaders to attach clarification notes to AI document findings so that contested or context-specific findings can be explained to reviewers without removing the AI signal.",
])

heading(doc, "AI integration — real-time support during upload", 3)
add_bullet_list(doc, [
    ("Pre-upload guidance. ",
     "Before the user uploads a document, the system shall display "
     "what the donor is asking for: the periods, signatures, sections, "
     "and structure that document type needs to satisfy the requirement. "
     "The NGO sees the success criteria before they spend time uploading."),
    ("Inline scoring as soon as extraction completes. ",
     "As soon as the document is parsed, the system shall display a "
     "score against the donor's specific requirements with classified "
     "findings (positive, neutral, concern) and a specific recommendation "
     "for each concern. The score shall update in place as the NGO "
     "uploads supplementary documents or adds clarification notes."),
    ("Donor-aware analysis. ",
     "The system shall analyse each document against the donor's "
     "stated requirements for that document type, not a generic "
     "completeness rubric. A financial statement is evaluated for the "
     "periods, signatures, and breakdowns this donor asked for; an MoU "
     "for the parties, dates, and scope this donor named."),
    ("One-click clarifications. ",
     "Where a finding is contested or context-specific (\"only one "
     "year of audit available because the organisation was founded in "
     "2024\"), the NGO shall be able to attach a clarification note "
     "in-place. The system shall incorporate the clarification into "
     "the displayed score without removing the underlying signal, so "
     "reviewers see both the finding and the explanation."),
    ("Reporting requirements extraction. ",
     "Where a donor uploads a draft grant agreement, the system shall "
     "extract the reporting cadence, indicators, narrative sections, "
     "and budget reporting format from the document and use them to "
     "populate the grant's reporting configuration in the wizard."),
])


# --- 6.6 AI-Powered Assistance ----------------------------------------
heading(doc, "6.6 AI-Powered Assistance Across the Workflow", 2)

add_callout(doc, "USE CASE",
    "Across every workspace, the system surfaces an action-oriented "
    "AI co-pilot. For an NGO programme officer it drafts narrative "
    "responses grounded in their evidence. For a donor it predicts how "
    "the median applicant will respond and warns where criteria are "
    "vague. For a reviewer it produces a one-screen summary with "
    "evidence cited per criterion and a draft rationale. For an "
    "administrator it explains compliance findings in plain language. "
    "In every case the assistance is grounded in the actual data the "
    "user is looking at, every claim is sourced, and every output is "
    "an editable starting point rather than a fait accompli."
)

heading(doc, "6.6.1 Design principles for AI use", 3)
add_body(doc,
    "The system applies AI assistance under a small number of explicit "
    "design principles, which together distinguish a working co-pilot "
    "from a surface-level chatbot. The system shall enforce these "
    "principles for every AI feature it ships."
)
add_bullet_list(doc, [
    ("Grounded, not generic. ",
     "AI output shall draw on the user's actual data — their "
     "organisation profile, capacity passport, prior submissions, "
     "uploaded documents, the grant they are looking at — rather than "
     "producing generic prose. Where the system cannot find evidence, "
     "it shall say so rather than invent."),
    ("Action-oriented. ",
     "AI output shall be an editable starting point or a structured "
     "decision aid (gap list, evidence list, rewrite, score with "
     "rationale) rather than a paragraph the user must re-read and "
     "interpret."),
    ("Traceable. ",
     "Every AI claim shall be linked to its source. Reviewers and "
     "applicants alike shall be able to see which document, which "
     "memory item, which prior application a claim drew on."),
    ("Honest about uncertainty. ",
     "AI output shall surface its own confidence — high, medium, low — "
     "and shall fall back to deterministic templates or rule-based "
     "logic when the AI service is unavailable or returns malformed "
     "output. Users shall always see a usable interface."),
    ("Multilingual and role-aware. ",
     "AI output shall be produced in the user's preferred language with "
     "tone appropriate to their role: warm, coaching, supportive for "
     "NGO-facing surfaces; precise, decision-oriented for donor-facing "
     "and reviewer-facing surfaces."),
    ("Auditable. ",
     "Every AI call shall produce a telemetry record (endpoint, model, "
     "input tokens, output tokens, latency, success or failure, user, "
     "language, role) retrievable by administrators."),
    ("Cost-aware. ",
     "Heavy AI work shall execute asynchronously so that the user's "
     "request returns immediately and the model is given the time it "
     "needs without blocking the user interface. Costs shall be "
     "forecastable on a thirty-day horizon."),
])

heading(doc, "6.6.2 AI surfaces the system shall provide", 3)
add_formatted_table(doc,
    ["Surface", "Workspace", "What the system does"],
    [
        ["Match scoring",
         "NGO and donor",
         "Score every NGO-grant pair on eligibility, sector, geography, "
         "capacity, and track record; surface top strength and top "
         "blocker."],
        ["Application co-authoring",
         "NGO",
         "Draft a complete first-pass application by drawing on the "
         "NGO's passport, prior submissions, and uploaded documents; "
         "track provenance on every claim."],
        ["Submission readiness",
         "NGO",
         "Run a pre-submission gap analysis identifying missing "
         "evidence, overclaims, generic answers, and weak passages with "
         "one-click rewrites."],
        ["Compliance pre-empt",
         "NGO",
         "Scan the application for compliance risks before submission "
         "(eligibility gaps, missing documents, financial inconsistencies, "
         "narrative-data contradictions); propose specific fixes."],
        ["Reviewer summary",
         "Reviewer",
         "Generate a one-screen summary per application with evidence "
         "for and against each criterion, decision-changers, "
         "comparable signal across the cohort, and a draft rationale."],
        ["Evidence extraction",
         "Reviewer",
         "Pull verbatim supporting, contradicting, and neutral quotes "
         "from the application against each rubric criterion so the "
         "reviewer cites evidence rather than writing from memory."],
        ["Criteria suggestion",
         "Reviewer and donor",
         "When a grant has no evaluation criteria defined, propose a "
         "calibrated rubric with labels, descriptions, weights summing "
         "to 100, and per-criterion rationale."],
        ["Grant brief generator",
         "Donor",
         "Draft a complete grant scaffold from a two-line donor prompt: "
         "title, description, criteria, eligibility, document and "
         "reporting requirements, recommended deadline."],
        ["Median-NGO preview",
         "Donor",
         "Predict how the median qualifying NGO will respond per "
         "criterion and rate criterion discrimination strength, so the "
         "donor can tighten the rubric pre-publish."],
        ["Burden estimator",
         "Donor",
         "Critique a grant draft for applicant burden; surface vague "
         "criteria, over-demanding evidence asks, and propose "
         "simplifications."],
        ["Report co-authoring",
         "NGO",
         "Draft a progress report narrative from prior reports, the "
         "current period's data, and uploaded evidence; surface gaps."],
        ["Report pre-flight",
         "NGO",
         "Predict donor concerns about the report before submission; "
         "identify vague claims, unexplained budget variance, missing "
         "evidence, with the donor concern each fix resolves."],
        ["Compliance explanation",
         "Donor and administrator",
         "Translate verification and sanctions findings into plain "
         "language with concrete follow-up actions."],
        ["Portfolio insights",
         "Donor",
         "Summarise the donor's portfolio with headline, anomalies, "
         "and next decisions tied to specific applications and grants."],
        ["Cross-grant patterns",
         "Donor",
         "Identify patterns across declined applications — common "
         "shortfalls, systematic misalignment — that suggest rubric "
         "tightening or eligibility refinement."],
        ["Insight narration",
         "All",
         "Generate plain-language captions for dashboard charts so "
         "numbers carry their meaning to non-analyst readers."],
        ["Page-aware suggestions",
         "All",
         "Surface a small list of context-appropriate next actions on "
         "every workspace page, scoped to the user's role and current "
         "scope (a grant, an application, an organisation)."],
        ["Ask Kuja",
         "All",
         "Conversational agent that answers operational questions by "
         "querying a registry of read-only data tools, role-checked and "
         "organisation-scoped, with three-step maximum iteration."],
    ],
    col_widths=[1.6, 1.1, 3.7]
)

doc.add_paragraph()
add_diagram(doc, "ai_integration_map",
    "Figure 4 — AI integration map across the three primary user "
    "workspaces. Every surface is grounded, traceable, action-oriented, "
    "and falls back gracefully when AI is unavailable.")


# --- 6.7 Compliance and sanctions --------------------------------------
heading(doc, "6.7 Compliance and Sanctions Screening", 2)

add_callout(doc, "USE CASE",
    "When an NGO is added to the platform, the system screens the "
    "organisation and its named principals against UN, OFAC, and EU "
    "sanctions lists. The screening runs again automatically when the "
    "organisation applies for a grant, when a grant is awarded, and on "
    "a scheduled cadence for organisations with active grants. A "
    "donor reviewing an application sees a clear status: clear, "
    "flagged, or pending — and on flag, a plain-language explanation "
    "of what was matched and the suggested follow-up."
)

add_numbered_list(doc, [
    "screen every organisation and its named principals against UN, OFAC, and EU sanctions lists at onboarding, at application submission, at award, and on a recurring schedule for organisations with active grants.",
    "use a live screening service as the primary feed and shall gracefully fall back to direct downloads of the underlying UN, OFAC, and EU lists when the primary service is unavailable, so that screening is never silently skipped.",
    "compute a structured screening result per organisation including overall status (clear, flagged, error), check-type breakdown, and matched entity details where applicable.",
    "persist screening results with timestamps for audit, and shall raise a notification to administrators when a new flag is detected on an organisation with active grants.",
    "translate technical screening findings into plain-language explanations with suggested follow-up actions.",
    "allow administrators to trigger an immediate rescreening run across all organisations with active grants.",
    "deduplicate against persona matches using country, name variants, and date-of-birth when supplied, to reduce false positives on common names.",
])


# --- 6.8 Registry verification ------------------------------------------
heading(doc, "6.8 Government Registry Verification", 2)

add_callout(doc, "USE CASE",
    "An NGO claims to be registered in Kenya under registration number "
    "OP.218/051/2017/0/3001. The system queries Kenya's NGO Coordination "
    "Board registry, verifies the registration is current and the legal "
    "name matches, and stamps the organisation profile with a verified "
    "badge dated today. If the registry returns an inconsistency, the "
    "system raises a verification finding for administrative review "
    "rather than silently passing."
)

add_numbered_list(doc, [
    "support live registry verification for organisations claiming registration in Kenya, Nigeria, South Africa, Uganda, Tanzania, Somalia, and Ethiopia.",
    "compare the registered legal name, status, and (where available) principal officer names against the values claimed by the organisation on the platform.",
    "stamp organisation profiles with a verification status (verified, unverified, AI-reviewed, expired) and timestamp.",
    "raise a verification finding for administrative review where registry data and claimed data diverge, rather than treating divergence as a soft warning the user can ignore.",
    "support periodic re-verification for organisations with active grants, with configurable expiry windows for previous verification stamps.",
    "expose a verification dashboard for administrators showing organisations grouped by verification status and expiry urgency.",
])


# --- 6.9 Reporting -----------------------------------------------------
heading(doc, "6.9 Reporting and Outcome Tracking", 2)

add_callout(doc, "USE CASE",
    "Six months into a grant, Amani Foundation's first quarterly report "
    "is due. The system reminds Fatima two weeks in advance, drafts the "
    "narrative from the prior quarter's report and the indicator data "
    "Amani has logged, and runs a pre-flight check from the donor's "
    "perspective. The pre-flight flags two unexplained budget variances "
    "and an impact claim with no supporting data. Fatima addresses both "
    "and submits."
)

add_numbered_list(doc, [
    "represent a progress report as a structured entity with grant reference, reporting period, report type, narrative sections, indicator values, budget actuals, milestone progress, uploaded evidence, and lifecycle status.",
    "track reports through a documented lifecycle of draft, submitted, accepted, and revision-requested.",
    "support the cadence and structure each donor specifies — quarterly, semi-annual, annual, custom — and the narrative and indicator template each donor specifies.",
    "remind reporters in advance of due dates, escalating reminders as deadlines approach.",
    "allow donors to request revisions with structured feedback, returning the report to the NGO with the donor's concerns clearly attached to specific sections.",
    "support evidence uploads attached to reports (photographs, beneficiary lists, financial documents, third-party verifications), analysed against the relevant requirement type.",
    "compute a compliance score for each accepted report combining narrative completeness, indicator coverage, budget reconciliation, and evidence sufficiency.",
])

heading(doc, "AI integration", 3)
add_bullet_list(doc, [
    ("Report co-authoring. ",
     "The system shall draft a progress report narrative by drawing "
     "on prior reports, the indicator data the NGO has already "
     "captured, and uploaded evidence, producing a first-pass "
     "narrative the reporter refines."),
    ("Donor-perspective pre-flight. ",
     "Before submission the system shall predict donor concerns about "
     "the report — vague claims without data, unexplained budget "
     "variance, missing evidence — and surface each fix mapped to the "
     "donor concern it resolves."),
])


# --- 6.10 Compliance — Flagship Section --------------------------------
heading(doc, "6.10 Compliance and Risk — Simplified for Both Sides", 2)

add_body(doc,
    "Compliance is the most expensive friction in the grant relationship "
    "today. For NGOs, it means tracking obligations across donors, "
    "interpreting reporting requirements written for someone else, "
    "drafting reports under deadline, and proving evidence under audit. "
    "For donors, it means trusting that funds are well-spent, that "
    "risks are surfaced early, and that the portfolio's compliance "
    "health is legible at a glance rather than hidden in PDFs."
)

add_body(doc,
    "The system shall simplify compliance and risk for both sides of "
    "the relationship — not by adding more rules but by carrying the "
    "work for the NGO and presenting the result, pre-assessed and "
    "scored, to the donor. This is the platform's central commitment."
)

add_diagram(doc, "ngo_compliance_journey",
    "Figure 6 — How the system supports the NGO at every stage of the "
    "compliance journey, from application pre-flight to ongoing "
    "delivery, reporting, and health monitoring.")

heading(doc, "6.10.1 For NGOs — AI as a compliance co-pilot", 3)

add_callout(doc, "USE CASE",
    "Fatima's first quarterly report is due in three weeks. The system "
    "has already extracted the donor's reporting requirements from the "
    "grant agreement (cadence, indicators, narrative sections, budget "
    "format), added each one to her compliance calendar, and surfaced "
    "the upcoming due date on her dashboard. Two weeks out, the system "
    "drafts a first-pass narrative from her prior reports and the "
    "indicator data she has logged, flags one missing evidence file, "
    "and sends her a gentle web push reminder. Fatima opens the report, "
    "refines two paragraphs, uploads the missing audit, runs the "
    "donor pre-flight (which scores her at 87 percent and surfaces one "
    "vague claim with a sharper rewrite), accepts the rewrite, and "
    "submits. The whole report took 90 minutes; the previous donor's "
    "equivalent report would have taken twelve."
)

add_body(doc,
    "The system shall act as the NGO's compliance co-pilot across the "
    "entire grant — from the moment the requirements are agreed to the "
    "moment the final report is accepted. The system shall:"
)

add_numbered_list(doc, [
    "extract structured compliance obligations from the grant agreement and any supporting documents the donor uploaded — reporting cadence, indicator definitions, narrative sections, budget format, document-of-record requirements — and present them to the NGO as a structured obligation list, not a PDF to re-read.",
    "translate donor-specific language into clear NGO-side action items: \"submit a Q1 progress report by 30 June including indicator data on visits completed and CHW retention, with a budget reconciliation showing variance against approved\".",
    "maintain a compliance calendar per NGO showing every upcoming obligation across every active grant, with date, type, status, and grant attribution.",
    "send proactive reminders ahead of each obligation, escalating as deadlines approach — in-app notifications, email, and (when subscribed) web push to the NGO's mobile devices.",
    "draft a first-pass narrative for each progress report by drawing on prior reports, the indicator data the NGO has captured, the documents they have uploaded, and their organisational memory.",
    "prompt the NGO to upload the right evidence at the right time: \"your indicator data shows 88 percent CHW retention this quarter — attach the training attendance records that support this number\".",
    "score every uploaded document against the donor's specific requirements in real time, with classified findings and one-click clarifications where context is needed.",
    "run a donor-perspective pre-flight check before submission, predicting the specific concerns the donor will raise (vague claims without data, unexplained budget variance, missing evidence) and offering specific fixes for each — each fix mapped to the donor concern it resolves.",
    "guide the NGO through donor revision requests by translating the donor's structured feedback into specific section-level actions the NGO can take, with AI-drafted updates the NGO can accept, edit, or override.",
    "score the report as a whole and the NGO's overall compliance health, so the NGO sees what the donor will see before the donor sees it.",
])

heading(doc, "6.10.2 For donors — pre-assessed, scored, monitored", 3)

add_callout(doc, "USE CASE",
    "Sarah, a donor managing twelve active grants, opens her compliance "
    "dashboard. Eleven grants show as on-track. One has a slips badge: "
    "\"Slips in 14 days\", meaning the system's trajectory model predicts "
    "this grant will drop below the at-risk threshold within two weeks "
    "if nothing changes. Sarah opens the grant, sees that report "
    "completion is trending down and one report is approaching overdue. "
    "She raises a risk, assigns it to her programme officer with a due "
    "date, and sends a structured comment to the NGO. The NGO is "
    "notified by web push and responds in-band. The whole exchange "
    "takes five minutes and is preserved in the grant's audit chain."
)

add_body(doc,
    "The system shall present every donor with a portfolio-level "
    "compliance command centre. Reports arrive pre-assessed and scored; "
    "compliance health is visible at a glance; risks are tracked with "
    "owners and due dates; trajectory forecasts surface degradation "
    "before grants slip. The donor shall not be surprised."
)

add_numbered_list(doc, [
    "deliver every progress report to the donor with an AI-pre-assessment already attached: section-level summary, indicator validation, budget reconciliation status, evidence coverage, and a composite quality score.",
    "compute a four-pillar compliance health score for every active grant combining completion (30 percent weight), timeliness (30 percent), workflow status (20 percent), and importance (20 percent).",
    "classify each grant as on-track, at-risk, or high-risk based on the composite score, with each pillar inspectable so the donor can see why the score is what it is.",
    "write a daily compliance health snapshot per active grant so that change over time is observable.",
    "compute a linear-regression forecast over the trailing 60 days of snapshots and project when (if ever) the grant is likely to drop below the at-risk threshold.",
    "surface a slips badge on grants forecast to slip within 30 days, colour-coded by urgency (within 7 days, within 14 days, within 30 days).",
    "maintain a risk register as a structured entity per organisation, application, or grant, with kind, severity, title, description, status, response, owner, due date, source (AI detection or human-raised), and lifecycle.",
    "track each risk through a documented lifecycle of open, mitigating, mitigated, accepted, and dismissed.",
    "surface risks awaiting donor response on the donor dashboard, prioritised by severity and age.",
    "summarise the donor's portfolio with an AI-generated insights panel: headline, anomalies (grants drifting, organisations under-performing), and the next decisions the donor owes.",
])

heading(doc, "6.10.3 Shared mechanics", 3)
add_bullet_list(doc, [
    ("Plain-language explanations of every finding. ",
     "Whether the audience is an NGO programme officer or a donor "
     "programme manager, the system shall translate technical "
     "findings (sanctions matches, document gaps, indicator "
     "variances) into plain language with specific follow-up actions."),
    ("Structured collaboration in-band. ",
     "Donor and NGO shall collaborate on every clarification, "
     "revision request, and risk response through threaded comments "
     "with @-mentions and (where subscribed) web push, rather than "
     "through email threads that escape the record."),
    ("Tamper-evident audit chain. ",
     "Every compliance-relevant event — sanctions screening result, "
     "verification stamp, document analysis, report submission, "
     "decision, risk lifecycle transition — shall be recorded in a "
     "hash-chained audit log so that the sequence of events is "
     "independently verifiable."),
    ("AI-driven risk detection. ",
     "The system shall use AI to detect risks during compliance "
     "screening, document analysis, application review, and report "
     "pre-flight, raising risk register entries with proposed "
     "severity and rationale rather than only surfacing them to the "
     "user who happens to be looking."),
    ("Health narrative. ",
     "The system shall layer a one- or two-sentence plain-language "
     "explanation on top of the rule-based compliance score, "
     "summarising why the score is what it is and what the donor "
     "should attend to first."),
])


# --- 6.11 Collaboration & Notifications --------------------------------
heading(doc, "6.11 Collaboration and Notifications", 2)

add_callout(doc, "USE CASE",
    "A donor reviewer wants the NGO to clarify a discrepancy between "
    "their reported beneficiaries and their financial statement. She "
    "writes a comment on the application detail page, @-mentioning "
    "the NGO's programme officer. The officer receives an in-app "
    "notification and, if subscribed, a web push notification on her "
    "phone. She replies in the thread; the conversation is preserved "
    "with the application rather than scattered across email."
)

add_numbered_list(doc, [
    "support threaded comments attached to any of: applications, grants, reports, risks, and organisations.",
    "resolve @mentions by email local-part across visible users, preferring same-organisation members when the local-part is ambiguous.",
    "raise an in-app notification to each mentioned user, with a deep link to the source comment.",
    "deliver a web push notification to each mentioned user's subscribed devices when web push is configured for the deployment, gracefully degrading to in-app-only when web push is not yet configured.",
    "allow comment authors to edit and delete their own comments, with edit timestamps preserved for audit.",
    "support administrator deletion of any comment for content moderation, with a documented record of the deletion.",
    "surface unread notification counts on the user's navigation, with a notifications page listing recent notifications grouped by kind.",
])


# --- 6.12 Organisational Memory ----------------------------------------
heading(doc, "6.12 Organisational Memory and Provenance", 2)

add_callout(doc, "USE CASE",
    "Amani has applied to nine grants over the past two years. Across "
    "those applications, the team has answered the same questions about "
    "their theory of change, their MEL approach, and their financial "
    "controls many times. The system has stored each consistent answer "
    "in the organisation's memory. When Fatima starts her next "
    "application, the system surfaces 'Drew on 11 facts from your "
    "memory' and shows exactly which facts informed which response. "
    "Fatima can review, accept, edit, or remove the underlying memory "
    "items, and her changes ripple to future applications."
)

add_numbered_list(doc, [
    "maintain an organisational memory store per organisation containing structured facts — beneficiary counts, geographic scope, key partnerships, methodology statements, financial controls — derived from completed applications, reports, and assessments.",
    "tag each memory item with its source (which application, which report), its kind (claim, policy, indicator, partnership), and a usage counter incremented each time the item informs a future draft.",
    "rank memory items by usage and recency when surfacing them to AI drafting, so that the most-validated facts shape new outputs.",
    "allow organisation users to view, edit, and remove memory items, with edits and removals taking effect on subsequent drafts but not retroactively altering historical submissions.",
    "track provenance on every AI-drafted claim, linking the claim to the specific source kind (organisation profile, prior application, prior report, uploaded document, organisational memory, or AI-general where no source applies) and recording the locator and confidence band.",
    "expose the provenance ledger so reviewers can inspect the sources behind a claim and applicants can see exactly what the system drew on.",
])


# --- 6.13 Saved searches -----------------------------------------------
heading(doc, "6.13 Saved Searches and Personalisation", 2)

add_callout(doc, "USE CASE",
    "A donor reviewer often filters open applications for the maternal "
    "health portfolio in East Africa, sorted by AI score. She saves "
    "this view as 'East Africa maternal health, top scoring'. Tomorrow "
    "she restores it with one click and reorders her saved views so "
    "this one is at the top."
)

add_numbered_list(doc, [
    "allow users to save the current filter and sort state of a list view as a named saved search, scoped to that user.",
    "support saved searches across grants, applications, reports, organisations, reviews, and risks.",
    "allow users to reorder their saved searches, delete them, and rename them.",
    "preserve the captured filter shape exactly, restoring the same view on next click without re-prompting the user.",
])


# --- 6.14 Administration -----------------------------------------------
heading(doc, "6.14 Administration and System Health", 2)

add_callout(doc, "USE CASE",
    "An administrator opens the system-health dashboard on Monday "
    "morning. Every row is green: API keys configured, scheduled-job "
    "secret set, AI failure rate at 0.4 percent over 24 hours, "
    "registry verification queue empty, rate-limit backend connected. "
    "She glances at the daily AI surface health probe — all seven "
    "flagship AI surfaces passed against synthetic fixtures overnight. "
    "She moves on to her day."
)

add_numbered_list(doc, [
    "provide a system health surface for administrators reporting on the readiness and operability of every critical subsystem: AI provider configuration, sanctions provider configuration, scheduled-job authentication, AI failure rate over the last 24 hours, document extraction queue depth, stuck extractions, orphaned risks, rate-limit backend, and AI cost budget status.",
    "provide an AI cost telemetry surface showing daily input and output token spend per endpoint, with a 30-day projection compared to a configurable budget threshold.",
    "provide a demo-readiness scanner surface that identifies categories of sparse data likely to make AI surfaces read as broken (grants without criteria, applications without documents or with empty responses, reports missing submitted timestamps, organisations missing profiles, administrators without two-factor enrolled).",
    "provide a flagship AI surface health probe that exercises every primary AI surface against synthetic fixtures on a daily cadence, reporting pass, fail, or skipped for each surface with latency and source. Failed surfaces shall raise an administrator notification.",
    "provide a configurable audit retention window, with the current window editable by administrators and the next prune count visible.",
    "provide a failed-login surface grouping recent failed authentication attempts by source and target, with rolling counts at 24-hour, 7-day, and 30-day windows.",
    "provide a feature flag administration surface exposing every flag with its current state, default state, and per-scope overrides.",
    "expose a synthesised OpenAPI document for the platform's HTTP surface and a rendered HTML viewer, for partner integration reference.",
])


# --- 6.15 Internationalisation -----------------------------------------
heading(doc, "6.15 Internationalisation and Localisation", 2)

add_callout(doc, "USE CASE",
    "An NGO programme officer in Mogadishu signs in. The interface, "
    "every AI-generated draft, every notification, and every error "
    "message arrive in Somali. When she @-mentions a donor reviewer in "
    "Geneva, the donor sees the comment in English with the original "
    "Somali available on hover. Neither party has to translate manually."
)

add_numbered_list(doc, [
    "support six interface languages: English, Arabic, French, Swahili, Somali, and Spanish — at parity across every user-facing surface.",
    "right-to-left layout shall be correctly applied for Arabic.",
    "support per-user language preference, persisted on the user profile and applied immediately on language change.",
    "produce AI output in the user's preferred language using a native-language directive applied at generation time, so that the model writes in the target language rather than translating from English.",
    "apply language-specific register adjustments (formality, idiom, numeric conventions) so that AI output reads naturally in each language rather than carrying a translated-from-English feel.",
    "apply role-aware tone: warm, coaching, and supportive for NGO-facing surfaces; precise and decision-oriented for donor-facing and reviewer-facing surfaces; clinical for administrator-facing surfaces.",
    "be available for native-speaker review of translation quality, with priority namespaces identified for review and a workflow for landing approved revisions.",
])


# --- 6.16 Field Operations --------------------------------------------
heading(doc, "6.16 Field Operations: Mobile and Offline Work", 2)

add_callout(doc, "USE CASE",
    "A field officer for a women's protection NGO is documenting a "
    "beneficiary intake in a rural area with no mobile signal. She "
    "opens the platform on her phone (the same login she uses at the "
    "office), drafts the report narrative, attaches three photographs "
    "of consent forms taken with her camera, and saves. The platform "
    "shows the work as queued for sync. Twenty minutes later, back in "
    "the town centre, her phone reconnects; the platform syncs the "
    "draft and uploads the photographs automatically in the "
    "background. When she opens her laptop that evening, the draft is "
    "already there, waiting for her to refine and submit."
)

add_body(doc,
    "The Global South operating environment includes intermittent "
    "connectivity, limited bandwidth, and a primary device that is "
    "typically a mobile phone. The system shall treat this as a "
    "first-class scenario rather than a degraded mode."
)

add_numbered_list(doc, [
    "be delivered as a progressive web application installable on a mobile device's home screen, with the same authentication, workspaces, and feature scope as the desktop experience.",
    "present responsive layouts tuned for mobile portrait, mobile landscape, tablet, and desktop viewports, so every workspace is fully usable on a phone rather than requiring a separate mobile app.",
    "operate in offline-first mode: users shall be able to draft applications, complete capacity assessments, write progress reports, fill out indicator data, and capture supporting evidence without an active internet connection.",
    "auto-save drafts continuously to local device storage during editing so that interrupted sessions, browser crashes, or sudden connectivity loss never lose work.",
    "queue file uploads (photographs, documents) for background sync when the device is offline, transmitting them automatically once connectivity is restored without the user re-uploading.",
    "synchronise queued edits and uploads on reconnection with appropriate conflict handling — if two devices edited the same draft offline, the system shall surface the divergence to the user rather than silently overwriting one side.",
    "operate acceptably over low-bandwidth and high-latency connections by minimising payload sizes, deferring heavy AI work to asynchronous background calls, and progressively enhancing rich features so that the core flow works on a slow 3G connection.",
    "treat photographs taken on a mobile camera as first-class evidence uploads, applying the same vision-based extraction and donor-aware analysis available to other document types.",
    "preserve the user's session across device restarts and prolonged disconnection, so a field officer returning to network after a day in a remote area can resume exactly where she left off.",
])

doc.add_page_break()


# ===========================================================================
# 7. NON-FUNCTIONAL REQUIREMENTS
# ===========================================================================

heading(doc, "7. Non-Functional Requirements", 1)

heading(doc, "7.1 Performance and Scalability", 2)
add_numbered_list(doc, [
    "return user-interactive responses (page navigations, list views, simple writes) within 500 milliseconds at the 95th percentile under normal load.",
    "execute heavy AI calls (application drafting, report drafting, reviewer summary, grant brief generation) asynchronously, returning the user's request within 50 milliseconds with a job identifier and surfacing the result via a polling endpoint once ready.",
    "cap the maximum runtime of any single AI call at 300 seconds and shall classify each AI call by expected work envelope (light up to 60 seconds, medium up to 120 seconds, heavy up to 240 seconds) so that frontends can present appropriate progress indication.",
    "cap concurrent in-flight AI calls per authenticated user, gracefully falling back to a template path when the cap is exceeded so that one client cannot saturate the worker pool.",
    "support horizontal scaling of the application tier such that adding workers proportionally increases serving capacity.",
])

heading(doc, "7.2 Reliability and Availability", 2)
add_numbered_list(doc, [
    "target 99.5 percent monthly availability for the user-facing surface, with documented degraded-mode behaviour for each dependent service (AI provider, sanctions provider, registry providers, push provider).",
    "fall back to deterministic rule-based or template output when an AI service is unavailable, returning a meaningful response in every case rather than an empty error.",
    "preserve user input across connection drops via automatic draft saving on application and report editing surfaces.",
    "retry transient infrastructure failures on idempotent read requests with exponential backoff, masking momentary worker restarts from the user.",
    "isolate scheduled-job state across workers via a shared backend so that scheduled work (compliance snapshots, rescreening, audit prune) executes exactly once per cycle rather than once per worker.",
])

heading(doc, "7.3 Security and Privacy", 2)
add_numbered_list(doc, [
    "encrypt all data in transit via TLS, and at rest via the database provider's default encryption.",
    "hash user credentials with a per-credential salt and a vetted hashing function with adjustable work factor.",
    "scope every data access by the requesting user's organisation, with cross-organisation access restricted to administrators and logged.",
    "maintain a tamper-evident audit chain for security-relevant events (authentication, authorisation, permission changes, sanctions findings, administrative actions), with every entry cryptographically linked to its predecessor so that retroactive tampering is detectable.",
    "support two-factor authentication via time-based one-time passwords with recovery codes, and shall support administrator-level hard enforcement of two-factor authentication for write actions.",
    "rate-limit authentication attempts per source address and per email account, with per-account lockout durations distinct from per-address rate limits to distinguish targeted brute force from shared-NAT noise.",
    "support a right-to-be-forgotten workflow allowing administrators to expunge personal data for a specific user while preserving the audit chain and any data the contract or programme record requires.",
    "apply a strict Content Security Policy, with no third-party script origins and explicit denials of object embedding, frame ancestors, and mixed-content loading.",
    "scope sensitive operations behind per-operation rate-limit policies separate from the global throttle, with policies named by operation class so that limits can be tuned per surface without touching the rest of the system.",
])

heading(doc, "7.4 Observability and Operability", 2)
add_numbered_list(doc, [
    "emit structured logs for every request including request identifier, route, user identifier, duration, and outcome, suitable for ingestion by a log analytics platform.",
    "emit a telemetry record for every AI call including endpoint, model, input tokens, output tokens, latency, success or failure, user, language, and role.",
    "compute and surface AI cost telemetry on a daily basis with a 30-day projection compared to a configurable budget threshold.",
    "execute scheduled jobs (compliance snapshots, audit prune, sanctions rescreening, AI surface health probe, demo readiness scan) on a documented cadence, with their last-run results visible to administrators.",
    "support graceful operation when optional infrastructure (shared rate-limit backend, push provider, primary sanctions feed) is not configured, surfacing the missing-but-recommended state in the system-health view without raising errors to end users.",
    "expose a synthesised OpenAPI document and a rendered HTML viewer for partner integration reference.",
])

heading(doc, "7.5 Usability, Accessibility, and Inclusivity", 2)
add_body(doc,
    "Usability is treated as a quality attribute, not a feature. The "
    "system shall be optimised for users whose primary device is a "
    "phone, whose connectivity is intermittent, whose first language "
    "may not be English, and whose time is scarce. Every surface shall "
    "answer the question \"what should I do next?\" before the user "
    "has to ask."
)
add_numbered_list(doc, [
    "meet WCAG 2.1 Level AA conformance for keyboard navigation, focus visibility, contrast, and screen-reader compatibility on every user-facing surface.",
    "present action-driven workspaces: every dashboard surface shall lead with what the user owes attention to (to-dos, deadlines, blockers, opportunities) rather than asking the user to navigate to discover their next action.",
    "minimise time-to-first-action: a user signing in shall be one click away from their highest-priority pending item.",
    "support touch interaction with mobile-portrait layouts on every primary workspace, treating mobile as a first-class delivery surface rather than a degraded fallback.",
    "operate offline-first: critical user actions (draft applications, complete assessments, write reports, capture evidence, save indicator data) shall function without an active connection and shall sync automatically on reconnection.",
    "operate acceptably over low-bandwidth and high-latency connections by deferring heavy AI work to asynchronous background calls and progressively enhancing rich features.",
    "auto-save user input continuously to local storage so that connectivity loss, browser crash, or device restart never causes loss of work.",
    "treat scanned and image-only documents as first-class inputs, falling back to vision-based extraction rather than rejecting field-collected evidence.",
    "render in right-to-left layout for Arabic without manual configuration on the part of the user.",
    "show clear status feedback for every long-running operation (document analysis, AI report drafting, background upload sync) so the user never wonders whether the system is working.",
    "make every AI output editable: AI is a starting point, never the final word.",
])

doc.add_page_break()


# ===========================================================================
# 8. USE CASES
# ===========================================================================

heading(doc, "8. Use Cases", 1)

add_body(doc,
    "This chapter presents end-to-end use cases connecting the "
    "functional requirements into the workflows the platform shall "
    "support. Each use case is written as a sequence of events with "
    "actors and pre- and post-conditions; the requirements in chapter 6 "
    "constitute the basis for each step."
)


def use_case(doc, uid, title, actor, pre, steps, post):
    heading(doc, f"{uid}: {title}", 2)
    add_formatted_table(doc,
        ["Actor", "Pre-conditions", "Post-conditions"],
        [[actor, pre, post]],
        col_widths=[1.0, 2.7, 2.7]
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Main flow")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{i}.  ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run_n = p.add_run(step)
        run_n.font.size = Pt(10)
        run_n.font.name = "Calibri"
    doc.add_paragraph()


use_case(doc, "UC-001", "Donor publishes a new grant",
    "Donor",
    "Donor user is signed in and belongs to an organisation authorised to publish grants.",
    [
        "Donor opens the grant creation wizard and provides a two-line prompt describing intent.",
        "The system drafts a complete grant scaffold (title, description, criteria with weights, eligibility, document requirements, reporting configuration, recommended deadline).",
        "The system runs an applicant-burden critique on the draft and flags vague criteria and over-demanding evidence asks.",
        "The system runs a median-NGO preview and identifies criteria that may not discriminate well across applicants.",
        "Donor refines the rubric, eligibility, and document requirements based on the system's feedback.",
        "Donor publishes the grant; the system transitions it from draft to open and notifies matching NGOs.",
    ],
    "Grant is open and discoverable; eligible NGOs are notified of the match.",
)

use_case(doc, "UC-002", "NGO completes capacity passport",
    "NGO",
    "NGO user is signed in with a registered organisation profile.",
    [
        "NGO opens the capacity assessment workspace and selects the Kuja framework as the baseline.",
        "The system drafts initial responses for each question by drawing on the organisation profile, uploaded documents, and organisational memory; each draft is shown with its source.",
        "NGO refines responses, uploads supporting evidence, and submits the assessment.",
        "The system scores the assessment against the framework rubric and surfaces gaps where evidence is missing or weak.",
        "The completed assessment is stored as part of the NGO's capacity passport, attached to the organisation profile.",
    ],
    "NGO has a versioned capacity passport that can be passported to future applications.",
)

use_case(doc, "UC-003", "NGO passports the capacity assessment to a new framework",
    "NGO",
    "NGO has a completed Kuja capacity passport. Applies to a grant requiring a different framework.",
    [
        "NGO opens the application and reaches the capacity section.",
        "The system detects the target framework and prefills responses from the passport using maintained translation tables.",
        "Where no curated mapping exists the system uses AI to propose a mapping and flags it for the NGO's confirmation.",
        "Each prefilled response is annotated with provenance showing the original passport answer, framework, and date.",
        "NGO reviews, refines, and submits the assessment portion of the application.",
    ],
    "Application's capacity section is complete; the NGO's effort is bounded by review rather than re-entry.",
)

use_case(doc, "UC-004", "Donor publishes a grant requiring a custom capacity framework",
    "Donor",
    "Donor needs a framework tailored to their portfolio.",
    [
        "Donor opens the custom framework editor.",
        "The system proposes a complete question bank, weights, evidence requirements, and scoring rubric derived from the grant's focus area.",
        "Donor refines questions, adjusts weights, and saves the framework.",
        "Donor attaches the framework to the grant.",
        "On submission, NGOs see the custom framework prefilled where their passport answers map to its questions, and clearly distinguished where they do not.",
    ],
    "Grant carries a donor-specific framework; applicants face the donor's questions but with passport prefill where possible.",
)

use_case(doc, "UC-005", "NGO drafts and submits an application",
    "NGO",
    "Open grant matches the NGO's profile.",
    [
        "NGO opens the grant and views the match score with top strength and top blocker.",
        "NGO opens the application wizard; the system prefills eligibility from the organisation profile and drafts each criterion response from the passport, prior applications, uploaded documents, and organisational memory.",
        "NGO refines responses; the system surfaces inline coaching where a response reads generic or makes an unsupported claim.",
        "NGO uploads required documents; the system extracts text, analyses against donor requirements, and scores each document.",
        "NGO runs the pre-submission readiness check; the system identifies missing evidence, overclaims, and weak passages, each with a one-click rewrite.",
        "NGO addresses the findings and submits; the system runs a compliance pre-empt scan and confirms no blockers.",
        "Application transitions from draft to submitted; the donor is notified.",
    ],
    "Application is submitted with traceable provenance on every drafted claim and a clean pre-flight check.",
)

use_case(doc, "UC-006", "Reviewer evaluates an application",
    "Reviewer",
    "Application has been submitted and assigned to the reviewer.",
    [
        "Reviewer opens the application; the system generates a one-screen summary with who the NGO is, what they propose, why-strong and why-weak summaries, and a draft rationale.",
        "Reviewer opens the evidence extraction view; the system has pulled verbatim supporting, contradicting, and neutral quotes per criterion.",
        "Reviewer scores each criterion, citing pulled evidence; the system records each score and rationale.",
        "Reviewer submits the score; the system aggregates against the rubric weights and produces a final AI-reference score for comparison.",
    ],
    "Application is scored with cited evidence; the reviewer's rationale is complete and defensible.",
)

use_case(doc, "UC-007", "Donor decides on an application",
    "Donor",
    "Application has been scored by the assigned reviewer.",
    [
        "Donor opens the application; the system presents the reviewer's score, the AI summary, the documents with their analysis, the applicant's compliance status, and the registry verification status.",
        "Donor reviews and either accepts, requests revision, or rejects.",
        "On rejection or revision, the donor adds a structured rationale that becomes part of the audit chain and is shared with the applicant.",
        "On acceptance, the system transitions the grant to awarded and initiates the contracting workflow.",
    ],
    "Application has a final decision with audit trail; on acceptance the grant becomes a live tracked engagement.",
)

use_case(doc, "UC-008", "NGO submits a progress report",
    "NGO",
    "Active grant has a reporting deadline approaching.",
    [
        "NGO opens the reporting workspace; the system has drafted a first-pass narrative from the prior report and the indicator data the NGO has captured.",
        "NGO refines the narrative and uploads supporting evidence (photographs, beneficiary data, financial statements).",
        "NGO runs the report pre-flight; the system identifies vague claims, unexplained budget variance, and missing evidence, each mapped to the donor concern it resolves.",
        "NGO addresses the findings and submits the report.",
        "The system computes a compliance score for the report and updates the grant's compliance health.",
    ],
    "Report is submitted with strong evidence coverage; donor is notified.",
)

use_case(doc, "UC-009", "Donor monitors compliance health and trajectory",
    "Donor",
    "Donor has one or more active grants.",
    [
        "Donor opens the compliance dashboard.",
        "The system displays each grant with its 4-pillar health score, compliance band, and (where applicable) a slips badge predicting how soon the grant will drop below the at-risk threshold.",
        "Donor opens a slipping grant; the system shows the trajectory chart over the last 60 days and identifies which pillar is degrading.",
        "Donor raises a risk on the grant; the system creates a risk register entry with severity, title, and proposed owner.",
        "Donor assigns the risk to a programme officer with a due date.",
        "The system surfaces the risk on the donor's awaiting-response queue until it is mitigated or accepted.",
    ],
    "At-risk grants are actively managed before they slip; risks have owners and deadlines.",
)

use_case(doc, "UC-010", "Donor and NGO collaborate on a clarification",
    "Donor and NGO",
    "Donor has a question about an application or report; NGO is reachable.",
    [
        "Donor writes a structured comment on the application or report, @-mentioning the relevant NGO contact.",
        "The system raises an in-app notification and, if subscribed, a web push notification on the NGO contact's devices.",
        "NGO contact opens the comment in context and replies in the thread.",
        "The conversation persists with the entity, preserved for audit and review.",
    ],
    "Clarification is resolved without leaving the platform; the exchange is part of the record.",
)

use_case(doc, "UC-011", "Administrator handles a flagged sanctions match",
    "Administrator",
    "Sanctions screening has raised a new flag on an organisation with an active grant.",
    [
        "Administrator receives a notification of the new flag.",
        "Administrator opens the organisation's compliance view; the system presents the matched entity, the screening source, and a plain-language explanation of the match.",
        "Administrator inspects the match against the organisation's known principals and beneficial owners.",
        "Administrator marks the match as a true positive (triggering account suspension) or a false positive (recording the rationale and continuing).",
        "The system records the determination in the audit chain.",
    ],
    "Sanctions flag is resolved with a documented determination.",
)

use_case(doc, "UC-012", "Administrator monitors AI surface health",
    "Administrator",
    "Daily scheduled probe has run; one flagship AI surface has reported failure.",
    [
        "Administrator opens the system-health surface and sees the AI surface health row in warn state.",
        "Administrator opens the AI surface health detail; the system shows the failing surface, the latency observed, and the failure reason.",
        "Administrator investigates the cause (model availability, prompt drift, schema change).",
        "Once resolved, the next daily probe transitions the row back to ok and the administrator notification is cleared.",
    ],
    "AI surface drift is detected and addressed before users encounter degraded behaviour.",
)

use_case(doc, "UC-013",
    "Donor creates a grant by uploading an existing brief document",
    "Donor",
    "Donor has drafted a grant brief in Word and wants to publish it on the platform.",
    [
        "Donor opens the grant creation wizard and chooses the upload path.",
        "Donor uploads the existing grant brief (PDF, DOCX, or plain text).",
        "The system extracts the structured content, mapping it across the wizard's six steps: basics, eligibility, criteria with weights, document requirements, reporting configuration, and review.",
        "Donor reviews each prefilled step; the system clearly distinguishes AI-extracted content from empty fields and from donor edits.",
        "Donor modifies, adds, and removes content where the extraction is incomplete or where the donor wants different language.",
        "Donor runs the applicant-burden critique and median-NGO preview before publishing.",
        "Donor publishes the grant; the system transitions it from draft to open and notifies matching NGOs.",
    ],
    "Grant is open with the donor's intended content preserved; the donor saved 90 minutes of manual data entry.",
)

use_case(doc, "UC-014",
    "NGO opens their This-Week action dashboard",
    "NGO",
    "NGO programme officer signs in on Monday morning.",
    [
        "The system presents the This-Week surface as the home page.",
        "Section: upcoming reports due — three reports due in the next four weeks, each with the draft state and an indication that two have AI-prepared narratives waiting.",
        "Section: applications in progress — two draft applications, each with the next recommended action (\"complete the M&E section\", \"upload last year's audit\").",
        "Section: matched grants — four newly published grants where the NGO's match score exceeds 75 percent, each with top strength and top blocker.",
        "Section: compliance to-dos — one open donor clarification request, one risk awaiting NGO response, and a capacity passport gap that would lift four match scores if closed.",
        "Programme officer clicks the highest-priority item (the report draft due in 14 days) and is taken directly to the report editor with the AI-prepared narrative loaded.",
    ],
    "NGO knows what to do before they have to ask.",
)

use_case(doc, "UC-015",
    "NGO uploads a document and receives real-time AI scoring",
    "NGO",
    "NGO is filling out an application that requires audited financial statements.",
    [
        "Before upload, the system shows the donor's specific requirement: two years of audited statements with auditor signature and breakdown of grants received.",
        "NGO uploads a PDF; the system begins extraction.",
        "Within seconds the system displays a score of 78 percent with three findings: \"auditor signature on page 4 ✓\", \"prior-year comparable included ✓\", \"three years requested; two provided\".",
        "The system offers two paths: upload the missing year (button), or attach a clarification note if unavailable (inline text field).",
        "NGO chooses the clarification path and types two sentences explaining the organisation was founded in 2023.",
        "The system incorporates the clarification, lifting the visible score to 92 percent without removing the underlying \"two of three years\" finding.",
        "Reviewer later sees both the score and the clarification in context.",
    ],
    "Real-time scoring helps the NGO close gaps while uploading rather than after rejection.",
)

use_case(doc, "UC-016",
    "NGO responds to a donor revision request with AI assistance",
    "NGO",
    "Donor has reviewed a progress report and returned it with revision requests on three sections.",
    [
        "NGO opens the report; the donor's structured feedback is attached to specific sections.",
        "The system translates each donor concern into a concrete NGO-side action and drafts the corresponding update by drawing on the report's existing content, prior reports, and uploaded evidence.",
        "For each donor concern, the system surfaces the proposed update and explains how it addresses the concern.",
        "NGO reviews, accepts, edits, or overrides each proposed update.",
        "NGO runs the donor pre-flight again before resubmitting; the system confirms each donor concern is resolved.",
        "NGO resubmits; the report transitions back to submitted state.",
    ],
    "Revision cycle closes in one round rather than three.",
)

use_case(doc, "UC-017",
    "NGO works offline in a low-connectivity area",
    "NGO field officer",
    "Field officer is documenting a beneficiary intake in an area with no mobile signal.",
    [
        "Field officer opens the platform as a progressive web app on her phone (already installed during onboarding).",
        "She drafts a report narrative; the system auto-saves to local device storage every few seconds.",
        "She attaches three photographs taken with her phone camera as evidence; the system queues them for upload.",
        "She closes the app; the work is preserved on the device.",
        "Twenty minutes later, on her way back to town, the phone reconnects to mobile data.",
        "The system synchronises the draft and uploads the photographs in the background, applying vision-based extraction to each image.",
        "That evening the officer opens her laptop; the draft and the analysed evidence are already waiting in the report editor.",
    ],
    "Field-collected evidence is captured at the point of collection without requiring connectivity, and synced when the network returns.",
)

doc.add_page_break()


# ===========================================================================
# 9. FEATURE FLAGS AND PHASED ROLLOUT
# ===========================================================================

heading(doc, "9. Feature Flags and Phased Rollout", 1)

add_body(doc,
    "The system shall support feature flags as a first-class mechanism "
    "for phased rollout, per-tenant configuration, and rapid rollback. "
    "Every flag has a global default, optional per-organisation and "
    "per-user overrides, and an environment-variable escape hatch. The "
    "flag set covers AI surfaces, UI surfaces, and operational gates."
)

add_body(doc,
    "The system shall:"
)

add_numbered_list(doc, [
    "support boolean and percentage flag types, with percentage flags bucketing each user stably so a given user receives a consistent answer across requests.",
    "expose the full flag inventory to administrators, showing each flag's current state, default, kind, description, and any per-scope overrides.",
    "evaluate flags in a documented order: environment override, per-user override, per-organisation override, global database row, in-code default — with the first match winning.",
    "support migration of default state without disturbing administrators' explicit overrides; flips of in-code defaults shall not silently override explicit administrative choices.",
])

add_body(doc,
    "Initial flag inventory at version 5.0 includes (non-exhaustive):"
)
add_bullet_list(doc, [
    "AI grant brief generator · AI compliance pre-empt · AI cross-grant patterns · AI compliance health narrative",
    "UI submission readiness · UI report readiness · UI reviewer summary · UI burden estimator",
    "UI this-week home · UI preview-as-reviewer · UI live-drafters pill · UI audit trail tab · UI compliance four-state · UI decision audit",
])

doc.add_page_break()


# ===========================================================================
# APPENDICES
# ===========================================================================

heading(doc, "Appendices", 1)

heading(doc, "Appendix A: Status Definitions", 2)

heading(doc, "Grant statuses", 3)
add_formatted_table(doc,
    ["Status", "Meaning"],
    [
        ["draft", "Donor is preparing the grant; not yet visible to applicants."],
        ["open", "Grant is published and accepting applications until deadline."],
        ["review", "Submission window closed; applications under review."],
        ["closed", "Window closed and no further submissions accepted."],
        ["awarded", "Grant has been awarded to one or more applicants."],
    ],
    col_widths=[1.2, 5.2]
)

doc.add_paragraph()
heading(doc, "Application statuses", 3)
add_formatted_table(doc,
    ["Status", "Meaning"],
    [
        ["draft", "Applicant is preparing; not yet submitted."],
        ["submitted", "Submitted; awaiting reviewer assignment or initial review."],
        ["under_review", "Assigned to a reviewer who is actively evaluating."],
        ["scored", "Reviewer has submitted scores; awaiting donor decision."],
        ["accepted", "Donor has accepted the application and the grant is awarded."],
        ["rejected", "Donor has declined the application."],
    ],
    col_widths=[1.5, 4.9]
)

doc.add_paragraph()
heading(doc, "Assessment statuses", 3)
add_formatted_table(doc,
    ["Status", "Meaning"],
    [
        ["in_progress", "Assessment started; responses being captured."],
        ["completed", "Assessment scored and added to the capacity passport."],
        ["expired", "Assessment is older than the donor's accepted recency window."],
    ],
    col_widths=[1.5, 4.9]
)

doc.add_paragraph()
heading(doc, "Report statuses", 3)
add_formatted_table(doc,
    ["Status", "Meaning"],
    [
        ["draft", "Report being prepared by the NGO."],
        ["submitted", "Report submitted; awaiting donor review."],
        ["accepted", "Report accepted by the donor."],
        ["revision_requested", "Donor has requested revisions with structured feedback."],
    ],
    col_widths=[1.7, 4.7]
)

doc.add_paragraph()
heading(doc, "Compliance statuses", 3)
add_formatted_table(doc,
    ["Status", "Meaning"],
    [
        ["clear", "No matches found across screening sources."],
        ["flagged", "One or more screening sources returned a match awaiting administrative review."],
        ["pending", "Screening is in progress or queued."],
        ["error", "Screening could not be completed; will be retried."],
    ],
    col_widths=[1.2, 5.2]
)

doc.add_paragraph()
heading(doc, "Compliance health bands", 3)
add_formatted_table(doc,
    ["Band", "Score range", "Meaning"],
    [
        ["on_track", "80–100", "Grant is meeting expectations across all pillars."],
        ["at_risk", "60–79", "Grant shows early-warning signals; intervention recommended."],
        ["high_risk", "0–59", "Grant is materially off-track on one or more pillars; immediate attention required."],
    ],
    col_widths=[1.2, 1.2, 4.0]
)

doc.add_paragraph()
heading(doc, "Risk statuses", 3)
add_formatted_table(doc,
    ["Status", "Meaning"],
    [
        ["open", "Risk identified; no response in flight."],
        ["mitigating", "Owner is actively working to mitigate; due date set."],
        ["mitigated", "Mitigation completed and verified; risk closed."],
        ["accepted", "Risk acknowledged and accepted by the donor as a known condition."],
        ["dismissed", "Risk determined to be a false positive or no longer applicable."],
    ],
    col_widths=[1.4, 5.0]
)

doc.add_page_break()


heading(doc, "Appendix B: Glossary", 2)

glossary = [
    ("Application", "An NGO's response to a specific grant, comprising eligibility responses, criterion responses, attached documents, and submission metadata."),
    ("Capacity passport", "The persistent, versioned set of capacity assessment results attached to an organisation, reusable across applications regardless of framework."),
    ("Compliance health score", "A composite score per active grant on a 0–100 scale, combining completion, timeliness, workflow, and importance pillars."),
    ("Co-pilot rail", "The persistent assistive surface in the user interface offering context-appropriate AI actions and answers without leaving the current workspace."),
    ("Grant", "A donor-published funding opportunity comprising title, description, eligibility, evaluation criteria, document requirements, reporting requirements, and deadline."),
    ("Match score", "A composite score per NGO-grant pair indicating alignment on eligibility, sector, geography, capacity, and track record."),
    ("Organisational memory", "The store of consistent facts derived from an organisation's submissions, used to ground AI drafting and surface provenance."),
    ("Passporting", "The act of carrying an existing capacity assessment forward to a new application, with prefill across frameworks and clearly marked provenance."),
    ("Pre-flight check", "An AI-driven analysis run before submission of an application or report, identifying gaps, overclaims, and weak passages with proposed fixes."),
    ("Provenance", "The traceable link from an AI-drafted claim to its source: organisation profile, prior submission, uploaded document, organisational memory, or AI-general."),
    ("Reviewer summary", "A one-screen synthesis of an application produced by the system for the reviewer, with evidence per criterion and a draft rationale."),
    ("Risk register", "The structured list of risks identified across organisations, applications, and grants, each with severity, status, owner, due date, and response."),
    ("Sanctions screening", "The continuous matching of organisations and named principals against UN, OFAC, and EU sanctions lists."),
    ("Slips badge", "A visual indicator surfaced on grants forecast to drop below the at-risk threshold within a configurable window (default 30 days)."),
    ("Trajectory", "The series of daily compliance health snapshots for an active grant, used to detect deterioration before it crosses a band threshold."),
]

add_formatted_table(doc,
    ["Term", "Definition"],
    glossary,
    col_widths=[1.6, 4.8]
)


# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

out_path = os.path.join(
    os.path.dirname(__file__),
    "Kuja_Grant_v5.0_Business_Requirements.docx",
)
doc.save(out_path)
print(f"Generated: {out_path}")
print(f"Size: {os.path.getsize(out_path):,} bytes")
