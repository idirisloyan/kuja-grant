"""
Generate Kuja Grant Management System — One-Page System Overview.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── Page setup (narrow margins for one-pager) ──
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(9.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # Blue-700
    hs.font.name = 'Calibri'

BLUE = RGBColor(0x1E, 0x40, 0xAF)
DARK = RGBColor(0x1F, 0x2A, 0x37)
GRAY = RGBColor(0x6B, 0x72, 0x80)
GREEN = RGBColor(0x05, 0x96, 0x69)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._element.get_or_add_tcPr()
        elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): '1E40AF', qn('w:val'): 'clear'})
        shading.append(elm)

    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(9)
        p.add_run(text).font.size = Pt(9)
    else:
        p.text = text
    for run in p.runs:
        run.font.size = Pt(9)
    return p


def small_para(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(4)
    return p


# ═══════════════════════════════════════════════════════════════
# HEADER
# ═══════════════════════════════════════════════════════════════

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('KUJA GRANT MANAGEMENT SYSTEM')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = BLUE

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('System Overview & Testing Guide')
run.font.size = Pt(12)
run.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Version 3.3.4  |  {datetime.date.today().strftime("%B %d, %Y")}  |  Adeso — African Development Solutions')
run.font.size = Pt(9)
run.font.color.rgb = GRAY

# Divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 80)
run.font.size = Pt(6)
run.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)

# ═══════════════════════════════════════════════════════════════
# QUICK REFERENCE
# ═══════════════════════════════════════════════════════════════

add_table(
    ['Item', 'Details'],
    [
        ['Production URL', 'https://web-production-6f8a.up.railway.app'],
        ['GitHub Repository', 'https://github.com/idirisloyan/kuja-grant (private)'],
        ['Stack', 'Python/Flask + Vanilla JS SPA + PostgreSQL (prod) / SQLite (dev)'],
        ['AI Engine', 'Anthropic Claude API (document analysis, scoring, chat)'],
        ['Deployment', 'Railway PaaS with Gunicorn + PostgreSQL'],
        ['Password (all accounts)', 'pass123'],
    ],
    col_widths=[4.5, 12.5]
)

# ═══════════════════════════════════════════════════════════════
# SYSTEM PURPOSE
# ═══════════════════════════════════════════════════════════════

doc.add_heading('System Purpose', level=2)
small_para(
    'Kuja Grant is an AI-powered end-to-end grant management platform for NGOs and Donors '
    'operating in Africa. It covers the complete grant lifecycle from discovery through '
    'compliance, with built-in capacity assessments, live sanctions screening, government '
    'registry verification, and AI-driven document analysis across 5 languages (EN/FR/AR/SW/SO).'
)

# ═══════════════════════════════════════════════════════════════
# INTENDED USERS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Intended Users & Test Accounts', level=2)

add_table(
    ['Role', 'Email', 'Organization', 'Key Capabilities'],
    [
        ['Admin', 'admin@kuja.org', 'Kuja', 'Full system configuration and user management'],
        ['NGO', 'fatima@amani.org', 'Amani Community Dev (Kenya)', 'Assessments, applications, document upload, reporting'],
        ['NGO', 'ahmed@salamrelief.org', 'Salam Relief (Somalia)', 'Grant discovery, capacity building, AI chat'],
        ['NGO', 'thandi@ubuntu.org', 'Ubuntu Education (SA)', 'Highest-rated NGO (91%), full workflow testing'],
        ['NGO', 'peter@hopebridges.org', 'Hope Bridges (Uganda)', 'Lower capacity (55%), edge case testing'],
        ['NGO', 'aisha@sahelwomen.org', 'Sahel Women (Nigeria)', 'Lowest capacity (47%), improvement workflows'],
        ['Donor', 'sarah@globalhealth.org', 'Global Health Fund', 'Grant wizard, AI evaluation, report review'],
        ['Donor', 'david@eatrust.org', 'East Africa Dev Trust', 'Compliance dashboard, sanctions screening'],
        ['Reviewer', 'james@reviewer.org', '—', 'Application scoring, structured review rubrics'],
        ['Reviewer', 'maria@reviewer.org', '—', 'Dual scoring (AI + human reviewer)'],
    ],
    col_widths=[2, 4.5, 4.5, 6]
)

# ═══════════════════════════════════════════════════════════════
# CORE FEATURES & WORKFLOWS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Core Features & Workflows', level=2)

# NGO workflow
p = doc.add_paragraph()
run = p.add_run('NGO Workflow: ')
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = GREEN
p.add_run(
    'Register → Complete Capacity Assessment (5 frameworks: Kuja, STEP, UN-HACT, CHS, NUPAS) → '
    'Discover Grants → Apply (4-step wizard with AI guidance) → Upload Documents (AI scores each) → '
    'Submit Reports → Track Compliance'
).font.size = Pt(9)

# Donor workflow
p = doc.add_paragraph()
run = p.add_run('Donor Workflow: ')
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = BLUE
p.add_run(
    'Create Grant (5-step wizard) → Upload Agreement (AI extracts reporting requirements) → '
    'Set Evaluation Criteria per Document Type → Review Applications (AI-ranked) → '
    'Award Grants → Review NGO Reports (AI evaluates against donor-specific requirements)'
).font.size = Pt(9)

# Feature summary table
add_table(
    ['Feature', 'Description'],
    [
        ['5-Framework Assessment', 'Kuja, STEP, UN-HACT, CHS, NUPAS — generates capacity scores per NGO'],
        ['Live Sanctions Screening', 'OpenSanctions API + direct UN/OFAC/EU/World Bank list downloads'],
        ['Registry Verification', 'Live verification for 7 African countries (Kenya, Nigeria, SA, Uganda, Tanzania, Somalia, Ethiopia)'],
        ['AI Document Analysis', 'Upload PDF/DOCX/XLSX — Claude AI scores against donor-specific criteria with detailed findings'],
        ['AI Grant Agreement Parsing', 'AI extracts reporting requirements, deadlines, and compliance obligations from agreements'],
        ['AI Report Evaluation', 'Reports scored per-requirement with compliance breakdown and risk flags'],
        ['AI Chat Assistant', 'Role-aware contextual guidance (different responses for NGOs vs Donors)'],
        ['Multi-Language Support', 'Full UI and AI analysis in English, French, Arabic, Swahili, Somali'],
        ['Donor Grant Wizard', '5-step wizard: details → documents → requirements → evaluation criteria → review'],
        ['NGO Reporting', 'Upcoming deadlines, AI evaluation, compliance scores, revision workflows'],
    ],
    col_widths=[4.5, 12.5]
)

# ═══════════════════════════════════════════════════════════════
# AI INTEGRATION
# ═══════════════════════════════════════════════════════════════

doc.add_heading('AI Integration (Claude API)', level=2)
small_para(
    'The system uses Anthropic Claude for intelligent document processing. When the ANTHROPIC_API_KEY '
    'is configured, all AI features are live. Without the key, the system falls back to simulated '
    'responses with pre-built templates.'
)

add_table(
    ['AI Function', 'Trigger', 'Output'],
    [
        ['Document Analysis', 'NGO uploads a document to an application', 'Score (0-100), detailed findings, recommendations'],
        ['Grant Agreement Extraction', 'Donor uploads agreement in grant wizard', 'Structured list of reporting requirements with deadlines'],
        ['Report Evaluation', 'NGO submits report for donor review', 'Per-requirement scores, compliance %, risk flags'],
        ['Chat Assistance', 'User opens AI chat panel', 'Context-aware guidance based on user role and current page'],
        ['Registration Analysis', 'System verifies NGO registration certificate', 'Validity assessment, registration details extraction'],
    ],
    col_widths=[4, 5.5, 7.5]
)

# ═══════════════════════════════════════════════════════════════
# TEST CASES & TEST FILES
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Test Cases & Test Files', level=2)

small_para(
    'The TEST_PLAN.md contains 178 test cases across 24 sections. A ZIP file of 22 test files '
    '(test-files.zip) provides the documents needed to execute the test cases. Below is how '
    'the test files map to the testing workflow:'
)

add_table(
    ['Test Area', 'Test Cases', 'Test Files Used', 'What is Verified'],
    [
        ['Document Upload & AI Scoring', 'TC-DOCU series',
         'financial_report_q1_2026.txt, audit_report_2025.txt, project_proposal.txt, budget_template.txt',
         'Files upload correctly, AI returns scores in expected ranges (e.g., financial: 70-80, audit: 80-90)'],
        ['Grant Agreement AI Extraction', 'TC-LIVEAI series',
         'grant_agreement_sample.txt',
         'AI extracts reporting requirements, deadlines, and obligations from agreement text'],
        ['Report Evaluation', 'TC-RPTN, TC-RPTD series',
         'financial_report_q1_2026.txt, impact_report_annual_2025.txt',
         'AI evaluates reports against donor-defined requirements with per-requirement scoring'],
        ['Multi-Language Testing', 'TC-LANG series',
         'arabic_grant_agreement.txt, french_project_report.txt',
         'AI correctly processes and scores non-English documents'],
        ['Edge Cases & Error Handling', 'TC-EDGE series',
         'empty_template.txt, poor_quality_report.txt, large_budget_detailed.txt',
         'Empty files score very low (5-15), poor quality scores 20-40, system handles large files gracefully'],
        ['Live Sanctions Screening', 'TC-SANC series',
         'sanctions_screening_test_data.txt (reference)',
         'Real-time screening against UN/OFAC/EU/World Bank lists returns correct matches'],
        ['Registry Verification', 'TC-REGV series',
         'registration_verification_test_data.txt (reference)',
         'Live lookup of real NGO registration numbers for 7 countries returns valid/invalid results'],
        ['Compliance & Due Diligence', 'TC-COMP series',
         'compliance_checklist.txt, due_diligence_questionnaire_completed.txt',
         'Documents upload, AI analyzes compliance posture, scores reflect completeness'],
    ],
    col_widths=[3.5, 2.5, 4.5, 6.5]
)

# ═══════════════════════════════════════════════════════════════
# PRE-LOADED TEST DATA
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Pre-Loaded Test Data', level=2)

small_para(
    'The system comes seeded with realistic test data so testers can immediately explore all features '
    'without setup:'
)

add_table(
    ['Data', 'Count', 'Details'],
    [
        ['Users', '10', '1 admin, 5 NGOs (across Kenya, Somalia, SA, Uganda, Nigeria), 2 donors, 2 reviewers'],
        ['Grants', '4', 'Community Health ($500K), EdTech ($250K), Climate ($1M), Women ($350K)'],
        ['Applications', '5', 'Various statuses: submitted, draft, incomplete — with AI scores'],
        ['NGO Reports', '5', 'Submitted, draft, accepted, revision requested, under review'],
        ['Assessments', '5', 'Capacity scores ranging from 47% to 91% across all 5 frameworks'],
        ['Departments', '7', 'Kenya, Nigeria, SA, Uganda, Tanzania, Somalia, Ethiopia registry data'],
    ],
    col_widths=[3, 2, 12]
)

# ═══════════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 80)
run.font.size = Pt(6)
run.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    f'Kuja Grant v3.3.4  |  Adeso — African Development Solutions  |  {datetime.date.today().year}'
)
run.font.size = Pt(8)
run.font.color.rgb = GRAY
run.italic = True

# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════

output = r'C:\Users\IdirisLoyan\kuja-grant\docs\Kuja_Grant_System_Overview.docx'
doc.save(output)
print(f'Document saved to: {output}')
