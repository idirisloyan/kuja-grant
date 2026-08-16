#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
make_testfiles.py -- generates the Kuja Marketplace UAT test-file pack.

Produces realistic, fictional documents/media for every upload point in the
Kuja Grant + Kuja Trust journeys, plus a folder of deliberately-broken edge-case
files. All content is INVENTED for testing (fictional org "Amani Health
Initiative", fictional donor "Global Health Fund"). No real PII.

Run:  py -3 docs/uat/make_testfiles.py
Out:  docs/uat/testfiles/<category>/...
"""
import os, sys, textwrap, subprocess, io, struct, random

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass
random.seed(42)
ROOT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "testfiles")

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                TableStyle, Image as RLImage)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from PIL import Image, ImageDraw, ImageFont
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from PyPDF2 import PdfReader, PdfWriter
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def ensure(d):
    os.makedirs(d, exist_ok=True); return d

def cat(name):
    return ensure(os.path.join(ROOT, name))

def _font(size):
    for p in (r"C:\Windows\Fonts\arialbd.ttf", r"C:\Windows\Fonts\arial.ttf"):
        try: return ImageFont.truetype(p, size)
        except Exception: pass
    return ImageFont.load_default()

# ---------------------------------------------------------------- PDF helpers
styles = getSampleStyleSheet()
H = ParagraphStyle('H', parent=styles['Heading1'], fontSize=15, spaceAfter=6,
                   textColor=colors.HexColor('#0f3d5c'))
H2 = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=11.5,
                    textColor=colors.HexColor('#146'))
BODY = ParagraphStyle('B', parent=styles['Normal'], fontSize=9.5, leading=13.5)
SMALL = ParagraphStyle('S', parent=styles['Normal'], fontSize=8, textColor=colors.grey)
CENTER = ParagraphStyle('C', parent=BODY, alignment=TA_CENTER)

def letterhead(org="Amani Health Initiative",
               tag="Community health for mothers &amp; children — Garissa, Kenya"):
    return [
        Paragraph(f"<b>{org}</b>", ParagraphStyle('org', parent=styles['Title'],
                  fontSize=17, textColor=colors.HexColor('#0f3d5c'))),
        Paragraph(tag, SMALL),
        Paragraph("Reg. No. NGO/KE/2018/04471 &nbsp;•&nbsp; P.O. Box 233-70100, "
                  "Garissa &nbsp;•&nbsp; info@amanihealth.org.ke", SMALL),
        Spacer(1, 8),
        Table([[""]], colWidths=[170*mm],
              style=TableStyle([('LINEBELOW',(0,0),(-1,-1),1.2,
                                 colors.HexColor('#0f3d5c'))])),
        Spacer(1, 10),
    ]

def build_pdf(path, flowables, head=True):
    doc = SimpleDocTemplate(path, pagesize=A4, topMargin=18*mm, bottomMargin=18*mm,
                            leftMargin=20*mm, rightMargin=20*mm,
                            title=os.path.basename(path))
    story = (letterhead() if head else []) + flowables
    doc.build(story)
    print("  PDF ", os.path.relpath(path, ROOT))

def P(t): return Paragraph(t, BODY)
def para_block(text):
    return [P(x.strip()) for x in text.strip().split("\n\n")]

def money_table(rows, header=("Budget line","Unit","Qty","Unit cost (USD)","Total (USD)")):
    data = [list(header)] + rows
    t = Table(data, colWidths=[62*mm,20*mm,14*mm,32*mm,32*mm])
    t.setStyle(TableStyle([
        ('BACKGROUND',(0,0),(-1,0),colors.HexColor('#0f3d5c')),
        ('TEXTCOLOR',(0,0),(-1,0),colors.white),
        ('FONTSIZE',(0,0),(-1,-1),8.5),
        ('GRID',(0,0),(-1,-1),0.4,colors.HexColor('#bbb')),
        ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white,colors.HexColor('#eef4f8')]),
        ('ALIGN',(2,0),(-1,-1),'RIGHT'),
        ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
    ]))
    return t

# ================================================================ 01 REGISTRATION
def registration():
    d = cat("01_registration")
    build_pdf(os.path.join(d,"Amani_Certificate_of_Registration.pdf"), [
        Paragraph("REPUBLIC OF KENYA", CENTER),
        Paragraph("<b>NGO Co-ordination Board — Certificate of Registration</b>", CENTER),
        Spacer(1,10),
        *para_block("""This is to certify that AMANI HEALTH INITIATIVE has been duly registered
        under the Non-Governmental Organizations Co-ordination Act (1990) and is authorised to
        operate in the Republic of Kenya.

        Registration Number: NGO/KE/2018/04471
        Date of Registration: 14 March 2018
        Registered Office: P.O. Box 233-70100, Garissa
        Area of Operation: Garissa, Wajir and Tana River Counties
        Sectors: Health; Water, Sanitation &amp; Hygiene; Nutrition"""),
        Spacer(1,20),
        Paragraph("____________________________<br/>Executive Director, NGO Co-ordination Board",
                  BODY),
        Spacer(1,6), Paragraph("[TEST DOCUMENT — fictional, for Kuja UAT only]", SMALL),
    ])
    build_pdf(os.path.join(d,"Amani_Constitution.pdf"), [
        Paragraph("Constitution &amp; Governing Document", H),
        *para_block("""ARTICLE 1 — NAME. The organisation shall be known as Amani Health Initiative
        (hereinafter "the Organisation").

        ARTICLE 2 — OBJECTS. To improve maternal, newborn and child health outcomes among
        underserved and displacement-affected communities in north-eastern Kenya through
        community health services, health education and referral support.

        ARTICLE 3 — MEMBERSHIP. Membership is open to individuals who subscribe to the objects
        of the Organisation and are approved by the Board.

        ARTICLE 4 — GOVERNANCE. The Organisation shall be governed by a Board of not fewer than
        five (5) and not more than nine (9) members, meeting at least quarterly.

        ARTICLE 5 — FINANCE. The financial year shall run 1 January to 31 December. Accounts
        shall be independently audited annually.

        ARTICLE 6 — SAFEGUARDING. The Organisation maintains a zero-tolerance policy toward
        sexual exploitation, abuse and harassment, and toward fraud and corruption."""),
    ])
    build_pdf(os.path.join(d,"Amani_Audited_Financials_2025.pdf"), [
        Paragraph("Independent Auditor's Report &amp; Financial Statements — Year ended 31 Dec 2025", H2),
        *para_block("""We have audited the financial statements of Amani Health Initiative, which
        comprise the statement of financial position and the statement of income and expenditure.
        In our opinion the financial statements present fairly, in all material respects, the
        financial position of the Organisation."""),
        Spacer(1,8), Paragraph("Statement of Income &amp; Expenditure (USD)", H2),
        money_table([
            ["Grant income — restricted","","","","842,150"],
            ["Grant income — unrestricted","","","","96,400"],
            ["Programme expenditure","","","","(731,980)"],
            ["Support &amp; admin (11.2%)","","","","(105,120)"],
            ["Surplus for the year","","","","101,450"],
        ], header=("Line","","","","Amount (USD)")),
        Spacer(1,8),
        Paragraph("Prepared by Ngoni &amp; Associates, Certified Public Accountants (K). "
                  "[TEST DOCUMENT — fictional].", SMALL),
    ])
    build_pdf(os.path.join(d,"Amani_Board_of_Directors.pdf"), [
        Paragraph("Board of Directors &amp; Key Personnel", H),
        Table([["Name","Role","Since"],
               ["Dr. Halima Abdi","Board Chair","2018"],
               ["Joseph Kimani","Treasurer","2019"],
               ["Amina Yusuf","Secretary","2020"],
               ["Peter Otieno","Director","2021"],
               ["Fatima Noor","Executive Director","2018"]],
              colWidths=[60*mm,60*mm,30*mm],
              style=TableStyle([('BACKGROUND',(0,0),(-1,0),colors.HexColor('#0f3d5c')),
                                ('TEXTCOLOR',(0,0),(-1,0),colors.white),
                                ('GRID',(0,0),(-1,-1),0.4,colors.grey),
                                ('FONTSIZE',(0,0),(-1,-1),9)])),
        Spacer(1,8), Paragraph("[TEST DOCUMENT — fictional persons, for Kuja UAT only]", SMALL),
    ])
    build_pdf(os.path.join(d,"Amani_Tax_Exemption_Letter.pdf"), [
        Paragraph("Kenya Revenue Authority — Income Tax Exemption", H2),
        *para_block("""Ref: KRA/EX/2019/2231. This letter confirms that Amani Health Initiative
        is granted income tax exemption under Paragraph 10, First Schedule of the Income Tax Act,
        valid and renewable every five (5) years. Current validity: 1 Jan 2024 – 31 Dec 2028."""),
        Spacer(1,10), Paragraph("[TEST DOCUMENT — fictional]", SMALL),
    ])
    # org chart image
    org_chart(os.path.join(d,"Amani_Org_Chart.png"))

def org_chart(path):
    W,Hh = 1000,620; im = Image.new("RGB",(W,Hh),"white"); dr = ImageDraw.Draw(im)
    def box(x,y,w,h,txt,fill="#0f3d5c"):
        dr.rounded_rectangle([x,y,x+w,y+h],12,fill=fill)
        f=_font(20); tb=dr.textbbox((0,0),txt,font=f)
        dr.text((x+(w-(tb[2]-tb[0]))/2, y+(h-(tb[3]-tb[1]))/2-4),txt,font=f,fill="white")
    dr.text((30,20),"Amani Health Initiative — Organisation Chart",font=_font(26),fill="#0f3d5c")
    box(400,90,200,60,"Board of Directors")
    box(400,200,200,60,"Executive Director")
    for i,(t) in enumerate(["Programmes","Finance &\nGrants","M&E","Community\nHealth Teams"]):
        x=70+i*230; box(x,340,180,70,t.replace("\n"," / "),fill="#2a6f97")
        dr.line([500,260,x+90,340],fill="#888",width=2)
    im.save(path); print("  IMG ", os.path.relpath(path,ROOT))

# ================================================================ 02 CAPACITY / TRUST
def capacity_trust():
    d = cat("02_capacity_trust_evidence")
    policies = {
        "Amani_Safeguarding_PSEA_Policy.pdf": ("Safeguarding &amp; PSEA Policy",
            """Amani Health Initiative maintains zero tolerance for sexual exploitation, abuse and
            harassment (SEAH). All staff, volunteers and partners sign a code of conduct on
            induction. A designated Safeguarding Focal Point receives and manages concerns through
            a confidential reporting channel. Survivors are referred to appropriate services.
            Recruitment includes reference and background checks."""),
        "Amani_Financial_Management_Policy.pdf": ("Financial Management Policy",
            """Segregation of duties is enforced between initiation, approval and payment.
            All expenditure above USD 500 requires dual authorisation. Bank reconciliations are
            performed monthly. Procurement follows a three-quote rule above USD 1,000. Assets are
            tagged and inventoried annually."""),
        "Amani_AntiFraud_AntiCorruption_Policy.pdf": ("Anti-Fraud &amp; Anti-Corruption Policy",
            """The Organisation prohibits bribery, kickbacks, and conflicts of interest. Staff must
            declare related-party interests. A whistle-blowing channel is available and protected
            from retaliation. Confirmed fraud is reported to the Board and relevant donors."""),
        "Amani_HR_Policy.pdf": ("Human Resources Policy",
            """Recruitment is merit-based and non-discriminatory. All staff have written contracts,
            job descriptions and annual performance reviews. Grievance and disciplinary procedures
            are documented. Statutory deductions (NSSF, SHIF, PAYE) are remitted monthly."""),
        "Amani_Data_Protection_Policy.pdf": ("Data Protection Policy",
            """Personal data of beneficiaries is collected for defined purposes, stored securely and
            retained only as long as necessary, consistent with the Kenya Data Protection Act 2019.
            Access is role-restricted. Data-sharing with donors is aggregated and de-identified."""),
    }
    for fn,(title,body) in policies.items():
        build_pdf(os.path.join(d,fn), [Paragraph(title,H), *para_block(body),
                  Spacer(1,10), Paragraph("Approved by the Board — v3.1, Jan 2025. "
                  "[TEST DOCUMENT — fictional].", SMALL)])
    build_pdf(os.path.join(d,"Amani_Bank_Confirmation_Letter.pdf"), [
        Paragraph("Kenya Commercial Bank — Account Confirmation", H2),
        *para_block("""We confirm that Amani Health Initiative maintains account number
        11xxxxxx88 (USD) and 22xxxxxx01 (KES) at our Garissa branch, in good standing since 2018.
        Authorised signatories operate on an any-two basis."""),
        Spacer(1,10), Paragraph("[TEST DOCUMENT — fictional, masked account numbers]", SMALL)])
    build_pdf(os.path.join(d,"Amani_Prior_Donor_Reference_Letter.pdf"), [
        Paragraph("Reference Letter — Horn Relief Foundation", H2),
        *para_block("""Amani Health Initiative implemented our 2023–2024 maternal health grant
        (USD 210,000) to a high standard, submitting timely narrative and financial reports with
        clean audit findings. We recommend them without reservation."""),
        Spacer(1,10), Paragraph("Grants Director, Horn Relief Foundation. [TEST — fictional].", SMALL)])

# ================================================================ 03 APPLICATION
def application():
    d = cat("03_grant_application")
    proposal_body = """Project title: Strengthening Maternal &amp; Newborn Health in Garissa County

    Summary. Amani Health Initiative requests USD 180,000 over 18 months to reduce maternal and
    newborn mortality in three sub-counties through community health promoter (CHP) networks,
    emergency referral support and facility mentorship.

    Problem. Garissa's maternal mortality ratio remains far above the national average, driven by
    low skilled-birth attendance, distance to facilities and delayed referral. Displacement from
    the wider region has increased pressure on services.

    Objectives. (1) Increase skilled birth attendance from 41% to 60% in the target area.
    (2) Train and equip 120 CHPs. (3) Establish an emergency transport voucher scheme covering
    six health facilities.

    Approach. CHP recruitment and training; monthly facility mentorship; a referral voucher
    scheme with local transporters; community dialogue sessions; routine data review.

    Beneficiaries. 24,000 women of reproductive age and 6,500 newborns, directly.

    Monitoring. Against the attached logframe; monthly CHP reporting; quarterly data review with
    the county health team.

    Sustainability. CHPs are integrated into the county community-health strategy; the voucher
    scheme transitions to county co-financing by month 15."""
    build_pdf(os.path.join(d,"Amani_Project_Proposal_MCH.pdf"),
              [Paragraph("Grant Application — Full Proposal", H), *para_block(proposal_body)])
    # docx variant
    doc = docx.Document()
    doc.add_heading("Grant Application — Full Proposal (Word variant)", level=1)
    for para in proposal_body.strip().split("\n\n"):
        doc.add_paragraph(para.strip())
    doc.save(os.path.join(d,"Amani_Project_Proposal_MCH.docx"))
    print("  DOCX", "03_grant_application/Amani_Project_Proposal_MCH.docx")

    # Budget xlsx
    budget_rows = [
        ("Personnel — Project Coordinator (18m)","month",18,1200,None),
        ("Personnel — M&E Officer (18m)","month",18,900,None),
        ("CHP training (120 x 3 days)","person",120,45,None),
        ("CHP kits &amp; equipment","kit",120,60,None),
        ("Emergency transport vouchers","voucher",1500,12,None),
        ("Facility mentorship visits","visit",108,80,None),
        ("Community dialogue sessions","session",36,150,None),
        ("Monitoring &amp; data review","quarter",6,700,None),
        ("Travel &amp; fuel","month",18,650,None),
        ("Indirect / support cost (11%)","lump",1,17820,None),
    ]
    budget_xlsx(os.path.join(d,"Amani_Detailed_Budget.xlsx"), budget_rows)
    # Budget pdf mirror
    pdf_rows=[]
    for name,unit,qty,uc,_ in budget_rows:
        tot = qty*uc; pdf_rows.append([name,unit,str(qty),f"{uc:,}",f"{tot:,}"])
    total = sum(q*u for _,_,q,u,_ in budget_rows)
    pdf_rows.append(["TOTAL","","","",f"{total:,}"])
    build_pdf(os.path.join(d,"Amani_Detailed_Budget.pdf"),
              [Paragraph("Detailed Budget (USD)",H2), money_table(pdf_rows)])
    # Logframe xlsx
    logframe_xlsx(os.path.join(d,"Amani_Logframe.xlsx"))
    # Workplan / gantt pdf
    build_pdf(os.path.join(d,"Amani_Workplan.pdf"), [
        Paragraph("18-Month Workplan",H2),
        Table([["Activity","Q1","Q2","Q3","Q4","Q5","Q6"],
               ["CHP recruitment","X","","","","",""],
               ["CHP training","X","X","","","",""],
               ["Voucher scheme","","X","X","X","X","X"],
               ["Facility mentorship","","X","X","X","X","X"],
               ["Data review","X","X","X","X","X","X"],
               ["Reporting","","X","","X","","X"]],
              colWidths=[55*mm]+[18*mm]*6,
              style=TableStyle([('BACKGROUND',(0,0),(-1,0),colors.HexColor('#0f3d5c')),
                                ('TEXTCOLOR',(0,0),(-1,0),colors.white),
                                ('GRID',(0,0),(-1,-1),0.4,colors.grey),
                                ('ALIGN',(1,0),(-1,-1),'CENTER'),
                                ('FONTSIZE',(0,0),(-1,-1),9)])),
    ])
    build_pdf(os.path.join(d,"Amani_Partnership_MOU.pdf"), [
        Paragraph("Memorandum of Understanding — County Health Team",H2),
        *para_block("""This MOU between Amani Health Initiative and the Garissa County Department of
        Health sets out roles for joint implementation of the maternal health project, including
        facility access, data sharing and co-supervision. Valid for the project period."""),
        Spacer(1,10), Paragraph("[TEST DOCUMENT — fictional]", SMALL)])
    build_pdf(os.path.join(d,"Amani_Letter_of_Support_County.pdf"), [
        Paragraph("Letter of Support — Office of the County Health Director",H2),
        *para_block("""The County Health Management Team endorses this application and confirms
        alignment with the county community-health strategy and integrated work plan."""),
        Spacer(1,10), Paragraph("[TEST DOCUMENT — fictional]", SMALL)])

def budget_xlsx(path, rows):
    wb=openpyxl.Workbook(); ws=wb.active; ws.title="Budget"
    hdr=["Budget line","Unit","Qty","Unit cost (USD)","Total (USD)"]
    ws.append(hdr)
    for c in ws[1]:
        c.font=Font(bold=True,color="FFFFFF"); c.fill=PatternFill("solid",fgColor="0F3D5C")
    for name,unit,qty,uc,_ in rows:
        r=ws.max_row+1
        ws.append([name.replace("&amp;","&"),unit,qty,uc,None])
        ws.cell(r,5).value=f"=C{r}*D{r}"
    tot=ws.max_row+1
    ws.append(["TOTAL","","","",f"=SUM(E2:E{tot-1})"])
    for c in ws[tot]: c.font=Font(bold=True)
    ws.column_dimensions["A"].width=42
    for col in "BCDE": ws.column_dimensions[col].width=16
    wb.save(path); print("  XLSX", os.path.relpath(path,ROOT))

def logframe_xlsx(path):
    wb=openpyxl.Workbook(); ws=wb.active; ws.title="Logframe"
    hdr=["Level","Indicator","Baseline","Target","Means of verification"]
    ws.append(hdr)
    for c in ws[1]:
        c.font=Font(bold=True,color="FFFFFF"); c.fill=PatternFill("solid",fgColor="0F3D5C")
    for row in [
        ["Goal","Maternal mortality ratio (target area)","362/100k","<300/100k","County HIS"],
        ["Outcome","Skilled birth attendance","41%","60%","Facility registers"],
        ["Output 1","CHPs trained &amp; active","0","120","Training records"],
        ["Output 2","Emergency referrals supported","0","1,500","Voucher log"],
        ["Output 3","Facilities mentored monthly","0","6","Mentorship reports"],
    ]:
        ws.append([x.replace("&amp;","&") for x in row])
    ws.column_dimensions["A"].width=12; ws.column_dimensions["B"].width=38
    for col in "CDE": ws.column_dimensions[col].width=18
    wb.save(path); print("  XLSX", os.path.relpath(path,ROOT))

# ================================================================ 04 COMPLIANCE / REPORTING
def compliance():
    d = cat("04_compliance_reporting")
    # narrative report docx
    doc=docx.Document()
    doc.add_heading("Quarterly Narrative Progress Report — Q1",0)
    doc.add_paragraph("Grant: Strengthening Maternal & Newborn Health in Garissa County")
    doc.add_paragraph("Reporting period: Jan–Mar 2026    Prepared by: Fatima Noor, ED")
    for h,b in [
        ("Progress against objectives",
         "84 of 120 CHPs recruited and trained (70%). The emergency transport voucher scheme "
         "launched at four of six facilities. Skilled birth attendance in the target area rose "
         "from 41% to 48% this quarter."),
        ("Activities delivered",
         "Three 3-day CHP training cohorts; 512 vouchers issued; monthly mentorship at four "
         "facilities; two community dialogue sessions."),
        ("Challenges & mitigation",
         "Seasonal flooding delayed access to two facilities; mitigated by rescheduling "
         "mentorship and using boat transport for referrals."),
        ("Finances",
         "USD 41,200 of USD 60,000 tranche 1 expended (69%); full reconciliation attached."),
        ("Next quarter",
         "Complete CHP recruitment; extend vouchers to all six facilities; first data review "
         "with the county team."),
    ]:
        doc.add_heading(h,level=1); doc.add_paragraph(b)
    doc.add_paragraph("[TEST DOCUMENT — fictional, for Kuja UAT only]")
    doc.save(os.path.join(d,"Amani_Q1_Narrative_Progress_Report.docx"))
    print("  DOCX", "04_compliance_reporting/Amani_Q1_Narrative_Progress_Report.docx")

    # financial report xlsx
    wb=openpyxl.Workbook(); ws=wb.active; ws.title="Q1 Financial Report"
    ws.append(["Budget line","Budget (USD)","Spent Q1 (USD)","Variance"])
    for c in ws[1]:
        c.font=Font(bold=True,color="FFFFFF"); c.fill=PatternFill("solid",fgColor="0F3D5C")
    data=[["Personnel",37800,9450],["CHP training",16200,11340],
          ["Vouchers",18000,6144],["Mentorship",8640,2880],
          ["Travel & fuel",11700,3900],["Support cost",17820,5940]]
    for name,bud,spent in data:
        r=ws.max_row+1; ws.append([name,bud,spent,None]); ws.cell(r,4).value=f"=B{r}-C{r}"
    tr=ws.max_row+1; ws.append(["TOTAL",f"=SUM(B2:B{tr-1})",f"=SUM(C2:C{tr-1})",f"=SUM(D2:D{tr-1})"])
    for c in ws[tr]: c.font=Font(bold=True)
    ws.column_dimensions["A"].width=26
    for col in "BCD": ws.column_dimensions[col].width=16
    wb.save(os.path.join(d,"Amani_Q1_Financial_Report.xlsx"))
    print("  XLSX", "04_compliance_reporting/Amani_Q1_Financial_Report.xlsx")

    # receipts bundle pdf
    build_pdf(os.path.join(d,"Amani_Expense_Receipts_Bundle.pdf"), [
        Paragraph("Expense Receipts — Q1 (sample)", H2),
        money_table([
            ["Garissa Pharmacy Ltd — MCH supplies","","","","1,240"],
            ["Northern Transporters — referral fuel","","","","615"],
            ["Sahara Stationers — training materials","","","","210"],
            ["Nomad Catering — training refreshments","","","","480"],
        ], header=("Vendor","","","","Amount (USD)")),
        Spacer(1,8), Paragraph("Originals held on file. [TEST — fictional].", SMALL)])

    # photo/receipt/attendance images
    field_photo(os.path.join(d,"field_photo_CHP_training.jpg"),
                "CHP Training Cohort 2","Garissa • 12 Feb 2026", (46,101,140))
    field_photo(os.path.join(d,"field_photo_clinic_site.jpg"),
                "Referral Facility — Sankuri","Garissa • 03 Mar 2026", (33,110,97))
    field_photo(os.path.join(d,"field_photo_beneficiaries.jpg"),
                "Community Dialogue Session","Garissa • 21 Mar 2026", (120,72,40))
    receipt_image(os.path.join(d,"receipt_pharmacy_supplies.jpg"))
    attendance_image(os.path.join(d,"attendance_sheet_scanned.jpg"))
    # WebP variant (accepted by photo-evidence + Trust evidence room, not by /documents)
    Image.open(os.path.join(d,"field_photo_CHP_training.jpg")).save(
        os.path.join(d,"field_photo_CHP_training.webp"),"WEBP",quality=80)
    print("  IMG  04_compliance_reporting/field_photo_CHP_training.webp")

def field_photo(path,title,sub,rgb):
    W,Hh=1280,860; im=Image.new("RGB",(W,Hh),rgb); dr=ImageDraw.Draw(im)
    for i in range(Hh):  # simple vertical gradient
        k=1-i/Hh*0.55
        dr.line([(0,i),(W,i)],fill=(int(rgb[0]*k),int(rgb[1]*k),int(rgb[2]*k)))
    # fake horizon / ground blocks so it reads as a scene
    dr.rectangle([0,int(Hh*0.72),W,Hh],fill=(int(rgb[0]*0.4)+40,int(rgb[1]*0.4)+35,int(rgb[2]*0.3)+25))
    for _ in range(9):
        x=random.randint(40,W-140); h=random.randint(60,180)
        dr.rectangle([x,int(Hh*0.72)-h,x+random.randint(40,90),int(Hh*0.72)],
                     fill=(230,220,200))
    dr.rectangle([0,Hh-120,W,Hh],fill=(0,0,0))
    dr.text((30,Hh-108),title,font=_font(40),fill="white")
    dr.text((30,Hh-56),sub+"   [Kuja UAT SAMPLE — fictional]",font=_font(24),fill=(210,210,210))
    im.save(path,quality=82); print("  IMG ", os.path.relpath(path,ROOT))

def receipt_image(path):
    W,Hh=680,900; im=Image.new("RGB",(W,Hh),"white"); dr=ImageDraw.Draw(im)
    dr.rectangle([0,0,W,Hh],outline=(180,180,180),width=2)
    y=40
    def line(t,f=22,c=(20,20,20),dx=40):
        nonlocal y; dr.text((dx,y),t,font=_font(f),fill=c); y+=f+12
    line("GARISSA PHARMACY LTD",30,(0,0,0)); line("Tax PIN P051xxxxxxA",18,(90,90,90))
    line("VAT Receipt  No. 0091827",18,(90,90,90)); y+=10
    dr.line([30,y,W-30,y],fill=(150,150,150)); y+=16
    for item,amt in [("Oxytocin inj. x50","420.00"),("Delivery kits x30","560.00"),
                     ("Chlorhexidine x40","160.00"),("Gloves (box) x10","100.00")]:
        dr.text((40,y),item,font=_font(20),fill=(20,20,20))
        dr.text((W-160,y),amt,font=_font(20),fill=(20,20,20)); y+=34
    y+=6; dr.line([30,y,W-30,y],fill=(150,150,150)); y+=16
    dr.text((40,y),"TOTAL (USD)",font=_font(24),fill=(0,0,0))
    dr.text((W-190,y),"1,240.00",font=_font(24),fill=(0,0,0)); y+=50
    line("Paid: M-PESA  •  12 Feb 2026",18,(90,90,90))
    line("[Kuja UAT SAMPLE — fictional receipt]",16,(150,150,150))
    im.save(path,quality=85); print("  IMG ", os.path.relpath(path,ROOT))

def attendance_image(path):
    W,Hh=1000,760; im=Image.new("RGB",(W,Hh),(248,248,244)); dr=ImageDraw.Draw(im)
    dr.text((30,24),"CHP Training — Attendance Sheet (Cohort 2)",font=_font(28),fill=(20,20,20))
    dr.text((30,64),"Facility: Sankuri  •  Date: 12 Feb 2026",font=_font(20),fill=(80,80,80))
    cols=[30,360,620,860]; y=110
    for c,t in zip(cols,["Name","Village","Phone (masked)","Sign"]):
        dr.text((c,y),t,font=_font(20),fill=(0,0,0))
    y+=34
    names=["A. Mohamed","H. Abdi","F. Noor","J. Kimani","S. Ali","M. Hassan","Z. Omar","K. Yusuf"]
    for n in names:
        dr.line([30,y,W-30,y],fill=(210,210,210))
        dr.text((30,y+6),n,font=_font(18),fill=(20,20,20))
        dr.text((360,y+6),"Sankuri",font=_font(18),fill=(20,20,20))
        dr.text((620,y+6),"07xx xxx "+str(random.randint(100,999)),font=_font(18),fill=(20,20,20))
        dr.text((860,y+6),"____",font=_font(18),fill=(120,120,120)); y+=52
    dr.text((30,Hh-40),"[Kuja UAT SAMPLE — fictional, masked numbers]",font=_font(16),fill=(150,150,150))
    im.save(path,quality=85); print("  IMG ", os.path.relpath(path,ROOT))

# ================================================================ 05 DONOR
def donor():
    d = cat("05_donor")
    build_pdf(os.path.join(d,"GlobalHealthFund_Call_for_Proposals_MCH.pdf"), [
        Paragraph("GLOBAL HEALTH FUND", ParagraphStyle('t',parent=styles['Title'],
                  fontSize=18,textColor=colors.HexColor('#7a2e2e'))),
        Paragraph("Call for Proposals — Maternal &amp; Newborn Health (Round 3)", H2),
        Spacer(1,6),
        *para_block("""Deadline for applications: 30 September 2026, 23:59 EAT.
        Grant ceiling: USD 200,000 per award. Duration: up to 24 months.
        Eligible applicants: nationally-registered NGOs operating in Kenya, Somalia or Ethiopia
        with at least three years of health programming and audited accounts.

        Thematic focus: reducing maternal and newborn mortality among underserved and
        displacement-affected populations.

        Evaluation criteria (weighted): Technical approach and evidence base (30%);
        Organisational capacity and past performance (25%); Value for money and budget (20%);
        Monitoring, evaluation and learning (15%); Safeguarding and risk management (10%).

        Required attachments: full proposal, detailed budget, logframe, latest audited accounts,
        registration certificate, and safeguarding policy.

        Applications are submitted through the Kuja Marketplace. A completed Trust &amp; Capacity
        assessment (Kuja Trust) is required before an application can be submitted."""),
        Spacer(1,10), Paragraph("[TEST DOCUMENT — fictional donor &amp; call, for Kuja UAT only]", SMALL),
    ])
    build_pdf(os.path.join(d,"GlobalHealthFund_Scoring_Rubric.pdf"), [
        Paragraph("Reviewer Scoring Rubric — MCH Round 3", H2),
        Table([["Criterion","Weight","0–5 scale anchors"],
               ["Technical approach","30%","0 weak / 3 adequate / 5 strong, evidence-based"],
               ["Organisational capacity","25%","0 unproven / 3 some track record / 5 strong"],
               ["Value for money","20%","0 unclear / 3 reasonable / 5 excellent"],
               ["MEL","15%","0 absent / 3 basic / 5 robust"],
               ["Safeguarding & risk","10%","0 none / 3 present / 5 comprehensive"]],
              colWidths=[45*mm,20*mm,95*mm],
              style=TableStyle([('BACKGROUND',(0,0),(-1,0),colors.HexColor('#7a2e2e')),
                                ('TEXTCOLOR',(0,0),(-1,0),colors.white),
                                ('GRID',(0,0),(-1,-1),0.4,colors.grey),
                                ('FONTSIZE',(0,0),(-1,-1),9),('VALIGN',(0,0),(-1,-1),'MIDDLE')])),
    ])

# ================================================================ 06 MEDIA (audio/video)
def media():
    d = cat("04_compliance_reporting")
    arial = r"C\:/Windows/Fonts/arial.ttf"
    mp4 = os.path.join(d,"field_evidence_clip.mp4")
    draw = (f"drawtext=fontfile='{arial}':text='Kuja UAT — sample field evidence':"
            "fontcolor=white:fontsize=26:x=(w-text_w)/2:y=h-60")
    cmds = [
        ["ffmpeg","-y","-f","lavfi","-i","testsrc=size=640x360:rate=15:duration=5",
         "-vf",draw,"-pix_fmt","yuv420p",mp4],
        ["ffmpeg","-y","-f","lavfi","-i","testsrc=size=640x360:rate=15:duration=5",
         "-pix_fmt","yuv420p",mp4],  # fallback without drawtext
    ]
    for c in cmds:
        try:
            r=subprocess.run(c,capture_output=True,timeout=90)
            if r.returncode==0 and os.path.exists(mp4): print("  MP4 ",os.path.relpath(mp4,ROOT)); break
        except Exception as e: print("   mp4 attempt failed",e)
    # voice memo (m4a + mp3) — tone stand-in for upload/format tests
    for ext,extra in (("m4a",["-c:a","aac"]),("mp3",[])):
        out=os.path.join(d,f"field_voice_memo.{ext}")
        try:
            subprocess.run(["ffmpeg","-y","-f","lavfi","-i",
                "sine=frequency=350:duration=6",*extra,out],capture_output=True,timeout=90)
            if os.path.exists(out): print("  AUD ",os.path.relpath(out,ROOT))
        except Exception as e: print("   audio failed",e)

# ================================================================ 99 EDGE CASES
def edge_cases(max_bytes=None):
    d = cat("99_edge_cases")
    # empty file with pdf extension
    open(os.path.join(d,"empty_file.pdf"),"wb").close()
    # corrupt pdf (valid-ish header then garbage)
    with open(os.path.join(d,"corrupt_truncated.pdf"),"wb") as f:
        f.write(b"%PDF-1.5\n%\xe2\xe3\xcf\xd3\n1 0 obj<< /Type /Catalog >>endobj\n")
        f.write(os.urandom(1500))  # no xref/trailer -> broken
    # text file masquerading as pdf
    with open(os.path.join(d,"not_really_a_pdf.pdf"),"w",encoding="utf-8") as f:
        f.write("This is plain text with a .pdf extension. A real parser should reject it.\n")
    # wrong format where a document is expected
    with open(os.path.join(d,"proposal_wrong_format.txt"),"w",encoding="utf-8") as f:
        f.write("Proposal submitted as .txt — should be rejected if only PDF/DOCX allowed.\n")
    # harmless executable (NOT malware — plain text renamed .exe)
    with open(os.path.join(d,"disguised_program.exe"),"wb") as f:
        f.write(b"MZ")  # DOS stub magic only; not a functional program
        f.write(b"\x00"*40)
        f.write(b"This is a harmless UAT stub, not a real executable.")
    # double extension
    with open(os.path.join(d,"report.pdf.exe"),"wb") as f:
        f.write(b"MZ"+b"\x00"*20+b"harmless UAT double-extension test")
    # SVG with script (stored-XSS-on-upload probe) — payload only shows an alert
    with open(os.path.join(d,"image_with_script.svg"),"w",encoding="utf-8") as f:
        f.write('<svg xmlns="http://www.w3.org/2000/svg" width="200" height="80">'
                '<rect width="200" height="80" fill="#eee"/>'
                '<text x="10" y="45">UAT SVG</text>'
                '<script>/*UAT XSS probe*/alert("uat-xss")</script></svg>')
    # HTML upload
    with open(os.path.join(d,"webpage_upload.html"),"w",encoding="utf-8") as f:
        f.write("<!doctype html><html><body><h1>UAT HTML upload test</h1>"
                "<script>console.log('uat')</script></body></html>")
    # very long filename
    longname = "very_long_filename_"+("a"*180)+".pdf"
    build_pdf(os.path.join(d,longname[:150]+".pdf"),
              [Paragraph("Long filename edge case.",BODY)], head=False)
    # Arabic (RTL) filename + Arabic content via docx (Word shapes Arabic natively)
    adoc=docx.Document()
    p=adoc.add_paragraph("تقرير ميداني تجريبي لاختبار قبول المستخدم — مبادرة أماني الصحية")
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    adoc.add_paragraph("هذا مستند تجريبي باللغة العربية للتحقق من عرض النص من اليمين إلى اليسار.")
    adoc.save(os.path.join(d,"تقرير_ميداني_عربي.docx"))
    print("  DOCX", "99_edge_cases/تقرير_ميداني_عربي.docx  (Arabic filename + RTL content)")
    # PDF carrying active content (JavaScript/OpenAction) -> Trust evidence room
    # rejects this (415 rejected_by_screening); Grant magic-byte check passes header.
    with open(os.path.join(d,"pdf_with_active_javascript.pdf"),"wb") as f:
        f.write(b"%PDF-1.4\n")
        f.write(b"1 0 obj << /Type /Catalog /Pages 2 0 R /OpenAction 4 0 R >> endobj\n")
        f.write(b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n")
        f.write(b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 300 200] >> endobj\n")
        f.write(b"4 0 obj << /Type /Action /S /JavaScript /JS (app.alert\\('UAT active-content probe'\\);) >> endobj\n")
        f.write(b"trailer << /Root 1 0 R >>\n%%EOF\n")
    print("  PDF  99_edge_cases/pdf_with_active_javascript.pdf  (Trust evidence security)")
    # password-protected PDF
    try:
        src=os.path.join(cat("03_grant_application"),"Amani_Project_Proposal_MCH.pdf")
        if os.path.exists(src):
            reader=PdfReader(src); writer=PdfWriter()
            for pg in reader.pages: writer.add_page(pg)
            writer.encrypt("uat-locked-123")
            with open(os.path.join(d,"password_protected.pdf"),"wb") as f: writer.write(f)
            print("  PDF  99_edge_cases/password_protected.pdf  (password: uat-locked-123)")
    except Exception as e: print("   encrypted pdf failed",e)
    # oversized file (size depends on the app limit; default target ~30MB)
    target = max_bytes if max_bytes else 30*1024*1024
    with open(os.path.join(d,f"oversized_{target//(1024*1024)}MB.pdf"),"wb") as f:
        f.write(b"%PDF-1.5\n");
        f.write(b"0"*(target))
    print(f"  BIN  99_edge_cases/oversized_{target//(1024*1024)}MB.pdf")
    for fn in os.listdir(d):
        pass
    print("  (edge-case files written)")

# ---------------------------------------------------------------- manifest + readme
def manifest_and_readme():
    import csv
    rows=[]
    for base,_,files in os.walk(ROOT):
        for fn in sorted(files):
            if fn in ("MANIFEST.csv",): continue
            fp=os.path.join(base,fn); rel=os.path.relpath(fp,ROOT)
            rows.append((rel.replace("\\","/"), os.path.getsize(fp)))
    with open(os.path.join(ROOT,"MANIFEST.csv"),"w",newline="",encoding="utf-8") as f:
        w=csv.writer(f); w.writerow(["file","bytes"])
        for r in rows: w.writerow(r)
    with open(os.path.join(ROOT,"_READ_ME.txt"),"w",encoding="utf-8") as f:
        f.write("KUJA MARKETPLACE — UAT TEST FILE PACK\n")
        f.write("="*40+"\n\n")
        f.write("All files are FICTIONAL and for testing only. Org 'Amani Health Initiative'\n")
        f.write("and donor 'Global Health Fund' do not exist; numbers/persons are invented.\n\n")
        f.write("Folders:\n")
        f.write("  01_registration        - org onboarding documents (NGO)\n")
        f.write("  02_capacity_trust_evidence - policies/letters for the Kuja Trust assessment\n")
        f.write("  03_grant_application    - proposal, budget, logframe, workplan, MOU\n")
        f.write("  04_compliance_reporting - narrative/financial reports, receipts, photos,\n")
        f.write("                            video + voice memo (post-award reporting)\n")
        f.write("  05_donor                - grant Call-for-Proposals PDF (AI-extraction test)\n")
        f.write("                            + reviewer scoring rubric\n")
        f.write("  99_edge_cases           - deliberately broken/unsafe files (see workbook)\n\n")
        f.write("Map each file to a test case using the 'Test Data / Files' column in the\n")
        f.write("UAT workbook (Kuja_UAT_Test_Plan.xlsx).\n")
    print("  MANIFEST + _READ_ME written;", len(rows), "files total")

if __name__ == "__main__":
    print("Generating Kuja UAT test-file pack ->", ROOT)
    registration(); capacity_trust(); application(); compliance(); donor(); media()
    edge_cases(max_bytes=20*1024*1024)  # app MAX_CONTENT_LENGTH=16MB -> 20MB triggers 413
    manifest_and_readme()
    print("DONE")
