"""
Kuja Grant Management — Client Demo Guide Generator
Creates a professionally formatted Word document with demo scenarios and steps.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# Branding colors
PRIMARY = RGBColor(0x1E, 0x40, 0xAF)
SECONDARY = RGBColor(0x0D, 0x94, 0x88)
ACCENT = RGBColor(0xF5, 0x9E, 0x0B)
DARK = RGBColor(0x1F, 0x2A, 0x37)
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF0, 0xF9, 0xFF)

FONT = "Calibri"


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def add_styled_paragraph(doc, text, size=11, bold=False, color=DARK, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT
        run.font.color.rgb = PRIMARY
    return h


def add_step_table(doc, steps):
    """Create a numbered steps table with alternating row colors."""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for cell in table.columns[0].cells:
        cell.width = Cm(1.5)
    for cell in table.columns[1].cells:
        cell.width = Cm(4.0)
    for cell in table.columns[2].cells:
        cell.width = Cm(11.0)

    # Header row
    hdr = table.rows[0]
    for i, txt in enumerate(["Step", "Action", "Details"]):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.name = FONT
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, "1E40AF")

    # Data rows
    for idx, (action, details) in enumerate(steps):
        row = table.add_row()
        # Step number
        c0 = row.cells[0]
        c0.text = ""
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(str(idx + 1))
        r0.font.name = FONT
        r0.font.size = Pt(11)
        r0.font.bold = True
        r0.font.color.rgb = PRIMARY
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Action
        c1 = row.cells[1]
        c1.text = ""
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(action)
        r1.font.name = FONT
        r1.font.size = Pt(10)
        r1.font.bold = True
        r1.font.color.rgb = DARK

        # Details
        c2 = row.cells[2]
        c2.text = ""
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(details)
        r2.font.name = FONT
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK

        # Alternating row colors
        bg = "F0F9FF" if idx % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            set_cell_shading(cell, bg)

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): 'D1D5DB',
        })
        borders.append(element)
    tblPr.append(borders)

    return table


def add_callout(doc, text, icon="💡"):
    """Add a highlighted callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # Use a table for the callout background
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(f"{icon}  {text}")
    run.font.name = FONT
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = PRIMARY
    set_cell_shading(cell, "EFF6FF")
    return table


def add_file_info(doc, filename, description):
    """Add file reference with icon."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"   📄  {filename}")
    run.font.name = FONT
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = SECONDARY
    run2 = p.add_run(f"  —  {description}")
    run2.font.name = FONT
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY


def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ─── COVER / TITLE ─────────────────────────────────────────────────
    doc.add_paragraph()  # spacer
    add_styled_paragraph(doc, "KUJA GRANT MANAGEMENT", size=28, bold=True,
                         color=PRIMARY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=40)
    add_styled_paragraph(doc, "Client Demo Guide", size=20, bold=False,
                         color=SECONDARY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=12)
    add_styled_paragraph(doc, "AI-Powered Grant Management for Impact",
                         size=13, italic=True, color=GRAY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    # Separator line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 60)
    run.font.color.rgb = SECONDARY
    run.font.size = Pt(8)

    add_styled_paragraph(doc, "Prepared by Adeso  |  March 2026", size=11,
                         color=GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=6)
    add_styled_paragraph(doc, "Live Platform: https://web-production-6f8a.up.railway.app",
                         size=10, color=GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ─── OVERVIEW ───────────────────────────────────────────────────────
    add_heading(doc, "Demo Overview", level=1)

    add_styled_paragraph(doc,
        "This guide walks through five demo scenarios that showcase Kuja's core "
        "capabilities. Each scenario is self-contained and can be presented "
        "independently. Total demo time: approximately 25-30 minutes.",
        size=11, color=DARK, space_after=12)

    # Scenario summary table
    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = summary_table.rows[0]
    for i, txt in enumerate(["#", "Scenario", "Login As", "Time"]):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(txt)
        r.font.name = FONT
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1E40AF")

    scenarios = [
        ("1", "Donor Creates a Grant with AI Document Requirements",
         "sarah@globalhealth.org (Donor)", "~8 min"),
        ("2", "NGO Capacity Assessment with AI Analysis",
         "fatima@amani.org (NGO)", "~5 min"),
        ("3", "Live Due Diligence: Sanctions + Registry Verification",
         "sarah@globalhealth.org (Donor)", "~5 min"),
        ("4", "NGO Grant Reporting with AI Compliance Scoring",
         "fatima@amani.org (NGO)", "~5 min"),
        ("5", "Multi-Language Demo (Arabic, French, Spanish)",
         "Any account", "~3 min"),
    ]
    for idx, (num, scenario, login, time) in enumerate(scenarios):
        row = summary_table.add_row()
        vals = [num, scenario, login, time]
        for i, v in enumerate(vals):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(v)
            r.font.name = FONT
            r.font.size = Pt(10)
            r.font.color.rgb = DARK
            if i == 0:
                r.font.bold = True
                r.font.color.rgb = PRIMARY
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bg = "F0F9FF" if idx % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            set_cell_shading(cell, bg)

    # Add borders to summary table
    tbl = summary_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single', qn('w:sz'): '4',
            qn('w:space'): '0', qn('w:color'): 'D1D5DB',
        })
        borders.append(element)
    tblPr.append(borders)

    doc.add_paragraph()  # spacer

    add_styled_paragraph(doc, "Demo Accounts (all passwords: pass123)", size=11,
                         bold=True, color=PRIMARY, space_before=8)
    accounts = [
        ("Donor", "sarah@globalhealth.org", "Sarah Mitchell — Global Health Foundation"),
        ("NGO", "fatima@amani.org", "Fatima Hassan — Amani Community Development (Kenya)"),
        ("Reviewer", "james@reviewer.org", "James Wilson — Independent Reviewer"),
        ("Admin", "admin@kuja.org", "System Administrator"),
    ]
    for role, email, desc in accounts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"  {role}: ")
        r1.font.name = FONT
        r1.font.size = Pt(10)
        r1.font.bold = True
        r1.font.color.rgb = SECONDARY
        r2 = p.add_run(f"{email}  ")
        r2.font.name = FONT
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK
        r3 = p.add_run(f"({desc})")
        r3.font.name = FONT
        r3.font.size = Pt(9)
        r3.font.color.rgb = GRAY

    add_styled_paragraph(doc, "Files Included in Demo Kit (demo-files.zip)", size=11,
                         bold=True, color=PRIMARY, space_before=12)
    demo_files = [
        ("grant_agreement_sample.txt", "Sample $500K grant agreement with reporting requirements"),
        ("project_proposal.txt", "MCH project proposal for Wajir County, Kenya ($350K)"),
        ("registration_certificate.txt", "Kenya NGO Board registration certificate"),
        ("financial_report_q1_2026.txt", "Q1 2026 quarterly financial report"),
        ("excellent_narrative_report.txt", "High-quality annual narrative report (Somalia health program)"),
        ("strategic_plan_2025_2030.txt", "5-year organizational strategic plan"),
    ]
    for fname, desc in demo_files:
        add_file_info(doc, fname, desc)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # SCENARIO 1: DONOR GRANT WIZARD
    # ═══════════════════════════════════════════════════════════════════
    add_heading(doc, "Scenario 1: Donor Creates a Grant with AI", level=1)

    add_styled_paragraph(doc,
        "Demonstrate the donor experience: creating a new grant opportunity with "
        "AI-generated evaluation criteria and document requirements. This showcases "
        "the 5-step grant wizard and AI integration.",
        size=11, color=DARK, space_after=8)

    add_callout(doc,
        "Talking Point: \"Donors configure grants once. Our AI generates tailored "
        "evaluation criteria for each required document type — proposals, budgets, "
        "certifications — so every application is scored consistently.\"")

    doc.add_paragraph()
    add_styled_paragraph(doc, "Pre-requisite: Log in as the Donor account",
                         size=10, bold=True, color=SECONDARY)

    add_step_table(doc, [
        ("Log in as Donor",
         "Navigate to the platform. Enter email: sarah@globalhealth.org and "
         "password: pass123. Click Sign In. You will see the Donor Dashboard "
         "showing Total Grants, Applications, Pending Review, and Total Awarded."),

        ("Open Grant Wizard",
         "Click '+ Create Grant' in the left sidebar. The 5-step wizard opens. "
         "Step 1 shows Basic Information fields."),

        ("Fill Basic Info (Step 1)",
         "Grant Title: Community Health Workers Scale-Up — East Africa\n"
         "Description: Scale-up of CHW programs across Kenya, Somalia, and Uganda...\n"
         "Total Funding: 500000\n"
         "Currency: USD\n"
         "Deadline: 2026-06-30\n"
         "Sectors: Click 'Health'\n"
         "Countries: Click 'Kenya' and 'Somalia'\n"
         "Click 'Next'."),

        ("Configure Requirements (Step 2)",
         "Step 2: Application Requirements. Select which document types applicants "
         "must submit. Toggle ON: Project Proposal, Budget, Registration Certificate, "
         "Organizational Capacity Statement. Each document type will get AI-generated "
         "evaluation criteria in the next step."),

        ("AI Evaluation Criteria (Step 3)",
         "Step 3: AI generates evaluation criteria for each document type. "
         "WAIT for AI to process — you'll see a loading spinner. Once complete, "
         "review the criteria. Each document type shows weighted scoring criteria "
         "(e.g., 'Problem Statement Clarity: 20%', 'Budget Justification: 15%'). "
         "You can edit weights or add custom criteria."),

        ("Upload Grant Agreement (Step 4)",
         "Step 4: Upload the grant agreement document. Use the file "
         "'grant_agreement_sample.txt' from the demo kit. "
         "AI will extract reporting requirements, deadlines, and disbursement schedules "
         "automatically. Review the extracted data."),

        ("Review & Publish (Step 5)",
         "Step 5 shows a summary of everything configured. Review all sections. "
         "Click 'Publish Grant' to make it live. The grant now appears on the "
         "NGO marketplace and accepting applications. You can also 'Save as Draft' "
         "to return later."),

        ("Show Grant Details",
         "Navigate to 'My Grants' in the sidebar. Click the newly created grant. "
         "Show the grant details page with requirements, AI criteria, and the "
         "published status."),
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # SCENARIO 2: NGO CAPACITY ASSESSMENT
    # ═══════════════════════════════════════════════════════════════════
    add_heading(doc, "Scenario 2: NGO Capacity Assessment with AI", level=1)

    add_styled_paragraph(doc,
        "Demonstrate the NGO experience: running an AI-powered capacity assessment "
        "using one of five industry-standard frameworks. The AI analyzes uploaded "
        "organizational documents and generates scored results with improvement roadmaps.",
        size=11, color=DARK, space_after=8)

    add_callout(doc,
        "Talking Point: \"NGOs complete ONE assessment and share it with any donor. "
        "No more filling out 4-12 duplicate questionnaires per year. Five frameworks "
        "today — Kuja, STEP, UN-HACT, CHS, NUPAS — and our system is designed to "
        "add new frameworks on demand.\"")

    doc.add_paragraph()
    add_styled_paragraph(doc, "Pre-requisite: Log out and log in as the NGO account",
                         size=10, bold=True, color=SECONDARY)

    add_step_table(doc, [
        ("Log in as NGO",
         "Log out from the donor account. Enter email: fatima@amani.org and "
         "password: pass123. Click Sign In. The NGO Dashboard shows a different "
         "layout: available grants, assessment status, applications, and reports."),

        ("Navigate to Assessments",
         "Click 'Capacity Assessment' in the left sidebar. You'll see the list of "
         "available assessment frameworks and any previous assessments."),

        ("Select a Framework",
         "Click 'Start New Assessment'. Choose the 'Kuja Framework' (the proprietary "
         "comprehensive framework). Other options visible: STEP (TechSoup), UN-HACT, "
         "CHS Alliance, NUPAS. Point out that all 5 are available in one platform."),

        ("Complete Assessment Questions",
         "The assessment questionnaire appears with sections covering governance, "
         "financial management, program delivery, M&E, and HR. Answer the questions "
         "(select ratings or enter text). For the demo, fill a few sections to show "
         "the interface."),

        ("Upload Documents for AI Analysis",
         "In the document upload section, upload 'strategic_plan_2025_2030.txt' from "
         "the demo kit. The AI will analyze the document against the framework criteria. "
         "WAIT for AI processing (10-15 seconds). The AI returns per-criterion scores "
         "(0-100) with specific findings."),

        ("Review AI Results",
         "Once AI analysis completes, review the results: overall score, per-section "
         "breakdown, identified gaps, and a 30/60/90-day improvement roadmap. Point "
         "out that these results are 'passportable' — shareable with any donor from "
         "a single assessment."),

        ("Show Donor Readiness",
         "Scroll to the Donor Readiness section. The assessment maps to readiness "
         "levels for different donor types: private philanthropy, bilateral, "
         "multilateral, and UN system. This helps NGOs understand which funding "
         "opportunities they qualify for."),
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # SCENARIO 3: LIVE DUE DILIGENCE
    # ═══════════════════════════════════════════════════════════════════
    add_heading(doc, "Scenario 3: Live Due Diligence", level=1)

    add_styled_paragraph(doc,
        "Demonstrate real-time compliance verification: sanctions screening against "
        "5 international databases and government registry verification. This replaces "
        "manual due diligence processes that cost donors $5,000-$15,000 per organization.",
        size=11, color=DARK, space_after=8)

    add_callout(doc,
        "Talking Point: \"This runs against LIVE databases — UN, OFAC, EU, World Bank "
        "sanctions lists and government registries in 7 African countries today. We're "
        "designed to expand to any Global South country with a verification portal. "
        "No other platform offers this.\"")

    doc.add_paragraph()
    add_styled_paragraph(doc, "Pre-requisite: Log in as the Donor account",
                         size=10, bold=True, color=SECONDARY)

    add_step_table(doc, [
        ("Log in as Donor",
         "Navigate to the platform. Enter email: sarah@globalhealth.org and "
         "password: pass123. Click Sign In."),

        ("Navigate to Compliance",
         "Click 'Compliance' in the left sidebar. This opens the compliance "
         "screening dashboard."),

        ("Run Sanctions Screening",
         "Click 'Screen Organization'. Enter organization name: "
         "'Amani Community Development' and country: 'Kenya'. "
         "Click 'Screen'. The system queries 5 databases in real-time:\n"
         "• UN Security Council Consolidated List\n"
         "• US OFAC SDN List\n"
         "• EU Financial Sanctions List\n"
         "• World Bank Debarment List\n"
         "• OpenSanctions (aggregated)\n"
         "Results show per-database status (Clear/Flagged) with match scores."),

        ("Review Screening Results",
         "The results panel shows an overall status. Each database check displays: "
         "source name, match score, matched entity (if any), and status. "
         "A clean organization shows 'CLEAR' across all databases. "
         "Point out that false positives show match details for manual review."),

        ("Navigate to Registration Checks",
         "Click 'Registration Checks' in the sidebar. This shows the "
         "government registry verification module."),

        ("Run Registry Verification",
         "Select an organization from the list (e.g., Amani Community Development). "
         "Click 'Verify Registration'. The system checks against the Kenya NGO "
         "Coordination Board registry using the org's registration number. "
         "Results show verification status, registry source, and extracted details."),

        ("Show Supported Countries",
         "Point out the list of supported registries: Kenya (NGO Board/BRS), "
         "Nigeria (CAC), South Africa (DSD NPO), Uganda (NGO Bureau), "
         "Tanzania (NiS), Somalia (MOIFAR), Ethiopia (ACSO). "
         "Emphasize: designed to expand to all Global South countries."),

        ("Upload Registration Certificate",
         "For deeper verification, upload 'registration_certificate.txt' from the "
         "demo kit. AI analyzes the certificate, extracting: organization name, "
         "registration number, date, validity, registered purposes. Cross-references "
         "with registry data."),
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # SCENARIO 4: NGO REPORTING + AI EVALUATION
    # ═══════════════════════════════════════════════════════════════════
    add_heading(doc, "Scenario 4: NGO Grant Reporting with AI", level=1)

    add_styled_paragraph(doc,
        "Demonstrate the reporting workflow: NGO submits grant reports (narrative "
        "and financial), AI evaluates compliance against grant agreement requirements, "
        "and generates per-requirement scores with risk flags.",
        size=11, color=DARK, space_after=8)

    add_callout(doc,
        "Talking Point: \"Our AI doesn't just store reports — it reads them, scores "
        "each requirement from the grant agreement, flags risks, and gives both the "
        "NGO and donor a compliance score. This turns weeks of manual review into "
        "minutes.\"")

    doc.add_paragraph()
    add_styled_paragraph(doc, "Pre-requisite: Log in as the NGO account",
                         size=10, bold=True, color=SECONDARY)

    add_step_table(doc, [
        ("Log in as NGO",
         "Log out from the donor account. Enter email: fatima@amani.org and "
         "password: pass123. Click Sign In."),

        ("Navigate to Reports",
         "Click 'Grant Reports' in the left sidebar. The reports dashboard shows "
         "upcoming deadlines, submitted reports, and compliance status per grant."),

        ("Select a Grant for Reporting",
         "You should see the 'Community Health Workers Scale-Up Program' grant. "
         "Click 'Submit Report' or 'New Report'. Select the reporting period."),

        ("Upload Narrative Report",
         "In the report submission form, upload 'excellent_narrative_report.txt' "
         "from the demo kit as the narrative report. This is a comprehensive annual "
         "report from a Somalia health program with detailed results, context "
         "analysis, and financial summary."),

        ("Upload Financial Report",
         "Upload 'financial_report_q1_2026.txt' from the demo kit as the financial "
         "report. This shows a quarterly financial report with budget vs. actual, "
         "variance analysis, and cash flow forecast."),

        ("Trigger AI Evaluation",
         "Click 'Submit Report'. The AI evaluates both documents against the grant "
         "agreement's specific requirements. WAIT for AI processing (15-20 seconds). "
         "The system generates:\n"
         "• Overall compliance score (0-100)\n"
         "• Per-requirement breakdown with scores\n"
         "• Risk flags for under-performing areas\n"
         "• Specific findings with quotes from the report"),

        ("Review AI Compliance Results",
         "The results page shows a compliance dashboard: overall score, per-section "
         "scores (financial reporting, narrative quality, M&E data, etc.), and "
         "any flagged risks. The AI provides specific recommendations for "
         "improvement."),

        ("Show Donor View (Optional)",
         "Log out and back in as the donor (sarah@globalhealth.org). Navigate to "
         "'Grant Reports'. You'll see the submitted report with AI scores visible "
         "to the donor. The donor can review the AI analysis and add their own "
         "review comments."),
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # SCENARIO 5: MULTI-LANGUAGE
    # ═══════════════════════════════════════════════════════════════════
    add_heading(doc, "Scenario 5: Multi-Language Demo", level=1)

    add_styled_paragraph(doc,
        "Demonstrate Kuja's multi-language support — a key differentiator that "
        "no other grant management platform offers. The platform currently supports "
        "English, French, Arabic, and Spanish, with more languages coming.",
        size=11, color=DARK, space_after=8)

    add_callout(doc,
        "Talking Point: \"No other grant management platform offers Arabic, French, "
        "or Spanish. Our users are in the Global South — they need tools in their "
        "languages. We have 5 languages live today and are adding more every week.\"")

    doc.add_paragraph()
    add_styled_paragraph(doc, "Pre-requisite: Be logged into any account",
                         size=10, bold=True, color=SECONDARY)

    add_step_table(doc, [
        ("Show Language Selector",
         "Point out the language selector in the top navigation bar. It shows "
         "the current language (e.g., 'English') with a globe icon."),

        ("Switch to Arabic",
         "Click the language selector and choose 'العربية' (Arabic). "
         "The entire interface switches to Arabic with right-to-left (RTL) layout. "
         "Navigation, labels, buttons, dashboard — everything is translated. "
         "Point out that even the AI assistant responds in Arabic."),

        ("Switch to French",
         "Click the language selector and choose 'Français' (French). "
         "The interface switches to French. Navigate to a few pages to show "
         "consistent translation across the platform — dashboard, grants list, "
         "forms, and reports all render in French."),

        ("Switch to Spanish",
         "Click the language selector and choose 'Español' (Spanish). "
         "Show the Spanish interface. Emphasize that this is live today — not "
         "a mockup or roadmap item."),

        ("Return to English",
         "Switch back to English to continue with any follow-up discussion. "
         "Emphasize: 675 translation keys per language, covering the entire "
         "platform. More languages (Swahili, Somali, Portuguese) in the pipeline."),
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # TIPS & TROUBLESHOOTING
    # ═══════════════════════════════════════════════════════════════════
    add_heading(doc, "Demo Tips & Troubleshooting", level=1)

    add_styled_paragraph(doc, "Before the Demo", size=13, bold=True, color=PRIMARY,
                         space_before=8)
    tips_before = [
        "Open the platform URL in Chrome (recommended) and verify you can log in",
        "Have the demo-files.zip extracted and files ready to upload",
        "Test your internet connection — sanctions screening and AI features require connectivity",
        "Pre-load the platform in a browser tab so there's no loading delay when starting",
        "If using the live Railway deployment, verify it's accessible at the URL above",
    ]
    for tip in tips_before:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"  ✓  {tip}")
        r.font.name = FONT
        r.font.size = Pt(10)
        r.font.color.rgb = DARK

    add_styled_paragraph(doc, "During the Demo", size=13, bold=True, color=PRIMARY,
                         space_before=12)
    tips_during = [
        "AI features take 10-20 seconds to process — use this time to explain what's happening",
        "If sanctions screening shows a 'flagged' result, explain it's a real-time check and "
        "false positives are expected (the system shows match scores for manual review)",
        "Emphasize the 8-stage lifecycle: Discovery → Assessment → Due Diligence → "
        "Matching → Application → Review → Reporting → ERP",
        "Key differentiators to highlight: only end-to-end platform, AI throughout, "
        "multi-language (no competitor has this), live due diligence, 5+ frameworks",
    ]
    for tip in tips_during:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"  ✓  {tip}")
        r.font.name = FONT
        r.font.size = Pt(10)
        r.font.color.rgb = DARK

    add_styled_paragraph(doc, "If Something Goes Wrong", size=13, bold=True, color=PRIMARY,
                         space_before=12)
    fallbacks = [
        ("AI is slow or times out", "Explain that AI processing happens in real-time via "
         "Anthropic's Claude API. If slow, skip to next scenario and return later."),
        ("Sanctions screening error", "The system has fallback: if OpenSanctions API is "
         "down, it falls back to direct UN XML, OFAC CSV, and EU CSV file parsing."),
        ("Session expires", "Simply log back in — all saved data persists. Sessions "
         "expire after inactivity for security."),
        ("Platform is down", "Use the local development server as backup. Run: "
         "python run.py from the kuja-grant directory."),
    ]
    for issue, fix in fallbacks:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"  ⚠  {issue}: ")
        r1.font.name = FONT
        r1.font.size = Pt(10)
        r1.font.bold = True
        r1.font.color.rgb = ACCENT
        r2 = p.add_run(fix)
        r2.font.name = FONT
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK

    # ─── FOOTER ─────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 60)
    run.font.color.rgb = SECONDARY
    run.font.size = Pt(8)
    add_styled_paragraph(doc,
        "Kuja Grant Management  |  Built by Adeso  |  Confidential",
        size=9, color=GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


def main():
    doc = build_document()
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, "Kuja_Demo_Guide.docx")
    doc.save(output_path)
    file_size = os.path.getsize(output_path)
    print(f"Generated: {output_path}")
    print(f"File size: {file_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
