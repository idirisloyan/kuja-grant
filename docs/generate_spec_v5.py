"""Kuja Grant — v5.0 Technical Design Specification generator.

A comprehensive technical spec that covers every aspect of the system:
architecture, design, database schema, tech stack, security stack,
external integrations, AI usage, scoring methodologies, calculations,
deployment, and observability.

Output: Kuja_Grant_v5.0_Technical_Design_Specification.docx
"""

import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Cm

# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------
KUJA_NAVY = RGBColor(26, 35, 126)
KUJA_NAVY_HEX = "1A237E"
KUJA_CLAY = RGBColor(194, 65, 12)
KUJA_CLAY_HEX = "C2410C"
INK = RGBColor(33, 33, 33)
MUTED = RGBColor(96, 96, 96)


def _shade(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    if level == 1:
        for run in h.runs:
            run.font.color.rgb = KUJA_NAVY
            run.font.size = Pt(20)
    elif level == 2:
        for run in h.runs:
            run.font.color.rgb = KUJA_NAVY
            run.font.size = Pt(15)
    elif level == 3:
        for run in h.runs:
            run.font.color.rgb = KUJA_CLAY
            run.font.size = Pt(12.5)


def body(doc, text, *, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = INK
    if italic:
        run.italic = True
    return p


def bullets(doc, items):
    for item in items:
        if isinstance(item, tuple):
            label, rest = item
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(label)
            r.bold = True
            r.font.size = Pt(10.5)
            r2 = p.add_run(" " + rest)
            r2.font.size = Pt(10.5)
        else:
            p = doc.add_paragraph(item, style="List Bullet")
            for r in p.runs:
                r.font.size = Pt(10.5)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for r in p.runs:
            r.font.size = Pt(10.5)


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = INK
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)


def callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label + ":  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = KUJA_CLAY
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.italic = True
    r2.font.color.rgb = MUTED


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        _shade(cell, KUJA_NAVY_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.bold = True
        r.font.size = Pt(10)
    # Body rows
    for ri, row in enumerate(rows, start=1):
        for ci, value in enumerate(row):
            cell = t.rows[ri].cells[ci]
            p = cell.paragraphs[0]
            p.add_run(str(value)).font.size = Pt(9.5)
    # Column widths
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)
    return t


def page_break(doc):
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # Page setup
    for s in doc.sections:
        s.page_width = Cm(21)
        s.page_height = Cm(29.7)
        s.top_margin = Cm(2.2)
        s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # COVER
    for _ in range(6):
        doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("KUJA GRANT MANAGEMENT SYSTEM")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = KUJA_NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Technical Design Specification — v5.0")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = KUJA_CLAY
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Complete architectural, security, AI, data, and "
                     "integration reference for the Kuja platform.")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    for _ in range(10):
        doc.add_paragraph("")
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d')}  ·  "
                     f"Reflects code at commit time of generation")
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED
    page_break(doc)

    # =================================================================
    # TABLE OF CONTENTS
    # =================================================================
    heading(doc, "Table of Contents", 1)
    toc = [
        "1. Executive Technical Summary",
        "2. Technology Stack",
        "    2.1 Backend",
        "    2.2 Frontend",
        "    2.3 Data layer",
        "    2.4 Infrastructure + deployment",
        "    2.5 Frontend libraries (full inventory)",
        "    2.6 Python dependencies (full inventory)",
        "3. System Architecture",
        "    3.1 Layered architecture",
        "    3.2 Request lifecycle",
        "    3.3 Blueprint catalogue",
        "    3.4 Service catalogue",
        "    3.5 Frontend route map",
        "4. Database Design",
        "    4.1 ER overview",
        "    4.2 Entity reference (31 models)",
        "    4.3 Key indexes + constraints",
        "    4.4 JSON-typed columns",
        "    4.5 Migration strategy",
        "5. Security Architecture",
        "    5.1 Authentication stack",
        "    5.2 Authorisation + role gating",
        "    5.3 Session + CSRF + headers",
        "    5.4 Rate limiting + lockouts",
        "    5.5 File upload security",
        "    5.6 GDPR + data deletion",
        "    5.7 Hash-chained audit + provenance",
        "6. AI Integration Architecture",
        "    6.1 AI service stack",
        "    6.2 Sync vs async dispatch",
        "    6.3 Tool-use agent + read-only registry",
        "    6.4 Cost budget + rate limits",
        "    6.5 Anti-hallucination discipline",
        "    6.6 AI surface inventory (28 surfaces)",
        "    6.7 Sustained chat threads",
        "    6.8 AI claim provenance ledger",
        "7. Scoring + Calculation Methodologies",
        "    7.1 Application AI scoring",
        "    7.2 5-framework capacity assessment scoring",
        "    7.3 Trust Profile aggregation",
        "    7.4 Compliance Health (4 pillars)",
        "    7.5 Reviewer match ranking",
        "    7.6 Peer benchmarks",
        "    7.7 Donor cohort analytics",
        "    7.8 NPS calculation",
        "    7.9 Funnel + retention math",
        "    7.10 A/B bucketing",
        "8. External Integrations + APIs",
        "    8.1 Anthropic Claude",
        "    8.2 OpenSanctions + fallback (UN/OFAC/EU)",
        "    8.3 Government registries (7 African countries)",
        "    8.4 Sentry",
        "    8.5 Web Push",
        "    8.6 Email transport",
        "    8.7 SMS + WhatsApp adapters",
        "    8.8 Railway platform",
        "    8.9 GitHub Actions (CI + crons)",
        "9. Real-User Metrics + Feedback Infrastructure",
        "10. Internationalisation (i18n)",
        "11. Frontend Architecture",
        "12. Background Jobs + Crons",
        "13. Observability + Operations",
        "14. Deployment Pipeline",
        "15. Performance + Scalability",
        "16. Testing Architecture",
        "17. Tech Debt + Future Hardening",
        "Appendix A: Environment Variables",
        "Appendix B: API endpoint census",
        "Appendix C: Tools inventory (what + why)",
    ]
    for line in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        r.font.size = Pt(10)
        if not line.startswith(" "):
            r.bold = True
            r.font.color.rgb = KUJA_NAVY
    page_break(doc)

    # =================================================================
    # 1. EXECUTIVE TECHNICAL SUMMARY
    # =================================================================
    heading(doc, "1. Executive Technical Summary", 1)

    body(doc,
        "Kuja is a multi-tenant, role-aware grant-management platform "
        "delivered as a single web application with a static-export "
        "Next.js frontend served by a Python/Flask backend. The platform "
        "is built around a relational PostgreSQL data model "
        "(SQLite for local dev), an Anthropic Claude AI layer woven "
        "into 28 product surfaces, a continuous compliance and due-"
        "diligence engine that integrates with OpenSanctions and seven "
        "African government registries, and a real-user metrics layer "
        "that captures behavioural events + NPS feedback to drive "
        "data-driven product decisions."
    )

    body(doc,
        "The system runs in production on Railway behind their managed "
        "PostgreSQL service, with deploys driven by `railway up` and "
        "smoke + browser regression gates running on GitHub Actions. "
        "Error tracking is wired through Sentry with Flask + SQLAlchemy "
        "integrations. The deployed unit is a single Gunicorn worker — "
        "AI generation calls dispatch to an in-process task runner "
        "(returning 202 + job id) so user requests never block on "
        "long-running model calls."
    )

    body(doc,
        "Every meaningful platform action writes to a tamper-evident "
        "hash-chained audit log, every AI-generated claim is linked "
        "back to its source documents through an AIProvenance ledger, "
        "and every behavioural event is captured by a lightweight "
        "UserEvent table that powers DAU/WAU/MAU + 6 funnels + per-"
        "language adoption + A/B outcome splits in the admin metrics "
        "dashboard. The platform supports six interface languages with "
        "role-appropriate tone and ships as a PWA installable on "
        "mobile devices."
    )

    body(doc,
        "This document is the authoritative technical reference. "
        "Sections 1–5 cover stack + architecture + database + security. "
        "Section 6 covers AI integration end-to-end. Section 7 details "
        "every scoring algorithm + calculation. Section 8 enumerates "
        "every external integration. Sections 9–17 cover metrics, i18n, "
        "frontend, jobs, observability, deployment, performance, testing, "
        "and tech debt. The appendices provide an environment-variable "
        "table, API endpoint census, and a tools inventory."
    )

    page_break(doc)

    # =================================================================
    # 2. TECHNOLOGY STACK
    # =================================================================
    heading(doc, "2. Technology Stack", 1)

    heading(doc, "2.1 Backend", 2)
    table(doc,
        ["Layer", "Choice", "Why"],
        [
            ["Language", "Python 3.11+", "Mature ecosystem for AI SDK, scientific libs, ops tooling. Used by Anthropic SDK natively."],
            ["Web framework", "Flask 3.1.0", "Lightweight, blueprint-based, well-understood patterns for role-gated APIs. Fits the per-feature boundary we want."],
            ["ORM", "Flask-SQLAlchemy 3.1.1 + SQLAlchemy 2.0.36", "Declarative models with the new 2.x typed Mapped API; supports both Postgres + SQLite from the same code."],
            ["Auth session", "Flask-Login 0.6.3", "Server-side sessions with secure cookies; simple integration with role checks."],
            ["DB migrations", "Flask-Migrate 4.0.7 (Alembic)", "Versioned schema migrations; supports auto-generation from model changes."],
            ["CORS", "Flask-CORS 5.0.1", "Cross-origin policy management for the SPA frontend served at the same origin in prod but separate in dev."],
            ["WSGI server", "Gunicorn 23.0.0", "Battle-tested HTTP server; chosen over uvicorn because Flask is sync-style. Single worker by design (see §15)."],
            ["Templating", "Jinja2 (via Flask)", "Used only for the small handful of legacy server-rendered surfaces + the email template."],
            ["AI SDK", "anthropic 0.42.0", "Official Python SDK; supports tool-use, structured output, vision."],
            ["HTTP client", "requests 2.32.3", "External API calls (OpenSanctions, government registries, email transport)."],
            ["PDF generation", "reportlab 4.2.5", "Portfolio bundle PDFs, calendar PDFs, reviewer take-home single-app PDFs."],
            ["PDF parsing", "PyPDF2 3.0.1", "Extract text from uploaded grant agreements, applications, supporting docs."],
            ["DOCX parsing", "python-docx 1.1.2", "Extract narrative from uploaded .docx files; also used to generate BRD + test cases."],
            ["XLSX parsing + generation", "openpyxl 3.1.5", "Read uploaded budget files; generate the test-cases workbook + bundle fixtures."],
            ["Postgres driver", "psycopg2-binary 2.9.10", "Production driver; pinned to binary distribution to avoid build deps on Railway."],
            ["Redis client", "redis 5.2.1", "Reserved for distributed lock + future cache; currently not on the critical path (single worker)."],
            ["Error tracking", "sentry-sdk[flask] 2.0.0", "Production error capture with Flask + SQLAlchemy integrations; releases tagged by build hash."],
            ["2FA TOTP", "pyotp 2.9.0", "RFC 6238 TOTP secrets + verification; nag-then-enforce model in app/routes/totp_routes.py."],
            ["Web push", "pywebpush 2.0.0", "VAPID-keyed browser push for the service worker."],
            ["WebAuthn", "webauthn 2.5.0", "FIDO2 enrolment + assertion + sign-count clone detection for biometric re-auth."],
            ["Environment loader", "python-dotenv 1.0.1", "Loads .env files in dev; prod uses Railway-injected env vars."],
        ],
        widths=[1.8, 2.0, 3.0]
    )

    heading(doc, "2.2 Frontend", 2)
    table(doc,
        ["Layer", "Choice", "Why"],
        [
            ["Language", "TypeScript 5+", "Type safety across 200+ components; the AI APIs return shaped JSON we want to model precisely."],
            ["Framework", "Next.js 14.2.35 (App Router, static export)", "Static-export build copied into /static/nextjs/ and served by Flask. Single deploy unit. App-Router for layout primitives."],
            ["UI components", "shadcn/ui + @base-ui/react 1.3.0", "Headless + style-it-yourself; fully migrated off MUI (Kuja Studio design tokens in globals.css). Tailwind for utility classes."],
            ["Styling", "Tailwind CSS 3.4.1 + tailwind-merge 3.5.0 + tailwindcss-animate 1.0.7 + tw-animate-css 1.4.0 + class-variance-authority 0.7.1 + clsx 2.1.1", "Utility-first with cva for component variants. Kuja Studio palette via CSS custom properties (--kuja-clay, --kuja-savanna, --kuja-spark)."],
            ["Icons", "lucide-react 0.577.0", "Open-source icon set with consistent stroke weights; works in RTL."],
            ["Charts", "recharts 3.8.0", "Composable React charts; chosen over D3 for ease of integration and SSR-safety."],
            ["Calendar", "@fullcalendar/react + daygrid + interaction 6.1.20", "Cross-entity calendar surface."],
            ["Forms", "react-hook-form 7.71.2 + @hookform/resolvers 5.2.2 + zod 4.3.6", "Declarative forms with zod schemas for validation; resolvers bridge react-hook-form to zod."],
            ["Data tables", "@tanstack/react-table 8.21.3", "Headless table primitive; works with shadcn styling."],
            ["State", "zustand 5.0.12 + swr 2.4.1", "zustand for client-only stores (auth, UI); swr for cached server data."],
            ["Theming", "next-themes 0.4.6", "Light/dark theme management."],
            ["Toasts", "sonner 2.0.7", "Toast notifications; integrated with the api helper for error surfacing."],
            ["Command palette", "cmdk 1.1.1", "Cmd+K palette for cross-entity navigation + search."],
            ["Date utilities", "date-fns 4.1.0 + react-day-picker 9.14.0", "Locale-aware formatting + the calendar day picker."],
            ["Fonts", "Fraunces (editorial display) + Inter (UI)", "Self-hosted via /frontend/src/app/fonts. Fraunces for kuja-display class, Inter for body."],
            ["Build", "Next.js 14 build + scripts/copy-build.js → /static/nextjs/", "Build outputs are committed to the repo so prod doesn't run node on the server."],
            ["Linter", "eslint 8 + eslint-config-next 14.2.35", "Standard Next + a11y rules."],
        ],
        widths=[1.6, 2.5, 2.7]
    )

    heading(doc, "2.3 Data layer", 2)
    table(doc,
        ["Layer", "Choice", "Why"],
        [
            ["Production database", "PostgreSQL (Railway-managed)", "ACID, full-text search, JSON support, mature pooling. Multi-tenant friendly via row-level scoping in app code."],
            ["Development database", "SQLite", "Zero-config local dev; the model layer is portable because we avoid Postgres-only types except where noted (TIMESTAMP WITH TIME ZONE handled in models)."],
            ["Connection pool", "SQLAlchemy default pool (size 5, overflow 10)", "Suitable for single-worker production; can be raised when scaling out."],
            ["Migrations", "Alembic (via Flask-Migrate)", "Generates from model diffs; manual review before apply."],
            ["JSON storage", "Text columns serialised with json.dumps", "Portable across SQLite + Postgres. Helpers in app/utils/helpers.py (_json_load / _json_dump)."],
            ["File storage", "Local filesystem during dev; Railway persistent volume in prod", "Uploaded documents land in /app/uploads/{org_id}/; rotated nightly into S3 if KUJA_S3_BUCKET env is set (future)."],
            ["Audit chain", "Append-only AuditChainEntry table with hash of prior row", "Tamper-evident provenance log; verify endpoint computes the full chain."],
        ],
        widths=[2.0, 2.4, 2.7]
    )

    heading(doc, "2.4 Infrastructure + deployment", 2)
    table(doc,
        ["Component", "Choice", "Why"],
        [
            ["Hosting", "Railway (project: clever-cooperation)", "Managed PaaS with built-in Postgres + automatic HTTPS + env-var management. Container deploys from `railway up`."],
            ["Domain", "https://web-production-6f8a.up.railway.app", "Railway-provided HTTPS endpoint. Custom domain pending DNS."],
            ["Container", "Single Gunicorn worker, sync mode", "Sufficient for current load; AI heavy lifting is async via in-process task runner (see §6.2)."],
            ["Deploy command", "railway up --detach", "From local; CLI uploads source, Railway builds + redeploys. Rolling restart takes 4-5 min."],
            ["CI", "GitHub Actions", "Python smoke (smoke.yml) + browser regression (browser-smoke.yml) on every push to main + every PR touching app/, frontend/, static/nextjs/."],
            ["Crons", "GitHub Actions schedule triggers", ".github/workflows/cron-*.yml hit /api/cron/* endpoints with Bearer CRON_SECRET. Daily for UAT fixtures + compliance rescreen + reviewer auto-assign sweep."],
            ["DNS", "Railway-managed", ""],
            ["TLS", "Railway-managed (Let's Encrypt)", ""],
            ["Source control", "GitHub (idirisloyan/kuja-grant)", "main branch is the prod deploy source; all commits pushed."],
        ],
        widths=[1.6, 2.5, 2.7]
    )

    heading(doc, "2.5 Frontend libraries (full inventory)", 2)
    body(doc,
        "Every dependency declared in frontend/package.json with a one-"
        "line role description. Curated to the minimum set sufficient "
        "to deliver the platform; removed MUI entirely in favour of "
        "shadcn/Tailwind in Phase 10."
    )
    table(doc,
        ["Package", "Version", "Role"],
        [
            ["next", "14.2.35", "App Router framework"],
            ["react / react-dom", "18.x", "UI runtime"],
            ["typescript", "5.x", "Type system"],
            ["tailwindcss", "3.4.1", "Utility CSS"],
            ["tailwind-merge", "3.5.0", "Merge conflicting Tailwind classes"],
            ["tailwindcss-animate", "1.0.7", "Animations preset"],
            ["tw-animate-css", "1.4.0", "Additional animation primitives"],
            ["class-variance-authority", "0.7.1", "Component variant API"],
            ["clsx", "2.1.1", "Class composition"],
            ["@base-ui/react", "1.3.0", "shadcn primitive headless components"],
            ["shadcn", "4.1.0", "Component CLI"],
            ["lucide-react", "0.577.0", "Icons"],
            ["recharts", "3.8.0", "Charts (funnel, bars, heatmap, score rings)"],
            ["@fullcalendar/react / daygrid / interaction", "6.1.20", "Calendar surface"],
            ["react-hook-form", "7.71.2", "Form state"],
            ["@hookform/resolvers", "5.2.2", "Bridge to zod"],
            ["zod", "4.3.6", "Schema validation"],
            ["@tanstack/react-table", "8.21.3", "Headless tables"],
            ["zustand", "5.0.12", "Client-only state (auth, UI)"],
            ["swr", "2.4.1", "Server data fetching + caching"],
            ["next-themes", "0.4.6", "Light/dark theme"],
            ["sonner", "2.0.7", "Toast notifications"],
            ["cmdk", "1.1.1", "Cmd+K palette"],
            ["date-fns", "4.1.0", "Date utilities"],
            ["react-day-picker", "9.14.0", "Date picker"],
        ],
        widths=[2.4, 1.0, 3.4]
    )

    heading(doc, "2.6 Python dependencies (full inventory)", 2)
    table(doc,
        ["Package", "Version", "Role"],
        [
            ["flask", "3.1.0", "Web framework"],
            ["flask-sqlalchemy", "3.1.1", "ORM integration"],
            ["flask-login", "0.6.3", "Session auth"],
            ["flask-cors", "5.0.1", "CORS"],
            ["flask-migrate", "4.0.7", "Alembic migrations"],
            ["sqlalchemy", "2.0.36", "ORM"],
            ["werkzeug", "3.1.3", "WSGI + password hashing (bcrypt-equivalent)"],
            ["python-dotenv", "1.0.1", ".env loader"],
            ["anthropic", "0.42.0", "Claude SDK (tool-use + structured output + vision)"],
            ["requests", "2.32.3", "External API HTTP client"],
            ["gunicorn", "23.0.0", "Production WSGI server"],
            ["PyPDF2", "3.0.1", "PDF text extraction"],
            ["python-docx", "1.1.2", "DOCX read + write"],
            ["openpyxl", "3.1.5", "XLSX read + write"],
            ["psycopg2-binary", "2.9.10", "Postgres driver"],
            ["redis", "5.2.1", "Redis client (reserved)"],
            ["sentry-sdk[flask]", "2.0.0", "Error tracking"],
            ["pyotp", "2.9.0", "TOTP 2FA"],
            ["pywebpush", "2.0.0", "VAPID web push"],
            ["reportlab", "4.2.5", "PDF generation"],
            ["webauthn", "2.5.0", "FIDO2 biometric re-auth"],
        ],
        widths=[2.4, 1.0, 3.4]
    )

    page_break(doc)

    # =================================================================
    # 3. SYSTEM ARCHITECTURE
    # =================================================================
    heading(doc, "3. System Architecture", 1)

    heading(doc, "3.1 Layered architecture", 2)
    body(doc,
        "The backend follows a layered architecture with clear "
        "boundaries between transport (Flask blueprints), business "
        "logic (services), data (SQLAlchemy models), and integrations "
        "(adapters to external APIs)."
    )
    code(doc, """
┌─────────────────────────────────────────────────────────────────┐
│  CLIENT TIER                                                    │
│  Next.js 14 (App Router) static export → /static/nextjs/        │
│  shadcn/ui + Tailwind · TypeScript · Recharts · SWR · Zustand   │
│  Service Worker (offline-first cache) · PWA manifest            │
└─────────────────────────────────────────────────────────────────┘
                          ↓  HTTPS  (Railway-managed)
┌─────────────────────────────────────────────────────────────────┐
│  TRANSPORT TIER  (Flask + Gunicorn)                             │
│  43 blueprints under /api/* + /admin/* + /trust/* + static      │
│  middleware: rate limit · CSRF · CORS · session · sentry        │
└─────────────────────────────────────────────────────────────────┘
                          ↓
┌─────────────────────────────────────────────────────────────────┐
│  SERVICE TIER  (72 services in app/services/)                   │
│  Lifecycle · AI (ai_service + ai_chat + ai_jobs + ai_agent)    │
│  Compliance · Trust · Reviewer · Donor portfolio · Metrics      │
│  Notifications · Search · Audit · Bundle · Calendar             │
└─────────────────────────────────────────────────────────────────┘
                          ↓
┌─────────────────────────────────────────────────────────────────┐
│  DATA TIER  (SQLAlchemy 2.x + Alembic)                          │
│  31 models  ·  PostgreSQL (prod)  ·  SQLite (dev)               │
│  AuditChainEntry (append-only) · UserEvent · AIProvenance       │
└─────────────────────────────────────────────────────────────────┘
                          ↓
┌─────────────────────────────────────────────────────────────────┐
│  INTEGRATIONS                                                   │
│  Anthropic Claude · OpenSanctions · UN/OFAC/EU sanctions lists  │
│  Kenya/Nigeria/SA/Uganda/Tanzania/Somalia/Ethiopia registries   │
│  Sentry · Web Push (VAPID) · SMTP · SMS · WhatsApp              │
└─────────────────────────────────────────────────────────────────┘
""".strip())

    heading(doc, "3.2 Request lifecycle", 2)
    numbered(doc, [
        "Browser issues an HTTPS request to web-production-6f8a.up.railway.app. Railway terminates TLS and forwards to the Gunicorn worker.",
        "Flask middleware runs in order: session loader (flask-login resolves user from cookie) → CORS → endpoint rate limit (per IP, per pattern) → CSRF check (X-Requested-With header required on mutating verbs) → route dispatch.",
        "Route handler validates role + ownership, then calls the relevant service function.",
        "Service does its work synchronously, OR enqueues a long-running AI call via ai_jobs.submit_ai_job() and returns a job_id (HTTP 202).",
        "Service writes to models (SQLAlchemy session), optionally writes an AuditChainEntry + UserEvent, commits.",
        "Response serialised to JSON via flask.jsonify; status code returned. Errors are caught by api_errors.error_response() into a uniform shape and reported to Sentry if env DSN is set.",
        "Client polls /api/ai/jobs/<job_id> with backoff until status='completed' or 'failed' (max 30 polls / ~50s).",
    ])

    heading(doc, "3.3 Blueprint catalogue", 2)
    body(doc,
        "All 43 blueprints registered in app/routes/__init__.py. Each "
        "blueprint owns a slice of the API surface and is independently "
        "testable. Url prefixes are stable: /api/auth, /api/grants, "
        "/api/applications, etc."
    )
    table(doc,
        ["Blueprint", "Prefix", "Role"],
        [
            ["auth_bp", "/api/auth", "Login/logout, language switch, session check"],
            ["totp_bp", "/api/auth/totp", "TOTP 2FA enrolment + verify"],
            ["webauthn_bp", "/api/auth/webauthn", "WebAuthn registration + assertion (Phase 26C)"],
            ["dashboard_bp", "/api/dashboard", "Per-role dashboard stats + benchmarks + cohort + heatmap"],
            ["organizations_bp", "/api/organizations", "Org CRUD + profile + summary + search"],
            ["grants_bp", "/api/grants", "Grant lifecycle + broadcast + Q&A + fit-compare"],
            ["applications_bp", "/api/applications", "Application CRUD + submit + decision + auto-assign + signals + timeline + comments"],
            ["assessments_bp", "/api/assessments", "Capacity assessment + framework selection"],
            ["documents_bp", "/api/documents", "Upload + extract + AI score"],
            ["ai_bp", "/api/ai", "Free-text chat + tool-use agent + threads + jobs + 20+ AI surfaces"],
            ["copilot_bp", "/api/ai (copilot subset)", "Streaming co-pilot rail (Now/Ask/Insights tabs)"],
            ["compliance_bp", "/api/compliance", "Compliance findings + health + risk register"],
            ["reviews_bp", "/api/reviews", "Reviewer assignment + scoring + completion + briefing"],
            ["reports_bp", "/api/reports", "Report draft + pre-flight + submit + review"],
            ["admin_bp", "/api/admin + /api/health + /api/version + /api/ready", "Admin tools + system health + metrics + clear-lockout"],
            ["notifications_bp", "/api/notifications", "In-app notifications + unread counts"],
            ["notif_pref_bp + notif_pref_alias_bp", "/api/notification-preferences + /api/notifications/preferences", "Per-user channel + cadence prefs"],
            ["match_bp", "/api/match", "Match engine + suggestions"],
            ["questions_bp", "/api/grant-questions", "Per-grant Q&A surface"],
            ["diligence_bp", "/api/diligence", "Due-diligence questionnaire"],
            ["org_memory_bp", "/api/org-memory", "Reusable item store"],
            ["risks_bp", "/api/risks", "Risk register CRUD + lifecycle transitions"],
            ["admin_health_bp", "/api/admin/system-health + /api/admin/ai-surface-health + /api/admin/demo-readiness", "Self-service ops endpoints"],
            ["comments_bp", "/api/comments", "Threaded comments on entities"],
            ["test_bp", "/api/test (env-gated)", "Test-only endpoints"],
            ["saved_searches_bp", "/api/saved-searches", "Persisted search queries"],
            ["push_bp", "/api/push", "Web push subscription"],
            ["trust_bp", "/api/trust-profile + /api/adverse-media + /api/bank-verification + /api/passport", "Two-pillar trust profile + diligence"],
            ["watchlist_bp", "/api/watchlist", "Personal star toggle"],
            ["signals_bp", "/api/signals", "ASK/RISK/DECISION rail"],
            ["preemption_bp", "/api/preemption", "AI compliance pre-emption scanner"],
            ["calendar_bp", "/api/calendar", "Deadline calendar across entities"],
            ["messaging_bp", "/api/messaging", "SMS + WhatsApp connector"],
            ["ai_budget_bp", "/api/ai-budget", "Per-org AI cost cap"],
            ["preflight_bp", "/api/preflight", "Donor pre-flight against application"],
            ["audit_chain_bp", "/api/audit-chain", "Audit chain read + verify"],
            ["report_bundle_bp", "/api/reports/<id>/bundle", "Report bundle assembly"],
            ["doc_search_bp", "/api/documents/search + /api/documents/search/global", "Document + cross-entity search"],
            ["search_alias_bp", "/api/search", "Phase 28A alias for global search"],
            ["digest_bp", "/api/digest", "Notification digest cron + cadence settings"],
            ["ai_compare_bp", "/api/applications/compare + /api/applications/autofill", "Compare + autofill"],
            ["phase11_bp", "/api/grants/agreement-unpack + /api/patterns", "Agreement smart-unpack + cross-grant patterns"],
            ["portfolio_bp", "/api/portfolio", "Donor + NGO portfolio bundle PDF + ZIP"],
            ["cron_bp", "/api/cron", "All scheduled task endpoints (Bearer CRON_SECRET or admin session)"],
            ["tags_bp", "/api/tags", "Tags + tag assignment"],
            ["exports_bp", "/api/exports/grants.csv + applications.csv + reviews.csv", "Role-scoped CSV exports"],
            ["feedback_bp", "/api/feedback + /api/feedback/my", "NPS micro-survey (Phase 31A)"],
        ],
        widths=[2.0, 2.4, 2.6]
    )

    heading(doc, "3.4 Service catalogue", 2)
    body(doc,
        "72 service modules under app/services/, grouped by capability "
        "pillar. Each service is a stateless module with classmethods + "
        "module-level functions; no service holds in-memory state across "
        "requests except for explicit caches (_dashboard_cache) and the "
        "in-process AI job registry (_AI_JOBS)."
    )
    callout(doc, "REFERENCE",
        "BRD Appendix C contains the full service-by-service table. The "
        "BRD is the canonical reference for what each service does; this "
        "spec covers HOW they interact (sync vs async dispatch, "
        "transaction boundaries, error handling).")

    heading(doc, "3.5 Frontend route map", 2)
    body(doc,
        "Frontend routes are organised under Next.js App Router groups: "
        "(app) for authenticated workspaces, (auth) for login + public "
        "trust verify, and top-level for /chat, /trust, /trust/verify. "
        "Static export generates /static/nextjs/ which Flask serves "
        "from app.send_from_directory."
    )
    table(doc,
        ["Route", "Role", "Purpose"],
        [
            ["/login", "any", "Email/password login"],
            ["/dashboard", "all roles", "Role-aware command center"],
            ["/chat", "ngo + donor", "Sustained AI chat (global scope)"],
            ["/grants", "all", "Browse + create + manage grants"],
            ["/grants/[id]", "all (gated)", "Grant detail + scoped chat"],
            ["/grants/new", "donor + admin", "Grant creation wizard"],
            ["/applications", "ngo + donor + admin", "Application list (NGO sees own; donor sees on own grants)"],
            ["/applications/[id]", "owner + donor + admin", "Application detail + timeline + chat + messaging"],
            ["/apply/[grantId]", "ngo", "Apply wizard with AI co-author"],
            ["/assessments", "ngo + admin", "Capacity assessment overview"],
            ["/assessments/wizard", "ngo + admin", "5-framework wizard"],
            ["/reports", "ngo + donor + admin", "Reports list"],
            ["/reports/[id]", "owner + donor + admin", "Report detail + scoped chat"],
            ["/reviews", "reviewer + donor + admin", "Reviewer assignments list"],
            ["/reviews/[id]", "reviewer + admin", "Score per criterion"],
            ["/reviews/completed", "reviewer + admin", "Completed reviews history"],
            ["/compliance", "donor + admin", "Compliance dashboard"],
            ["/verification", "donor + admin", "Registry verification dashboard"],
            ["/trust", "ngo + admin", "Trust profile + gap insights"],
            ["/trust/verify/[slug]", "public (token-gated)", "Capacity passport public verifier"],
            ["/calendar", "ngo + donor + admin", "Cross-entity deadline calendar"],
            ["/organizations/profile", "ngo + donor + admin", "Org profile editor"],
            ["/organizations/memory", "ngo + admin", "Org memory items"],
            ["/organizations/search", "donor + admin", "Find + verify orgs"],
            ["/donors/[id]", "any logged-in user", "Public donor profile (aggregate only)"],
            ["/ngo/[id]", "public (opt-in)", "Public NGO summary"],
            ["/settings/notifications", "all", "Per-channel + per-category prefs + digest cadence"],
            ["/settings/security", "all", "WebAuthn enrolment + 2FA"],
            ["/admin/audit-chain", "admin", "Chronological audit log viewer"],
            ["/admin/metrics", "admin", "DAU/WAU + funnels + NPS rollup"],
            ["/admin/security", "admin", "Security events + lockout management"],
            ["/observability", "admin", "System health + AI surface health"],
        ],
        widths=[2.2, 1.6, 3.2]
    )

    page_break(doc)

    # =================================================================
    # 4. DATABASE DESIGN
    # =================================================================
    heading(doc, "4. Database Design", 1)

    heading(doc, "4.1 ER overview", 2)
    body(doc,
        "31 models grouped into nine clusters. Foreign keys are explicit; "
        "many-to-many relationships are denormalised (e.g. tag → "
        "TagAssignment with polymorphic target). Cross-tenant scoping "
        "is enforced in application code (every list query filters by "
        "org_id from session); row-level security is not enabled at the "
        "Postgres level."
    )
    code(doc, """
LIFECYCLE        ID/AUTH           AI                 TRUST/DILIGENCE
─────────        ───────           ──                 ────────────────
Organization     User              AIThread           CapacityPassport
Grant            WebAuthnCredential AIMessage          AdverseMediaScreening
Application                        AICallLog          BankAccountVerification
Report                             AIProvenance       ComplianceCheck
Review                                                RegistrationVerification
Assessment       OBSERVABILITY     COMPLIANCE         ComplianceSnapshot
Document         ──────────────    ──────────         DiligenceItem
GrantQuestion    UserEvent         Risk
EntityComment    UserFeedback      AuditChainEntry    DISCOVERY
                                                      ─────────
COLLAB           METRICS           MISC               SavedSearch
──────           ───────                              WatchlistItem
Notification     UserEvent         OrgMemory          Tag
NotificationPreference             StatusSignal       TagAssignment
PushSubscription
""".strip())

    heading(doc, "4.2 Entity reference (31 models)", 2)
    body(doc,
        "Every model with its primary fields + purpose. Date fields "
        "default to UTC. JSON-typed columns are documented in §4.4."
    )

    # Models grouped
    entities = [
        # (cluster, model, fields, purpose)
        ("Identity", "User", "id PK, email UQ, password_hash, name, role (ngo/donor/reviewer/admin), org_id FK, language, avatar_url, totp_secret, totp_enabled, digest_cadence, failed_login_count, locked_until, last_failed_login",
         "Auth + role + 2FA + lockout + digest cadence."),
        ("Identity", "WebAuthnCredential", "id PK, user_id FK, credential_id UQ, public_key, sign_count, label, transport_hint, created_at, last_used_at",
         "Per-device biometric credential (Phase 26C)."),
        ("Lifecycle", "Organization", "id PK, name, org_type (ngo/donor), country, sector, language, website, mission, year_established, staff_count, annual_budget, ai_monthly_budget_usd, stage_labels_json, summary_published",
         "Multi-tenant root."),
        ("Lifecycle", "Grant", "id PK, donor_org_id FK, title, description, total_funding NUMERIC, currency, deadline DATE, status, sectors JSON, countries JSON, eligibility JSON, criteria JSON, doc_requirements JSON, reporting_requirements JSON, grant_document, report_template, reporting_frequency, published_at",
         "Donor-published opportunity."),
        ("Lifecycle", "Application", "id PK, grant_id FK, ngo_org_id FK, status, responses JSON, eligibility_responses JSON, ai_score FLOAT, human_score FLOAT, final_score FLOAT, submitted_at, decision_reason_code, decision_notes, decision_recorded_at, decision_recorded_by_user_id FK",
         "NGO response to grant + decision."),
        ("Lifecycle", "Report", "id PK, grant_id FK, application_id FK, submitted_by_org_id FK, report_type, reporting_period, title, content JSON, attachments JSON, status, due_date DATE, submitted_at, reviewed_at, reviewer_notes, ai_analysis JSON, revision_history JSON, revision_number",
         "Progress + final reports."),
        ("Lifecycle", "Review", "id PK, application_id FK, reviewer_user_id FK, status (assigned/in_progress/completed), scores JSON, overall_score FLOAT, comments JSON, completed_at",
         "Reviewer scoring per application."),
        ("Lifecycle", "Assessment", "id PK, org_id FK, framework, version, total_score FLOAT, framework_scores JSON, completed_sections JSON, answers JSON, status, completed_at",
         "Capacity assessment instance."),
        ("Lifecycle", "Document", "id PK, application_id FK (nullable), org_id FK, doc_type, original_filename, stored_path, content_text, ai_score INT, ai_analysis JSON, file_size, mime_type, uploaded_by_user_id FK",
         "Uploaded file with extracted text + AI score."),
        ("AI", "AIThread", "id PK, user_id FK, scope_kind, scope_id, title, created_at, updated_at",
         "Sustained chat thread metadata."),
        ("AI", "AIMessage", "id PK, thread_id FK, role (user/assistant), content TEXT, model, created_at",
         "Individual chat message."),
        ("AI", "AICallLog", "id PK, user_id FK, endpoint, prompt_tokens INT, completion_tokens INT, total_tokens INT, cost_usd FLOAT, latency_ms INT, source, language, created_at",
         "Cost-tagged AI call ledger."),
        ("AI", "AIProvenance", "id PK, ai_call_id FK, claim TEXT, source_doc_id FK, source_excerpt, confidence",
         "Per-claim source linkage."),
        ("Trust", "CapacityPassport", "id PK, org_id FK, slug UQ, status, public_token, expires_at, scoring_snapshot JSON, published_at, revoked_at",
         "Public passport URL + verify token."),
        ("Trust", "AdverseMediaScreening", "id PK, org_id FK, screened_at, status (clear/flagged/error), summary_json (matched counts + per-list breakdown)",
         "Sanctions + adverse media result row."),
        ("Trust", "BankAccountVerification", "id PK, org_id FK, account_holder, iban, swift, country, status, error_reason, verified_at",
         "IBAN checksum + connector verification."),
        ("Trust", "ComplianceCheck", "id PK, org_id FK, check_type (sanctions/registration/tax/bo/adverse), status, findings JSON, last_checked_at, expires_at",
         "Per-org compliance check result."),
        ("Trust", "RegistrationVerification", "id PK, org_id FK, country, registration_number, status (verified/unverified/expired/ai_reviewed), legal_name_match, source, verified_at, expires_at",
         "Registry verification result."),
        ("Trust", "ComplianceSnapshot", "id PK, grant_id FK, snapshot_date DATE, score INT, band, pillars_json",
         "Daily compliance trajectory snapshot."),
        ("Trust", "DiligenceItem", "id PK, org_id FK, question_key, answer_json, signed_by_user_id FK, signed_at",
         "Due-diligence questionnaire response."),
        ("Compliance", "Risk", "id PK, subject_kind, subject_id, severity (info/low/med/high), status (open/mitigated/resolved), title, description, owner_user_id FK, due_date DATE, created_at, resolved_at",
         "Polymorphic risk row."),
        ("Audit", "AuditChainEntry", "id PK, action, actor_email, subject_kind, subject_id, details JSON, prev_hash, hash, created_at",
         "Tamper-evident append-only log."),
        ("Audit", "EntityComment", "id PK, entity_kind, entity_id, author_user_id FK, body, mentioned_user_ids JSON, created_at",
         "Threaded comments on apps + reports."),
        ("Discovery", "GrantQuestion", "id PK, grant_id FK, asker_user_id FK, question, answer, answered_at",
         "Per-grant Q&A."),
        ("Discovery", "SavedSearch", "id PK, user_id FK, name, scope_kind, query_json, sort_order, created_at",
         "Persisted discovery query."),
        ("Discovery", "WatchlistItem", "id PK, user_id FK, target_kind, target_id, created_at",
         "Personal star per (user, target)."),
        ("Discovery", "Tag", "id PK, org_id FK, key, label, color",
         "Per-org tag definition."),
        ("Discovery", "TagAssignment", "id PK, tag_id FK, target_kind, target_id, assigned_by_user_id FK",
         "Polymorphic tag → target."),
        ("Collab", "Notification", "id PK, user_id FK, category, title, body, deep_link_url, related_kind, related_id, read_at, created_at",
         "In-app notification."),
        ("Collab", "NotificationPreference", "id PK, user_id FK, category, channels JSON, phone_e164, whatsapp_e164",
         "Per-user channel preferences."),
        ("Collab", "PushSubscription", "id PK, user_id FK, endpoint, keys JSON, created_at, last_seen_at",
         "Web push subscription."),
        ("Collab", "StatusSignal", "id PK, entity_kind, entity_id, signal_type (ASK/RISK/DECISION), severity, owner_user_id FK, status, summary, created_at, resolved_at",
         "Lightweight signal rail."),
        ("Collab", "OrgMemory", "id PK, org_id FK, item_kind, content_text, source_application_id FK, confidence, created_at",
         "Auto-extracted reusable items."),
        ("Metrics", "UserEvent", "id PK, user_id FK, org_id FK (denorm), role, language, event_name, event_props JSON, ab_arm, occurred_at",
         "Behavioural event (Phase 29A)."),
        ("Metrics", "UserFeedback", "id PK, user_id FK, org_id FK, role, language, surface, related_kind, related_id, score (0-10), comment (≤500ch), created_at; UNIQUE (user_id, surface, related_kind, related_id)",
         "NPS micro-survey response (Phase 31A)."),
    ]
    # Group by cluster
    clusters = {}
    for cluster, model, fields, purpose in entities:
        clusters.setdefault(cluster, []).append((model, fields, purpose))
    for cluster, rows in clusters.items():
        heading(doc, f"4.2.{['Identity','Lifecycle','AI','Trust','Compliance','Audit','Discovery','Collab','Metrics'].index(cluster) + 1} {cluster}", 3)
        for model, fields, purpose in rows:
            p = doc.add_paragraph()
            r1 = p.add_run(f"{model}.  ")
            r1.bold = True
            r1.font.color.rgb = KUJA_CLAY
            r1.font.size = Pt(10.5)
            r2 = p.add_run(purpose)
            r2.font.size = Pt(10.5)
            r2.italic = True
            code(doc, fields)

    heading(doc, "4.3 Key indexes + constraints", 2)
    bullets(doc, [
        ("Per-user time-series queries: ", "Index (user_id, created_at) on UserEvent, UserFeedback, AIMessage."),
        ("Lifecycle list views: ", "Composite (ngo_org_id, status) on Application; (donor_org_id, status) on Grant; (submitted_by_org_id, status) on Report."),
        ("Audit verification: ", "AuditChainEntry has UNIQUE on (prev_hash, hash) and an index on created_at for ordered scan."),
        ("Uniqueness rules: ", "Application UNIQUE (grant_id, ngo_org_id) prevents duplicate applications; UserFeedback UNIQUE (user_id, surface, related_kind, related_id) prevents duplicate survey responses; WebAuthnCredential UNIQUE on credential_id."),
        ("Soft FK strategy: ", "Polymorphic targets (Risk.subject_kind/subject_id, EntityComment.entity_kind/entity_id, TagAssignment) intentionally NOT enforced at the DB level. App code is responsible for integrity; trade-off accepted because the polymorphism crosses 4-5 entity tables."),
    ])

    heading(doc, "4.4 JSON-typed columns", 2)
    body(doc,
        "Stored as TEXT and (de)serialised in app code via _json_load + _json_dump "
        "(app/utils/helpers.py). Portable across SQLite + Postgres. Each model "
        "exposes get_X() / set_X() helpers so callers never touch the raw column."
    )
    table(doc,
        ["Column", "Shape", "Why JSON not relational"],
        [
            ["Grant.sectors / countries", "list[str]", "Variable size; only used as a tag set, never joined."],
            ["Grant.eligibility / criteria / doc_requirements / reporting_requirements / report_template", "list[dict] / dict", "Highly variable shape per donor; relational normalisation would explode tables."],
            ["Application.responses / eligibility_responses", "dict[criterion_key → text]", "Free-form per criterion; AI-extracted; doesn't drive joins."],
            ["Review.scores / comments", "dict[criterion_key → value]", "Variable per grant rubric; AI-extracted."],
            ["Assessment.framework_scores / answers", "dict", "5 frameworks × ~40 sub-components × variable; relational would be 200+ tables."],
            ["Document.ai_analysis / Report.ai_analysis", "dict (AI output structure)", "Variable per AI surface."],
            ["AuditChainEntry.details", "dict", "Per-action variable; hashed as a unit."],
            ["UserEvent.event_props", "dict (event-specific)", "Per event_name; loose by design."],
            ["NotificationPreference.channels", "list[str]", "Tag set."],
            ["TagAssignment + others (polymorphic)", "polymorphic kind+id", "Cross-entity tagging without joining 5 tables."],
            ["AIThread / Risk / StatusSignal scope_kind / subject_kind", "string discriminator", "Polymorphic target."],
        ],
        widths=[2.6, 1.8, 2.6]
    )

    heading(doc, "4.5 Migration strategy", 2)
    body(doc,
        "Flask-Migrate (Alembic) manages schema versions. New models or "
        "field changes generate an alembic revision under migrations/. "
        "In dev, `flask db migrate -m 'msg'` then `flask db upgrade`. "
        "In prod, deploys apply pending migrations on first boot via "
        "the Railway start command. Backward-incompatible changes "
        "(rename + drop) are deferred until all reads have moved off "
        "the old column; the Phase 5 SOPs deletion in adeso-pmo-v2 "
        "demonstrates the pattern (dual-write → cutover → drop)."
    )

    page_break(doc)

    # =================================================================
    # 5. SECURITY ARCHITECTURE
    # =================================================================
    heading(doc, "5. Security Architecture", 1)

    heading(doc, "5.1 Authentication stack", 2)
    bullets(doc, [
        ("Primary: ", "Email + password via Flask-Login session. Password hashed with werkzeug.security.generate_password_hash (scrypt-based; ~150ms hash cost)."),
        ("Second factor (admins enforced, all users opt-in): ", "TOTP 2FA via pyotp; QR-code provisioning at /api/auth/totp/enroll/start; recovery codes (bcrypt-hashed single-use) generated at enrolment."),
        ("Re-auth for sensitive actions: ", "WebAuthn (Phase 26C) — Touch ID / Face ID / Windows Hello / hardware security key. Multiple devices per user. Short-lived (5min) single-use re-auth tokens returned by /api/auth/webauthn/authenticate/finish, consumed via X-Reauth-Token header on the gated request."),
        ("Brute force protection: ", "Per-email failed_login_count + locked_until on the User row. After 5 failures in 15 min, account locks; lockout duration 15 min default. Admin can clear via /api/admin/clear-lockout (single email) or /api/admin/clear-all-lockouts (bulk)."),
        ("Session: ", "HTTP-only secure cookie (SameSite=Lax). Sessions are permanent with 30-day expiry. Server-side state in cookie via flask.session encryption."),
        ("Sign-count clone detection: ", "WebAuthn sign counts must strictly increase. A regression triggers hard auth failure (treated as a clone attempt) and logs a warning."),
    ])

    heading(doc, "5.2 Authorisation + role gating", 2)
    bullets(doc, [
        ("Decorators: ", "@login_required (flask-login) + custom @role_required(*allowed) in app/utils/decorators.py."),
        ("Per-row scoping: ", "Every list query filters by current_user.org_id (NGO) or current_user.id (reviewer)."),
        ("Cross-org access: ", "Returns 403 OR 404 (preferred for privacy — 'not yours' is indistinguishable from 'does not exist')."),
        ("Admin-only endpoints: ", "Manually checked via current_user.role == 'admin' inside the handler (vs decorator) because some admin endpoints also accept Bearer CRON_SECRET for unattended runs."),
    ])

    heading(doc, "5.3 Session + CSRF + headers", 2)
    bullets(doc, [
        ("CSRF strategy: ", "Custom 'X-Requested-With: XMLHttpRequest' header required on every mutating API call (POST/PUT/PATCH/DELETE) except file uploads. Standard browsers cannot set this header cross-origin without preflight, so it functions as a same-origin gate without per-form CSRF tokens."),
        ("File uploads: ", "X-Requested-With cleared on multipart so the browser controls the boundary. CSRF check is skipped for uploads (rationale: file-pick interaction is already a same-origin user gesture)."),
        ("CORS: ", "Flask-CORS configured to allow the Next.js dev origin (localhost:3000) + the prod origin; credentials are included so the session cookie travels."),
        ("Security headers: ", "Content-Security-Policy, X-Frame-Options DENY, X-Content-Type-Options nosniff, Referrer-Policy strict-origin-when-cross-origin, Permissions-Policy strict — applied via middleware on every response."),
    ])

    heading(doc, "5.4 Rate limiting + lockouts", 2)
    body(doc,
        "Two-tier rate limit. Per-user AI limiter (ai_limiter in "
        "app/utils/rate_limiter.py) tracks failure-count per "
        "(user, key) and short-circuits with 429 when exceeded. "
        "Per-IP endpoint limiter in app/middleware.py runs as a "
        "before_request hook."
    )
    table(doc,
        ["Endpoint pattern", "Cap", "Window"],
        [
            ["^/api/auth/login$", "30 (env: RATE_LIMIT_LOGIN_PER_IP_PER_MIN)", "60 s"],
            ["^/api/ai/jobs/", "600", "60 s (status polls excluded from generation budget)"],
            ["^/api/ai/", "40", "60 s"],
            ["^/api/documents/upload$", "10", "60 s"],
            ["^/api/grants/\\d+/upload-grant-doc$", "10", "60 s"],
            ["^/api/compliance/screen$", "5", "60 s"],
        ],
        widths=[3.0, 1.5, 2.5]
    )
    body(doc,
        "Walked back from 60→40 on /api/ai/ in Phase 27 after saturation "
        "testing wedged the single Gunicorn worker (AI calls take 5-26s "
        "each). The 40/min cap protects worker pressure while still "
        "supporting normal cross-page navigation (a reviewer touching "
        "dashboard + reviews + an application can fire 8-12 legitimate "
        "AI calls)."
    )

    heading(doc, "5.5 File upload security", 2)
    bullets(doc, [
        ("Size limit: ", "MAX_CONTENT_LENGTH = 16 MB. Rejected with 413 by Flask before hitting the handler."),
        ("Type whitelist: ", "Allowed extensions enforced in app code; .exe + .zip + others rejected with clean error."),
        ("Filename sanitisation: ", "werkzeug.utils.secure_filename for storage path."),
        ("Storage: ", "/app/uploads/{org_id}/ — outside the web-served static directory; downloads go through a Flask handler that re-checks ownership."),
        ("Content extraction: ", "PyPDF2 for PDF, python-docx for DOCX, openpyxl for XLSX. Failed extraction surfaces 'no text extracted — pre-flight against requirements not possible'."),
        ("Malicious payloads: ", "HTML uploads sanitised + rendered in iframe with sandbox attribute; JavaScript in narrative content is escaped on render (React XSS protection)."),
    ])

    heading(doc, "5.6 GDPR + data deletion", 2)
    body(doc,
        "Right-to-erasure endpoint accepts an admin-triggered deletion "
        "for a specific user. User row marked as anonymised (email "
        "scrubbed, name redacted to 'Deleted User N'), related records "
        "(applications, reports, messages) have PII fields nulled. The "
        "audit chain retains the erasure event itself (action: "
        "'gdpr.erased') so the deletion is recorded without re-storing "
        "the PII. Public profiles + benchmarks recalculate without the "
        "deleted user on next aggregation."
    )

    heading(doc, "5.7 Hash-chained audit + provenance", 2)
    body(doc,
        "Every meaningful action writes an AuditChainEntry row. The "
        "hash for row N is SHA-256(prev_hash + canonical_json(row N)). "
        "Append-only: no UPDATE or DELETE endpoint exists for the table. "
        "The verify endpoint walks every row and reports 'integrity ok' "
        "or 'hash mismatch at row X' — tamper detection by construction."
    )
    body(doc,
        "AI claim provenance lives in a separate AIProvenance table. "
        "Each row maps an AI-generated claim to the AICallLog row that "
        "produced it AND the source document excerpt that the AI was "
        "shown. Donors questioning a number can ask 'where did this "
        "come from?' and get the exact source span back."
    )

    page_break(doc)

    # =================================================================
    # 6. AI INTEGRATION ARCHITECTURE
    # =================================================================
    heading(doc, "6. AI Integration Architecture", 1)

    heading(doc, "6.1 AI service stack", 2)
    body(doc,
        "AI is woven across 28 product surfaces (see §6.6 inventory + "
        "BRD §6.6.2 for narrative). All surfaces flow through one "
        "service layer: app/services/ai_service.py wraps the Anthropic "
        "SDK with cost-tagged endpoint names, budget enforcement, "
        "rate-limit retry, deterministic mock fallback, and a "
        "canonical tool-use schema."
    )
    code(doc, """
ai_service.py (the wrapper)
├── _call_claude(system, user_message, max_tokens, endpoint, language, role)
│     • Free-text generation; returns string
│     • Cost-tagged via endpoint param
│     • Falls back to deterministic mock (ai_mock.py) when key missing
│     • Logged via ai_budget_service.enforce_budget(endpoint, org_id)
├── _call_claude_tool(system, user_message, tool_name, tool_description, tool_schema)
│     • Forced tool-use; returns the tool's input dict
│     • Schema-validated; always-structured output
│     • Used for: scoring rubrics, gap analysis, criteria suggestion,
│       grant brief generation, etc.
├── chat() / guidance() / generate_application() / score_application() …
│     • ~50 public methods wrapping specific surfaces with their prompts
├── analyze_document() / extract_evidence() / check_submission_readiness() …
│     • Vision-capable for scanned PDFs
└── classmethod _anthropic_client (lazy singleton)

ai_chat_service.py (Phase 24B — sustained conversations)
├── open_or_resume()    — opens AIThread per (user, scope_kind, scope_id)
├── list_messages()
├── post_message()      — appends user msg, calls _call_claude with last 12
└── reset_thread()      — wipes messages, keeps thread row

ai_jobs.py (Phase 13.42 — async dispatch)
├── submit_ai_job(task_type, work_fn) → job_id
│     • Spawns daemon thread; stores result in in-process _AI_JOBS dict
├── get_ai_job(job_id) → {status, result, error, created_at, completed_at}
└── maybe_async_jsonify(request, task_type, work_fn)
      • If async=true in query, returns 202 with job_id; else runs sync

ai_agent.py (tool-use agent for /api/ai/agent)
├── Read-only data registry with role + org scope checks
└── Max 3 iterations; each tool call returns structured data

ai_budget_service.py
├── per_org cost cap + monthly forecast
└── enforce_budget() raises if org has exceeded their cap
""".strip())

    heading(doc, "6.2 Sync vs async dispatch", 2)
    body(doc,
        "AI calls fall into two latency classes. SYNCHRONOUS calls "
        "(chat reply, single-doc scoring, guidance) typically return "
        "in 2-10s and are called inline by the route handler. "
        "ASYNCHRONOUS calls (donor portfolio insights, reviewer summary, "
        "cross-grant patterns, bundle generation, report co-author) "
        "can take 30-120s; the route dispatches via "
        "maybe_async_jsonify which returns HTTP 202 + job_id "
        "immediately, then the client polls /api/ai/jobs/<id> with "
        "exponential backoff (250ms / 500ms / 1s / 1.5s / 2s, max 30 "
        "polls)."
    )
    body(doc,
        "Critical property: a slow AI call NEVER blocks other endpoints. "
        "/api/health and /api/version always return within 200ms even "
        "under sustained AI load."
    )

    heading(doc, "6.3 Tool-use agent + read-only registry", 2)
    body(doc,
        "The conversational 'Ask Kuja' agent at /api/ai/agent is a "
        "Claude tool-use loop with a read-only tool registry. Each "
        "tool wraps a single data query (e.g. list_applications_for_user, "
        "get_grant_summary, search_organisations) that role-checks + "
        "org-scopes before returning data. The agent runs at most 3 "
        "tool-use iterations to bound latency + cost."
    )

    heading(doc, "6.4 Cost budget + rate limits", 2)
    body(doc,
        "Every AI call carries an endpoint label (e.g. 'ai.chat', "
        "'ai.report.preflight'). ai_budget_service.enforce_budget() "
        "checks the per-org monthly cap before each call; over budget "
        "returns 429 with a friendly message. ai_surface_health probes "
        "every flagship surface on /api/admin/ai-surface-health and "
        "reports per-surface latency + success rate."
    )

    heading(doc, "6.5 Anti-hallucination discipline", 2)
    body(doc,
        "Every chat thread system prompt includes a footer instructing "
        "the assistant to NEVER invent grant amounts, status values, or "
        "scores it cannot see in the scope context. When asked about "
        "data outside scope, it must say 'I don't have that loaded in "
        "this thread — try the corresponding page.' This is enforced by "
        "prompt design; we also audit chat replies in test cases "
        "TC-409 (anti-hallucination probe) and TC-409 sub-cases."
    )
    body(doc,
        "Grounding pattern across all generation: the user's actual "
        "data (organisation profile, prior submissions, uploaded "
        "documents, the grant criteria) is included in the system "
        "prompt OR fetched via the agent's tool registry. The model is "
        "never asked to imagine details about an organisation."
    )

    heading(doc, "6.6 AI surface inventory (28 surfaces)", 2)
    body(doc,
        "See BRD §6.6.2 for the canonical table with workspace + "
        "behaviour per surface. Spec-side, the 28 surfaces span:"
    )
    bullets(doc, [
        ("NGO drafting + readiness (10): ", "application co-author, draft section, strengthen against criterion, guidance, submission readiness, past wins suggester, grant fit compare, readiness check, NGO readiness console, trust gap insights."),
        ("Donor pre + post publish (7): ", "grant brief generator, burden estimator, median-NGO preview, donor portfolio insights, portfolio diagnostics, cross-grant patterns, donor cohort analytics."),
        ("Reviewer workflow (4): ", "reviewer summary, evidence extraction, criteria suggestion, reviewer briefing."),
        ("Compliance + diligence (4): ", "compliance pre-empt, compliance explanation, adverse media disambiguation, grant agreement smart-unpack."),
        ("Conversation (2): ", "Ask Kuja (streaming, one-shot, tool-use) AND Sustained chat threads (3 scopes, persistent)."),
        ("Cross-cutting (1): ", "insight narration (captions for dashboard charts)."),
    ])

    heading(doc, "6.7 Sustained chat threads", 2)
    body(doc,
        "Phase 24B introduced a real conversation surface — not a "
        "one-shot chatbot. Architecture: AIThread + AIMessage tables; "
        "ai_chat_service.open_or_resume() returns the latest thread "
        "matching (user_id, scope_kind, scope_id). post_message() "
        "appends the user message, fetches the last 12 messages, "
        "flattens into [USER] / [ASSISTANT] markers, and calls "
        "_call_claude with a scope-aware system prompt. Per-user "
        "isolation is enforced by user_id keying on the thread. "
        "Cost cap: 12-message history sent per turn, regardless of "
        "thread length."
    )

    heading(doc, "6.8 AI claim provenance ledger", 2)
    body(doc,
        "The AIProvenance table records, for each AI-generated claim "
        "shown in the UI, the AICallLog row that produced it AND the "
        "source document excerpt the AI was shown. This means donors "
        "questioning a number can ask 'where did this come from?' and "
        "get a verifiable answer. Without provenance, AI claims are "
        "indistinguishable from confabulation; with it, every claim is "
        "auditable."
    )

    page_break(doc)

    # =================================================================
    # 7. SCORING + CALCULATION METHODOLOGIES
    # =================================================================
    heading(doc, "7. Scoring + Calculation Methodologies", 1)

    heading(doc, "7.1 Application AI scoring", 2)
    body(doc,
        "ScoringEngine.score_application() runs on every application "
        "submit. For each criterion in the grant's rubric:"
    )
    numbered(doc, [
        "Find the corresponding response in application.responses (matched by criterion id OR by index — supports both legacy and modern grant shapes).",
        "Call AIService._call_claude_tool with a per-criterion scoring schema (score 0-100 + 1-2 sentence rationale + per-keyword evidence count).",
        "Capture the score + rationale + evidence count.",
        "Compute overall_score as weighted average: sum(score_i × weight_i) / sum(weight_i). Weights are taken from grant.criteria[].weight; if no weights provided, equal-weight average.",
        "Persist into application.ai_score + application.final_score (until human reviews land — then final_score = 40% × ai + 60% × human).",
    ])
    body(doc,
        "Fallback: when ANTHROPIC_API_KEY is missing OR the AI call "
        "fails, ai_mock.score_application returns a deterministic "
        "heuristic score (40-85) based on response length + keyword "
        "matches. Tests and offline dev runs use this path."
    )

    heading(doc, "7.2 5-framework capacity assessment scoring", 2)
    body(doc,
        "5 capacity frameworks supported: Kuja Internal, STEP, UN-HACT, "
        "CHS, NUPAS. Each framework has its own section weighting + "
        "sub-component weighting. Scoring proceeds per framework:"
    )
    numbered(doc, [
        "For each section in the framework, sum sub-component scores (each on its own scale, typically 0-4 or 0-5) and normalise to 0-100.",
        "Apply section weights to compute framework_score = sum(section_score_i × section_weight_i).",
        "Persist into Assessment.framework_scores[framework_key] = framework_score.",
        "Compute total_score as weighted average across frameworks (default equal weight, donor-overridable).",
        "AI review (optional, on completion): _call_claude_tool with the answers + extracted strengths/gaps; persists to assessment.ai_analysis.",
    ])
    body(doc,
        "Capacity passport derives from the latest Assessment row per "
        "organisation. Passport URL + verify token are minted by "
        "capacity_passport_service.publish()."
    )

    heading(doc, "7.3 Trust Profile aggregation", 2)
    body(doc,
        "trust_profile_service.build(org_id) returns a two-pillar "
        "structure:"
    )
    code(doc, """
{
  "capacity_pillar": {
    "score": int,                # latest Assessment.total_score, or null
    "framework": str,
    "sub_components": [           # rolled up from framework_scores
        {"key": "governance", "score": int, "max": 100},
        ...
    ],
    "last_updated": iso,
    "trend": "improving|flat|declining"  # vs prior assessment
  },
  "diligence_pillar": {
    "score": int,                 # composite, see formula below
    "sub_components": [
        {"key": "sanctions",       "status": "clear|flagged|error"},
        {"key": "registration",    "status": "verified|expired|...", "country": str},
        {"key": "tax_exempt",      "status": ...},
        {"key": "beneficial_owners","disclosed": bool},
        {"key": "adverse_media",   "active_flags": int},
        {"key": "bank_account",    "status": "verified|unverified"}
    ],
    "last_updated": iso
  },
  "overall_score": int,           # avg of two pillars (50/50)
  "overall_band": "excellent|good|attention|critical"
}
""".strip())
    body(doc,
        "Diligence pillar composite: sanctions clear = 100 OR flagged = 0 "
        "(hard gate). When sanctions are clear, the remaining 5 "
        "sub-components contribute equally (20% each) to the diligence "
        "score. A failed registration or expired status drops 30 points "
        "from the registration sub-component. Active adverse media flags "
        "drop 10 points per flag (capped at 50)."
    )

    heading(doc, "7.4 Compliance Health (4 pillars)", 2)
    body(doc,
        "compliance_health.calculate_grant_compliance_health(grant_id) "
        "returns a 0-100 composite + 4 pillars (each 0-100):"
    )
    table(doc,
        ["Pillar", "Weight", "What it measures"],
        [
            ["completion", "0.30", "Required documents uploaded + required fields filled across applications + reports on the grant"],
            ["timeliness", "0.30", "% of reports submitted by due_date on this grant (rolling)"],
            ["workflow", "0.20", "% of applications progressing through expected statuses without stalling > 14 days"],
            ["importance", "0.20", "Grant priority signal (deadline urgency + funding amount + number of active applications)"],
        ],
        widths=[1.6, 1.0, 4.0]
    )
    body(doc,
        "Health bands: excellent (≥85), good (70-84), attention (50-69), "
        "critical (<50). Daily snapshots written to ComplianceSnapshot by "
        "write_daily_snapshots() cron. Trajectory = last 30 daily "
        "snapshots; 'slips-in-N-days' forecast = linear regression on "
        "the trajectory projecting when score will cross 70."
    )

    heading(doc, "7.5 Reviewer match ranking", 2)
    body(doc,
        "reviewer_match_service.suggest_for_application(app_id, top_n) "
        "ranks reviewers by:"
    )
    numbered(doc, [
        "Sector fit: +30 points if reviewer.sectors intersects application.grant.sectors.",
        "Country fit: +20 if reviewer.country matches application.ngo_org.country; +10 if same region.",
        "Capacity fit: +15 if reviewer.expertise covers the application's grant_type.",
        "Throughput health: +25 if median_turnaround_days ≤ 7 ('healthy'); 0 if 8-14 ('normal'); -10 if >14 ('slipping').",
        "Workload balance: -1 point per current open assignment (caps at -20).",
        "Total score per reviewer; sort descending; return top_n with score + reasons array (e.g. ['sector fit: education', 'country fit: Kenya', 'busy queue — assign carefully']).",
    ])
    body(doc,
        "Auto-assignment (Phase 25A) takes the top 3 by default, prefers "
        "'healthy' over 'slipping', and falls back to slipping reviewers "
        "if the healthy pool is exhausted (rather than leaving the panel "
        "empty)."
    )

    heading(doc, "7.6 Peer benchmarks", 2)
    body(doc,
        "peer_benchmark_service has two surfaces: for_ngo (compares "
        "against same-country NGOs on capacity, win rate, submission "
        "count) and for_donor (compares against other donors on decision "
        "speed, decline rate, portfolio size). Sample-size guard: "
        "MIN_PEERS=3; below threshold returns source='sparse' and the UI "
        "renders 'not enough peers' instead of fake confidence."
    )
    body(doc,
        "Math: median + percentile rank. Median = standard middle "
        "value. Percentile = (count_below + count_equal/2) / N * 100. "
        "Verdict bucket: above (≥75th percentile, higher-is-better) / "
        "around (33-75) / below (<33). Inverted for lower-is-better "
        "metrics like decision speed and decline rate."
    )

    heading(doc, "7.7 Donor cohort analytics", 2)
    body(doc,
        "donor_cohort_analytics_service compares a donor's portfolio "
        "QUALITY (the NGOs they fund) against cohort medians. Six "
        "metrics per donor:"
    )
    bullets(doc, [
        ("avg_grantee_capacity_score: ", "mean Assessment.total_score across the donor's grantees."),
        ("avg_ai_score_at_award: ", "mean Application.ai_score across the donor's awarded applications."),
        ("country_diversity_pct: ", "distinct grantee countries / total grantees × 100."),
        ("sector_diversity_pct: ", "distinct sectors across awarded grants / total grants × 100."),
        ("small_org_funding_share_pct: ", "fraction of grantees with capacity_score < 60 × 100."),
        ("report_on_time_rate_pct: ", "% of grantee reports submitted ≤ due_date."),
    ])
    body(doc,
        "Cohort = all other donors on the platform (caller excluded). "
        "Sample guards: MIN_COHORT=3 donors required for any comparison; "
        "MIN_SAMPLE=2 grantees required per metric (otherwise the metric "
        "is omitted entirely from the response). Percentile rank + "
        "verdict computed identically to §7.6. Anonymity: NEVER returns "
        "NGO names, donor names, grant ids, or specific dollar amounts."
    )

    heading(doc, "7.8 NPS calculation", 2)
    body(doc,
        "UserFeedbackService.nps_summary(days=30) returns:"
    )
    code(doc, """
overall_nps = ((promoters - detractors) / total_responses) * 100

  promoters  = count where score >= 9
  passives   = count where 7 <= score <= 8
  detractors = count where score <= 6

  per_surface: same formula scoped to (surface, window)
  per_language: same formula scoped to (language, window)
  histogram: count per integer score 0-10
""".strip())
    body(doc,
        "Sample guards: when total_responses == 0, overall_nps = null "
        "(not 0) and the histogram is all-zeros. Per-surface and "
        "per-language rollups are unbounded — even a single response is "
        "shown because they're triage signals, not statistical claims."
    )

    heading(doc, "7.9 Funnel + retention math", 2)
    body(doc,
        "UserEventService.funnel(stages, days) computes a naive funnel: "
        "for each stage in the ordered list, count distinct users who "
        "recorded that event_name in the window. NOT a strict ordered "
        "funnel (would require per-session tracking); enough to answer "
        "'how many people who opened chat actually sent a message?'."
    )
    code(doc, """
for each stage in stages:
    count = distinct(user_id) WHERE event_name = stage AND occurred_at >= since
    rate_vs_base_pct = (count / first_stage_count) * 100
""".strip())
    body(doc,
        "Active users: DAU/WAU/MAU = distinct(user_id) where "
        "occurred_at >= now - {1, 7, 30} days. Breakdowns by role + "
        "language returned in parallel from the same query for the "
        "per-cohort view."
    )

    heading(doc, "7.10 A/B bucketing", 2)
    body(doc,
        "ab_arm(experiment, org_id|user_id, arms=('A','B')) in "
        "app/utils/feature_flags.py:"
    )
    code(doc, """
def ab_arm(experiment, *, org_id=None, user_id=None, arms=('A','B')):
    subject = str(org_id) if org_id is not None else (str(user_id) if user_id else None)
    if subject is None:
        return None
    h = sha256(f"{experiment}|{subject}".encode()).digest()
    idx = h[0] % len(arms)
    return arms[idx]
""".strip())
    body(doc,
        "Stability: SHA-256 ensures the same subject always lands in "
        "the same arm for a given experiment name. Switching the arms "
        "tuple (e.g. ('control','variant','holdout')) gives a 3-way "
        "split. Subject-less calls return None and metrics surface "
        "those as '(unbucketed)'."
    )

    page_break(doc)

    # =================================================================
    # 8. EXTERNAL INTEGRATIONS + APIs
    # =================================================================
    heading(doc, "8. External Integrations + APIs", 1)

    heading(doc, "8.1 Anthropic Claude", 2)
    bullets(doc, [
        ("SDK: ", "anthropic 0.42.0 (Python)."),
        ("Model: ", "claude-sonnet-4-20250514 (pinned)."),
        ("Auth: ", "ANTHROPIC_API_KEY env var; lazy singleton client in AIService._anthropic_client."),
        ("Modes used: ", "Free-text generation (most chat surfaces), forced tool-use (scoring, gap analysis, criteria suggestion — structured output), vision (scanned PDFs)."),
        ("Cost tracking: ", "Every call writes an AICallLog row with prompt + completion tokens + cost. ai_budget_service.enforce_budget guards per-org caps."),
        ("Retry: ", "On HTTP 429 or 5xx, retry once with 2s backoff. On second failure, return success=False to caller (UI shows friendly fallback)."),
        ("Fallback: ", "ai_mock.py provides deterministic stub outputs for every method when key is missing — used by tests and offline dev."),
    ])

    heading(doc, "8.2 OpenSanctions + fallback (UN/OFAC/EU)", 2)
    bullets(doc, [
        ("Primary feed: ", "OpenSanctions API at https://api.opensanctions.org/match/sanctions + /match/peps. Authenticated with OPENSANCTIONS_API_KEY (Bearer 'ApiKey {key}' header)."),
        ("Match shape: ", "POST with { queries: { id: { schema: 'Person'|'Organization', properties: { name, country, birthDate? } } } } → returns matches with scores."),
        ("Dedup: ", "App-side filtering using country + name variants + DOB when available. Surfaces the dedup logic transparently so compliance officers see why a match was discarded."),
        ("Fallback (when API down or key missing): ", "Direct downloads of UN XML, OFAC SDN CSV, EU CSV files (cached locally + refreshed daily)."),
        ("Coverage: ", "UN consolidated, OFAC SDN, EU consolidated, plus AML/CTF lists where OpenSanctions has them."),
        ("Latency budget: ", "10s timeout; logs each check + outcome in ComplianceCheck. AdverseMediaScreening rows preserve the screening result with timestamp + matched-counts."),
        ("Continuous rescreening: ", "compliance_rerun_service runs nightly via /api/cron/compliance-rerun and re-checks orgs with active grants whose last screen is > 30 days old."),
    ])

    heading(doc, "8.3 Government registries (7 African countries)", 2)
    body(doc,
        "registry_service routes per-country verification via "
        "_verify_{country} methods. Coverage:"
    )
    table(doc,
        ["Country", "Registry", "Method"],
        [
            ["Kenya", "NGO Coordination Board", "API + HTML scrape (registry surface varies)"],
            ["Nigeria", "Corporate Affairs Commission (CAC)", "API"],
            ["South Africa", "Companies & Intellectual Property Commission (CIPC) + DSD NPO Register", "API"],
            ["Uganda", "Uganda Registration Services Bureau (URSB) + NGO Bureau", "Scrape"],
            ["Tanzania", "Registrar of NGOs", "Scrape"],
            ["Somalia", "Ministry of Planning (NGO Register)", "Scrape + manual where needed"],
            ["Ethiopia", "Federal Civil Society Agency (FCSA)", "Scrape"],
        ],
        widths=[1.8, 3.4, 1.5]
    )
    body(doc,
        "Roadmap expansion: registry coverage extends to additional "
        "Sub-Saharan jurisdictions through 2027 H1 per the GTM doc. "
        "Each new country = a new _verify_{country} method + an entry "
        "in the dispatch table."
    )

    heading(doc, "8.4 Sentry", 2)
    bullets(doc, [
        ("SDK: ", "sentry-sdk[flask] 2.0.0 with FlaskIntegration + SqlalchemyIntegration."),
        ("Initialised: ", "app/__init__.py when SENTRY_DSN env is set; release tagged to the current build hash."),
        ("Captures: ", "Unhandled exceptions, request transactions (sampled), SQL queries (sampled)."),
        ("PII strategy: ", "Default scrubbing on; user email + IP attached only when explicitly opted in."),
        ("Alerts: ", "Configured on Sentry dashboard (not in repo); error spike alerts → Slack."),
    ])

    heading(doc, "8.5 Web Push", 2)
    bullets(doc, [
        ("Library: ", "pywebpush 2.0.0."),
        ("Protocol: ", "VAPID keys (PUSH_VAPID_PUBLIC, PUSH_VAPID_PRIVATE env)."),
        ("Subscription: ", "Client requests permission, posts subscription endpoint + keys to /api/push/subscribe; stored in PushSubscription model."),
        ("Dispatch: ", "notification_dispatcher.dispatch() iterates a user's subscriptions and POSTs through pywebpush.webpush()."),
        ("Service worker: ", "/sw.js handles incoming push events + displays notification with deep-link click handler."),
    ])

    heading(doc, "8.6 Email transport", 2)
    bullets(doc, [
        ("Service: ", "email_service.send_email() wraps an SMTP client. Production uses Postmark-style transactional provider via SMTP_HOST + SMTP_USER + SMTP_PASS env vars."),
        ("Fallback: ", "Logs to console when SMTP env is unset (dev mode)."),
        ("Templates: ", "Jinja2 templates in app/templates/email/; subject + body rendered per category + user language."),
        ("Categories: ", "deadlines, decisions, compliance, reviews — match the NotificationPreference category vocab."),
        ("Bounce handling: ", "Logged but not currently retried; future hardening."),
    ])

    heading(doc, "8.7 SMS + WhatsApp adapters", 2)
    bullets(doc, [
        ("Service: ", "messaging_service abstracts both channels behind a single send() method."),
        ("SMS adapter: ", "Africa's Talking-compatible (env: AT_USERNAME + AT_API_KEY) OR Twilio (env: TWILIO_SID + TWILIO_TOKEN + TWILIO_FROM)."),
        ("WhatsApp adapter: ", "Twilio WhatsApp API (env: TWILIO_WA_FROM)."),
        ("Phone resolution: ", "Each NotificationPreference row carries phone_e164 + whatsapp_e164 fields; kept in sync across the user's category rows."),
    ])

    heading(doc, "8.8 Railway platform", 2)
    bullets(doc, [
        ("Project: ", "clever-cooperation; service id c47233a9-9828-4430-b75c-f11df546764a."),
        ("Deployment: ", "`railway up --detach` uploads source, Railway builds Docker image, rotates container."),
        ("Database: ", "Railway-managed PostgreSQL; DATABASE_URL auto-injected."),
        ("Env vars: ", "Set via Railway dashboard or CLI; loaded at container start. See Appendix A for canonical list."),
        ("Logs: ", "`railway logs --tail N` for live; persistent logs visible in dashboard."),
        ("Domain: ", "web-production-6f8a.up.railway.app (Railway-provided); custom domain pending."),
    ])

    heading(doc, "8.9 GitHub Actions (CI + crons)", 2)
    bullets(doc, [
        ("Workflows: ", ".github/workflows/smoke.yml + browser-smoke.yml + cron-compliance-rerun.yml + cron-uat-fixtures.yml + cron-reviewer-auto-assign.yml + e2e-regression.yml."),
        ("Triggers: ", "smoke runs on every push to main + every PR; browser smoke runs on app/, frontend/, static/nextjs/ changes; crons fire on schedule (UAT 03:15 UTC, compliance 02:30 UTC, reviewer sweep 02:45 UTC)."),
        ("Secrets: ", "CRON_SECRET (matches Railway env), PROD_BASE (overrides prod URL for staging tests)."),
        ("Cron auth: ", "POST with Authorization: Bearer $CRON_SECRET against /api/cron/* endpoints."),
        ("Step summary: ", "Each cron workflow writes a summary to $GITHUB_STEP_SUMMARY so failures show in the run UI."),
    ])

    page_break(doc)

    # =================================================================
    # 9. REAL-USER METRICS + FEEDBACK
    # =================================================================
    heading(doc, "9. Real-User Metrics + Feedback Infrastructure", 1)

    body(doc,
        "Phase 29-31 added a lightweight behavioural event store + NPS "
        "micro-survey infrastructure (BRD §6.20 covers the spec). "
        "Architecture details:"
    )
    bullets(doc, [
        ("UserEvent table: ", "id + user_id + org_id (denorm) + role + language + event_name + event_props JSON + ab_arm + occurred_at. Indexes: (user_id, occurred_at), (org_id, occurred_at), (event_name, occurred_at)."),
        ("Stable event vocab (12 events instrumented): ", "session.start, application.start_draft/.submit, report.start_draft/.submit/.preflight_used, readiness_check.used, chat.thread_open/.message_sent, search.query, donor.decision_recorded/.broadcast_sent, reviewer.assignment_opened/.review_submitted, trust_profile.viewed."),
        ("Non-blocking discipline: ", "UserEventService.record() catches all exceptions; a failed event write NEVER raises in the user-visible request path."),
        ("Aggregation API: ", "active_users(days), event_counts(days), funnel(stages, days), feature_usage_by_language(event_name, days), ab_outcome(outcome_event, days)."),
        ("A/B helper: ", "ab_arm() in app/utils/feature_flags.py — see §7.10 for math."),
        ("Generic event ingest: ", "POST /api/ai/events/track with whitelisted event_name + arbitrary props (capped at 4KB). Frontend uses this for client-only events like ai_assist.suggestion_accepted."),
        ("UserFeedback table: ", "Same shape with score (0-10) + comment (≤500 chars). UNIQUE constraint on (user_id, surface, related_kind, related_id) prevents duplicates."),
        ("NPS dashboard: ", "/admin/metrics surfaces overall NPS + by_surface + by_language + recent_comments stream."),
    ])

    page_break(doc)

    # =================================================================
    # 10. INTERNATIONALISATION
    # =================================================================
    heading(doc, "10. Internationalisation (i18n)", 1)
    bullets(doc, [
        ("Supported languages: ", "English (en), French (fr), Arabic (ar), Swahili (sw), Somali (so), Spanish (es)."),
        ("Frontend dictionaries: ", "frontend/src/i18n/{en,fr,ar,sw,so,es}.json — flat key/value, ~2100 keys each (parity across locales). Loaded once at module init; switching language re-renders without reload."),
        ("Translation function: ", "i18n/index.ts translate(key, lang, params) — returns translations[lang]?.[key] ?? translations.en?.[key] ?? key. Param interpolation via {placeholder} replacement."),
        ("Backend dictionaries: ", "app/translations/{lang}.json — same shape, used for server-rendered messages (email subjects, notification body templates)."),
        ("Language preference: ", "User.language column; PUT /api/auth/language to switch. Persisted across sessions."),
        ("Tone discipline: ", "NGO surfaces (apply, dashboard, reports) lean coaching + supportive; donor surfaces (dashboard, decisions, broadcast) lean precise + executive. Maintained per locale."),
        ("RTL support: ", "Arabic + Somali. <html dir='rtl'> applied via useTranslation hook on language switch. Tailwind RTL plugin handles padding/margin mirroring."),
        ("AI replies: ", "Every chat + AI surface receives a 'User language: {lang}' line in the system prompt so model replies in the user's locale."),
        ("Tests: ", "TC-550-557 (RTL + i18n edge cases), TC-406-407 (chat in Arabic + Swahili), TC-553 (cohort verdict pills in French)."),
    ])

    page_break(doc)

    # =================================================================
    # 11. FRONTEND ARCHITECTURE
    # =================================================================
    heading(doc, "11. Frontend Architecture", 1)
    bullets(doc, [
        ("Build: ", "Next.js 14 App Router with output: 'export'. Static HTML + JS generated by `next build`, copied to /static/nextjs/ by scripts/copy-build.js. Flask serves /static/nextjs/* under the root path."),
        ("Routing: ", "App-Router file-based. Authenticated pages under (app)/ group share a layout with sidebar + header; public pages under (auth)/."),
        ("State: ", "Zustand stores for auth (auth-store.ts) + UI (ui-store.ts). SWR for server data fetching with cache + revalidate-on-focus."),
        ("API helper: ", "frontend/src/lib/api.ts — apiFetch(method, path, body) handles JSON serialisation, CSRF X-Requested-With header, 401 redirect, transient 5xx retry, /api/api/ normalisation (Phase 27 fix). Exposes api.get/.post/.put/.patch/.delete/.upload."),
        ("Auth gate: ", "401 responses trigger window.location.href = '/login' (handled in api.ts). Session check via /api/auth/me on app boot."),
        ("Component library: ", "shadcn/ui primitives in frontend/src/components/ui; Kuja-specific components in frontend/src/components/{dashboards,copilot,grants,applications,etc.}."),
        ("Design tokens: ", "Kuja Studio palette in globals.css as CSS custom properties (--kuja-clay, --kuja-savanna, --kuja-spark, etc.). All components reference tokens, not hex codes."),
        ("PWA: ", "manifest.webmanifest + sw.js in /public/. Service worker registered on first paint by app-registry.tsx. Install banner (shared/pwa-install-banner.tsx) captures beforeinstallprompt + persists dismiss in localStorage."),
        ("Bundle size: ", "Tree-shaken; common chunks split. Main bundle ~250 KB gz after Tailwind purge + Next.js code-splitting."),
    ])

    page_break(doc)

    # =================================================================
    # 12. BACKGROUND JOBS + CRONS
    # =================================================================
    heading(doc, "12. Background Jobs + Crons", 1)
    body(doc,
        "Three classes of background work: in-process async AI jobs, "
        "GitHub-Actions-driven crons, and the per-process scheduler."
    )

    heading(doc, "In-process async AI jobs", 3)
    body(doc,
        "ai_jobs.py spawns a daemon thread per submitted job and "
        "stores the result in an in-memory _AI_JOBS dict. Single-worker "
        "design means cross-request job lookup works without external "
        "state. Acceptable for current load (3-30 concurrent jobs); "
        "when scaling out, _AI_JOBS would move to Redis."
    )

    heading(doc, "GitHub Actions crons", 3)
    table(doc,
        ["Cron", "Schedule", "Purpose"],
        [
            ["cron-compliance-rerun", "02:30 UTC daily", "Re-screen NGOs whose adverse media check is > 30 days old"],
            ["cron-reviewer-auto-assign", "02:45 UTC daily", "Sweep applications in {submitted, under_review} with zero reviews; auto-assign top-3"],
            ["cron-uat-fixtures", "03:15 UTC daily", "Ensure each donor has open + awarded + rejected apps + a published report bundle"],
        ],
        widths=[2.4, 1.6, 3.0]
    )
    body(doc,
        "All cron endpoints require Bearer CRON_SECRET (matches Railway "
        "env) OR an admin session. Each runs idempotently — re-running "
        "is safe + reports zero-deltas in its drift summary."
    )

    heading(doc, "Per-process scheduler", 3)
    body(doc,
        "Rescreening scheduler (app/__init__.py) runs daily compliance "
        "snapshots into ComplianceSnapshot if RESCREENING_SCHEDULER=true. "
        "Disabled by default in prod (the cron handles it); enabled in "
        "single-instance dev to keep dashboards populated."
    )

    page_break(doc)

    # =================================================================
    # 13. OBSERVABILITY + OPS
    # =================================================================
    heading(doc, "13. Observability + Operations", 1)
    bullets(doc, [
        ("System health: ", "/api/health returns 200 with status='healthy' if DB ping succeeds; /api/ready also checks DB + critical envs."),
        ("AI surface health: ", "/api/admin/ai-surface-health probes 7 flagship surfaces (chat, scoring, briefing, gap insights, etc.) with timeout + reports per-surface pass/fail + latency."),
        ("System health detail: ", "/api/admin/system-health checks OpenSanctions + Anthropic + DB + email transport + Sentry connection + storage + recent error rate. Returns overall='ok'/'degraded'/'down'."),
        ("Metrics dashboard: ", "/admin/metrics covered in §9."),
        ("Audit chain viewer: ", "/admin/audit-chain — chronological table with action + actor + subject + timestamp + hash verify status."),
        ("Failed login analytics: ", "/api/admin/security-events lists recent failed logins + lockouts."),
        ("Demo readiness: ", "/api/admin/demo-readiness scans for empty fixtures / missing data that would break a demo."),
        ("Logging: ", "Structured logging via Python logging; every request logged with method + path + status + duration + user + request_id."),
        ("Sentry: ", "Unhandled exceptions captured; sampled performance traces; release tagged per build."),
    ])

    page_break(doc)

    # =================================================================
    # 14. DEPLOYMENT PIPELINE
    # =================================================================
    heading(doc, "14. Deployment Pipeline", 1)
    numbered(doc, [
        "Developer makes changes locally. Runs `py -3 smoke_test.py` — must be 137/137 PASS.",
        "If frontend changed: `cd frontend && npm run build` — copies output to /static/nextjs/.",
        "git commit + push to main.",
        "GitHub Actions runs smoke + browser regression on the PR. Required to pass for merge (when PR mode is used; in solo mode pushed directly).",
        "Developer runs `railway up --detach` from the repo root. Railway uploads source, builds Docker image, rotates container.",
        "Container takes 4-5 min to be fully serving traffic (old container drains, new one comes up).",
        "Verify by hitting /api/version (build hash) + /api/health.",
        "Run smoke from outside (curl + cookies) against prod to confirm critical paths work.",
        "If issue: roll back via Railway dashboard (one-click revert to prior deploy).",
    ])
    body(doc,
        "Frontend build IDs change on every `npm run build`, so the "
        "/api/version response's frontend_build field is a strong deploy "
        "fingerprint. Backend build is shorter — first 12 chars of the "
        "git commit hash."
    )

    page_break(doc)

    # =================================================================
    # 15. PERFORMANCE + SCALABILITY
    # =================================================================
    heading(doc, "15. Performance + Scalability", 1)
    bullets(doc, [
        ("Current scale: ", "Single Gunicorn worker; ~600 active users; ~200 organisations. Sufficient for current load."),
        ("Concurrency model: ", "Sync handlers + threaded async AI jobs. Long AI calls don't block other endpoints because the route returns 202 + job_id immediately."),
        ("DB connection pool: ", "SQLAlchemy default (5 pool + 10 overflow). Sufficient; haven't observed pool exhaustion in prod."),
        ("Caching: ", "_dashboard_cache (30s TTL, in-process) for dashboard aggregations. AI job results cached in-process via ai_jobs._AI_JOBS until cleanup."),
        ("Rate limits: ", "Covered in §5.4."),
        ("AI cost cap: ", "Per-org monthly budget via ai_budget_service. Forecasted spend visible on /admin/ai-budget."),
        ("Scale-out plan: ", "When > 1 worker needed: (a) move _AI_JOBS + ai_limiter to Redis; (b) bump worker count from 1 to 3-4; (c) verify rate limits still scope correctly per IP (they do — middleware reads X-Forwarded-For); (d) confirm rescreening scheduler runs on exactly one worker (use distributed_lock)."),
        ("SLOs: ", "First chat reply p95 ≤ 15s; follow-up chat reply p95 ≤ 10s; dashboard TTI ≤ 5s; verdict card AI synthesis ≤ 15s; report pre-flight p95 ≤ 25s; search response p95 ≤ 1s; /api/health < 200ms always. Covered in test cases TC-780-789."),
    ])

    page_break(doc)

    # =================================================================
    # 16. TESTING ARCHITECTURE
    # =================================================================
    heading(doc, "16. Testing Architecture", 1)
    bullets(doc, [
        ("Python smoke tests: ", "smoke_test.py — 153/153 cases, runs in-process Flask app on a free port + executes against it. Covers every blueprint + critical regression points."),
        ("Browser UAT tests: ", "browser_test.py + browser_test_strict.py (Playwright). 100+ specs across 23 categories incl. Phase 27 strict user-flow tests that catch the 2026-05-16 double-prefix wiring class of bug."),
        ("Test cases (manual UAT): ", "354 cases across 37 categories in Kuja_Grant_v5.0_Test_Cases.xlsx (covered in detail in the test-cases doc)."),
        ("Test fixtures: ", "53 files (22 txt + 31 binary including PDF/XLSX/CSV/HTML/EXE/Unicode/AI-adversarial) in kuja-test-files-v5.zip."),
        ("CI gates: ", "smoke must pass on every push; browser smoke must pass when frontend/backend changes."),
        ("Mocks: ", "ai_mock.py for AI calls when ANTHROPIC_API_KEY missing; deterministic outputs so tests pass offline."),
        ("UAT fixture cron: ", "Daily cron ensures demo state stays meaningful — each donor has open + awarded + rejected apps + published bundle."),
    ])

    page_break(doc)

    # =================================================================
    # 17. TECH DEBT + FUTURE HARDENING
    # =================================================================
    heading(doc, "17. Tech Debt + Future Hardening", 1)
    body(doc,
        "Known limitations documented honestly. None block current scale; "
        "all have concrete paths forward."
    )
    bullets(doc, [
        ("Single Gunicorn worker: ", "Wedges under sustained AI load (Phase 27 wall — 60/min AI rate limit hit a hard ceiling). Walked back to 40/min; long-term: scale workers + move state to Redis."),
        ("In-process AI job store: ", "_AI_JOBS dict doesn't survive restart + doesn't share across workers. Acceptable today; move to Redis when scaling out."),
        ("require_reauth() wired but not gated: ", "WebAuthn helper ready but no specific route currently uses it. Pending team input on which sensitive actions to gate (likely candidates: decision recording, document downloads, notification preference changes)."),
        ("Payment integration: ", "Only deferred high-priority backlog item; everything around money — award, debrief, audit, bundles — exists; the actual disbursement is the gap. Blocking on finance ops sign-off + real Stripe/Flutterwave keys."),
        ("Outcome metrics A/B: ", "Phase 29-31 infrastructure is live + data is accumulating; first experiment not yet wired. Awaiting 30 days of baseline data to design the first experiment."),
        ("Deep i18n per locale: ", "Translations are present + parity is enforced. Native-speaker QA per locale waits for usage data — invest where users actually are, not across all 5 non-English locales upfront."),
        ("Soft FK polymorphism: ", "Risk.subject_kind/id + similar are app-enforced, not DB-enforced. Trade-off accepted because the polymorphism crosses 4-5 entities; tightening would require explicit join tables per (subject, target) pair."),
        ("Single SMTP provider: ", "No retry / bounce-handling logic. Future: connect to a service like Postmark + handle bounces + maintain reputation."),
        ("Web push: ", "Endpoint deduplication is naive — same browser may register multiple PushSubscription rows over time. Future: dedupe on endpoint URL + reap stale subscriptions."),
    ])

    page_break(doc)

    # =================================================================
    # APPENDIX A: ENV VARS
    # =================================================================
    heading(doc, "Appendix A: Environment Variables", 1)
    table(doc,
        ["Variable", "Required", "Purpose"],
        [
            ["DATABASE_URL", "yes", "Postgres connection string (Railway auto-injects)"],
            ["SECRET_KEY", "yes", "Flask session signing key (generate with secrets.token_hex(32))"],
            ["ANTHROPIC_API_KEY", "yes for AI", "Anthropic Claude API key"],
            ["OPENSANCTIONS_API_KEY", "yes for sanctions", "OpenSanctions API key"],
            ["SENTRY_DSN", "no (recommended)", "Sentry DSN; disables Sentry when unset"],
            ["CRON_SECRET", "yes for crons", "Bearer secret for /api/cron/* endpoints"],
            ["RATE_LIMIT_LOGIN_PER_IP_PER_MIN", "no (default 30)", "Per-IP login attempts per minute"],
            ["PUSH_VAPID_PUBLIC + PUSH_VAPID_PRIVATE", "yes for web push", "VAPID key pair"],
            ["SMTP_HOST + SMTP_USER + SMTP_PASS + SMTP_FROM", "yes for email", "SMTP transport config"],
            ["AT_USERNAME + AT_API_KEY", "yes for SMS via Africa's Talking", "SMS auth"],
            ["TWILIO_SID + TWILIO_TOKEN + TWILIO_FROM + TWILIO_WA_FROM", "yes for SMS/WhatsApp via Twilio", "Twilio auth + sender"],
            ["WEBAUTHN_ORIGIN + WEBAUTHN_RP_ID", "yes for WebAuthn", "Origin + RP id for FIDO2 verification"],
            ["RESCREENING_SCHEDULER", "no (default false)", "Enable per-process daily snapshot scheduler"],
            ["KUJA_FF_<KEY>", "no", "Per-feature-flag overrides (env > DB > default)"],
            ["FLASK_ENV", "no (default production)", "Set to 'development' for verbose errors"],
        ],
        widths=[2.6, 1.4, 3.0]
    )

    page_break(doc)

    # =================================================================
    # APPENDIX B: API ENDPOINT CENSUS
    # =================================================================
    heading(doc, "Appendix B: API Endpoint Census", 1)
    body(doc,
        "Sampled enumeration; full enumeration available via `flask "
        "routes` against a running app. See §3.3 for the canonical "
        "blueprint catalogue. Roughly 220 routes across 45 blueprints."
    )
    table(doc,
        ["Category", "Sample endpoints", "Methods"],
        [
            ["Auth + 2FA", "/api/auth/{login,logout,me,language}, /api/auth/totp/{status,enroll/start,enroll/confirm,verify,disable}, /api/auth/webauthn/{credentials,register/begin,register/finish,authenticate/begin,authenticate/finish}", "POST/GET/PUT/DELETE"],
            ["Dashboard", "/api/dashboard/{today,stats,benchmarks,portfolio-risk-heatmap,donor-cohort-analytics,onboarding,reviewer-throughput}", "GET"],
            ["Grants", "/api/grants/, /api/grants/<id>, /api/grants/<id>/broadcast, /api/grants/fit-compare, /api/grants/<id>/upload-grant-doc", "GET/POST/PUT/DELETE"],
            ["Applications", "/api/applications/, /api/applications/<id>, /api/applications/<id>/{submit,debrief,auto-assign-reviewers,status,timeline,signals,comments,compare}", "GET/POST/PATCH"],
            ["Reviews", "/api/reviews/, /api/reviews/<id>, /api/reviews/<id>/complete, /api/reviews/<id>/briefing", "GET/POST/PUT"],
            ["Reports", "/api/reports/, /api/reports/<id>, /api/reports/<id>/{submit,precheck,review,bundle,attachments}", "GET/POST/PUT"],
            ["AI", "/api/ai/{chat,guidance,strengthen-section,jobs/<id>,extract-evidence,burden-estimate,grant-brief,fit-compare,submission-readiness,threads/{open,/<id>/messages,/<id>/reset},events/track,calls/<id>/feedback}", "POST/GET/PATCH"],
            ["Trust", "/api/trust-profile/<id>, /api/trust-profile/<id>/gap-insights, /api/adverse-media/<id>, /api/adverse-media/screen, /api/bank-verification/<id>, /api/bank-verification/verify, /api/passport/<id>, /api/passport/publish, /api/passport/<id>/revoke", "GET/POST"],
            ["Compliance + Risk", "/api/compliance/screen, /api/risks/, /api/signals/, /api/preemption/me", "GET/POST/PATCH/DELETE"],
            ["Discovery", "/api/search?q=, /api/documents/search, /api/documents/search/global, /api/saved-searches, /api/watchlist/toggle, /api/match", "GET/POST/PATCH/DELETE"],
            ["Notifications", "/api/notifications/, /api/notification-preferences, /api/notifications/preferences, /api/push/subscribe", "GET/PUT/POST"],
            ["Admin", "/api/admin/{stats,canary,system-health,ai-surface-health,demo-readiness,security-events,perf-budgets,clear-lockout,clear-all-lockouts,orgs/merge,metrics,trigger-rescreening}, /api/audit-chain/{recent,verify}", "GET/POST"],
            ["Cron", "/api/cron/{compliance-rerun,uat-fixtures,reviewer-auto-assign-sweep}", "POST"],
            ["Exports", "/api/exports/{grants,applications,reviews}.csv", "GET"],
            ["Bundles", "/api/portfolio/audit-timeline, /api/portfolio/bundle, /api/reports/<id>/bundle, /api/calendar/deadlines, /api/calendar/pdf", "GET"],
            ["Feedback", "/api/feedback, /api/feedback/my", "POST/GET"],
            ["Health", "/api/health, /api/ready, /api/version", "GET"],
        ],
        widths=[1.4, 4.4, 1.2]
    )

    page_break(doc)

    # =================================================================
    # APPENDIX C: TOOLS INVENTORY
    # =================================================================
    heading(doc, "Appendix C: Tools Inventory (what + why)", 1)
    body(doc,
        "Every external service + library currently in use, with the "
        "specific Kuja capability it enables. Cross-reference with §2 "
        "for version pins."
    )
    table(doc,
        ["Tool", "Category", "What it does for Kuja", "Why this choice"],
        [
            ["Anthropic Claude (claude-sonnet-4-20250514)", "AI", "Powers all 28 AI surfaces — drafting, scoring, briefing, agent tool-use, vision OCR on scanned PDFs", "Best-in-class reasoning, strong vision, official Python SDK with native tool-use schema, cost-effective at our scale"],
            ["OpenSanctions API", "Compliance", "Sanctions + PEP + adverse media screening of orgs, principals, beneficial owners", "Live, well-maintained, single endpoint covering UN+OFAC+EU+regional; saves us building + maintaining 5+ separate scrape pipelines"],
            ["UN/OFAC/EU direct downloads", "Compliance fallback", "Sanctions screening fallback when OpenSanctions API is down or key missing", "Authoritative lists; no API dependency for the critical-path compliance check"],
            ["Government registries (7 countries)", "Verification", "Verify NGO registration + standing in Kenya/Nigeria/SA/Uganda/Tanzania/Somalia/Ethiopia", "Direct authoritative source; no third-party intermediary"],
            ["Sentry", "Observability", "Captures unhandled exceptions, sampled traces, release-tagged error attribution", "De-facto industry standard with mature Flask + SQLAlchemy integrations"],
            ["Railway", "Hosting", "Managed PaaS with built-in Postgres + HTTPS + env-var management", "Single-command deploys (`railway up`); fits a small team without ops overhead; cheaper than AWS for our scale"],
            ["PostgreSQL (Railway-managed)", "Data", "Production database for all 31 models + audit chain + event ledger", "ACID, mature JSON support, well-understood scaling path, Railway-managed"],
            ["SQLite", "Dev data", "Local development + test runs", "Zero-config; same SQL surface as Postgres (with helpers for TZ-aware timestamps)"],
            ["Anthropic SDK (Python)", "AI integration", "Wraps Claude API calls (free-text, tool-use, vision)", "Official SDK, kept current with model updates"],
            ["pyotp", "Security", "TOTP 2FA secrets + verification", "RFC 6238 standard; works with Google Authenticator / Authy / 1Password"],
            ["webauthn (Python)", "Security", "WebAuthn registration + assertion + sign-count clone detection", "Spec-compliant; handles platform + cross-platform authenticators"],
            ["pywebpush + VAPID", "Notifications", "Browser push notifications via service worker", "Open-standard; no vendor lock-in (vs FCM/APNs SDKs)"],
            ["Africa's Talking OR Twilio", "SMS + WhatsApp", "SMS + WhatsApp notification fan-out", "Africa's Talking has best African mobile-network coverage + pricing; Twilio is the fallback + WhatsApp Business connector"],
            ["SMTP (Postmark-class)", "Email", "Transactional email transport", "Pluggable via SMTP_HOST env; supports any SMTP provider"],
            ["PyPDF2", "Document", "Extract text from uploaded PDFs", "Pure-Python, no native deps"],
            ["python-docx", "Document", "Read + write .docx files", "Used both for parsing uploads + generating BRD/test-cases docs"],
            ["openpyxl", "Document", "Read + write .xlsx files", "Used both for parsing uploads + generating test-cases workbook"],
            ["reportlab", "Document", "Generate PDF bundles + calendars + reviewer take-home PDFs", "Mature; supports complex layouts (cover pages, tables, page numbers)"],
            ["GitHub Actions", "CI + crons", "Smoke + browser tests; daily cron triggers for compliance + reviewer sweep + UAT fixtures", "Free for our scale; tight integration with repo"],
            ["Flask + Flask-Login + Flask-SQLAlchemy + Flask-Migrate + Flask-CORS", "Backend", "Web framework + session auth + ORM + migrations + CORS", "Mature, well-understood, blueprint pattern fits our per-feature boundary"],
            ["Gunicorn", "Server", "Production WSGI server", "Battle-tested with Flask; sync workers match our sync-style handlers"],
            ["Next.js 14 (App Router, static export)", "Frontend", "React app shell + routing", "Static export → no server-side JS in prod → single deploy unit served by Flask"],
            ["shadcn/ui + Tailwind", "Frontend UI", "Component library + utility CSS", "Headless = full design control; Kuja Studio palette implementable; tree-shakeable"],
            ["zustand + swr", "Frontend state", "Auth/UI store + cached server data", "Minimal API surface; better than Redux for our scope"],
            ["Recharts", "Frontend charts", "Funnel + bar + heatmap + score-ring charts", "React-native; composable; SSR-safe"],
            ["Playwright", "Testing", "Browser UAT against in-process Flask + against prod URL", "Cross-browser, fast, robust selectors; Chromium-default keeps CI lean"],
        ],
        widths=[1.7, 1.0, 2.7, 1.8]
    )

    # =================================================================
    # SAVE
    # =================================================================
    out_path = os.path.join(
        os.path.dirname(__file__),
        "Kuja_Grant_v5.0_Technical_Design_Specification.docx",
    )
    doc.save(out_path)
    print(f"Generated: {out_path}")
    print(f"Size: {os.path.getsize(out_path):,} bytes")


if __name__ == "__main__":
    build()
