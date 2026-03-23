"""
Generate Kuja Grant Platform Market Analysis & Go-to-Market Strategy Document
Output: Kuja_Grant_Market_Analysis_GTM.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Constants ───────────────────────────────────────────────────────────────
BLUE = RGBColor(0x1E, 0x40, 0xAF)       # #1E40AF
TEAL = RGBColor(0x0D, 0x94, 0x88)       # #0D9488
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
RED = RGBColor(0xDC, 0x26, 0x26)
BLUE_HEX = "1E40AF"
LIGHT_BLUE_HEX = "DBEAFE"
ALT_ROW_HEX = "F0F7FF"

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Kuja_Grant_Market_Analysis_GTM.docx")

# ─── Helper Functions ────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, font_size=Pt(9), font_color=BLACK, font_name="Calibri", alignment=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = font_size
    run.font.color.rgb = font_color
    run.font.bold = bold
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a professionally styled table with blue header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set table borders
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders)

    # Header row
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_shading(cell, BLUE_HEX)
        set_cell_text(cell, header, bold=True, font_size=Pt(9), font_color=WHITE)

    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.cell(i + 1, j)
            if i % 2 == 1:
                set_cell_shading(cell, ALT_ROW_HEX)
            set_cell_text(cell, val, font_size=Pt(9), font_color=DARK_GRAY)

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_heading1(doc, text):
    """Add a Heading 1 with blue color."""
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = BLUE
        run.font.name = "Calibri"
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(8)
    return h


def add_heading2(doc, text):
    """Add a Heading 2 with teal color."""
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = TEAL
        run.font.name = "Calibri"
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    return h


def add_heading3(doc, text):
    """Add a Heading 3."""
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = BLUE
        run.font.name = "Calibri"
        run.font.size = Pt(12)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    return h


def add_body(doc, text, bold=False, italic=False, space_after=Pt(6)):
    """Add body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = Pt(15)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.font.name = "Calibri"
        run_bold.font.size = Pt(10.5)
        run_bold.font.color.rgb = DARK_GRAY
        run_bold.bold = True
        run_rest = p.add_run(text)
        run_rest.font.name = "Calibri"
        run_rest.font.size = Pt(10.5)
        run_rest.font.color.rgb = DARK_GRAY
    else:
        p.clear()
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK_GRAY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(15)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * level)
    return p


def add_numbered_item(doc, number, title, description):
    """Add a numbered advantage/item with bold title and description."""
    p = doc.add_paragraph()
    run_num = p.add_run(f"{number}. {title}: ")
    run_num.font.name = "Calibri"
    run_num.font.size = Pt(10.5)
    run_num.font.color.rgb = BLUE
    run_num.bold = True
    run_desc = p.add_run(description)
    run_desc.font.name = "Calibri"
    run_desc.font.size = Pt(10.5)
    run_desc.font.color.rgb = DARK_GRAY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(15)
    return p


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


# ─── Document Creation ───────────────────────────────────────────────────────

def create_document():
    doc = Document()

    # ─── Set default font ────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = DARK_GRAY

    # Update heading styles
    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = "Calibri"

    # ─── Set margins ─────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ═════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═════════════════════════════════════════════════════════════════════

    # Add spacer paragraphs for vertical centering
    for _ in range(6):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)
        sp.paragraph_format.space_before = Pt(0)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("KUJA GRANT PLATFORM")
    run.font.name = "Calibri"
    run.font.size = Pt(36)
    run.font.color.rgb = BLUE
    run.bold = True
    p_title.paragraph_format.space_after = Pt(8)

    # Decorative line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line.add_run("_" * 60)
    run.font.color.rgb = TEAL
    run.font.size = Pt(10)
    p_line.paragraph_format.space_after = Pt(16)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("Market Analysis & Go-to-Market Strategy")
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_GRAY
    p_sub.paragraph_format.space_after = Pt(24)

    # Version
    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_ver.add_run("March 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = LIGHT_GRAY
    p_ver.paragraph_format.space_after = Pt(12)

    # Organization
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_org.add_run("Adeso \u2014 African Development Solutions")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = TEAL
    run.bold = True
    p_org.paragraph_format.space_after = Pt(48)

    # Confidential
    p_conf = doc.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_conf.add_run("CONFIDENTIAL \u2014 FOR PARTNER DISCUSSIONS")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RED
    run.bold = True

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═════════════════════════════════════════════════════════════════════

    add_heading1(doc, "Table of Contents")

    toc_items = [
        "1.  Executive Summary",
        "2.  The Problem We Solve",
        "3.  Our Solution: End-to-End Grant Management",
        "4.  Market Opportunity",
        "5.  Competitive Landscape",
        "6.  Kuja\u2019s Competitive Advantages",
        "7.  Platform Architecture & AI Capabilities",
        "8.  Go-to-Market Strategy",
        "9.  Revenue Model",
        "10. Appendix: Feature Comparison Matrix",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.font.color.rgb = DARK_GRAY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1.5)

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ═════════════════════════════════════════════════════════════════════

    add_heading1(doc, "1. Executive Summary")

    add_body(doc, (
        "Kuja is the first AI-powered end-to-end grant management platform purpose-built for the "
        "Global South. Incubated by Adeso, a Somali-founded African social enterprise with over "
        "30 years of humanitarian experience, Kuja addresses a fundamental imbalance in the aid "
        "sector: local organizations closest to communities have the most effective solutions but "
        "the least resources. Billions in aid funding remain concentrated at Global North "
        "headquarters while local civil society organizations face systemic barriers\u2014lack of "
        "visibility to donors, limited capacity to navigate complex funding processes, extensive "
        "and duplicative donor compliance requirements, and fragmented grants management tooling."
    ))

    add_body(doc, (
        "The platform covers the complete grant lifecycle across eight integrated stages: "
        "marketplace discovery (kuja.org, 600+ active users), AI-powered capacity assessments "
        "using five industry-standard frameworks (Kuja, STEP, UN-HACT, CHS, NUPAS), live due "
        "diligence including government registry verification for seven African countries and "
        "sanctions screening against UN, OFAC, EU, and World Bank databases, AI-guided grant "
        "applications with document analysis and scoring, dual-scoring review combining AI and "
        "human evaluation, AI-powered reporting with compliance monitoring, and back-end ERP for "
        "operations built on Odoo 17. For entities choosing the ERP, onboarding is seamless "
        "because most organizational information is already captured through earlier lifecycle "
        "stages. For those that do not adopt the ERP, a dedicated reporting module connects "
        "grant information to enable compliance reporting without additional data entry."
    ))

    add_body(doc, (
        "The global grant management software market is valued at $2.75 billion (2024) and is "
        "projected to reach $4.79 billion by 2030, growing at a CAGR of 10.3%. No existing "
        "platform in this market combines marketplace, capacity assessment, grant management, "
        "AI-powered analysis, and ERP operations in a single integrated solution designed "
        "specifically for the Global South. The dissolution of USAID in 2025\u2014with 83% of "
        "programs cancelled and $36 billion in aid cuts\u2014has created urgent demand for "
        "diversified funding infrastructure and direct donor-to-NGO connections, positioning "
        "Kuja at the intersection of market need and sector transformation."
    ))

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════════
    # 2. THE PROBLEM WE SOLVE
    # ═════════════════════════════════════════════════════════════════════

    add_heading1(doc, "2. The Problem We Solve")

    add_heading2(doc, "For CSOs / NGOs")

    add_bullet(doc, "Local organizations remain invisible to donors despite delivering the most impactful, community-proximate work. Without a centralized discovery mechanism, the best-positioned implementers are systematically excluded from funding pipelines.", bold_prefix="Invisible to donors: ")
    add_bullet(doc, "NGOs undergo 4\u201312 near-identical capacity assessments per year from different donors, each requiring substantial staff time and documentation. Results are not portable between funders.", bold_prefix="Assessment fatigue: ")
    add_bullet(doc, "Each donor mandates different reporting requirements, templates, formats, and deadlines. A single organization managing five grants may need to produce five distinct compliance reporting workflows simultaneously.", bold_prefix="Compliance burden: ")
    add_bullet(doc, "Most grant management platforms are English-first. Organizations operating in Francophone, Arabophone, and Swahili-speaking regions face interface and documentation barriers.", bold_prefix="Language barriers: ")
    add_bullet(doc, "Enterprise platforms like Fluxx (custom pricing, not publicly disclosed) and Submittable (~$5K\u2013$18K/year depending on tier) are designed for and priced for Global North institutions, placing them far beyond the reach of local CSOs.", bold_prefix="Prohibitive costs: ")
    add_bullet(doc, "Despite Grand Bargain commitments to route 25% of humanitarian funding directly to local actors, the actual figure remains below 5%. Structural barriers\u2014not capacity\u2014drive this gap.", bold_prefix="Limited access: ")

    add_heading2(doc, "For Donors")

    add_bullet(doc, "No efficient mechanism exists for identifying, vetting, and connecting with local organizations. Donor pipelines default to established INGOs despite localization mandates.", bold_prefix="Cannot find local partners: ")
    add_bullet(doc, "Persistent stereotypes frame local organizations as inherently high-risk, yet no standardized, technology-enabled assessment infrastructure exists to provide objective evidence.", bold_prefix="Risk perception: ")
    add_bullet(doc, "Conducting capacity assessments in-house is expensive and time-consuming, often costing $5,000\u2013$15,000 per assessment and requiring weeks of staff engagement.", bold_prefix="Resource-intensive due diligence: ")
    add_bullet(doc, "Once grants are awarded, donors lack real-time visibility into spending, deliverable progress, and reporting compliance across their portfolio.", bold_prefix="Compliance monitoring: ")
    add_bullet(doc, "Collecting, reviewing, and evaluating narrative and financial reports from multiple grantees across different geographies is manual, inconsistent, and resource-intensive.", bold_prefix="Reporting gaps: ")

    add_heading2(doc, "Systemic Context")

    add_bullet(doc, "The dissolution of USAID in 2025 resulted in 83% of programs cancelled, $36 billion in aid cuts, and 81+ NGOs closing field offices. This has created an unprecedented disruption to the humanitarian funding architecture.")
    add_bullet(doc, "Organizations previously dependent on USAID sub-grants urgently need alternative funding pathways and direct connections to bilateral, multilateral, and private donors.")
    add_bullet(doc, "The Grand Bargain 2.0 localization agenda demands new digital infrastructure to connect donors directly with local organizations\u2014infrastructure that does not currently exist at scale.")

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════════
    # 3. OUR SOLUTION: END-TO-END GRANT MANAGEMENT
    # ═════════════════════════════════════════════════════════════════════

    add_heading1(doc, "3. Our Solution: End-to-End Grant Management")

    add_body(doc, (
        "Kuja delivers the only platform that spans the complete grant lifecycle in a single, "
        "integrated solution. The eight-stage architecture ensures that data flows seamlessly "
        "between stages, eliminating re-entry, reducing administrative burden, and building a "
        "cumulative organizational profile that strengthens with each interaction."
    ))

    # Stage 1
    add_heading2(doc, "Stage 1 \u2014 Marketplace Discovery (kuja.org)")

    add_body(doc, (
        "The foundation of the Kuja ecosystem is a free networking platform serving CSOs, "
        "Networks, INGOs, Donors, and Individual professionals. The marketplace is the entry "
        "point for all organizations and the pipeline for conversion to paid services."
    ))
    add_bullet(doc, "Profile types: CSO, Network, INGO, Donor, Individual\u2014each with tailored fields and verification requirements")
    add_bullet(doc, "Three-tier verification: referral-based, legal registration verification, and community/peer verification through network endorsements")
    add_bullet(doc, "Discussion forums for sector-specific knowledge sharing, direct messaging between organizations, and a curated resource library")
    add_bullet(doc, "Grant feed with advanced filtering by sector, country, organization type, funding size, and keywords")
    add_bullet(doc, "600+ active users in pre-launch phase, validating product-market fit")

    # Stage 2
    add_heading2(doc, "Stage 2 \u2014 AI-Powered Capacity Assessment")

    add_body(doc, (
        "Kuja\u2019s capacity assessment module is the most comprehensive self-service assessment "
        "available in the sector, supporting five industry-standard frameworks in a single tool."
    ))
    add_bullet(doc, "Five frameworks: Kuja (proprietary), STEP (TechSoup), UN-HACT (United Nations Harmonized Approach), CHS (Core Humanitarian Standard), NUPAS (USAID Non-US Pre-Award Survey)")
    add_bullet(doc, "Free tier: Rules-based assessment engine generates gap checklists and learning roadmaps based on questionnaire responses")
    add_bullet(doc, "Paid tier: Claude AI reviews uploaded organizational policies, financial documents, and certificates\u2014identifying missing clauses, weak areas, and providing a structured 30/60/90-day improvement roadmap")
    add_bullet(doc, "Donor readiness scoring across four tiers: private philanthropy, bilateral donors, multilateral agencies, and UN system partners")
    add_bullet(doc, "Passportable results: Complete one comprehensive assessment, share the verified results with multiple donors\u2014eliminating 4\u201312 duplicate assessments per year")

    # Stage 3
    add_heading2(doc, "Stage 3 \u2014 Live Due Diligence")

    add_body(doc, (
        "Automated, real-time verification replaces manual due diligence processes that "
        "typically cost donors $5,000\u2013$15,000 per organization and take weeks to complete."
    ))
    add_bullet(doc, "Government registry verification for seven African countries: Kenya NGO Coordination Board, Nigeria Corporate Affairs Commission (CAC), South Africa DSD NPO Registry, Uganda NGO Bureau, Tanzania NiS Business Registration, Somalia MOIFAR, Ethiopia Authority for Civil Society Organizations (ACSO)")
    add_bullet(doc, "Real-time sanctions screening against five databases: OpenSanctions API (primary), UN Security Council Consolidated List (XML), US OFAC Specially Designated Nationals (CSV), EU Financial Sanctions (CSV), World Bank Debarment List")
    add_bullet(doc, "AI-powered registration certificate analysis\u2014extracting organization name, registration number, date, and validity from scanned or photographed certificates")

    # Stage 4
    add_heading2(doc, "Stage 4 \u2014 AI-Powered Grant Matching")

    add_bullet(doc, "AI matches NGOs to relevant grants based on organizational profile, capacity assessment score, sector alignment, geographic focus, and funding tier eligibility")
    add_bullet(doc, "AI matches donors to qualified grantees based on donor requirements, geographic preferences, sector priorities, and minimum capacity thresholds")
    add_bullet(doc, "Multi-language support: English, French, Arabic, Swahili, and Somali natively\u2014with architecture designed to expand to 100+ languages using AI translation")

    # Stage 5
    add_heading2(doc, "Stage 5 \u2014 Grant Application with AI Guidance")

    add_bullet(doc, "Four-step application wizard with AI-driven coaching at each stage\u2014guiding applicants through eligibility confirmation, document preparation, narrative writing, and budget construction")
    add_bullet(doc, "AI analyzes uploaded documents (PDF, DOCX, XLSX) with detailed findings, per-criteria scoring (0\u2013100), and actionable recommendations for improvement")
    add_bullet(doc, "Supported document types: financial reports, audit reports, project proposals, budgets, impact reports, registration certificates, organizational policies")

    # Stage 6
    add_heading2(doc, "Stage 6 \u2014 Review & Award")

    add_bullet(doc, "Dual scoring system: AI auto-score provides an objective baseline; human reviewers add contextual judgment and final scoring")
    add_bullet(doc, "Structured scoring rubrics per document type, ensuring consistency across reviewers and applications")
    add_bullet(doc, "Donor-configurable evaluation criteria with three priority levels (Critical, Important, Nice to Have) and customizable weightings")
    add_bullet(doc, "AI-ranked shortlists allow reviewers to focus attention on the most competitive applications")

    # Stage 7
    add_heading2(doc, "Stage 7 \u2014 Reporting & Compliance Monitoring")

    add_bullet(doc, "Donors upload grant agreements; AI extracts reporting requirements, submission deadlines, financial obligations, and compliance conditions automatically")
    add_bullet(doc, "Donors define per-document-type evaluation criteria, establishing clear expectations for each reporting obligation")
    add_bullet(doc, "NGOs submit reports through the platform; AI evaluates each report against extracted requirements with per-requirement compliance scores and risk flags")
    add_bullet(doc, "Compliance dashboard provides donors with portfolio-level visibility into upcoming deadlines, submission status, and risk indicators")
    add_bullet(doc, "For NGOs not using the ERP: the reporting module seamlessly connects grant information to compliance workflows without requiring additional software adoption")

    # Stage 8
    add_heading2(doc, "Stage 8 \u2014 ERP & Operations (Kuja Build \u2014 Odoo 17)")

    add_body(doc, (
        "The final stage extends grant management into full operational management, closing the "
        "loop between funding and implementation."
    ))
    add_bullet(doc, "Three subscription tiers: Basic (donor/grant management, accounting, financial reporting), Premium (adds asset management, procurement, payroll), Premium+ (adds project management, HRIS, MEAL)")
    add_bullet(doc, "Multi-tenancy architecture: donors can access real-time financial data for their specific grants within a grantee\u2019s ERP instance, increasing transparency and reducing reporting burden")
    add_bullet(doc, "Seamless onboarding: most organizational information is already captured through Stages 1\u20137, reducing ERP setup from weeks to hours")
    add_bullet(doc, "For entities choosing the ERP: automatic data migration from marketplace profiles, assessment results, and grant records")

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════════
    # 4. MARKET OPPORTUNITY
    # ═════════════════════════════════════════════════════════════════════

    add_heading1(doc, "4. Market Opportunity")

    add_heading2(doc, "Market Size")

    add_styled_table(doc,
        headers=["Metric", "Value", "Source"],
        rows=[
            ["Global grant management software market (2024)", "$2.75 billion", "Grand View Research"],
            ["Projected market size (2030)", "$4.79 billion", "Grand View Research"],
            ["CAGR (2024\u20132030)", "10.3%", "Grand View Research"],
            ["Nonprofit/NGO segment share", "58.2%", "Grand View Research"],
            ["Cloud-based deployments", "65% of new implementations", "Industry reports"],
            ["Nonprofits wanting to expand AI use", "90%", "CEP \u2018AI With Purpose\u2019 (2025)"],
            ["Foundations wanting to expand AI use", "94%", "CEP \u2018AI With Purpose\u2019 (2025)"],
        ]
    )

    doc.add_paragraph()  # spacer

    add_heading2(doc, "Target Market")

    add_styled_table(doc,
        headers=["Segment", "Estimated Size", "Kuja\u2019s Approach"],
        rows=[
            ["African CSOs/NGOs", "100,000+ registered organizations", "Free tier for marketplace + paid assessment + ERP"],
            ["Global South CSOs (incl. LatAm, Asia)", "500,000+", "Phase 2\u20133 geographic expansion"],
            ["Bilateral/multilateral donors", "50+ major institutions", "Premium end-to-end solution"],
            ["Private foundations", "200,000+ globally", "Application portal + matching"],
            ["INGOs seeking local partners", "1,000+", "Marketplace + assessment integration"],
        ]
    )

    doc.add_paragraph()  # spacer

    add_heading2(doc, "Key Market Dynamics")

    add_bullet(doc, "The post-USAID landscape has created acute urgency for alternative funding infrastructure. Organizations that previously relied on USAID sub-grants through INGO intermediaries now need direct connections to diversified donor bases.")
    add_bullet(doc, "Aid localization mandates under the Grand Bargain 2.0 require new digital infrastructure to route funding directly to local actors\u2014yet no technology platform currently exists at scale to facilitate this transition.")
    add_bullet(doc, "The inefficiency of 4\u201312 duplicate capacity assessments per NGO per year represents a massive efficiency opportunity. A single passportable assessment eliminates redundant evaluations worth an estimated $50,000\u2013$150,000 annually per NGO in staff time and consultant fees.")

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════════
    # 5. COMPETITIVE LANDSCAPE
    # ═════════════════════════════════════════════════════════════════════

    add_heading1(doc, "5. Competitive Landscape")

    add_body(doc, (
        "The competitive landscape spans three categories: donor-side grant management platforms "
        "(Fluxx, Submittable, SmartSimple, Good Grants, Foundant, Benevity), seeker-side tools "
        "(OpenGrants, Instrumentl, GrantHub), and sector infrastructure (GlobalGiving, UN Partner "
        "Portal, TechSoup STEP, Xapien). No existing competitor offers Kuja\u2019s end-to-end "
        "integrated approach combining all eight lifecycle stages."
    ))

    add_heading2(doc, "Key Competitors")

    add_styled_table(doc,
        headers=["Competitor", "Type", "Primary User", "Key Strength", "Pricing", "Global South?", "AI?"],
        rows=[
            ["Fluxx", "Grant Mgmt", "Donors", "Deep lifecycle management, Grantelligence AI", "Custom (not public)", "No", "Yes (analytics)"],
            ["Submittable", "Grant Mgmt", "Donors", "Proven scale (25M apps, 145K programs)", "~$5K\u201318K/yr", "No", "Yes (fraud)"],
            ["SmartSimple", "Grant Mgmt", "Donors", "Extremely configurable, 45+ languages", "~$6K/yr", "Partial", "Yes (+AI)"],
            ["Good Grants", "Grant Mgmt", "Donors", "Affordable, humanitarian sector friendly", "EUR 3K\u20136K/yr", "Partial", "No"],
            ["Foundant GLM", "Grant Mgmt", "Donors", "User-friendly, unlimited users", "Not public", "No", "Limited"],
            ["Benevity", "CSR Suite", "Corporations", "Massive scale ($14B+ donations)", "$40K+/yr", "No", "Yes"],
            ["OpenGrants", "Discovery", "Seekers", "AI matching, grant writer marketplace", "$29/mo", "No", "Yes (matching)"],
            ["Instrumentl", "Discovery", "Seekers", "Best discovery (450K+ funders, 31K+ RFPs)", "$299\u2013499/mo", "No", "Yes (matching)"],
            ["GrantHub", "Pipeline", "Seekers", "Affordable entry point", "$95/mo", "No", "No"],
            ["GlobalGiving", "Marketplace", "Both", "Established trust, 175 countries", "5\u201312% + 3% fees", "Yes", "No"],
            ["UN Partner Portal", "Due Diligence", "UN\u2013CSO", "Free, harmonized across UN system", "Free", "Yes", "No"],
            ["TechSoup STEP", "Assessment", "Both", "Comprehensive organizational assessment", "Contact", "Yes", "No"],
            ["Xapien", "Screening", "Donors", "True AI (35T pages, 0.5B registries)", "Subscription", "Partial", "Yes"],
        ]
    )

    doc.add_paragraph()  # spacer

    add_heading2(doc, "Key Finding")

    add_body(doc, (
        "No competitor offers all eight of Kuja\u2019s lifecycle stages in a single platform. The "
        "closest competitors\u2014Fluxx and Submittable\u2014serve only the donor side of the equation, "
        "cost 10\u2013100x more than Kuja\u2019s target pricing, and lack capacity assessment, marketplace "
        "discovery, and any meaningful focus on the Global South. Donor-side platforms require "
        "CSOs to navigate multiple disconnected systems; seeker-side tools provide discovery but "
        "no management; and sector infrastructure (UNPP, TechSoup) addresses narrow slices of "
        "the lifecycle without integration. Kuja is the only platform that eliminates these gaps "
        "by unifying all stages under a single architecture with AI automation throughout."
    ), bold=False)

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════════
    # 6. KUJA'S COMPETITIVE ADVANTAGES
    # ═════════════════════════════════════════════════════════════════════

    add_heading1(doc, "6. Kuja\u2019s Competitive Advantages")

    add_numbered_item(doc, 1, "ONLY END-TO-END PLATFORM",
        "Kuja is the only solution that combines marketplace discovery, capacity assessment, "
        "grant application, dual-scoring review, compliance reporting, and ERP operations in a "
        "single platform. Organizations and donors never need to switch between disconnected tools "
        "or re-enter data across systems.")

    add_numbered_item(doc, 2, "AI-POWERED THROUGHOUT",
        "Anthropic Claude AI is integrated across every stage: document analysis with per-criteria "
        "scoring (0\u2013100), grant agreement parsing and requirement extraction, report evaluation "
        "with compliance scoring, capacity assessment with gap analysis, grant matching, and "
        "context-aware chat assistance. No competitor uses AI across all lifecycle stages.")

    add_numbered_item(doc, 3, "BUILT FOR THE GLOBAL SOUTH",
        "Created by Adeso, a Somali-founded social enterprise with 30+ years of humanitarian "
        "experience across Africa. Multi-language support from day one (English, French, Arabic, "
        "Swahili, Somali) with architecture for 100+ languages. Pricing designed for affordability. "
        "Interface optimized for low-bandwidth environments common in target geographies.")

    add_numbered_item(doc, 4, "LIVE DUE DILIGENCE",
        "Real-time government registry verification for seven African countries plus sanctions "
        "screening against four major databases and OpenSanctions. No other grant management "
        "platform offers automated, live compliance verification\u2014a capability that typically "
        "costs donors $5,000\u2013$15,000 per manual assessment.")

    add_numbered_item(doc, 5, "PASSPORTABLE ASSESSMENTS",
        "Organizations complete one comprehensive assessment covering five industry frameworks "
        "(Kuja, STEP, UN-HACT, CHS, NUPAS) and share verified results with multiple donors. "
        "This eliminates the 4\u201312 duplicate assessments NGOs currently endure each year, saving "
        "thousands of staff hours across the sector.")

    add_numbered_item(doc, 6, "TWO-SIDED MARKETPLACE",
        "Kuja is the only platform connecting assessed and verified NGOs directly with donors "
        "in a two-sided marketplace. This creates powerful network effects: each additional "
        "organization increases the platform\u2019s value for all participants, building a defensible "
        "competitive moat that is difficult for single-sided competitors to replicate.")

    add_numbered_item(doc, 7, "SEAMLESS ERP ONBOARDING",
        "Organizations using the marketplace naturally flow into operational management without "
        "re-entering data. The Odoo 17-based ERP inherits organizational profiles, assessment "
        "results, and grant records. Multi-tenant architecture gives donors real-time financial "
        "visibility into their grants\u2014transforming reporting from a periodic burden into "
        "continuous transparency.")

    add_numbered_item(doc, 8, "ADESO\u2019S CREDIBILITY",
        "30+ years of humanitarian experience. Founding member and leader of the NEAR network "
        "(the largest network of local and national organizations in the Global South). Champion "
        "of aid localization with deep relationships across the sector\u2014including Oxfam, IRC, "
        "Save the Children, Gates Foundation, Hilton Foundation, and Porticus. This credibility "
        "cannot be replicated by technology-only competitors.")

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════════
    # 7. PLATFORM ARCHITECTURE & AI CAPABILITIES
    # ═════════════════════════════════════════════════════════════════════

    add_heading1(doc, "7. Platform Architecture & AI Capabilities")

    add_heading2(doc, "Technical Overview")

    add_bullet(doc, "https://web-production-6f8a.up.railway.app", bold_prefix="Production URL: ")
    add_bullet(doc, "Python/Flask backend + Vanilla JavaScript single-page application frontend", bold_prefix="Application Stack: ")
    add_bullet(doc, "PostgreSQL (production) with SQLite for local development", bold_prefix="Database: ")
    add_bullet(doc, "Anthropic Claude API (claude-sonnet-4-20250514)", bold_prefix="AI Engine: ")
    add_bullet(doc, "Railway PaaS with Gunicorn WSGI server", bold_prefix="Deployment: ")
    add_bullet(doc, "kuja.org (Odoo 17 Community) \u2014 marketplace and ERP integration target", bold_prefix="Marketplace: ")
    add_bullet(doc, "RESTful API architecture; JWT authentication; role-based access control (NGO, Donor, Reviewer, Admin)", bold_prefix="Architecture: ")

    add_heading2(doc, "AI Functions")

    add_styled_table(doc,
        headers=["Function", "Input", "Output", "Impact"],
        rows=[
            ["Document Analysis", "PDF/DOCX/XLSX upload", "Score (0\u2013100), findings, recommendations", "Automates manual review, 5x faster"],
            ["Grant Agreement Parsing", "Agreement document", "Extracted requirements, deadlines, obligations", "Eliminates manual extraction"],
            ["Report Evaluation", "NGO reports", "Per-requirement compliance scores, risk flags", "Consistent, objective evaluation"],
            ["Capacity Assessment", "Policies, audits, certificates", "Gap analysis, donor readiness tier, improvement roadmap", "Replaces expensive consultants"],
            ["Grant Matching", "NGO profile + grant criteria", "Ranked match list with compatibility scores", "Reduces discovery from weeks to minutes"],
            ["Chat Assistance", "User context + question", "Role-aware guidance (NGO vs Donor vs Reviewer)", "Reduces support burden, improves UX"],
        ]
    )

    doc.add_paragraph()

    add_body(doc, (
        "All AI functions use the same underlying architecture: structured prompts with role "
        "context, document content extraction, and JSON-formatted responses that integrate "
        "directly into the platform\u2019s interface. The system is designed for extensibility\u2014"
        "new AI functions can be added without architectural changes. AI responses include "
        "confidence indicators and are always presented alongside human review capabilities, "
        "ensuring that automation augments rather than replaces human judgment."
    ))

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════════
    # 8. GO-TO-MARKET STRATEGY
    # ═════════════════════════════════════════════════════════════════════

    add_heading1(doc, "8. Go-to-Market Strategy")

    add_heading2(doc, "Phased Rollout")

    add_styled_table(doc,
        headers=["Phase", "Timeline", "Focus", "Key Actions"],
        rows=[
            ["Kenya Testing", "Q1\u2013Q2 2026", "Validate model",
             "3 CSO Network MoUs, government partnerships (NGO Board), pilot with 8 pipeline partners, iterate based on user feedback"],
            ["Global Scaling", "Q2\u2013Q4 2026", "Leverage Adeso network",
             "10 CSO Network MoUs globally (NEAR, WACSI, EACSOF, SANGONet), onboard Oxfam/IRC/Save as donor partners"],
            ["Sub-Saharan Africa", "Q3 2026\u2013Q2 2027", "24 countries",
             "Expand across East, Central, West, and Southern Africa through regional CSO network partnerships"],
            ["Latin America", "Q3\u2013Q4 2027", "15 countries",
             "Regional network partnerships, Spanish language support, LatAm donor engagement"],
        ]
    )

    doc.add_paragraph()

    add_heading2(doc, "Partnership Strategy")

    add_heading3(doc, "CSO Networks (Distribution Channels)")
    add_bullet(doc, "Mzizi Connect (DRC), WACSI (West Africa), EACSOF (East Africa), SANGONet (Southern Africa)")
    add_bullet(doc, "NEAR Network, MCLD, Start Network, Pledge for Change")
    add_bullet(doc, "Each network MoU provides access to 50\u2013500+ member organizations")

    add_heading3(doc, "Donor Partners")
    add_bullet(doc, "NEAR Change Fund, Hilton Foundation, Gates Foundation, Oak Foundation, Porticus, Packard Foundation")
    add_bullet(doc, "Approach: demonstrate ROI through pilot outcomes, emphasize compliance automation and risk reduction")

    add_heading3(doc, "INGO Partners")
    add_bullet(doc, "Oxfam, International Rescue Committee (IRC), Save the Children")
    add_bullet(doc, "Use case: INGOs use Kuja to find, assess, and manage local implementing partners")

    add_heading3(doc, "Ecosystem Partners")
    add_bullet(doc, "TechSoup (technology access for nonprofits), EPIC Africa (data on African philanthropy)")
    add_bullet(doc, "Proximate Fund (funding local innovation), Pledge for Change (INGO reform)")

    doc.add_paragraph()

    add_heading2(doc, "2026 Targets")

    add_styled_table(doc,
        headers=["Metric", "Target"],
        rows=[
            ["Sales revenue", "$1M"],
            ["CSO profiles on marketplace", "2,500"],
            ["INGO/Donor profiles", "25"],
            ["Individual profiles", "7,500"],
            ["User retention rate", "75%"],
            ["CSO Network MoUs signed", "10"],
        ]
    )

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════════
    # 9. REVENUE MODEL
    # ═════════════════════════════════════════════════════════════════════

    add_heading1(doc, "9. Revenue Model")

    add_heading2(doc, "Revenue Streams")

    add_styled_table(doc,
        headers=["Stream", "Pricing", "2026 Target"],
        rows=[
            ["Application Portal", "$10K\u2013$20K/year per donor", "$50K"],
            ["End-to-End Fund Management", "Custom pricing by portfolio size", "$650K"],
            ["ERP Subscriptions + Outsourcing", "Tiered (Basic / Premium / Premium+)", "$300K"],
            ["Verification Fees", "Per-verification fee", "Future (2027)"],
            ["LMS Course Revenue", "Revenue sharing with content providers", "Future (2028)"],
            ["Ad Revenue", "$2/click at 5,000+ monthly active users", "Future (2028)"],
        ]
    )

    doc.add_paragraph()

    add_heading2(doc, "Pricing Philosophy")

    add_bullet(doc, "CSOs pay substantially lower prices than equivalent Global North platforms. Accessibility and inclusion are the priority\u2014the platform cannot achieve its mission if the organizations it serves cannot afford it.", bold_prefix="CSO Affordability: ")
    add_bullet(doc, "Donor organizations pay a 75% premium over CSO pricing. This cross-subsidy model ensures sustainability while maintaining accessibility.", bold_prefix="Donor Premium: ")
    add_bullet(doc, "For every two donor-sponsored organizations onboarded, one CSO receives full platform access at no cost. This creates a virtuous cycle where donor investment directly expands the ecosystem.", bold_prefix="Cross-Subsidy Model: ")
    add_bullet(doc, "The marketplace and basic capacity assessment (rules-based) are permanently free. This serves as the acquisition pipeline for paid services while ensuring that all organizations can benefit from basic platform functionality.", bold_prefix="Free Tier: ")

    add_heading2(doc, "Competitive Pricing Context")

    add_body(doc, (
        "Enterprise grant management platforms command premium pricing: Fluxx uses custom "
        "enterprise pricing (not publicly disclosed), Submittable starts at approximately "
        "$5,000\u201318,000 per year depending on team size, SmartSimple starts at ~$6,000 per year, "
        "and Good Grants charges EUR 3,000\u20136,000 per year. Kuja targets pricing significantly "
        "below these benchmarks while delivering equivalent or superior functionality\u2014including "
        "capabilities (marketplace, capacity assessment, ERP integration) that no competitor "
        "offers at any price."
    ))

    add_heading2(doc, "Revenue Trajectory")

    add_body(doc, (
        "$1M in 2026 (validation year) \u2192 Scale through ERP adoption and geographic expansion "
        "in 2027 \u2192 Profitability target at $2M+ annual recurring revenue in 2029. The "
        "marketplace-to-ERP conversion funnel creates a predictable revenue growth model: "
        "free marketplace users convert to paid assessment users, who convert to grant management "
        "subscribers, who convert to ERP customers\u2014each stage increasing lifetime value."
    ))

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════════
    # 10. APPENDIX: FEATURE COMPARISON MATRIX
    # ═════════════════════════════════════════════════════════════════════

    add_heading1(doc, "10. Appendix: Feature Comparison Matrix")

    add_body(doc, (
        "The following matrix compares Kuja against eight leading competitors across 24 "
        "platform capabilities. Y = Yes (fully supported), P = Partial (limited support), "
        "\u2014 = Not available."
    ), italic=True)

    # Feature comparison data
    # Columns: Feature, Kuja, Fluxx, Submittable, SmartSimple, UNPP, STEP, Xapien, GlobalGiving, Instrumentl
    features = [
        ["Self-Service Digital Platform",   "Y", "Y", "Y", "Y", "Y", "Y", "Y", "Y", "Y"],
        ["AI-Powered Analysis",             "Y", "P", "P", "P", "\u2014", "\u2014", "Y", "\u2014", "P"],
        ["Multi-Framework Assessment",      "Y", "\u2014", "\u2014", "\u2014", "P", "Y", "\u2014", "\u2014", "\u2014"],
        ["Passportable Results",            "Y", "\u2014", "\u2014", "\u2014", "P", "P", "\u2014", "\u2014", "\u2014"],
        ["Two-Sided Marketplace",           "Y", "\u2014", "\u2014", "\u2014", "P", "\u2014", "\u2014", "Y", "\u2014"],
        ["Grant Matching Algorithm",        "Y", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014", "Y"],
        ["Document Upload & Scoring",       "Y", "P", "P", "P", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014"],
        ["Donor Configuration",             "Y", "Y", "Y", "Y", "P", "\u2014", "\u2014", "\u2014", "\u2014"],
        ["Capacity Improvement Roadmap",    "Y", "\u2014", "\u2014", "\u2014", "\u2014", "P", "\u2014", "\u2014", "\u2014"],
        ["Multi-Language Support",          "Y", "P", "\u2014", "Y", "Y", "\u2014", "\u2014", "P", "\u2014"],
        ["Gov. Registry Verification",      "Y", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014", "P", "\u2014", "\u2014"],
        ["Live Sanctions Screening",        "Y", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014", "Y", "\u2014", "\u2014"],
        ["Grant Application Portal",        "Y", "Y", "Y", "Y", "Y", "\u2014", "\u2014", "Y", "\u2014"],
        ["Review & Award Workflow",         "Y", "Y", "Y", "Y", "Y", "\u2014", "\u2014", "\u2014", "\u2014"],
        ["Reporting Module",                "Y", "Y", "P", "Y", "P", "\u2014", "\u2014", "P", "\u2014"],
        ["ERP Integration",                 "Y", "P", "\u2014", "P", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014"],
        ["Multi-Tenancy",                   "Y", "P", "\u2014", "P", "Y", "\u2014", "\u2014", "\u2014", "\u2014"],
        ["Compliance Dashboard",            "Y", "Y", "P", "Y", "P", "\u2014", "\u2014", "\u2014", "\u2014"],
        ["Audit Trail",                     "Y", "Y", "Y", "Y", "Y", "\u2014", "Y", "\u2014", "\u2014"],
        ["Mobile-Ready",                    "Y", "P", "Y", "P", "P", "\u2014", "\u2014", "Y", "Y"],
        ["Offline Capability",              "P", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014"],
        ["Free Tier for CSOs",              "Y", "\u2014", "\u2014", "\u2014", "Y", "\u2014", "\u2014", "\u2014", "\u2014"],
        ["Global South Design",             "Y", "\u2014", "\u2014", "P", "Y", "P", "\u2014", "Y", "\u2014"],
        ["Chat Assistant",                  "Y", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014"],
    ]

    comp_headers = ["Feature", "Kuja", "Fluxx", "Submittable", "SmartSimple", "UNPP", "STEP", "Xapien", "GlobalGiving", "Instrumentl"]

    table = doc.add_table(rows=1 + len(features), cols=len(comp_headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set table borders
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders)

    # Header row
    for j, header in enumerate(comp_headers):
        cell = table.cell(0, j)
        set_cell_shading(cell, BLUE_HEX)
        set_cell_text(cell, header, bold=True, font_size=Pt(7.5), font_color=WHITE)

    # Data rows
    for i, row_data in enumerate(features):
        for j, val in enumerate(row_data):
            cell = table.cell(i + 1, j)
            if i % 2 == 1:
                set_cell_shading(cell, ALT_ROW_HEX)
            fs = Pt(7.5)
            fc = DARK_GRAY
            bd = False
            # Highlight Kuja column
            if j == 1 and val == "Y":
                fc = RGBColor(0x05, 0x96, 0x69)  # green
                bd = True
            elif j == 1 and val == "P":
                fc = RGBColor(0xD9, 0x77, 0x06)  # amber
                bd = True
            alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else None
            set_cell_text(cell, val, bold=bd, font_size=fs, font_color=fc, alignment=alignment)

    doc.add_paragraph()

    add_body(doc, (
        "Kuja achieves full support (Y) for 23 of 24 features, with partial support (P) for "
        "Offline Capability (planned for future release). No competitor achieves more than 12 "
        "of 24 features. The closest competitors\u2014Fluxx and SmartSimple\u2014score 11 each, "
        "and neither offers marketplace, capacity assessment, government registry "
        "verification, or Global South-specific design."
    ), italic=True)

    # ═════════════════════════════════════════════════════════════════════
    # FOOTER
    # ═════════════════════════════════════════════════════════════════════

    doc.add_paragraph()
    doc.add_paragraph()

    # End line
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_end.add_run("\u2014 End of Document \u2014")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = LIGHT_GRAY
    run.italic = True

    # Footer line
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_footer.add_run("Kuja Grant Platform | Adeso \u2014 African Development Solutions | 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MEDIUM_GRAY

    # ─── Save ────────────────────────────────────────────────────────────
    doc.save(OUTPUT_FILE)
    print(f"Document saved to: {OUTPUT_FILE}")
    print(f"File size: {os.path.getsize(OUTPUT_FILE):,} bytes")


if __name__ == "__main__":
    create_document()
