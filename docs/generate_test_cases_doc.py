"""
Generate Kuja Grant Management System v3.0 Test Cases Document
Creates a professional Word document with comprehensive test cases.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, size=9, color=None, alignment=None):
    """Set text in a cell with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_test_case_table(doc, tc):
    """Add a single test case as a formatted table."""
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(14)

    fields = [
        ("Test ID", tc["id"]),
        ("Test Name", tc["name"]),
        ("Category", tc["category"]),
        ("Priority", tc["priority"]),
        ("Requirement", tc.get("requirement", "N/A")),
        ("Prerequisites", tc.get("prereqs", "None")),
        ("Test Steps", tc["steps"]),
        ("Test Data", tc.get("data", "N/A")),
        ("Expected Result", tc["expected"]),
        ("Pass/Fail Criteria", tc.get("criteria", "Test passes if expected result is observed without errors.")),
    ]

    for label, value in fields:
        row = table.add_row()
        # Label cell
        cell0 = row.cells[0]
        set_cell_shading(cell0, "E8EAF6")
        set_cell_text(cell0, label, bold=True, size=9)
        # Value cell
        cell1 = row.cells[1]
        set_cell_text(cell1, value, size=9)

    # Priority color coding in the priority row
    priority_row = table.rows[3]
    p_text = tc["priority"]
    if "P1" in p_text:
        set_cell_shading(priority_row.cells[1], "FFCDD2")
    elif "P2" in p_text:
        set_cell_shading(priority_row.cells[1], "FFF9C4")
    elif "P3" in p_text:
        set_cell_shading(priority_row.cells[1], "C8E6C9")

    doc.add_paragraph("")  # spacer


# ── Define ALL test cases ────────────────────────────────────────────────────

test_cases = []

# ============================================================================
# CATEGORY 1: AUTHENTICATION (TC-001 to TC-015)
# ============================================================================

auth_cases = [
    {
        "id": "TC-001", "name": "Valid NGO Login", "category": "Authentication",
        "priority": "P1 - Critical", "requirement": "FR-AUTH-001",
        "prereqs": "System is running. Seed data loaded. User fatima@amani.org exists.",
        "steps": (
            "1. Open the application URL (https://web-production-6f8a.up.railway.app)\n"
            "2. Verify the login page is displayed with email and password fields\n"
            "3. Enter 'fatima@amani.org' in the Email field\n"
            "4. Enter 'pass123' in the Password field\n"
            "5. Click the 'Sign In' button\n"
            "6. Wait for page to load completely"
        ),
        "data": "Email: fatima@amani.org | Password: pass123",
        "expected": (
            "1. User is redirected to the NGO Dashboard\n"
            "2. Dashboard displays stat cards: Capacity Score (82%), My Applications, Open Grants, Documents\n"
            "3. Sidebar shows NGO-specific menu items (Dashboard, Grants, Applications, Assessments, Reports, Documents)\n"
            "4. Welcome message shows 'Welcome, Fatima' or similar\n"
            "5. Session cookie is set in browser"
        ),
        "criteria": "Pass: NGO dashboard loads with correct stats. Fail: Login error or wrong dashboard displayed."
    },
    {
        "id": "TC-002", "name": "Valid Donor Login", "category": "Authentication",
        "priority": "P1 - Critical", "requirement": "FR-AUTH-001",
        "prereqs": "System is running. Seed data loaded. User sarah@globalhealth.org exists.",
        "steps": (
            "1. Open the application URL\n"
            "2. Enter 'sarah@globalhealth.org' in the Email field\n"
            "3. Enter 'pass123' in the Password field\n"
            "4. Click the 'Sign In' button\n"
            "5. Wait for page to load completely"
        ),
        "data": "Email: sarah@globalhealth.org | Password: pass123",
        "expected": (
            "1. User is redirected to the Donor Dashboard\n"
            "2. Dashboard displays stat cards: Total Grants, Applications, Pending Reviews, Funding Awarded, Reports Due\n"
            "3. Sidebar shows Donor-specific menu items (Dashboard, My Grants, Applications, Organizations, Reports, Compliance)\n"
            "4. Quick action buttons for creating grants, reviewing applications, etc. are visible"
        ),
        "criteria": "Pass: Donor dashboard loads with correct stats and menu. Fail: Login error or wrong role dashboard."
    },
    {
        "id": "TC-003", "name": "Valid Reviewer Login", "category": "Authentication",
        "priority": "P1 - Critical", "requirement": "FR-AUTH-001",
        "prereqs": "System is running. User james@reviewer.org exists.",
        "steps": (
            "1. Open the application URL\n"
            "2. Enter 'james@reviewer.org' in the Email field\n"
            "3. Enter 'pass123' in the Password field\n"
            "4. Click the 'Sign In' button"
        ),
        "data": "Email: james@reviewer.org | Password: pass123",
        "expected": (
            "1. User is redirected to the Reviewer Dashboard\n"
            "2. Dashboard displays stat cards: Assigned, In Progress, Completed, Average Score\n"
            "3. Sidebar shows Reviewer menu items (Dashboard, My Reviews)\n"
            "4. List of assigned applications is visible"
        ),
        "criteria": "Pass: Reviewer dashboard loads correctly. Fail: Login error or wrong dashboard."
    },
    {
        "id": "TC-004", "name": "Valid Admin Login", "category": "Authentication",
        "priority": "P1 - Critical", "requirement": "FR-AUTH-001",
        "prereqs": "System is running. User admin@kuja.org exists.",
        "steps": (
            "1. Open the application URL\n"
            "2. Enter 'admin@kuja.org' in the Email field\n"
            "3. Enter 'pass123' in the Password field\n"
            "4. Click the 'Sign In' button"
        ),
        "data": "Email: admin@kuja.org | Password: pass123",
        "expected": (
            "1. User is redirected to the Admin Dashboard\n"
            "2. Dashboard displays: Users by role breakdown, Organizations by type, Total users count, Flagged compliance items\n"
            "3. Sidebar shows Admin menu items (Dashboard, Users, Organizations, Grants, Compliance, System)"
        ),
        "criteria": "Pass: Admin dashboard loads with system-wide stats. Fail: Login error or restricted dashboard."
    },
    {
        "id": "TC-005", "name": "Invalid Credentials - Wrong Password", "category": "Authentication",
        "priority": "P1 - Critical", "requirement": "FR-AUTH-002",
        "prereqs": "System is running. User fatima@amani.org exists.",
        "steps": (
            "1. Open the application URL\n"
            "2. Enter 'fatima@amani.org' in the Email field\n"
            "3. Enter 'wrongpassword' in the Password field\n"
            "4. Click the 'Sign In' button"
        ),
        "data": "Email: fatima@amani.org | Password: wrongpassword",
        "expected": (
            "1. Login fails\n"
            "2. Error message displayed: 'Invalid email or password'\n"
            "3. User remains on the login page\n"
            "4. Password field is cleared\n"
            "5. No session cookie is set"
        ),
        "criteria": "Pass: Error message shown, no redirect. Fail: User logs in or no error message."
    },
    {
        "id": "TC-006", "name": "Empty Email Field", "category": "Authentication",
        "priority": "P2 - High", "requirement": "FR-AUTH-003",
        "prereqs": "System is running.",
        "steps": (
            "1. Open the application URL\n"
            "2. Leave the Email field empty\n"
            "3. Enter 'pass123' in the Password field\n"
            "4. Click the 'Sign In' button"
        ),
        "data": "Email: (empty) | Password: pass123",
        "expected": (
            "1. Login fails\n"
            "2. Error message displayed: 'Email and password are required'\n"
            "3. User remains on the login page"
        ),
        "criteria": "Pass: Validation error shown. Fail: Request sent without email."
    },
    {
        "id": "TC-007", "name": "Empty Password Field", "category": "Authentication",
        "priority": "P2 - High", "requirement": "FR-AUTH-003",
        "prereqs": "System is running.",
        "steps": (
            "1. Open the application URL\n"
            "2. Enter 'fatima@amani.org' in the Email field\n"
            "3. Leave the Password field empty\n"
            "4. Click the 'Sign In' button"
        ),
        "data": "Email: fatima@amani.org | Password: (empty)",
        "expected": (
            "1. Login fails\n"
            "2. Error message displayed: 'Email and password are required'\n"
            "3. User remains on the login page"
        ),
        "criteria": "Pass: Validation error shown. Fail: Request sent without password."
    },
    {
        "id": "TC-008", "name": "Rate Limiting - 5 Failed Attempts", "category": "Authentication",
        "priority": "P1 - Critical", "requirement": "FR-AUTH-004",
        "prereqs": "System is running. No prior failed attempts for the test email.",
        "steps": (
            "1. Open the application URL\n"
            "2. Enter 'fatima@amani.org' in the Email field\n"
            "3. Enter 'wrong1' in the Password field and click Sign In\n"
            "4. Repeat with passwords 'wrong2', 'wrong3', 'wrong4', 'wrong5'\n"
            "5. On the 6th attempt, enter the correct password 'pass123'\n"
            "6. Click Sign In"
        ),
        "data": "Email: fatima@amani.org | Passwords: wrong1 through wrong5, then pass123",
        "expected": (
            "1. First 4 attempts show 'Invalid email or password'\n"
            "2. After 5th failed attempt, message changes to account lockout with remaining time\n"
            "3. 6th attempt (even with correct password) is rejected with lockout message\n"
            "4. Message includes countdown or remaining lockout duration"
        ),
        "criteria": "Pass: Account locks after 5 failures, correct password rejected during lockout. Fail: No rate limiting or lockout bypassed."
    },
    {
        "id": "TC-009", "name": "Rate Limiting Recovery", "category": "Authentication",
        "priority": "P2 - High", "requirement": "FR-AUTH-004",
        "prereqs": "Account has been locked from TC-008.",
        "steps": (
            "1. Wait for the lockout period to expire (typically 5-15 minutes)\n"
            "2. Enter 'fatima@amani.org' in the Email field\n"
            "3. Enter 'pass123' in the Password field\n"
            "4. Click Sign In"
        ),
        "data": "Email: fatima@amani.org | Password: pass123",
        "expected": (
            "1. Login succeeds after lockout period expires\n"
            "2. User is redirected to NGO Dashboard\n"
            "3. Failed attempt counter is reset"
        ),
        "criteria": "Pass: Login succeeds after cooldown. Fail: Account remains locked indefinitely."
    },
    {
        "id": "TC-010", "name": "Session Persistence - Close and Reopen Tab", "category": "Authentication",
        "priority": "P2 - High", "requirement": "FR-AUTH-005",
        "prereqs": "None.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Verify NGO dashboard is displayed\n"
            "3. Close the browser tab (not the entire browser)\n"
            "4. Open a new tab and navigate to the application URL\n"
            "5. Observe the page"
        ),
        "data": "Email: fatima@amani.org | Password: pass123",
        "expected": (
            "1. After reopening, user is still logged in\n"
            "2. NGO dashboard is displayed without needing to re-login\n"
            "3. Session cookie persists across tabs"
        ),
        "criteria": "Pass: Session persists across tab close/reopen. Fail: User redirected to login."
    },
    {
        "id": "TC-011", "name": "Logout", "category": "Authentication",
        "priority": "P1 - Critical", "requirement": "FR-AUTH-006",
        "prereqs": "User is logged in as fatima@amani.org.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Verify NGO dashboard is displayed\n"
            "3. Click the Logout button (typically in sidebar or user menu)\n"
            "4. Observe the redirect"
        ),
        "data": "Email: fatima@amani.org | Password: pass123",
        "expected": (
            "1. User is redirected to the login page\n"
            "2. Session cookie is cleared\n"
            "3. Navigating back to a protected route shows login page, not cached dashboard"
        ),
        "criteria": "Pass: Session ended, login page shown. Fail: Session persists after logout."
    },
    {
        "id": "TC-012", "name": "Access Protected Route Without Login", "category": "Authentication",
        "priority": "P1 - Critical", "requirement": "FR-AUTH-007",
        "prereqs": "User is NOT logged in (clear cookies/incognito).",
        "steps": (
            "1. Open an incognito/private browser window\n"
            "2. Navigate directly to a protected API endpoint: GET /api/auth/me\n"
            "3. Observe the response"
        ),
        "data": "URL: /api/auth/me (no session cookie)",
        "expected": (
            "1. Response status code is 401 Unauthorized\n"
            "2. Response body contains error message\n"
            "3. No user data is returned"
        ),
        "criteria": "Pass: 401 returned. Fail: User data returned without authentication."
    },
    {
        "id": "TC-013", "name": "Language Preference - Set to French", "category": "Authentication",
        "priority": "P2 - High", "requirement": "FR-AUTH-008",
        "prereqs": "User is logged in.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Locate the language selector (typically in the sidebar or header)\n"
            "3. Click the language selector\n"
            "4. Select 'Francais' from the dropdown\n"
            "5. Wait for the UI to update"
        ),
        "data": "Email: fatima@amani.org | Language: fr (French)",
        "expected": (
            "1. UI text translates to French\n"
            "2. Dashboard labels, menu items, and buttons display in French\n"
            "3. Language preference is saved (persists on page refresh)\n"
            "4. API returns success for PUT /api/auth/language"
        ),
        "criteria": "Pass: UI displays in French. Fail: Text remains in English or errors occur."
    },
    {
        "id": "TC-014", "name": "Language Preference - Set to Arabic (RTL)", "category": "Authentication",
        "priority": "P2 - High", "requirement": "FR-AUTH-008",
        "prereqs": "User is logged in.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Click the language selector\n"
            "3. Select Arabic from the dropdown\n"
            "4. Wait for the UI to update\n"
            "5. Observe the layout direction"
        ),
        "data": "Email: fatima@amani.org | Language: ar (Arabic)",
        "expected": (
            "1. UI text translates to Arabic\n"
            "2. Layout switches to RTL (right-to-left)\n"
            "3. Sidebar moves to the right side\n"
            "4. Text alignment is right-aligned\n"
            "5. Language preference is saved"
        ),
        "criteria": "Pass: Arabic text displayed with RTL layout. Fail: LTR layout or English text."
    },
    {
        "id": "TC-015", "name": "Invalid Language Preference via API", "category": "Authentication",
        "priority": "P3 - Low", "requirement": "FR-AUTH-008",
        "prereqs": "User is logged in.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Using browser DevTools or curl, send:\n"
            "   PUT /api/auth/language\n"
            "   Body: {\"language\": \"xx\"}\n"
            "3. Observe the response"
        ),
        "data": "API: PUT /api/auth/language | Body: {\"language\": \"xx\"}",
        "expected": (
            "1. Response status code is 400 Bad Request\n"
            "2. Error message indicates invalid language code\n"
            "3. Language preference is not changed"
        ),
        "criteria": "Pass: 400 error returned. Fail: Invalid language accepted."
    },
]
test_cases.extend(auth_cases)

# ============================================================================
# CATEGORY 2: DASHBOARD (TC-020 to TC-028)
# ============================================================================

dashboard_cases = [
    {
        "id": "TC-020", "name": "NGO Dashboard - Stats Cards Display", "category": "Dashboard",
        "priority": "P1 - Critical", "requirement": "FR-DASH-001",
        "prereqs": "Logged in as fatima@amani.org. Seed data loaded.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Verify the NGO Dashboard loads\n"
            "3. Observe the stat cards at the top of the dashboard\n"
            "4. Note the values displayed in each card"
        ),
        "data": "Email: fatima@amani.org | Password: pass123",
        "expected": (
            "1. Dashboard displays 4 stat cards:\n"
            "   - Capacity Score: 82%\n"
            "   - My Applications: count of applications\n"
            "   - Open Grants: count of available grants\n"
            "   - Documents: count of uploaded documents\n"
            "2. Each card has an icon, label, and numeric value\n"
            "3. Quick action buttons are visible below stats"
        ),
        "criteria": "Pass: All 4 stat cards display with correct data. Fail: Missing cards or incorrect values."
    },
    {
        "id": "TC-021", "name": "Donor Dashboard - Stats Cards Display", "category": "Dashboard",
        "priority": "P1 - Critical", "requirement": "FR-DASH-002",
        "prereqs": "Logged in as sarah@globalhealth.org. Seed data loaded.",
        "steps": (
            "1. Login as sarah@globalhealth.org / pass123\n"
            "2. Verify the Donor Dashboard loads\n"
            "3. Observe the stat cards\n"
            "4. Note the values in each card"
        ),
        "data": "Email: sarah@globalhealth.org | Password: pass123",
        "expected": (
            "1. Dashboard displays stat cards:\n"
            "   - Total Grants: count of donor's grants\n"
            "   - Applications: count of received applications\n"
            "   - Pending Reviews: applications awaiting review\n"
            "   - Funding Awarded: total amount awarded\n"
            "   - Reports Due: pending report reviews\n"
            "2. Grant status breakdown visible\n"
            "3. Quick actions: Create Grant, Review Applications, etc."
        ),
        "criteria": "Pass: Donor stats display correctly. Fail: Missing or incorrect metrics."
    },
    {
        "id": "TC-022", "name": "Reviewer Dashboard - Stats Cards Display", "category": "Dashboard",
        "priority": "P1 - Critical", "requirement": "FR-DASH-003",
        "prereqs": "Logged in as james@reviewer.org.",
        "steps": (
            "1. Login as james@reviewer.org / pass123\n"
            "2. Verify the Reviewer Dashboard loads\n"
            "3. Observe the stat cards"
        ),
        "data": "Email: james@reviewer.org | Password: pass123",
        "expected": (
            "1. Dashboard displays stat cards:\n"
            "   - Assigned: number of assigned reviews\n"
            "   - In Progress: reviews currently being worked on\n"
            "   - Completed: finished reviews\n"
            "   - Average Score: mean of all review scores\n"
            "2. List of assigned applications visible"
        ),
        "criteria": "Pass: Reviewer stats shown correctly. Fail: Missing cards or no review list."
    },
    {
        "id": "TC-023", "name": "Admin Dashboard - System Stats Display", "category": "Dashboard",
        "priority": "P1 - Critical", "requirement": "FR-DASH-004",
        "prereqs": "Logged in as admin@kuja.org.",
        "steps": (
            "1. Login as admin@kuja.org / pass123\n"
            "2. Verify the Admin Dashboard loads\n"
            "3. Observe system-wide statistics"
        ),
        "data": "Email: admin@kuja.org | Password: pass123",
        "expected": (
            "1. Dashboard displays system stats:\n"
            "   - Users by role (NGO, Donor, Reviewer, Admin counts)\n"
            "   - Organizations by type\n"
            "   - Total users count\n"
            "   - Flagged compliance items\n"
            "2. Admin-specific menu items visible in sidebar"
        ),
        "criteria": "Pass: Admin stats with role breakdown visible. Fail: Missing system stats."
    },
    {
        "id": "TC-024", "name": "NGO Dashboard - Data Accuracy Verification", "category": "Dashboard",
        "priority": "P2 - High", "requirement": "FR-DASH-005",
        "prereqs": "Logged in as fatima@amani.org. Known seed data: 2 submitted applications for Amani.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Note the 'My Applications' count on dashboard\n"
            "3. Navigate to 'Applications' page from sidebar\n"
            "4. Count the actual applications listed\n"
            "5. Return to dashboard and compare"
        ),
        "data": "Email: fatima@amani.org | Password: pass123",
        "expected": (
            "1. Dashboard 'My Applications' count matches actual number of applications in the list\n"
            "2. 'Open Grants' count matches the number of grants with 'open' status\n"
            "3. 'Documents' count matches uploaded documents count\n"
            "4. Capacity Score matches the assessment score from the profile"
        ),
        "criteria": "Pass: All dashboard numbers match actual data. Fail: Any count mismatch."
    },
    {
        "id": "TC-025", "name": "NGO Dashboard - Quick Actions Navigation", "category": "Dashboard",
        "priority": "P2 - High", "requirement": "FR-DASH-006",
        "prereqs": "Logged in as fatima@amani.org.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. On the dashboard, locate the quick action buttons\n"
            "3. Click 'Browse Grants' quick action\n"
            "4. Verify navigation to Grants page\n"
            "5. Return to dashboard\n"
            "6. Click 'My Applications' quick action\n"
            "7. Verify navigation to Applications page\n"
            "8. Repeat for each available quick action"
        ),
        "data": "Email: fatima@amani.org | Password: pass123",
        "expected": (
            "1. Each quick action button navigates to the correct page\n"
            "2. Back navigation returns to dashboard\n"
            "3. No broken links or error pages"
        ),
        "criteria": "Pass: All quick actions navigate correctly. Fail: Broken link or wrong page."
    },
    {
        "id": "TC-026", "name": "Donor Dashboard - Quick Actions Navigation", "category": "Dashboard",
        "priority": "P2 - High", "requirement": "FR-DASH-006",
        "prereqs": "Logged in as sarah@globalhealth.org.",
        "steps": (
            "1. Login as sarah@globalhealth.org / pass123\n"
            "2. On the donor dashboard, locate quick actions\n"
            "3. Click 'Create Grant' quick action\n"
            "4. Verify grant creation wizard opens\n"
            "5. Return to dashboard\n"
            "6. Click 'Review Applications' quick action\n"
            "7. Verify applications list opens"
        ),
        "data": "Email: sarah@globalhealth.org | Password: pass123",
        "expected": (
            "1. 'Create Grant' opens the grant creation wizard\n"
            "2. 'Review Applications' shows applications for donor's grants\n"
            "3. Each quick action navigates to the correct page"
        ),
        "criteria": "Pass: All donor quick actions work correctly. Fail: Navigation errors."
    },
    {
        "id": "TC-027", "name": "Dashboard Responsive Layout", "category": "Dashboard",
        "priority": "P3 - Low", "requirement": "FR-DASH-007",
        "prereqs": "Logged in as any user.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. View dashboard at full desktop width (1920px)\n"
            "3. Resize browser to tablet width (~768px)\n"
            "4. Resize to mobile width (~375px)\n"
            "5. Observe stat cards and layout at each breakpoint"
        ),
        "data": "Email: fatima@amani.org | Password: pass123",
        "expected": (
            "1. Desktop: stat cards in a horizontal row, sidebar visible\n"
            "2. Tablet: stat cards may wrap to 2 columns, sidebar may collapse\n"
            "3. Mobile: stat cards stack vertically, sidebar is hamburger menu\n"
            "4. No horizontal scrollbar at any width\n"
            "5. All text remains readable"
        ),
        "criteria": "Pass: Layout adapts correctly at all breakpoints. Fail: Overflow or unreadable content."
    },
    {
        "id": "TC-028", "name": "Dashboard Auto-Refresh on Navigation", "category": "Dashboard",
        "priority": "P3 - Low", "requirement": "FR-DASH-008",
        "prereqs": "Logged in as fatima@amani.org.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Note dashboard stats\n"
            "3. Navigate to another section (e.g., Grants)\n"
            "4. Navigate back to Dashboard\n"
            "5. Verify stats are refreshed (not stale cached data)"
        ),
        "data": "Email: fatima@amani.org | Password: pass123",
        "expected": (
            "1. Dashboard stats refresh when returning to dashboard page\n"
            "2. API call is made to fetch fresh data\n"
            "3. Stats reflect any changes made during the session"
        ),
        "criteria": "Pass: Dashboard data refreshes on navigation. Fail: Stale data displayed."
    },
]
test_cases.extend(dashboard_cases)

# ============================================================================
# CATEGORY 3: GRANTS (TC-030 to TC-050)
# ============================================================================

grant_cases = [
    {
        "id": "TC-030", "name": "Browse Open Grants (NGO)", "category": "Grants",
        "priority": "P1 - Critical", "requirement": "FR-GRNT-001",
        "prereqs": "Logged in as NGO user. Open grants exist in system.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Click 'Grants' in the sidebar menu\n"
            "3. Observe the grants list page"
        ),
        "data": "Email: fatima@amani.org | Password: pass123",
        "expected": (
            "1. Grants page displays list of open grants\n"
            "2. Each grant card shows: Title, Funding Amount, Deadline, Sectors, Countries\n"
            "3. At least 3 open grants visible (Community Health Workers, Education Technology, Women's Protection)\n"
            "4. Draft grants (Climate Resilience) are NOT visible to NGOs\n"
            "5. 'Apply' button visible on eligible grants"
        ),
        "criteria": "Pass: Open grants listed with correct details, drafts hidden. Fail: Missing grants or drafts shown."
    },
    {
        "id": "TC-031", "name": "Grant Search Filter", "category": "Grants",
        "priority": "P2 - High", "requirement": "FR-GRNT-002",
        "prereqs": "Logged in as NGO user. Multiple grants exist.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Navigate to Grants page\n"
            "3. Locate the search/filter input\n"
            "4. Type 'Health' in the search box\n"
            "5. Observe the filtered results\n"
            "6. Clear search and type 'Education'\n"
            "7. Observe filtered results"
        ),
        "data": "Search terms: 'Health', 'Education'",
        "expected": (
            "1. Searching 'Health' shows: 'Community Health Workers Scale-Up Program'\n"
            "2. 'Education Technology' grant is filtered out when searching 'Health'\n"
            "3. Searching 'Education' shows: 'Education Technology for Rural Schools'\n"
            "4. Filter is case-insensitive\n"
            "5. Results update in real-time as user types"
        ),
        "criteria": "Pass: Search filters grants correctly. Fail: Wrong results or no filtering."
    },
    {
        "id": "TC-032", "name": "Grant Detail View", "category": "Grants",
        "priority": "P1 - Critical", "requirement": "FR-GRNT-003",
        "prereqs": "Logged in as NGO user. 'Community Health Workers' grant exists.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Navigate to Grants page\n"
            "3. Click on 'Community Health Workers Scale-Up Program'\n"
            "4. Review all sections on the detail page"
        ),
        "data": "Email: fatima@amani.org | Grant: Community Health Workers Scale-Up Program",
        "expected": (
            "1. Grant detail page displays:\n"
            "   - Title: Community Health Workers Scale-Up Program\n"
            "   - Funding: $500,000\n"
            "   - Donor: Global Health Fund\n"
            "   - Countries: Kenya, Somalia, Uganda\n"
            "   - Status: Open\n"
            "2. Eligibility section with requirements\n"
            "3. Evaluation criteria section with weights\n"
            "4. Required documents list\n"
            "5. Apply button visible at the bottom"
        ),
        "criteria": "Pass: All grant details shown correctly. Fail: Missing sections or wrong data."
    },
    {
        "id": "TC-033", "name": "Create Grant - Step 1 Basic Info (Donor)", "category": "Grants",
        "priority": "P1 - Critical", "requirement": "FR-GRNT-004",
        "prereqs": "Logged in as sarah@globalhealth.org (Donor).",
        "steps": (
            "1. Login as sarah@globalhealth.org / pass123\n"
            "2. Click 'Create Grant' button on dashboard or sidebar\n"
            "3. In Step 1 (Basic Information), fill in:\n"
            "   - Title: 'Water Security in Horn of Africa'\n"
            "   - Description: 'Program to improve water access in drought-affected regions'\n"
            "   - Funding Amount: 750000\n"
            "   - Application Deadline: 2026-06-30\n"
            "   - Sectors: WASH, Climate\n"
            "   - Countries: Kenya, Somalia, Ethiopia\n"
            "4. Click 'Next' to proceed to Step 2"
        ),
        "data": "Title: Water Security in Horn of Africa | Amount: $750,000 | Deadline: 2026-06-30",
        "expected": (
            "1. Grant creation wizard opens at Step 1\n"
            "2. All fields accept input correctly\n"
            "3. Multi-select works for sectors and countries\n"
            "4. Clicking Next advances to Step 2\n"
            "5. Data is preserved (going back shows entered data)"
        ),
        "criteria": "Pass: Step 1 data captured, advances to Step 2. Fail: Validation errors or data loss."
    },
    {
        "id": "TC-034", "name": "Create Grant - Step 2 Eligibility", "category": "Grants",
        "priority": "P1 - Critical", "requirement": "FR-GRNT-005",
        "prereqs": "Continuing from TC-033, on Step 2 of grant creation.",
        "steps": (
            "1. In Step 2 (Eligibility Requirements), add requirements:\n"
            "   - 'Registered NGO in target country' (Required: Yes)\n"
            "   - 'Minimum 3 years WASH experience' (Required: Yes)\n"
            "   - 'Annual budget exceeds $100,000' (Required: No)\n"
            "2. Click 'Add Requirement' for each\n"
            "3. Review the list of added requirements\n"
            "4. Click Next to proceed to Step 3"
        ),
        "data": "3 eligibility requirements with varying required status",
        "expected": (
            "1. Requirements are added to the list\n"
            "2. Each shows text and required/optional status\n"
            "3. Requirements can be reordered or deleted\n"
            "4. Clicking Next advances to Step 3"
        ),
        "criteria": "Pass: Requirements saved and visible. Fail: Requirements lost or validation error."
    },
    {
        "id": "TC-035", "name": "Create Grant - Step 3 Evaluation Criteria", "category": "Grants",
        "priority": "P1 - Critical", "requirement": "FR-GRNT-006",
        "prereqs": "Continuing from TC-034, on Step 3.",
        "steps": (
            "1. In Step 3 (Evaluation Criteria), add criteria:\n"
            "   - 'Technical Approach' - Weight: 30%\n"
            "   - 'Organizational Capacity' - Weight: 25%\n"
            "   - 'Budget Reasonableness' - Weight: 25%\n"
            "   - 'Impact & Sustainability' - Weight: 20%\n"
            "2. Verify total weight displays 100%\n"
            "3. Click Next to proceed"
        ),
        "data": "4 criteria: Technical (30%), Capacity (25%), Budget (25%), Impact (20%)",
        "expected": (
            "1. Criteria added with weights\n"
            "2. Running total shows 100%\n"
            "3. System accepts the configuration\n"
            "4. Advances to Step 4"
        ),
        "criteria": "Pass: Criteria with 100% total weight accepted. Fail: Weight calculation error."
    },
    {
        "id": "TC-036", "name": "Create Grant - Invalid Criteria Weight Total", "category": "Grants",
        "priority": "P2 - High", "requirement": "FR-GRNT-006",
        "prereqs": "On Step 3 of grant creation.",
        "steps": (
            "1. Add criteria with weights that do NOT sum to 100%:\n"
            "   - 'Technical Approach' - Weight: 40%\n"
            "   - 'Capacity' - Weight: 30%\n"
            "   (Total: 70%, not 100%)\n"
            "2. Try to click Next"
        ),
        "data": "2 criteria: Technical (40%), Capacity (30%) = 70% total",
        "expected": (
            "1. Validation error displayed: weights must sum to 100%\n"
            "2. Cannot proceed to next step\n"
            "3. Total weight indicator shows 70% with warning"
        ),
        "criteria": "Pass: Validation prevents proceeding. Fail: Allows advancement with non-100% weights."
    },
    {
        "id": "TC-037", "name": "Create Grant - Step 4 Documents", "category": "Grants",
        "priority": "P1 - Critical", "requirement": "FR-GRNT-007",
        "prereqs": "Continuing from TC-035, on Step 4.",
        "steps": (
            "1. In Step 4 (Required Documents), select document types:\n"
            "   - Financial Report (Required)\n"
            "   - Registration Certificate (Required)\n"
            "   - Audit Report (Required)\n"
            "   - PSEA Policy (Required)\n"
            "   - Strategic Plan (Optional)\n"
            "2. For each, set AI evaluation criteria if available\n"
            "3. Click Next to proceed"
        ),
        "data": "5 document types with required/optional flags",
        "expected": (
            "1. Document types selected and displayed\n"
            "2. Required/optional flags set correctly\n"
            "3. AI evaluation criteria can be configured per document type\n"
            "4. Advances to Step 5"
        ),
        "criteria": "Pass: Document requirements configured. Fail: Configuration lost or error."
    },
    {
        "id": "TC-038", "name": "Create Grant - Step 5 Reporting Requirements", "category": "Grants",
        "priority": "P1 - Critical", "requirement": "FR-GRNT-008",
        "prereqs": "Continuing from TC-037, on Step 5.",
        "steps": (
            "1. In Step 5 (Reporting), configure:\n"
            "   - Reporting Frequency: Quarterly\n"
            "   - Report Types: Financial, Narrative\n"
            "   - Add specific requirements for each report type\n"
            "2. Click Next to proceed to Review"
        ),
        "data": "Quarterly reporting with Financial and Narrative types",
        "expected": (
            "1. Reporting configuration saved\n"
            "2. Report template sections displayed\n"
            "3. Advances to final Review step"
        ),
        "criteria": "Pass: Reporting requirements configured. Fail: Configuration not saved."
    },
    {
        "id": "TC-039", "name": "Create Grant - Step 6 Review and Publish", "category": "Grants",
        "priority": "P1 - Critical", "requirement": "FR-GRNT-009",
        "prereqs": "Continuing from TC-038, on final Review step.",
        "steps": (
            "1. In Step 6 (Review), verify all entered information:\n"
            "   - Basic info, eligibility, criteria, documents, reporting\n"
            "2. Click 'Publish Grant'\n"
            "3. Wait for confirmation"
        ),
        "data": "All data from TC-033 through TC-038",
        "expected": (
            "1. Review page shows summary of all configuration\n"
            "2. After publishing, grant status = 'open'\n"
            "3. Success message displayed\n"
            "4. Grant appears in the grants list for NGOs\n"
            "5. Redirected to grant detail page or grants list"
        ),
        "criteria": "Pass: Grant published and visible to NGOs. Fail: Grant not created or not visible."
    },
    {
        "id": "TC-040", "name": "Create Grant as NGO (Unauthorized)", "category": "Grants",
        "priority": "P1 - Critical", "requirement": "FR-GRNT-010",
        "prereqs": "Logged in as NGO user.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Using browser DevTools or curl, send:\n"
            "   POST /api/grants/\n"
            "   Body: {\"title\": \"Test Grant\", \"description\": \"Test\", \"amount\": 100000}\n"
            "3. Observe the response"
        ),
        "data": "Email: fatima@amani.org | API: POST /api/grants/",
        "expected": (
            "1. Response status code is 403 Forbidden\n"
            "2. Error message: 'Donor role required' or similar\n"
            "3. No grant is created in the system"
        ),
        "criteria": "Pass: 403 returned, no grant created. Fail: Grant created by NGO user."
    },
    {
        "id": "TC-041", "name": "Grant Agreement Upload with AI Extraction", "category": "Grants",
        "priority": "P1 - Critical", "requirement": "FR-GRNT-011",
        "prereqs": "Logged in as donor, in grant creation wizard.",
        "steps": (
            "1. Login as sarah@globalhealth.org / pass123\n"
            "2. Start creating a new grant\n"
            "3. In the grant agreement section, click 'Upload Grant Agreement'\n"
            "4. Select the file 'grant_agreement_sample.txt' from test-files/\n"
            "5. Wait for AI analysis to complete\n"
            "6. Review extracted requirements"
        ),
        "data": "File: test-files/grant_agreement_sample.txt",
        "expected": (
            "1. File uploads successfully\n"
            "2. Loading indicator shown during AI analysis\n"
            "3. AI extracts reporting requirements:\n"
            "   - Report types (financial, narrative, impact)\n"
            "   - Frequencies (quarterly, annual)\n"
            "   - Specific deadlines or due dates\n"
            "4. Extracted requirements displayed for donor review\n"
            "5. Donor can modify or add to extracted requirements"
        ),
        "criteria": "Pass: AI extracts meaningful requirements from agreement. Fail: Extraction fails or returns empty."
    },
    {
        "id": "TC-042", "name": "Edit Existing Grant (Owner)", "category": "Grants",
        "priority": "P2 - High", "requirement": "FR-GRNT-012",
        "prereqs": "Logged in as sarah@globalhealth.org. Grant exists that she owns.",
        "steps": (
            "1. Login as sarah@globalhealth.org / pass123\n"
            "2. Navigate to My Grants\n"
            "3. Click on an existing grant\n"
            "4. Click Edit\n"
            "5. Change the title to 'Updated Grant Title'\n"
            "6. Save changes"
        ),
        "data": "Email: sarah@globalhealth.org | Modified field: title",
        "expected": (
            "1. Edit interface opens with current grant data\n"
            "2. Title field is editable\n"
            "3. After save, success message shown\n"
            "4. Grant detail shows updated title"
        ),
        "criteria": "Pass: Grant updated successfully. Fail: Edit fails or data not saved."
    },
    {
        "id": "TC-043", "name": "Edit Grant (Non-Owner Donor)", "category": "Grants",
        "priority": "P2 - High", "requirement": "FR-GRNT-013",
        "prereqs": "david@eatrust.org logged in. Grant owned by sarah exists.",
        "steps": (
            "1. Login as david@eatrust.org / pass123\n"
            "2. Using DevTools or curl, attempt:\n"
            "   PUT /api/grants/1 (owned by Global Health Fund)\n"
            "   Body: {\"title\": \"Hacked Grant\"}\n"
            "3. Observe the response"
        ),
        "data": "Email: david@eatrust.org | API: PUT /api/grants/1",
        "expected": (
            "1. Response status code is 403 Forbidden\n"
            "2. Grant title is NOT changed\n"
            "3. Only the grant owner can edit"
        ),
        "criteria": "Pass: 403 returned, no modification. Fail: Grant modified by non-owner."
    },
    {
        "id": "TC-044", "name": "Grant Sector Filter", "category": "Grants",
        "priority": "P2 - High", "requirement": "FR-GRNT-002",
        "prereqs": "Logged in as NGO user.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Navigate to Grants page\n"
            "3. Use sector filter to select 'Health'\n"
            "4. Observe filtered results\n"
            "5. Change filter to 'Education'\n"
            "6. Observe updated results"
        ),
        "data": "Sector filters: Health, Education",
        "expected": (
            "1. 'Health' filter shows Community Health Workers grant\n"
            "2. 'Education' filter shows Education Technology grant\n"
            "3. Filters can be combined or cleared"
        ),
        "criteria": "Pass: Sector filter works correctly. Fail: Wrong grants shown."
    },
    {
        "id": "TC-045", "name": "Grant Country Filter", "category": "Grants",
        "priority": "P2 - High", "requirement": "FR-GRNT-002",
        "prereqs": "Logged in as NGO user.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Navigate to Grants page\n"
            "3. Use country filter to select 'Kenya'\n"
            "4. Observe filtered results\n"
            "5. Change filter to 'Nigeria'\n"
            "6. Observe updated results"
        ),
        "data": "Country filters: Kenya, Nigeria",
        "expected": (
            "1. 'Kenya' filter shows grants targeting Kenya (Community Health Workers, Education Technology)\n"
            "2. 'Nigeria' filter shows Women's Protection grant\n"
            "3. Filter updates dynamically"
        ),
        "criteria": "Pass: Country filter works correctly. Fail: Wrong grants shown."
    },
]
test_cases.extend(grant_cases)

# ============================================================================
# CATEGORY 4: APPLICATIONS (TC-060 to TC-085)
# ============================================================================

application_cases = [
    {
        "id": "TC-060", "name": "Start New Application", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-001",
        "prereqs": "Logged in as fatima@amani.org. 'Education Technology' grant is open and Amani has not applied.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Navigate to Grants page\n"
            "3. Find 'Education Technology for Rural Schools'\n"
            "4. Click 'Apply' button\n"
            "5. Observe the application wizard"
        ),
        "data": "Email: fatima@amani.org | Grant: Education Technology for Rural Schools",
        "expected": (
            "1. Application wizard opens at Step 1 (Eligibility)\n"
            "2. Grant name displayed in wizard header\n"
            "3. Eligibility checklist items shown\n"
            "4. Step progress indicator shows Step 1 of 4"
        ),
        "criteria": "Pass: Application wizard opens at Step 1. Fail: Error or wrong page."
    },
    {
        "id": "TC-061", "name": "Eligibility Check - All Requirements Met", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-002",
        "prereqs": "In Step 1 of application wizard from TC-060.",
        "steps": (
            "1. Check all eligibility requirement checkboxes\n"
            "2. Observe the status indicators\n"
            "3. Click 'Next' to proceed to Step 2"
        ),
        "data": "All eligibility items checked",
        "expected": (
            "1. All items show green checkmarks\n"
            "2. 'Next' button becomes enabled\n"
            "3. Proceeds to Step 2 (Proposal Responses)"
        ),
        "criteria": "Pass: All items checked, advances to Step 2. Fail: Cannot proceed."
    },
    {
        "id": "TC-062", "name": "Eligibility Check - Required Item Unchecked", "category": "Applications",
        "priority": "P2 - High", "requirement": "FR-APPL-002",
        "prereqs": "In Step 1 of application wizard.",
        "steps": (
            "1. Check most eligibility items but leave one REQUIRED item unchecked\n"
            "2. Try to click 'Next'"
        ),
        "data": "One required eligibility item left unchecked",
        "expected": (
            "1. Warning message displayed about unmet requirement\n"
            "2. Cannot proceed to Step 2\n"
            "3. The unchecked required item is highlighted"
        ),
        "criteria": "Pass: Warning shown, cannot proceed. Fail: Proceeds with unmet requirements."
    },
    {
        "id": "TC-063", "name": "Proposal Response with AI Guidance", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-003",
        "prereqs": "In Step 2 of application wizard.",
        "steps": (
            "1. In Step 2 (Proposal), find the first criterion text area\n"
            "2. Click the 'AI Guidance' button next to it\n"
            "3. Wait for AI to generate guidance\n"
            "4. Read the guidance provided\n"
            "5. Write a response in the text area based on the guidance"
        ),
        "data": "Criterion: first evaluation criterion for the grant",
        "expected": (
            "1. AI Guidance button triggers API call\n"
            "2. Loading indicator shown while AI processes\n"
            "3. AI returns relevant writing suggestions\n"
            "4. Suggestions appear in a panel/tooltip near the text area\n"
            "5. Suggestions are in the user's selected language"
        ),
        "criteria": "Pass: AI guidance generated and relevant. Fail: No guidance or irrelevant suggestions."
    },
    {
        "id": "TC-064", "name": "Proposal Response Word Count", "category": "Applications",
        "priority": "P2 - High", "requirement": "FR-APPL-004",
        "prereqs": "In Step 2 of application wizard.",
        "steps": (
            "1. In a criterion response text area, type a very long response\n"
            "2. Continue typing past any word count limit\n"
            "3. Observe the word count indicator"
        ),
        "data": "Long text exceeding typical 500-word limit",
        "expected": (
            "1. Word count indicator shows current/max words\n"
            "2. When exceeding limit, indicator turns red/warning\n"
            "3. Warning message about exceeding word limit"
        ),
        "criteria": "Pass: Word count tracking works with warning. Fail: No word count or no warning."
    },
    {
        "id": "TC-065", "name": "Document Upload - Valid PDF", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-005",
        "prereqs": "In Step 3 of application wizard. test-files-v3/valid_financial_report.pdf available.",
        "steps": (
            "1. In Step 3 (Documents), find 'Financial Report' upload area\n"
            "2. Click 'Upload' or drag the file valid_financial_report.pdf\n"
            "3. Wait for upload and AI analysis to complete\n"
            "4. Review the AI analysis results"
        ),
        "data": "File: test-files-v3/valid_financial_report.pdf",
        "expected": (
            "1. File uploads successfully (progress bar shown)\n"
            "2. AI analysis runs automatically after upload\n"
            "3. AI score displayed (expected: 70-90)\n"
            "4. Analysis findings listed\n"
            "5. Document appears in the uploaded documents list with green status"
        ),
        "criteria": "Pass: File uploaded, AI score shown. Fail: Upload fails or no AI analysis."
    },
    {
        "id": "TC-066", "name": "Document Upload - Invalid File Type (EXE)", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-006",
        "prereqs": "In Step 3 of application wizard.",
        "steps": (
            "1. In Step 3, attempt to upload 'malicious_script.html'\n"
            "2. Observe the system response"
        ),
        "data": "File: test-files-v3/malicious_script.html",
        "expected": (
            "1. File is rejected before upload\n"
            "2. Error message: 'File type not allowed' or similar\n"
            "3. No file is stored on the server\n"
            "4. Only PDF, DOC, DOCX, XLS, XLSX, TXT, CSV formats accepted"
        ),
        "criteria": "Pass: File rejected with error. Fail: File accepted."
    },
    {
        "id": "TC-067", "name": "Document Upload - Oversized File", "category": "Applications",
        "priority": "P2 - High", "requirement": "FR-APPL-006",
        "prereqs": "In Step 3 of application wizard. A file >16MB available.",
        "steps": (
            "1. Create or use a PDF file larger than 16MB\n"
            "2. Attempt to upload it in Step 3\n"
            "3. Observe the system response"
        ),
        "data": "File: large file > 16MB (or 10MB per system config)",
        "expected": (
            "1. File is rejected with size error\n"
            "2. Error message: 'File too large. Max 10 MB' or 'Max 16 MB'\n"
            "3. No file is stored"
        ),
        "criteria": "Pass: File rejected with size error. Fail: Large file accepted."
    },
    {
        "id": "TC-068", "name": "Document Upload - Empty/Corrupt PDF", "category": "Applications",
        "priority": "P2 - High", "requirement": "FR-APPL-007",
        "prereqs": "In Step 3 of application wizard.",
        "steps": (
            "1. Upload the file 'invalid_empty.pdf' (minimal PDF structure, no content)\n"
            "2. Wait for upload and AI analysis\n"
            "3. Review the analysis results"
        ),
        "data": "File: test-files-v3/invalid_empty.pdf",
        "expected": (
            "1. File uploads (valid PDF structure)\n"
            "2. AI analysis returns very low score (5-15)\n"
            "3. Analysis notes: 'Document appears empty' or 'No meaningful content'\n"
            "4. Warning indicator shown"
        ),
        "criteria": "Pass: Low AI score with appropriate warning. Fail: High score for empty document."
    },
    {
        "id": "TC-069", "name": "Application Review Step - Summary", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-008",
        "prereqs": "Steps 1-3 completed in application wizard.",
        "steps": (
            "1. Navigate to Step 4 (Review & Submit)\n"
            "2. Review the summary of all entered information\n"
            "3. Check completion indicators for each section"
        ),
        "data": "All previous steps completed",
        "expected": (
            "1. Summary shows:\n"
            "   - Eligibility: all items checked\n"
            "   - Proposal: responses for each criterion with word counts\n"
            "   - Documents: uploaded files with AI scores\n"
            "2. Completion indicators show green for completed sections\n"
            "3. 'Submit Application' button is visible"
        ),
        "criteria": "Pass: Complete summary displayed. Fail: Missing sections or wrong data."
    },
    {
        "id": "TC-070", "name": "Submit Application", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-009",
        "prereqs": "In Step 4 of application wizard, all sections complete.",
        "steps": (
            "1. On the Review step, verify all sections are complete\n"
            "2. Click 'Submit Application'\n"
            "3. Wait for processing\n"
            "4. Observe the confirmation"
        ),
        "data": "Complete application ready for submission",
        "expected": (
            "1. Application submitted successfully\n"
            "2. Status changes from 'draft' to 'submitted'\n"
            "3. AI calculates overall application score\n"
            "4. Confirmation message displayed with score\n"
            "5. Application appears in 'My Applications' list with 'Submitted' status\n"
            "6. Timestamp recorded for submission"
        ),
        "criteria": "Pass: Status = submitted, score calculated. Fail: Error on submission."
    },
    {
        "id": "TC-071", "name": "Duplicate Application Prevention", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-010",
        "prereqs": "fatima@amani.org has already submitted application for 'Community Health Workers'.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Navigate to Grants page\n"
            "3. Find 'Community Health Workers Scale-Up Program'\n"
            "4. Try to click 'Apply' again\n"
            "5. Alternatively, via API: POST /api/applications/ with grant_id=1"
        ),
        "data": "Email: fatima@amani.org | Grant: Community Health Workers (already applied)",
        "expected": (
            "1. Apply button is disabled or hidden (already applied)\n"
            "2. If API call made: 409 Conflict\n"
            "3. Message: 'Already applied to this grant' or 'Application exists'\n"
            "4. No duplicate application created"
        ),
        "criteria": "Pass: Duplicate prevented (409 or disabled button). Fail: Duplicate application created."
    },
    {
        "id": "TC-072", "name": "View Application Detail", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-011",
        "prereqs": "fatima@amani.org has submitted application for Community Health Workers.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Navigate to 'My Applications' from sidebar\n"
            "3. Click on the 'Community Health Workers' application\n"
            "4. Review all detail sections"
        ),
        "data": "Email: fatima@amani.org | Application: Community Health Workers",
        "expected": (
            "1. Application detail page shows:\n"
            "   - Grant name and donor\n"
            "   - Status: Submitted\n"
            "   - AI Score: 78.5\n"
            "   - Eligibility responses\n"
            "   - Proposal responses for each criterion\n"
            "   - Uploaded documents with AI scores\n"
            "   - Status timeline with timestamps"
        ),
        "criteria": "Pass: Full application details visible. Fail: Missing data or wrong application."
    },
    {
        "id": "TC-073", "name": "Draft Application - Save and Return", "category": "Applications",
        "priority": "P2 - High", "requirement": "FR-APPL-012",
        "prereqs": "Logged in as NGO user. An open grant exists that user has not applied to.",
        "steps": (
            "1. Login as peter@hopebridges.org / pass123\n"
            "2. Start a new application for an open grant\n"
            "3. Complete Step 1 (Eligibility)\n"
            "4. Start Step 2, write partial responses\n"
            "5. Click 'Save as Draft' or navigate away\n"
            "6. Logout\n"
            "7. Login again as peter@hopebridges.org\n"
            "8. Go to My Applications\n"
            "9. Click on the draft application\n"
            "10. Verify data is preserved"
        ),
        "data": "Email: peter@hopebridges.org | Password: pass123",
        "expected": (
            "1. Application saved as draft\n"
            "2. After re-login, draft appears in 'My Applications' with 'Draft' status\n"
            "3. Opening the draft restores all entered data:\n"
            "   - Eligibility checkboxes still checked\n"
            "   - Partial proposal responses preserved\n"
            "4. Can continue editing from where left off"
        ),
        "criteria": "Pass: Draft preserved with all data. Fail: Data lost on re-login."
    },
    {
        "id": "TC-074", "name": "Application Status Transition - Under Review", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-013",
        "prereqs": "Logged in as sarah@globalhealth.org. Submitted application exists for her grant.",
        "steps": (
            "1. Login as sarah@globalhealth.org / pass123\n"
            "2. Navigate to Applications for 'Community Health Workers'\n"
            "3. Find Amani's submitted application\n"
            "4. Change status to 'Under Review'\n"
            "5. Observe the status update"
        ),
        "data": "Email: sarah@globalhealth.org | Application: Amani's submitted application",
        "expected": (
            "1. Status changes from 'submitted' to 'under_review'\n"
            "2. Timestamp recorded for the transition\n"
            "3. Status history updated\n"
            "4. NGO can see updated status when they view the application"
        ),
        "criteria": "Pass: Status updated with timestamp. Fail: Status unchanged or error."
    },
    {
        "id": "TC-075", "name": "Assign Reviewer to Application", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-014",
        "prereqs": "Logged in as sarah@globalhealth.org. Application is under review.",
        "steps": (
            "1. Login as sarah@globalhealth.org / pass123\n"
            "2. Open the application detail page\n"
            "3. Click 'Assign Reviewer' button\n"
            "4. Select 'james@reviewer.org' from the reviewer list\n"
            "5. Confirm assignment"
        ),
        "data": "Reviewer: james@reviewer.org | Application: Amani's application",
        "expected": (
            "1. Review assignment created\n"
            "2. Review status = 'assigned'\n"
            "3. James sees the application in his reviewer dashboard\n"
            "4. Assignment confirmation shown to donor"
        ),
        "criteria": "Pass: Reviewer assigned successfully. Fail: Assignment fails."
    },
    {
        "id": "TC-076", "name": "View Application as Reviewer", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-015",
        "prereqs": "james@reviewer.org assigned to review Amani's application.",
        "steps": (
            "1. Login as james@reviewer.org / pass123\n"
            "2. View the reviewer dashboard\n"
            "3. Find the assigned application\n"
            "4. Click to open it\n"
            "5. Review all sections"
        ),
        "data": "Email: james@reviewer.org | Password: pass123",
        "expected": (
            "1. Assigned application visible in reviewer dashboard\n"
            "2. Full application visible:\n"
            "   - Proposal responses\n"
            "   - Uploaded documents\n"
            "   - AI scores\n"
            "3. Scoring interface visible for each criterion\n"
            "4. Score input fields (0-100) and comment areas"
        ),
        "criteria": "Pass: Full app visible with scoring interface. Fail: Missing data or no scoring."
    },
    {
        "id": "TC-077", "name": "Score Application - Manual Scoring", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-016",
        "prereqs": "Logged in as james@reviewer.org. Assigned application open.",
        "steps": (
            "1. For each evaluation criterion, enter a score (0-100):\n"
            "   - Technical Approach: 85\n"
            "   - Organizational Capacity: 78\n"
            "   - Budget: 72\n"
            "   - Impact: 90\n"
            "2. Add comments for each criterion\n"
            "3. Observe the weighted score calculation in real-time"
        ),
        "data": "Scores: Technical=85, Capacity=78, Budget=72, Impact=90",
        "expected": (
            "1. Score inputs accept values 0-100\n"
            "2. Weighted score calculates in real-time as scores are entered\n"
            "3. Comment fields accept text\n"
            "4. Running total updates dynamically"
        ),
        "criteria": "Pass: Scores entered, weighted total calculates. Fail: Calculation error or input rejected."
    },
    {
        "id": "TC-078", "name": "Score Application - AI Auto-Score", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-017",
        "prereqs": "Logged in as james@reviewer.org. Assigned application open.",
        "steps": (
            "1. On the scoring interface, click 'AI Auto-Score' button\n"
            "2. Wait for AI to process\n"
            "3. Review the AI-generated scores for each criterion\n"
            "4. Optionally adjust scores\n"
            "5. Add or modify comments"
        ),
        "data": "Application: Amani's submitted application",
        "expected": (
            "1. AI generates scores for all criteria\n"
            "2. Loading indicator during processing\n"
            "3. AI scores populated in score fields\n"
            "4. AI-generated comments/justifications added\n"
            "5. Reviewer can manually adjust any AI score\n"
            "6. Weighted total recalculates after AI scoring"
        ),
        "criteria": "Pass: AI generates scores for all criteria. Fail: AI fails or returns empty scores."
    },
    {
        "id": "TC-079", "name": "Submit Review", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-018",
        "prereqs": "Logged in as james@reviewer.org. All criteria scored.",
        "steps": (
            "1. Verify all criteria have scores\n"
            "2. Click 'Submit Review'\n"
            "3. Wait for confirmation"
        ),
        "data": "Complete review with all scores and comments",
        "expected": (
            "1. Review submitted successfully\n"
            "2. Review status changes to 'completed'\n"
            "3. Scores saved permanently\n"
            "4. Review removed from 'pending' list\n"
            "5. Appears in 'completed' reviews"
        ),
        "criteria": "Pass: Review status = completed, scores saved. Fail: Submission error."
    },
    {
        "id": "TC-080", "name": "Final Score Calculation", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-019",
        "prereqs": "Reviewer has submitted scores. Application has AI document scores.",
        "steps": (
            "1. Login as sarah@globalhealth.org / pass123\n"
            "2. View the application that has been reviewed\n"
            "3. Check the final composite score\n"
            "4. Verify the calculation breakdown"
        ),
        "data": "Application with reviewer scores and AI document scores",
        "expected": (
            "1. Final score calculated as weighted average:\n"
            "   Final = (Criteria Score x 60%) + (Document Score x 20%) + (Eligibility x 20%)\n"
            "2. Score breakdown visible showing each component\n"
            "3. Percentage and numeric score displayed"
        ),
        "criteria": "Pass: Final score matches expected formula. Fail: Calculation mismatch."
    },
    {
        "id": "TC-081", "name": "View Applicant Rankings (Donor)", "category": "Applications",
        "priority": "P1 - Critical", "requirement": "FR-APPL-020",
        "prereqs": "Logged in as sarah@globalhealth.org. Multiple applications with scores.",
        "steps": (
            "1. Login as sarah@globalhealth.org / pass123\n"
            "2. Navigate to 'Community Health Workers' grant\n"
            "3. View applications/rankings\n"
            "4. Check the ordering"
        ),
        "data": "Email: sarah@globalhealth.org | Grant: Community Health Workers",
        "expected": (
            "1. Applications listed in descending score order\n"
            "2. Amani (78.5) ranked above Salam Relief (62.3)\n"
            "3. Each entry shows: NGO name, score, status, review status\n"
            "4. Award/reject action buttons available"
        ),
        "criteria": "Pass: Ranked list in descending order. Fail: Wrong order or missing scores."
    },
]
test_cases.extend(application_cases)

# ============================================================================
# CATEGORY 5: ASSESSMENTS (TC-090 to TC-100)
# ============================================================================

assessment_cases = [
    {
        "id": "TC-090", "name": "Start Kuja Framework Assessment", "category": "Assessments",
        "priority": "P1 - Critical", "requirement": "FR-ASMT-001",
        "prereqs": "Logged in as fatima@amani.org.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Click 'Assessments' in the sidebar\n"
            "3. Click 'Start New Assessment'\n"
            "4. Select 'Kuja Standard' framework\n"
            "5. Observe the checklist displayed"
        ),
        "data": "Email: fatima@amani.org | Framework: Kuja Standard",
        "expected": (
            "1. Kuja assessment checklist displayed\n"
            "2. 5 categories visible with total of 26 items\n"
            "3. Categories: Governance, Financial, Technical, HR, Compliance (or equivalent)\n"
            "4. Each item has a checkbox and description\n"
            "5. Score starts at 0%"
        ),
        "criteria": "Pass: Kuja checklist with 5 categories displayed. Fail: Wrong framework or missing items."
    },
    {
        "id": "TC-091", "name": "Complete Kuja Checklist", "category": "Assessments",
        "priority": "P1 - Critical", "requirement": "FR-ASMT-002",
        "prereqs": "Kuja assessment started from TC-090.",
        "steps": (
            "1. Check various items across all 5 categories\n"
            "2. For each category, check approximately 70-80% of items\n"
            "3. Observe the score calculations updating\n"
            "4. Click 'Complete Assessment'"
        ),
        "data": "Approximately 70-80% of items checked across all categories",
        "expected": (
            "1. Category scores calculate as (checked/total) x 100 x weight\n"
            "2. Overall score is sum of weighted category scores\n"
            "3. Score updates in real-time as items are checked/unchecked\n"
            "4. Assessment saved on completion\n"
            "5. Results page shows score, category breakdown, gaps"
        ),
        "criteria": "Pass: Correct score calculation and saved. Fail: Wrong score or save failure."
    },
    {
        "id": "TC-092", "name": "Start STEP Framework Assessment", "category": "Assessments",
        "priority": "P2 - High", "requirement": "FR-ASMT-001",
        "prereqs": "Logged in as NGO user.",
        "steps": (
            "1. Login as ahmed@salamrelief.org / pass123\n"
            "2. Navigate to Assessments\n"
            "3. Start new assessment\n"
            "4. Select 'STEP' framework\n"
            "5. Verify checklist content"
        ),
        "data": "Email: ahmed@salamrelief.org | Framework: STEP",
        "expected": (
            "1. STEP checklist displayed with 5 categories\n"
            "2. Items relevant to STEP framework\n"
            "3. Different from Kuja items"
        ),
        "criteria": "Pass: STEP checklist displayed correctly. Fail: Wrong framework items."
    },
    {
        "id": "TC-093", "name": "Start UN-HACT Framework Assessment", "category": "Assessments",
        "priority": "P2 - High", "requirement": "FR-ASMT-001",
        "prereqs": "Logged in as NGO user.",
        "steps": (
            "1. Login as thandi@ubuntu.org / pass123\n"
            "2. Navigate to Assessments\n"
            "3. Start new assessment, select 'UN-HACT'\n"
            "4. Verify checklist"
        ),
        "data": "Email: thandi@ubuntu.org | Framework: UN-HACT",
        "expected": (
            "1. UN-HACT checklist with 5 categories, 22 items\n"
            "2. Items reflect UN-HACT micro-assessment criteria"
        ),
        "criteria": "Pass: UN-HACT checklist correct. Fail: Wrong items or count."
    },
    {
        "id": "TC-094", "name": "Start CHS Framework Assessment", "category": "Assessments",
        "priority": "P2 - High", "requirement": "FR-ASMT-001",
        "prereqs": "Logged in as NGO user.",
        "steps": (
            "1. Login as aisha@sahelwomen.org / pass123\n"
            "2. Navigate to Assessments\n"
            "3. Start new assessment, select 'CHS'\n"
            "4. Verify checklist"
        ),
        "data": "Email: aisha@sahelwomen.org | Framework: CHS",
        "expected": (
            "1. CHS checklist with 7 categories, 27 items\n"
            "2. Items aligned to Core Humanitarian Standard commitments"
        ),
        "criteria": "Pass: CHS checklist with 7 categories. Fail: Wrong structure."
    },
    {
        "id": "TC-095", "name": "Start NUPAS Framework Assessment", "category": "Assessments",
        "priority": "P2 - High", "requirement": "FR-ASMT-001",
        "prereqs": "Logged in as NGO user.",
        "steps": (
            "1. Login as peter@hopebridges.org / pass123\n"
            "2. Navigate to Assessments\n"
            "3. Start new assessment, select 'NUPAS'\n"
            "4. Verify checklist"
        ),
        "data": "Email: peter@hopebridges.org | Framework: NUPAS",
        "expected": (
            "1. NUPAS checklist with 5 categories, 26 items\n"
            "2. Items reflect NUPAS assessment criteria"
        ),
        "criteria": "Pass: NUPAS checklist correct. Fail: Wrong items."
    },
    {
        "id": "TC-096", "name": "Assessment Score Calculation Accuracy", "category": "Assessments",
        "priority": "P1 - Critical", "requirement": "FR-ASMT-003",
        "prereqs": "In Kuja assessment with known items.",
        "steps": (
            "1. Start Kuja assessment\n"
            "2. In Category 1 (e.g., 5 items, weight 25%), check 4 out of 5\n"
            "3. In Category 2 (e.g., 6 items, weight 20%), check 3 out of 6\n"
            "4. Leave other categories at 0\n"
            "5. Calculate expected score manually\n"
            "6. Compare with displayed score"
        ),
        "data": "Category 1: 4/5 checked (80%), Category 2: 3/6 checked (50%)",
        "expected": (
            "1. Category 1 score = (4/5) x 100 x 0.25 = 20.0\n"
            "2. Category 2 score = (3/6) x 100 x 0.20 = 10.0\n"
            "3. Overall = 20.0 + 10.0 = 30.0%\n"
            "4. Displayed score matches manual calculation"
        ),
        "criteria": "Pass: Score matches manual calculation. Fail: Score mismatch."
    },
    {
        "id": "TC-097", "name": "Assessment Gaps Identification", "category": "Assessments",
        "priority": "P2 - High", "requirement": "FR-ASMT-004",
        "prereqs": "Assessment completed with some items unchecked.",
        "steps": (
            "1. Complete an assessment with some items unchecked\n"
            "2. Click 'Complete Assessment'\n"
            "3. View the results page\n"
            "4. Look for the 'Gaps' or 'Recommendations' section"
        ),
        "data": "Assessment with gaps in multiple categories",
        "expected": (
            "1. Unchecked items listed as capacity gaps\n"
            "2. Gaps organized by category\n"
            "3. Recommendations provided for addressing gaps\n"
            "4. Priority indicators for critical gaps"
        ),
        "criteria": "Pass: Gaps correctly identified. Fail: Missing gaps or wrong items listed."
    },
    {
        "id": "TC-098", "name": "Assessment Document Upload", "category": "Assessments",
        "priority": "P2 - High", "requirement": "FR-ASMT-005",
        "prereqs": "In assessment wizard.",
        "steps": (
            "1. During assessment, find the document upload area\n"
            "2. Upload a supporting document (e.g., valid_audit_report.pdf)\n"
            "3. Wait for upload to complete"
        ),
        "data": "File: test-files-v3/valid_audit_report.pdf",
        "expected": (
            "1. Document uploaded successfully\n"
            "2. Document linked to the assessment\n"
            "3. File name displayed in the uploaded list"
        ),
        "criteria": "Pass: Document uploaded and linked. Fail: Upload fails."
    },
    {
        "id": "TC-099", "name": "View Assessment Results", "category": "Assessments",
        "priority": "P1 - Critical", "requirement": "FR-ASMT-006",
        "prereqs": "Assessment completed.",
        "steps": (
            "1. After completing assessment, view results page\n"
            "2. Review all sections of the results"
        ),
        "data": "Completed assessment",
        "expected": (
            "1. Results page shows:\n"
            "   - Overall score (percentage)\n"
            "   - Category-by-category breakdown with scores\n"
            "   - Visual indicators (progress bars or charts)\n"
            "   - List of gaps (unchecked items)\n"
            "   - Recommendations for improvement"
        ),
        "criteria": "Pass: Comprehensive results displayed. Fail: Missing sections."
    },
    {
        "id": "TC-100", "name": "Assessment Updates Organization Profile", "category": "Assessments",
        "priority": "P1 - Critical", "requirement": "FR-ASMT-007",
        "prereqs": "Assessment completed for an NGO.",
        "steps": (
            "1. Complete a new assessment\n"
            "2. Navigate to the organization profile\n"
            "3. Verify the assessment score and date are updated"
        ),
        "data": "Newly completed assessment",
        "expected": (
            "1. Organization profile shows updated assess_score\n"
            "2. assess_date reflects the completion date\n"
            "3. Score visible to donors searching organizations"
        ),
        "criteria": "Pass: Profile updated with new score/date. Fail: Old score persists."
    },
]
test_cases.extend(assessment_cases)

# ============================================================================
# CATEGORY 6: AI SERVICES (TC-120 to TC-130)
# ============================================================================

ai_cases = [
    {
        "id": "TC-120", "name": "AI Chat - Basic Question", "category": "AI Services",
        "priority": "P1 - Critical", "requirement": "FR-AI-001",
        "prereqs": "Logged in as any user. ANTHROPIC_API_KEY configured.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Open the AI Chat panel (typically a button/icon in the header or sidebar)\n"
            "3. Type: 'What can Kuja help with?'\n"
            "4. Press Enter or click Send\n"
            "5. Wait for AI response"
        ),
        "data": "Question: 'What can Kuja help with?'",
        "expected": (
            "1. AI panel opens\n"
            "2. Loading indicator while processing\n"
            "3. AI responds with relevant platform information\n"
            "4. Response is contextual and accurate about Kuja features\n"
            "5. Response appears within reasonable time (< 15 seconds)"
        ),
        "criteria": "Pass: AI responds with relevant info. Fail: Error, timeout, or irrelevant response."
    },
    {
        "id": "TC-121", "name": "AI Chat - Role-Specific Guidance", "category": "AI Services",
        "priority": "P2 - High", "requirement": "FR-AI-002",
        "prereqs": "Logged in as NGO user.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Open AI Chat\n"
            "3. Ask: 'How do I strengthen my application?'\n"
            "4. Review the response"
        ),
        "data": "Question: 'How do I strengthen my application?'",
        "expected": (
            "1. AI provides NGO-specific advice\n"
            "2. Response references proposal writing, document quality, capacity scores\n"
            "3. Advice is actionable and relevant to the NGO context\n"
            "4. Does not provide donor-specific or admin-specific guidance"
        ),
        "criteria": "Pass: Role-appropriate advice given. Fail: Generic or wrong-role advice."
    },
    {
        "id": "TC-122", "name": "AI Chat in French", "category": "AI Services",
        "priority": "P2 - High", "requirement": "FR-AI-003",
        "prereqs": "Logged in. Language set to French.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Switch language to French\n"
            "3. Open AI Chat\n"
            "4. Ask: 'Comment renforcer ma candidature?'\n"
            "5. Review the response"
        ),
        "data": "Language: French | Question: 'Comment renforcer ma candidature?'",
        "expected": (
            "1. AI responds entirely in French\n"
            "2. Response is grammatically correct French\n"
            "3. Content is relevant and helpful\n"
            "4. No English text mixed in"
        ),
        "criteria": "Pass: Complete French response. Fail: English or mixed-language response."
    },
    {
        "id": "TC-123", "name": "AI Chat in Arabic", "category": "AI Services",
        "priority": "P2 - High", "requirement": "FR-AI-003",
        "prereqs": "Logged in. Language set to Arabic.",
        "steps": (
            "1. Login as ahmed@salamrelief.org / pass123\n"
            "2. Switch language to Arabic\n"
            "3. Open AI Chat\n"
            "4. Ask a question in Arabic\n"
            "5. Review the response"
        ),
        "data": "Language: Arabic",
        "expected": (
            "1. AI responds in Arabic\n"
            "2. RTL text rendering is correct\n"
            "3. Content is relevant\n"
            "4. Arabic characters display properly"
        ),
        "criteria": "Pass: Arabic response with correct rendering. Fail: Wrong language or broken rendering."
    },
    {
        "id": "TC-124", "name": "AI Document Analysis - Good Document", "category": "AI Services",
        "priority": "P1 - Critical", "requirement": "FR-AI-004",
        "prereqs": "Logged in as NGO. valid_financial_report.pdf available.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Navigate to Documents section\n"
            "3. Upload valid_financial_report.pdf\n"
            "4. Wait for AI analysis to complete\n"
            "5. Review the analysis results"
        ),
        "data": "File: test-files-v3/valid_financial_report.pdf",
        "expected": (
            "1. AI analysis completes within 30 seconds\n"
            "2. Score: 70-90 (good quality document)\n"
            "3. Findings list with specific observations\n"
            "4. Recommendations for improvement\n"
            "5. Score breakdown by evaluation criteria"
        ),
        "criteria": "Pass: Score 70+, meaningful findings. Fail: Very low score or empty analysis."
    },
    {
        "id": "TC-125", "name": "AI Document Analysis - Poor Document", "category": "AI Services",
        "priority": "P2 - High", "requirement": "FR-AI-004",
        "prereqs": "Logged in as NGO. invalid_financial_incomplete.pdf available.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Upload invalid_financial_incomplete.pdf\n"
            "3. Wait for AI analysis\n"
            "4. Review results"
        ),
        "data": "File: test-files-v3/invalid_financial_incomplete.pdf",
        "expected": (
            "1. AI returns lower score (30-50)\n"
            "2. Identifies missing sections (no auditor opinion, incomplete data)\n"
            "3. Specific findings about what is missing\n"
            "4. Recommendations to complete the document"
        ),
        "criteria": "Pass: Lower score with identified gaps. Fail: High score for incomplete document."
    },
    {
        "id": "TC-126", "name": "AI Application Scoring", "category": "AI Services",
        "priority": "P1 - Critical", "requirement": "FR-AI-005",
        "prereqs": "Application submitted with well-written responses.",
        "steps": (
            "1. Submit a complete application with detailed responses\n"
            "2. After submission, observe the AI score calculation\n"
            "3. Review the score breakdown"
        ),
        "data": "Complete application with detailed proposal responses",
        "expected": (
            "1. AI calculates score based on:\n"
            "   - Completeness of responses\n"
            "   - Relevance to criteria\n"
            "   - Depth and quality of writing\n"
            "2. Score is 0-100\n"
            "3. Per-criterion AI feedback available"
        ),
        "criteria": "Pass: AI score calculated with breakdown. Fail: No score or error."
    },
    {
        "id": "TC-127", "name": "AI Guidance for Proposal Writing", "category": "AI Services",
        "priority": "P2 - High", "requirement": "FR-AI-006",
        "prereqs": "In application wizard Step 2.",
        "steps": (
            "1. In the application wizard, navigate to Step 2 (Proposal)\n"
            "2. For any evaluation criterion, click 'AI Guidance'\n"
            "3. Wait for response\n"
            "4. Review the suggestions"
        ),
        "data": "Criterion: any evaluation criterion from the grant",
        "expected": (
            "1. AI generates writing suggestions\n"
            "2. Suggestions are specific to the criterion\n"
            "3. Include key points to address\n"
            "4. In the user's selected language\n"
            "5. Actionable and concrete"
        ),
        "criteria": "Pass: Relevant, actionable suggestions. Fail: Generic or missing suggestions."
    },
    {
        "id": "TC-128", "name": "AI Report Analysis", "category": "AI Services",
        "priority": "P1 - Critical", "requirement": "FR-AI-007",
        "prereqs": "NGO submits a report for a grant with defined requirements.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Submit a report for a grant\n"
            "3. Wait for AI analysis\n"
            "4. Review the analysis results"
        ),
        "data": "Report submitted for grant with defined requirements",
        "expected": (
            "1. AI analyzes report against grant-specific requirements\n"
            "2. Per-requirement scores (0-100)\n"
            "3. Risk flags for low-scoring requirements\n"
            "4. Overall compliance score\n"
            "5. Specific recommendations"
        ),
        "criteria": "Pass: Per-requirement scoring and risk flags. Fail: No analysis or generic feedback."
    },
    {
        "id": "TC-129", "name": "AI Requirement Extraction from Grant Agreement", "category": "AI Services",
        "priority": "P1 - Critical", "requirement": "FR-AI-008",
        "prereqs": "Logged in as donor.",
        "steps": (
            "1. Login as sarah@globalhealth.org / pass123\n"
            "2. Start creating a new grant\n"
            "3. Upload a grant agreement document\n"
            "4. Click 'Extract Requirements' or similar button\n"
            "5. Wait for AI processing\n"
            "6. Review extracted requirements"
        ),
        "data": "File: test-files/grant_agreement_sample.txt",
        "expected": (
            "1. AI extracts:\n"
            "   - Report types (financial, narrative, impact)\n"
            "   - Reporting frequencies\n"
            "   - Specific deadlines\n"
            "   - Key compliance requirements\n"
            "2. Extracted items displayed in editable format\n"
            "3. Donor can modify, add, or remove items"
        ),
        "criteria": "Pass: Meaningful requirements extracted. Fail: Empty or irrelevant extraction."
    },
    {
        "id": "TC-130", "name": "AI Fallback When API Unavailable", "category": "AI Services",
        "priority": "P2 - High", "requirement": "FR-AI-009",
        "prereqs": "Simulate API failure (e.g., invalid API key or rate limiting).",
        "steps": (
            "1. (For testing) Temporarily disable or simulate AI API failure\n"
            "2. Upload a document for analysis\n"
            "3. Observe the system behavior\n"
            "4. Check if fallback mechanism activates"
        ),
        "data": "Document upload with AI API unavailable",
        "expected": (
            "1. System does not crash or show raw error\n"
            "2. Template-based fallback response returned\n"
            "3. User informed that AI analysis is temporarily unavailable\n"
            "4. Document is still saved/uploaded\n"
            "5. User can retry analysis later"
        ),
        "criteria": "Pass: Graceful fallback without data loss. Fail: Crash or data loss."
    },
]
test_cases.extend(ai_cases)

# ============================================================================
# CATEGORY 7: COMPLIANCE & SANCTIONS (TC-150 to TC-156)
# ============================================================================

compliance_cases = [
    {
        "id": "TC-150", "name": "Sanctions Screening - Clean Organization", "category": "Compliance & Sanctions",
        "priority": "P1 - Critical", "requirement": "FR-COMP-001",
        "prereqs": "Logged in as admin@kuja.org. OpenSanctions API configured.",
        "steps": (
            "1. Login as admin@kuja.org / pass123\n"
            "2. Navigate to Compliance section\n"
            "3. Select 'Amani Community Development' for screening\n"
            "4. Click 'Screen' or 'Run Screening'\n"
            "5. Wait for results from all sanctions lists"
        ),
        "data": "Organization: Amani Community Development | Country: Kenya",
        "expected": (
            "1. Screening runs against: UN, OFAC, EU, World Bank\n"
            "2. All results return CLEAR (no matches)\n"
            "3. Green status indicators for each list\n"
            "4. Screening date/time recorded\n"
            "5. Result saved to compliance history"
        ),
        "criteria": "Pass: All clear across all lists. Fail: False positive or screening failure."
    },
    {
        "id": "TC-151", "name": "Sanctions Screening - Flagged Match", "category": "Compliance & Sanctions",
        "priority": "P1 - Critical", "requirement": "FR-COMP-002",
        "prereqs": "Logged in as admin. Ability to screen arbitrary names.",
        "steps": (
            "1. Login as admin@kuja.org / pass123\n"
            "2. Navigate to Compliance\n"
            "3. Screen entity name similar to a known sanctioned entity\n"
            "4. Wait for results"
        ),
        "data": "Entity name similar to sanctioned organization",
        "expected": (
            "1. Screening returns FLAGGED status\n"
            "2. Match score displayed (percentage similarity)\n"
            "3. Matched sanctioned entity details shown:\n"
            "   - Entity name\n"
            "   - List(s) where found\n"
            "   - Reason for listing\n"
            "4. Red status indicators\n"
            "5. Result saved to compliance history"
        ),
        "criteria": "Pass: Flagged with match details. Fail: False negative (cleared when should be flagged)."
    },
    {
        "id": "TC-152", "name": "View Compliance History", "category": "Compliance & Sanctions",
        "priority": "P2 - High", "requirement": "FR-COMP-003",
        "prereqs": "Previous screenings have been run.",
        "steps": (
            "1. Login as admin@kuja.org / pass123\n"
            "2. Navigate to Compliance\n"
            "3. Select an organization\n"
            "4. View the compliance history tab"
        ),
        "data": "Organization with previous screening records",
        "expected": (
            "1. All previous screenings listed chronologically\n"
            "2. Each entry shows: date, result (clear/flagged), lists checked\n"
            "3. Can click to see full details of any screening\n"
            "4. History is immutable (cannot be deleted)"
        ),
        "criteria": "Pass: Full history displayed. Fail: Missing records or incomplete data."
    },
    {
        "id": "TC-153", "name": "OpenSanctions API Integration", "category": "Compliance & Sanctions",
        "priority": "P1 - Critical", "requirement": "FR-COMP-004",
        "prereqs": "OPENSANCTIONS_API_KEY is configured.",
        "steps": (
            "1. Trigger a sanctions screening\n"
            "2. Monitor network requests (DevTools)\n"
            "3. Verify API call to OpenSanctions"
        ),
        "data": "Valid OpenSanctions API key configured",
        "expected": (
            "1. HTTP request made to https://api.opensanctions.org/\n"
            "2. Response status: 200 OK\n"
            "3. Response contains match results\n"
            "4. Results parsed and displayed in UI"
        ),
        "criteria": "Pass: Successful API call with parsed results. Fail: API error or no results."
    },
    {
        "id": "TC-154", "name": "Fallback Sanctions Check (API Down)", "category": "Compliance & Sanctions",
        "priority": "P2 - High", "requirement": "FR-COMP-005",
        "prereqs": "OpenSanctions API is unavailable or returns error.",
        "steps": (
            "1. (Simulate) When OpenSanctions API is down\n"
            "2. Trigger a sanctions screening\n"
            "3. Observe fallback behavior"
        ),
        "data": "OpenSanctions API unavailable",
        "expected": (
            "1. System falls back to direct downloads:\n"
            "   - UN XML from scsanctions.un.org\n"
            "   - OFAC CSV from treasury.gov\n"
            "   - EU CSV from webgate.ec.europa.eu\n"
            "2. Screening still completes using fallback data\n"
            "3. User informed about fallback mode\n"
            "4. Results may be less comprehensive but still functional"
        ),
        "criteria": "Pass: Fallback screening completes. Fail: Screening fails entirely."
    },
    {
        "id": "TC-155", "name": "Keyword Screening", "category": "Compliance & Sanctions",
        "priority": "P2 - High", "requirement": "FR-COMP-006",
        "prereqs": "Logged in as admin.",
        "steps": (
            "1. Login as admin@kuja.org / pass123\n"
            "2. Screen an organization with suspicious keywords in name\n"
            "3. Review screening results"
        ),
        "data": "Organization name with flagged keywords",
        "expected": (
            "1. Keyword screening detects suspicious terms\n"
            "2. Flagged matches listed with keyword context\n"
            "3. Severity level indicated"
        ),
        "criteria": "Pass: Keywords detected and flagged. Fail: No keyword screening."
    },
    {
        "id": "TC-156", "name": "Personnel Screening", "category": "Compliance & Sanctions",
        "priority": "P2 - High", "requirement": "FR-COMP-007",
        "prereqs": "Logged in as admin or donor.",
        "steps": (
            "1. Login as admin@kuja.org / pass123\n"
            "2. Navigate to Compliance\n"
            "3. Select personnel screening option\n"
            "4. Enter individual name for screening\n"
            "5. Review results"
        ),
        "data": "Individual name for sanctions screening",
        "expected": (
            "1. Person screened against all sanctions lists\n"
            "2. Results show clear or flagged status\n"
            "3. If flagged, matched entity details shown\n"
            "4. Result linked to the organization's compliance record"
        ),
        "criteria": "Pass: Person-level screening works. Fail: No person screening capability."
    },
]
test_cases.extend(compliance_cases)

# ============================================================================
# CATEGORY 8: REGISTRY VERIFICATION (TC-180 to TC-186)
# ============================================================================

registry_cases = [
    {
        "id": "TC-180", "name": "Verify South African NGO", "category": "Registry Verification",
        "priority": "P1 - Critical", "requirement": "FR-REG-001",
        "prereqs": "Logged in as admin. Live internet access.",
        "steps": (
            "1. Login as admin@kuja.org / pass123\n"
            "2. Navigate to Compliance > Registry Verification\n"
            "3. Select 'Ubuntu Education Trust' (South Africa)\n"
            "4. Click 'Verify Registration'\n"
            "5. Wait for live check to complete"
        ),
        "data": "Organization: Ubuntu Education Trust | Country: South Africa | Registry: DSD NPO",
        "expected": (
            "1. System queries DSD NPO Registry (https://www.npo.gov.za/)\n"
            "2. Verification result returned (registered/not found)\n"
            "3. Registration details shown if found\n"
            "4. Verification timestamp recorded\n"
            "5. Status updated on organization profile"
        ),
        "criteria": "Pass: Live verification completes with result. Fail: Verification fails or times out."
    },
    {
        "id": "TC-181", "name": "Verify Kenyan NGO", "category": "Registry Verification",
        "priority": "P1 - Critical", "requirement": "FR-REG-002",
        "prereqs": "Logged in as admin. Live internet access.",
        "steps": (
            "1. Login as admin@kuja.org / pass123\n"
            "2. Navigate to Registry Verification\n"
            "3. Select 'Amani Community Development' (Kenya)\n"
            "4. Click 'Verify Registration'\n"
            "5. Wait for result"
        ),
        "data": "Organization: Amani Community Development | Country: Kenya | Registry: BRS",
        "expected": (
            "1. System checks BRS portal (https://brs.go.ke/)\n"
            "2. Result returned: found/not found\n"
            "3. Verification details displayed\n"
            "4. Timestamp recorded"
        ),
        "criteria": "Pass: Kenya BRS check completes. Fail: Check fails."
    },
    {
        "id": "TC-182", "name": "Verify Nigerian NGO", "category": "Registry Verification",
        "priority": "P2 - High", "requirement": "FR-REG-003",
        "prereqs": "Logged in as admin.",
        "steps": (
            "1. Login as admin@kuja.org / pass123\n"
            "2. Select 'Sahel Women's Network' (Nigeria)\n"
            "3. Click 'Verify Registration'"
        ),
        "data": "Organization: Sahel Women's Network | Country: Nigeria | Registry: CAC",
        "expected": (
            "1. System checks CAC portal (https://search.cac.gov.ng/)\n"
            "2. Result returned\n"
            "3. Verification status updated"
        ),
        "criteria": "Pass: Nigeria CAC check completes. Fail: Check fails."
    },
    {
        "id": "TC-183", "name": "Verify Ugandan NGO", "category": "Registry Verification",
        "priority": "P2 - High", "requirement": "FR-REG-004",
        "prereqs": "Logged in as admin.",
        "steps": (
            "1. Login as admin@kuja.org / pass123\n"
            "2. Select 'Hope Bridges Initiative' (Uganda)\n"
            "3. Click 'Verify Registration'"
        ),
        "data": "Organization: Hope Bridges Initiative | Country: Uganda | Registry: NGO Bureau",
        "expected": (
            "1. System searches NGO Bureau register (https://ngobureau.go.ug/)\n"
            "2. Result returned\n"
            "3. Verification status updated"
        ),
        "criteria": "Pass: Uganda NGO Bureau check completes. Fail: Check fails."
    },
    {
        "id": "TC-184", "name": "Verify Somali NGO (No Public Registry)", "category": "Registry Verification",
        "priority": "P2 - High", "requirement": "FR-REG-005",
        "prereqs": "Logged in as admin.",
        "steps": (
            "1. Login as admin@kuja.org / pass123\n"
            "2. Select 'Salam Relief Foundation' (Somalia)\n"
            "3. Click 'Verify Registration'"
        ),
        "data": "Organization: Salam Relief Foundation | Country: Somalia",
        "expected": (
            "1. System detects no public registry available for Somalia\n"
            "2. Message: 'No public registry available. Manual verification required.'\n"
            "3. Guidance on alternative verification methods (MOIFAR contact info)\n"
            "4. Status set to 'manual_verification_needed'"
        ),
        "criteria": "Pass: Appropriate message about no registry. Fail: Error or false verification."
    },
    {
        "id": "TC-185", "name": "Verify Ethiopian NGO (No Public Registry)", "category": "Registry Verification",
        "priority": "P2 - High", "requirement": "FR-REG-005",
        "prereqs": "Logged in as admin.",
        "steps": (
            "1. Login as admin@kuja.org / pass123\n"
            "2. Select or create verification for Ethiopia-based org\n"
            "3. Attempt verification"
        ),
        "data": "Country: Ethiopia",
        "expected": (
            "1. Message about ACSO (Authority for Civil Society Organizations)\n"
            "2. Manual verification guidance provided\n"
            "3. No false positive from automated check"
        ),
        "criteria": "Pass: Manual verification guidance shown. Fail: Automated check attempted."
    },
    {
        "id": "TC-186", "name": "View All Verification Statuses (Admin)", "category": "Registry Verification",
        "priority": "P2 - High", "requirement": "FR-REG-006",
        "prereqs": "Logged in as admin. Some verifications completed.",
        "steps": (
            "1. Login as admin@kuja.org / pass123\n"
            "2. Navigate to Compliance > Verification Dashboard\n"
            "3. View the list of all organizations with verification status"
        ),
        "data": "Admin view of all organizations",
        "expected": (
            "1. All organizations listed with columns:\n"
            "   - Organization name\n"
            "   - Country\n"
            "   - Registry type\n"
            "   - Verification status (verified/pending/manual/not verified)\n"
            "   - Last verified date\n"
            "2. Filters available by status and country\n"
            "3. Can initiate verification from this view"
        ),
        "criteria": "Pass: All orgs listed with statuses. Fail: Missing orgs or no status info."
    },
]
test_cases.extend(registry_cases)

# ============================================================================
# CATEGORY 9: REPORTS (TC-200 to TC-209)
# ============================================================================

report_cases = [
    {
        "id": "TC-200", "name": "View Reports List (NGO)", "category": "Reports",
        "priority": "P1 - Critical", "requirement": "FR-RPT-001",
        "prereqs": "Logged in as fatima@amani.org. Reports exist for Amani.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Click 'Reports' in the sidebar\n"
            "3. Observe the reports list"
        ),
        "data": "Email: fatima@amani.org | Password: pass123",
        "expected": (
            "1. Reports page displays list of reports for Amani:\n"
            "   - Q1 2026 Financial Report (Submitted)\n"
            "   - Annual Progress Report H1 2026 (Draft)\n"
            "   - Test Q2 Financial Report (Accepted)\n"
            "   - Annual Impact Report 2025 (Revision Requested)\n"
            "   - Under Review Financial Report (Under Review)\n"
            "2. Each report shows: title, type, status, grant name, date\n"
            "3. Status indicators color-coded"
        ),
        "criteria": "Pass: All reports listed with correct statuses. Fail: Missing reports or wrong status."
    },
    {
        "id": "TC-201", "name": "Create New Financial Report", "category": "Reports",
        "priority": "P1 - Critical", "requirement": "FR-RPT-002",
        "prereqs": "Logged in as fatima@amani.org. Active grant exists.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Navigate to Reports\n"
            "3. Click 'New Report'\n"
            "4. Select Type: Financial\n"
            "5. Select Grant: Community Health Workers Scale-Up Program\n"
            "6. Select Period: Q2 2026\n"
            "7. Observe the report form"
        ),
        "data": "Type: Financial | Grant: Community Health Workers | Period: Q2 2026",
        "expected": (
            "1. Report creation form opens\n"
            "2. Financial report template sections displayed:\n"
            "   - Executive Summary\n"
            "   - Income/Revenue\n"
            "   - Expenditure Breakdown\n"
            "   - Budget vs. Actual\n"
            "   - Notes to Accounts\n"
            "3. Template matches donor-defined reporting requirements"
        ),
        "criteria": "Pass: Financial template with correct sections. Fail: Wrong template or missing sections."
    },
    {
        "id": "TC-202", "name": "Create New Narrative Report", "category": "Reports",
        "priority": "P1 - Critical", "requirement": "FR-RPT-002",
        "prereqs": "Logged in as NGO user.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Click New Report\n"
            "3. Select Type: Narrative\n"
            "4. Select appropriate grant and period"
        ),
        "data": "Type: Narrative",
        "expected": (
            "1. Narrative report template sections:\n"
            "   - Executive Summary\n"
            "   - Activities & Outputs\n"
            "   - Progress Against Indicators\n"
            "   - Challenges & Lessons Learned\n"
            "   - Plans for Next Period"
        ),
        "criteria": "Pass: Narrative template displayed. Fail: Wrong template."
    },
    {
        "id": "TC-203", "name": "Fill Report Template Sections", "category": "Reports",
        "priority": "P1 - Critical", "requirement": "FR-RPT-003",
        "prereqs": "Report form open from TC-201.",
        "steps": (
            "1. Fill in Executive Summary: 'This report covers Q2 2026 financial activities...'\n"
            "2. Fill in Revenue section with income data\n"
            "3. Fill in Expenditure section with spending data\n"
            "4. Add notes and explanations\n"
            "5. Click Save"
        ),
        "data": "Financial report content for all template sections",
        "expected": (
            "1. All sections accept text input\n"
            "2. Data saved successfully\n"
            "3. Report status remains 'Draft'\n"
            "4. Can return and edit later"
        ),
        "criteria": "Pass: Content saved in all sections. Fail: Data loss or save error."
    },
    {
        "id": "TC-204", "name": "Submit Report", "category": "Reports",
        "priority": "P1 - Critical", "requirement": "FR-RPT-004",
        "prereqs": "Report filled out and saved as draft.",
        "steps": (
            "1. Open the draft report\n"
            "2. Review all sections\n"
            "3. Click 'Submit Report'\n"
            "4. Confirm submission"
        ),
        "data": "Completed draft report ready for submission",
        "expected": (
            "1. Report submitted successfully\n"
            "2. Status changes from 'draft' to 'submitted'\n"
            "3. AI analysis triggered automatically\n"
            "4. Submission timestamp recorded\n"
            "5. Report visible to donor for review"
        ),
        "criteria": "Pass: Status = submitted, AI analysis runs. Fail: Submission error."
    },
    {
        "id": "TC-205", "name": "AI Report Analysis Results", "category": "Reports",
        "priority": "P1 - Critical", "requirement": "FR-RPT-005",
        "prereqs": "Report submitted (TC-204). AI analysis complete.",
        "steps": (
            "1. View the submitted report\n"
            "2. Check the AI analysis section\n"
            "3. Review per-requirement scores"
        ),
        "data": "Submitted report with AI analysis",
        "expected": (
            "1. AI analysis section shows:\n"
            "   - Overall compliance score\n"
            "   - Per-requirement scores (for each donor-defined requirement)\n"
            "   - Risk flags for low-scoring areas\n"
            "   - Specific recommendations\n"
            "2. Visual indicators (green/yellow/red) for each requirement\n"
            "3. Detailed findings with citations from the report"
        ),
        "criteria": "Pass: Per-requirement AI analysis with risk flags. Fail: No analysis or generic feedback."
    },
    {
        "id": "TC-206", "name": "Donor Reviews Report - Accept", "category": "Reports",
        "priority": "P1 - Critical", "requirement": "FR-RPT-006",
        "prereqs": "Logged in as sarah@globalhealth.org. Submitted report exists.",
        "steps": (
            "1. Login as sarah@globalhealth.org / pass123\n"
            "2. Navigate to Reports section\n"
            "3. Find submitted report from Amani\n"
            "4. Review the report and AI analysis\n"
            "5. Click 'Accept' button"
        ),
        "data": "Email: sarah@globalhealth.org | Report: submitted financial report",
        "expected": (
            "1. Report detail view shows full content and AI analysis\n"
            "2. Accept button available\n"
            "3. After clicking Accept:\n"
            "   - Status changes to 'accepted'\n"
            "   - reviewed_at timestamp set\n"
            "   - NGO can see accepted status\n"
            "4. Confirmation message displayed"
        ),
        "criteria": "Pass: Report accepted, status updated. Fail: Status not changed."
    },
    {
        "id": "TC-207", "name": "Donor Reviews Report - Request Revision", "category": "Reports",
        "priority": "P1 - Critical", "requirement": "FR-RPT-007",
        "prereqs": "Logged in as sarah@globalhealth.org. Submitted report exists.",
        "steps": (
            "1. Login as sarah@globalhealth.org / pass123\n"
            "2. Open a submitted report\n"
            "3. Click 'Request Revision'\n"
            "4. Enter revision notes: 'Please provide more detail on personnel expenditure'\n"
            "5. Confirm"
        ),
        "data": "Revision notes: 'Please provide more detail on personnel expenditure'",
        "expected": (
            "1. Status changes to 'revision_requested'\n"
            "2. Reviewer notes saved and visible to NGO\n"
            "3. NGO sees the report with revision request and notes\n"
            "4. NGO can edit and resubmit"
        ),
        "criteria": "Pass: Revision requested with notes saved. Fail: Notes not saved or status error."
    },
    {
        "id": "TC-208", "name": "View Upcoming Report Deadlines", "category": "Reports",
        "priority": "P2 - High", "requirement": "FR-RPT-008",
        "prereqs": "Logged in as NGO. Reports with deadlines exist.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Navigate to Reports page\n"
            "3. Check for upcoming deadlines section"
        ),
        "data": "Email: fatima@amani.org",
        "expected": (
            "1. Upcoming report deadlines listed\n"
            "2. Each deadline shows: report type, grant name, due date\n"
            "3. Deadlines sorted by date (nearest first)\n"
            "4. Visual indicators for urgency (days remaining)"
        ),
        "criteria": "Pass: Deadlines listed and sorted. Fail: No deadline information."
    },
    {
        "id": "TC-209", "name": "Overdue Report Detection", "category": "Reports",
        "priority": "P2 - High", "requirement": "FR-RPT-009",
        "prereqs": "Reports with past-due deadlines exist.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Navigate to Reports\n"
            "3. Look for overdue report indicators"
        ),
        "data": "Reports with due dates in the past",
        "expected": (
            "1. Overdue reports flagged with red indicator\n"
            "2. 'Overdue' label visible\n"
            "3. Days overdue count shown\n"
            "4. Overdue reports appear at top of list or in separate section"
        ),
        "criteria": "Pass: Overdue reports clearly flagged. Fail: No overdue indication."
    },
]
test_cases.extend(report_cases)

# ============================================================================
# CATEGORY 10: INTERNATIONALIZATION (TC-230 to TC-237)
# ============================================================================

i18n_cases = [
    {
        "id": "TC-230", "name": "English Translations Loaded", "category": "Internationalization",
        "priority": "P2 - High", "requirement": "FR-I18N-001",
        "prereqs": "Application accessible.",
        "steps": (
            "1. In browser, navigate to: /static/js/translations/en.json\n"
            "2. Verify the JSON file loads\n"
            "3. Count the translation keys"
        ),
        "data": "URL: /static/js/translations/en.json",
        "expected": (
            "1. JSON file loads successfully\n"
            "2. Contains 627+ translation keys\n"
            "3. All values are in English\n"
            "4. No missing or empty values"
        ),
        "criteria": "Pass: 627+ keys with English values. Fail: Missing file or fewer keys."
    },
    {
        "id": "TC-231", "name": "Arabic Translations Loaded", "category": "Internationalization",
        "priority": "P2 - High", "requirement": "FR-I18N-001",
        "prereqs": "Application accessible.",
        "steps": (
            "1. Navigate to: /static/js/translations/ar.json\n"
            "2. Verify the JSON file loads\n"
            "3. Verify Arabic text content"
        ),
        "data": "URL: /static/js/translations/ar.json",
        "expected": (
            "1. JSON file loads successfully\n"
            "2. Contains 627+ translation keys\n"
            "3. Values are in Arabic script\n"
            "4. Keys match the English file keys"
        ),
        "criteria": "Pass: Arabic translations present. Fail: Missing or incomplete translations."
    },
    {
        "id": "TC-232", "name": "French Translations Loaded", "category": "Internationalization",
        "priority": "P2 - High", "requirement": "FR-I18N-001",
        "prereqs": "Application accessible.",
        "steps": (
            "1. Navigate to: /static/js/translations/fr.json\n"
            "2. Verify the JSON file loads"
        ),
        "data": "URL: /static/js/translations/fr.json",
        "expected": (
            "1. JSON file loads with 627+ keys\n"
            "2. French text values"
        ),
        "criteria": "Pass: French translations present. Fail: Missing translations."
    },
    {
        "id": "TC-233", "name": "Swahili Translations Loaded", "category": "Internationalization",
        "priority": "P2 - High", "requirement": "FR-I18N-001",
        "prereqs": "Application accessible.",
        "steps": (
            "1. Navigate to: /static/js/translations/sw.json (or equivalent)\n"
            "2. Verify the JSON file loads"
        ),
        "data": "URL: /static/js/translations/sw.json",
        "expected": (
            "1. JSON file loads with translation keys\n"
            "2. Swahili text values"
        ),
        "criteria": "Pass: Swahili translations present. Fail: Missing translations."
    },
    {
        "id": "TC-234", "name": "Switch to Arabic - RTL Layout", "category": "Internationalization",
        "priority": "P1 - Critical", "requirement": "FR-I18N-002",
        "prereqs": "Logged in as any user.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Select Arabic language from language selector\n"
            "3. Wait for UI to update\n"
            "4. Inspect the page layout direction"
        ),
        "data": "Language: Arabic",
        "expected": (
            "1. HTML dir attribute set to 'rtl'\n"
            "2. Sidebar moves to right side of screen\n"
            "3. Text is right-aligned\n"
            "4. Icons and buttons mirror appropriately\n"
            "5. Arabic font (Noto Sans Arabic or similar) applied\n"
            "6. No overlapping elements or broken layout"
        ),
        "criteria": "Pass: Full RTL layout with Arabic text. Fail: LTR maintained or layout breaks."
    },
    {
        "id": "TC-235", "name": "Switch Back to English - LTR Layout", "category": "Internationalization",
        "priority": "P2 - High", "requirement": "FR-I18N-002",
        "prereqs": "Currently in Arabic RTL mode.",
        "steps": (
            "1. From Arabic mode, select English from language selector\n"
            "2. Wait for UI to update\n"
            "3. Verify layout returns to LTR"
        ),
        "data": "Language: English (from Arabic)",
        "expected": (
            "1. HTML dir attribute reverts to 'ltr'\n"
            "2. Sidebar returns to left side\n"
            "3. Text is left-aligned\n"
            "4. All English text displayed\n"
            "5. No layout artifacts from RTL mode"
        ),
        "criteria": "Pass: Clean LTR layout restored. Fail: RTL artifacts remain."
    },
    {
        "id": "TC-236", "name": "Assessment Checklist in French", "category": "Internationalization",
        "priority": "P2 - High", "requirement": "FR-I18N-003",
        "prereqs": "Logged in as NGO. Language set to French.",
        "steps": (
            "1. Login and switch language to French\n"
            "2. Navigate to Assessments\n"
            "3. Start a new assessment\n"
            "4. Observe checklist item text"
        ),
        "data": "Language: French | Section: Assessment Checklist",
        "expected": (
            "1. All checklist items displayed in French\n"
            "2. Category headers in French\n"
            "3. Button labels in French\n"
            "4. Score labels in French"
        ),
        "criteria": "Pass: Complete French checklist. Fail: English items mixed in."
    },
    {
        "id": "TC-237", "name": "Error Messages Translated", "category": "Internationalization",
        "priority": "P2 - High", "requirement": "FR-I18N-004",
        "prereqs": "Logged in. Language set to non-English.",
        "steps": (
            "1. Switch to French\n"
            "2. Trigger a validation error (e.g., submit empty form)\n"
            "3. Observe error message language"
        ),
        "data": "Language: French | Action: trigger validation error",
        "expected": (
            "1. Error message displayed in French\n"
            "2. All system messages in French\n"
            "3. No English error fallback"
        ),
        "criteria": "Pass: Error in selected language. Fail: English error shown."
    },
]
test_cases.extend(i18n_cases)

# ============================================================================
# CATEGORY 11: ACCESS CONTROL (TC-250 to TC-255)
# ============================================================================

access_cases = [
    {
        "id": "TC-250", "name": "NGO Cannot Access Admin Stats", "category": "Access Control",
        "priority": "P1 - Critical", "requirement": "FR-RBAC-001",
        "prereqs": "Logged in as NGO user.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Using DevTools or curl, send:\n"
            "   GET /api/admin/stats\n"
            "3. Observe the response"
        ),
        "data": "Email: fatima@amani.org | API: GET /api/admin/stats",
        "expected": (
            "1. Response status: 403 Forbidden\n"
            "2. Error message: role-based access denied\n"
            "3. No admin stats data returned"
        ),
        "criteria": "Pass: 403 returned. Fail: Admin stats accessible to NGO."
    },
    {
        "id": "TC-251", "name": "NGO Cannot Create Grants", "category": "Access Control",
        "priority": "P1 - Critical", "requirement": "FR-RBAC-002",
        "prereqs": "Logged in as NGO user.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Via API: POST /api/grants/\n"
            "   Body: {\"title\": \"NGO Grant\", \"amount\": 50000}\n"
            "3. Observe response"
        ),
        "data": "Email: fatima@amani.org | API: POST /api/grants/",
        "expected": (
            "1. Response: 403 Forbidden\n"
            "2. No grant created"
        ),
        "criteria": "Pass: 403. Fail: Grant created."
    },
    {
        "id": "TC-252", "name": "Reviewer Cannot Edit Applications", "category": "Access Control",
        "priority": "P1 - Critical", "requirement": "FR-RBAC-003",
        "prereqs": "Logged in as reviewer.",
        "steps": (
            "1. Login as james@reviewer.org / pass123\n"
            "2. Via API: PUT /api/applications/1\n"
            "   Body: {\"status\": \"approved\"}\n"
            "3. Observe response"
        ),
        "data": "Email: james@reviewer.org | API: PUT /api/applications/1",
        "expected": (
            "1. Response: 403 Forbidden\n"
            "2. Application not modified\n"
            "3. Only donors can change application status"
        ),
        "criteria": "Pass: 403 returned. Fail: Application modified by reviewer."
    },
    {
        "id": "TC-253", "name": "Donor Can Only See Own Grants", "category": "Access Control",
        "priority": "P1 - Critical", "requirement": "FR-RBAC-004",
        "prereqs": "Logged in as sarah@globalhealth.org.",
        "steps": (
            "1. Login as sarah@globalhealth.org / pass123\n"
            "2. Navigate to My Grants\n"
            "3. Verify only Global Health Fund grants are shown\n"
            "4. Check that East Africa Trust grants are NOT visible"
        ),
        "data": "Email: sarah@globalhealth.org (Global Health Fund)",
        "expected": (
            "1. Only grants owned by Global Health Fund shown:\n"
            "   - Community Health Workers Scale-Up Program\n"
            "   - Women's Protection and Empowerment\n"
            "2. East Africa Trust grants (Education Technology, Climate Resilience) NOT shown\n"
            "3. No data leakage from other donors"
        ),
        "criteria": "Pass: Only own grants visible. Fail: Other donor's grants shown."
    },
    {
        "id": "TC-254", "name": "NGO Can Only See Own Applications", "category": "Access Control",
        "priority": "P1 - Critical", "requirement": "FR-RBAC-005",
        "prereqs": "Logged in as fatima@amani.org.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Navigate to My Applications\n"
            "3. Verify only Amani's applications shown\n"
            "4. Check that other NGOs' applications are not visible"
        ),
        "data": "Email: fatima@amani.org (Amani Community Development)",
        "expected": (
            "1. Only Amani's applications shown (Community Health Workers)\n"
            "2. Salam Relief, Ubuntu, Sahel, Hope Bridges applications NOT visible\n"
            "3. No access to other organizations' application data"
        ),
        "criteria": "Pass: Only own applications visible. Fail: Other orgs' apps shown."
    },
    {
        "id": "TC-255", "name": "Cross-Organization Document Access Blocked", "category": "Access Control",
        "priority": "P1 - Critical", "requirement": "FR-RBAC-006",
        "prereqs": "Documents uploaded by multiple organizations.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Via API, try to access a document belonging to Salam Relief:\n"
            "   GET /api/documents/{salam_doc_id}\n"
            "3. Observe response"
        ),
        "data": "Email: fatima@amani.org | API: GET /api/documents/{other_org_doc_id}",
        "expected": (
            "1. Response: 403 Forbidden\n"
            "2. No document data returned\n"
            "3. Access restricted to own organization's documents"
        ),
        "criteria": "Pass: 403 returned. Fail: Other org's document accessible."
    },
]
test_cases.extend(access_cases)

# ============================================================================
# CATEGORY 12: SYSTEM HEALTH & VERSION (TC-270 to TC-275)
# ============================================================================

system_cases = [
    {
        "id": "TC-270", "name": "Health Check Endpoint", "category": "System Health",
        "priority": "P1 - Critical", "requirement": "FR-SYS-001",
        "prereqs": "System is running.",
        "steps": (
            "1. Send GET request to /api/health\n"
            "2. Observe the response"
        ),
        "data": "URL: GET /api/health",
        "expected": (
            "1. Response status: 200 OK\n"
            "2. Response body: {\"status\": \"healthy\", \"timestamp\": \"...\"}\n"
            "3. Timestamp is current"
        ),
        "criteria": "Pass: 200 with healthy status. Fail: Non-200 or unhealthy."
    },
    {
        "id": "TC-271", "name": "Version Endpoint", "category": "System Health",
        "priority": "P2 - High", "requirement": "FR-SYS-002",
        "prereqs": "System is running.",
        "steps": (
            "1. Send GET request to /api/version\n"
            "2. Observe the response"
        ),
        "data": "URL: GET /api/version",
        "expected": (
            "1. Response status: 200 OK\n"
            "2. Response body: {\"version\": \"3.0.0\", \"name\": \"Kuja Grant Management System\"}"
        ),
        "criteria": "Pass: Correct version returned. Fail: Wrong version or error."
    },
    {
        "id": "TC-272", "name": "Readiness Check", "category": "System Health",
        "priority": "P2 - High", "requirement": "FR-SYS-003",
        "prereqs": "System is running.",
        "steps": (
            "1. Send GET request to /api/ready\n"
            "2. Observe the response"
        ),
        "data": "URL: GET /api/ready",
        "expected": (
            "1. Response status: 200 OK\n"
            "2. Response body: {\"ready\": true, \"checks\": {\"database\": \"ok\", \"ai_service\": \"configured\"}}\n"
            "3. Both database and AI service checks pass"
        ),
        "criteria": "Pass: All readiness checks pass. Fail: Any check fails."
    },
    {
        "id": "TC-273", "name": "Telemetry Endpoint", "category": "System Health",
        "priority": "P3 - Low", "requirement": "FR-SYS-004",
        "prereqs": "System is running. User is logged in.",
        "steps": (
            "1. Login as any user\n"
            "2. Send POST /api/telemetry with event data:\n"
            "   {\"event\": \"page_view\", \"page\": \"dashboard\"}\n"
            "3. Observe response"
        ),
        "data": "API: POST /api/telemetry | Body: {\"event\": \"page_view\"}",
        "expected": (
            "1. Response status: 200 OK\n"
            "2. Event recorded (or acknowledged)\n"
            "3. No errors"
        ),
        "criteria": "Pass: 200 returned. Fail: Error response."
    },
    {
        "id": "TC-274", "name": "Security Headers Present", "category": "System Health",
        "priority": "P1 - Critical", "requirement": "FR-SYS-005",
        "prereqs": "System is running.",
        "steps": (
            "1. Send any API request (e.g., GET /api/health)\n"
            "2. Inspect the response headers\n"
            "3. Check for security headers"
        ),
        "data": "Any API endpoint response headers",
        "expected": (
            "1. Headers present:\n"
            "   - X-Content-Type-Options: nosniff\n"
            "   - X-Frame-Options: DENY or SAMEORIGIN\n"
            "   - Content-Security-Policy: defined\n"
            "   - X-XSS-Protection: 1; mode=block (or CSP equivalent)\n"
            "2. No Server header revealing technology details"
        ),
        "criteria": "Pass: All security headers present. Fail: Missing critical headers."
    },
    {
        "id": "TC-275", "name": "CORS Configuration", "category": "System Health",
        "priority": "P2 - High", "requirement": "FR-SYS-006",
        "prereqs": "System is running.",
        "steps": (
            "1. Send an OPTIONS preflight request from a different origin\n"
            "2. Check the CORS response headers\n"
            "3. Verify allowed origins"
        ),
        "data": "Cross-origin OPTIONS request",
        "expected": (
            "1. CORS headers present:\n"
            "   - Access-Control-Allow-Origin: appropriate value\n"
            "   - Access-Control-Allow-Methods: listed methods\n"
            "   - Access-Control-Allow-Headers: listed headers\n"
            "2. Not set to wildcard '*' in production"
        ),
        "criteria": "Pass: CORS properly configured. Fail: Missing CORS or too permissive."
    },
]
test_cases.extend(system_cases)

# ============================================================================
# CATEGORY 13: DOCUMENTS (TC-280 to TC-290)
# ============================================================================

doc_mgmt_cases = [
    {
        "id": "TC-280", "name": "View Documents List (NGO)", "category": "Documents",
        "priority": "P1 - Critical", "requirement": "FR-DOC-001",
        "prereqs": "Logged in as fatima@amani.org. Documents uploaded for Amani.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Navigate to Documents from sidebar\n"
            "3. Review the documents list"
        ),
        "data": "Email: fatima@amani.org",
        "expected": (
            "1. Documents page shows list of uploaded documents:\n"
            "   - amani_psea_policy_2024.pdf (AI Score: 88%)\n"
            "   - amani_audit_2024.pdf (AI Score: 78%)\n"
            "   - amani_registration_certificate.pdf (AI Score: 95%)\n"
            "   - amani_financial_2023-2025.pdf (AI Score: 82%)\n"
            "2. Each shows: filename, type, AI score, upload date\n"
            "3. AI score color-coded (green > 80, yellow 50-80, red < 50)"
        ),
        "criteria": "Pass: All documents listed with scores. Fail: Missing documents."
    },
    {
        "id": "TC-281", "name": "Upload New Document", "category": "Documents",
        "priority": "P1 - Critical", "requirement": "FR-DOC-002",
        "prereqs": "Logged in as NGO. Test file available.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Navigate to Documents\n"
            "3. Click 'Upload Document'\n"
            "4. Select document type: 'Strategic Plan'\n"
            "5. Choose file: valid_strategic_plan.pdf\n"
            "6. Click Upload\n"
            "7. Wait for upload and AI analysis"
        ),
        "data": "File: test-files-v3/valid_strategic_plan.pdf | Type: Strategic Plan",
        "expected": (
            "1. File uploads with progress indicator\n"
            "2. AI analysis runs automatically\n"
            "3. AI score displayed (expected 70-90)\n"
            "4. Document appears in documents list\n"
            "5. Type correctly categorized"
        ),
        "criteria": "Pass: Upload successful with AI analysis. Fail: Upload error."
    },
    {
        "id": "TC-282", "name": "Download Document", "category": "Documents",
        "priority": "P2 - High", "requirement": "FR-DOC-003",
        "prereqs": "Document uploaded and visible.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Navigate to Documents\n"
            "3. Click download icon on a document\n"
            "4. Verify file downloads"
        ),
        "data": "Any uploaded document",
        "expected": (
            "1. File downloads to browser\n"
            "2. Downloaded file matches uploaded content\n"
            "3. Correct filename and extension"
        ),
        "criteria": "Pass: File downloads correctly. Fail: Download fails or wrong file."
    },
    {
        "id": "TC-283", "name": "View Document AI Analysis Detail", "category": "Documents",
        "priority": "P2 - High", "requirement": "FR-DOC-004",
        "prereqs": "Document with AI analysis exists.",
        "steps": (
            "1. Click on a document with AI analysis\n"
            "2. View the analysis detail panel"
        ),
        "data": "Document with completed AI analysis",
        "expected": (
            "1. Analysis detail shows:\n"
            "   - Overall score with visual indicator\n"
            "   - Category-level scores\n"
            "   - Specific findings/observations\n"
            "   - Recommendations for improvement\n"
            "   - Analysis timestamp"
        ),
        "criteria": "Pass: Detailed analysis visible. Fail: Missing analysis data."
    },
    {
        "id": "TC-284", "name": "File Type Validation", "category": "Documents",
        "priority": "P1 - Critical", "requirement": "FR-DOC-005",
        "prereqs": "Logged in as NGO.",
        "steps": (
            "1. Try uploading each of these file types:\n"
            "   - .pdf (should accept)\n"
            "   - .doc (should accept)\n"
            "   - .docx (should accept)\n"
            "   - .xls (should accept)\n"
            "   - .xlsx (should accept)\n"
            "   - .txt (should accept)\n"
            "   - .csv (should accept)\n"
            "   - .exe (should reject)\n"
            "   - .html (should reject)\n"
            "   - .js (should reject)"
        ),
        "data": "Files with various extensions",
        "expected": (
            "1. PDF, DOC, DOCX, XLS, XLSX, TXT, CSV: accepted\n"
            "2. EXE, HTML, JS, and other types: rejected with 'File type not allowed'\n"
            "3. Validation happens client-side and server-side"
        ),
        "criteria": "Pass: Correct file types accepted/rejected. Fail: Wrong types accepted."
    },
    {
        "id": "TC-285", "name": "File Size Limit Enforcement", "category": "Documents",
        "priority": "P1 - Critical", "requirement": "FR-DOC-006",
        "prereqs": "Logged in as NGO.",
        "steps": (
            "1. Attempt to upload a file exceeding 10MB (or system max)\n"
            "2. Observe the response"
        ),
        "data": "File > 10MB",
        "expected": (
            "1. File rejected with size error\n"
            "2. Error message states maximum allowed size\n"
            "3. File not stored"
        ),
        "criteria": "Pass: Oversized file rejected. Fail: Large file accepted."
    },
]
test_cases.extend(doc_mgmt_cases)

# ============================================================================
# CATEGORY 14: ORGANIZATION PROFILE (TC-290 to TC-299)
# ============================================================================

org_cases = [
    {
        "id": "TC-290", "name": "View Organization Profile", "category": "Organization Profile",
        "priority": "P2 - High", "requirement": "FR-ORG-001",
        "prereqs": "Logged in as NGO user.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Navigate to Organization Profile (if available in sidebar)\n"
            "3. Review profile information"
        ),
        "data": "Email: fatima@amani.org",
        "expected": (
            "1. Profile displays:\n"
            "   - Organization name: Amani Community Development\n"
            "   - Country: Kenya\n"
            "   - Sectors: Health, WASH, Nutrition\n"
            "   - Assessment Score: 82%\n"
            "   - Registration status\n"
            "   - Contact information"
        ),
        "criteria": "Pass: Complete profile displayed. Fail: Missing information."
    },
    {
        "id": "TC-291", "name": "Donor Organization Search", "category": "Organization Profile",
        "priority": "P2 - High", "requirement": "FR-ORG-002",
        "prereqs": "Logged in as donor.",
        "steps": (
            "1. Login as sarah@globalhealth.org / pass123\n"
            "2. Navigate to Organizations\n"
            "3. Search for 'Amani'\n"
            "4. Review search results"
        ),
        "data": "Email: sarah@globalhealth.org | Search: 'Amani'",
        "expected": (
            "1. Search returns Amani Community Development\n"
            "2. Shows: name, country, sectors, capacity score\n"
            "3. Can click to view detailed profile\n"
            "4. Assessment scores visible to donor"
        ),
        "criteria": "Pass: Organization found with details. Fail: No results or missing data."
    },
    {
        "id": "TC-292", "name": "Organization Search by Country", "category": "Organization Profile",
        "priority": "P3 - Low", "requirement": "FR-ORG-003",
        "prereqs": "Logged in as donor.",
        "steps": (
            "1. Login as sarah@globalhealth.org / pass123\n"
            "2. Navigate to Organizations\n"
            "3. Filter by country: Kenya\n"
            "4. Review results"
        ),
        "data": "Filter: Country = Kenya",
        "expected": (
            "1. Only Kenyan organizations shown:\n"
            "   - Amani Community Development\n"
            "2. Non-Kenyan orgs filtered out"
        ),
        "criteria": "Pass: Only Kenyan orgs shown. Fail: Wrong country orgs included."
    },
]
test_cases.extend(org_cases)

# ============================================================================
# CATEGORY 15: END-TO-END WORKFLOWS (TC-300 to TC-310)
# ============================================================================

e2e_cases = [
    {
        "id": "TC-300", "name": "E2E: Complete Grant Lifecycle", "category": "End-to-End Workflows",
        "priority": "P1 - Critical", "requirement": "FR-E2E-001",
        "prereqs": "All seed data loaded. All test accounts available.",
        "steps": (
            "1. DONOR (sarah@globalhealth.org): Create a new grant\n"
            "   a. Fill basic info, eligibility, criteria (100% weight), documents, reporting\n"
            "   b. Upload grant agreement for AI extraction\n"
            "   c. Publish the grant\n\n"
            "2. NGO (fatima@amani.org): Apply for the grant\n"
            "   a. Complete eligibility checklist\n"
            "   b. Write proposal responses using AI Guidance\n"
            "   c. Upload required documents\n"
            "   d. Submit application\n\n"
            "3. DONOR (sarah@globalhealth.org): Assign reviewer\n"
            "   a. Change app status to 'under_review'\n"
            "   b. Assign james@reviewer.org\n\n"
            "4. REVIEWER (james@reviewer.org): Score application\n"
            "   a. Review all proposal responses\n"
            "   b. Use AI Auto-Score then adjust\n"
            "   c. Submit review\n\n"
            "5. DONOR: View rankings, award the grant\n\n"
            "6. NGO: Submit first report\n"
            "   a. Create financial report\n"
            "   b. Fill template sections\n"
            "   c. Submit -- AI analyzes against requirements\n\n"
            "7. DONOR: Review report\n"
            "   a. View AI analysis with per-requirement scores\n"
            "   b. Accept or request revision"
        ),
        "data": "Multiple accounts, grant agreement file, document files",
        "expected": (
            "1. Grant created and published\n"
            "2. Application submitted with AI score\n"
            "3. Reviewer assigned and scores submitted\n"
            "4. Final score calculated correctly\n"
            "5. Report submitted with AI analysis\n"
            "6. Donor can review and accept/revise report\n"
            "7. All status transitions recorded with timestamps"
        ),
        "criteria": "Pass: All steps complete without errors. Fail: Any step fails."
    },
    {
        "id": "TC-301", "name": "E2E: Due Diligence Workflow", "category": "End-to-End Workflows",
        "priority": "P1 - Critical", "requirement": "FR-E2E-002",
        "prereqs": "NGO with registration certificate. Admin access.",
        "steps": (
            "1. NGO (fatima@amani.org): Upload registration certificate\n"
            "2. SYSTEM: AI analyzes certificate content\n"
            "3. ADMIN (admin@kuja.org): View compliance dashboard\n"
            "4. ADMIN: Trigger registry verification for Kenya\n"
            "5. SYSTEM: Cross-checks against Kenya BRS portal\n"
            "6. ADMIN: Trigger sanctions screening\n"
            "7. SYSTEM: Screens against UN, OFAC, EU, World Bank\n"
            "8. DONOR (sarah@globalhealth.org): View compliance status"
        ),
        "data": "NGO: Amani (Kenya) | Registration cert | Admin and donor accounts",
        "expected": (
            "1. Certificate uploaded and AI analyzed\n"
            "2. Registry verification completes (Kenya BRS)\n"
            "3. Sanctions screening returns CLEAR\n"
            "4. Full compliance history visible\n"
            "5. Donor can see verification and screening results"
        ),
        "criteria": "Pass: Full due diligence workflow completes. Fail: Any check fails."
    },
    {
        "id": "TC-302", "name": "E2E: Assessment to Grant Matching", "category": "End-to-End Workflows",
        "priority": "P2 - High", "requirement": "FR-E2E-003",
        "prereqs": "NGO account. Open grants exist.",
        "steps": (
            "1. NGO (peter@hopebridges.org): Complete Kuja assessment\n"
            "2. Verify capacity score updates on profile\n"
            "3. Browse available grants\n"
            "4. Check if grants match org's sectors and capacity\n"
            "5. Apply for a matching grant\n"
            "6. Verify capacity score is factored into application"
        ),
        "data": "Email: peter@hopebridges.org | Framework: Kuja",
        "expected": (
            "1. Assessment completes, score saved to profile\n"
            "2. Grants filtered by relevance to org's sectors\n"
            "3. Application includes capacity score component\n"
            "4. Score visible to donors during review"
        ),
        "criteria": "Pass: Assessment feeds into grant application process. Fail: Score not reflected."
    },
    {
        "id": "TC-303", "name": "E2E: Multi-Language Workflow", "category": "End-to-End Workflows",
        "priority": "P2 - High", "requirement": "FR-E2E-004",
        "prereqs": "NGO account with French language preference.",
        "steps": (
            "1. Login as fatima@amani.org / pass123\n"
            "2. Switch language to French\n"
            "3. Navigate through:\n"
            "   - Dashboard (verify French text)\n"
            "   - Grants (verify French labels)\n"
            "   - Assessments (verify French checklist)\n"
            "   - AI Chat (ask question in French)\n"
            "4. Switch to Arabic\n"
            "5. Verify RTL layout throughout:\n"
            "   - Dashboard\n"
            "   - Grants\n"
            "   - Documents\n"
            "6. Switch back to English\n"
            "7. Verify clean LTR restoration"
        ),
        "data": "Languages: French, Arabic, English",
        "expected": (
            "1. French: all text translated, labels, buttons, messages\n"
            "2. Arabic: RTL layout throughout, all text Arabic\n"
            "3. English: clean return to LTR\n"
            "4. AI Chat responds in selected language\n"
            "5. No mixed-language artifacts"
        ),
        "criteria": "Pass: Clean language switching throughout. Fail: Mixed languages or layout issues."
    },
]
test_cases.extend(e2e_cases)


# ============================================================================
# CATEGORY 16: SUSTAINED AI CHAT  (TC-320 to TC-326)   — Phase 24B, 25B, 26A
# ============================================================================

chat_cases = [
    {
        "id": "TC-320", "name": "Open Chat With Kuja From Sidebar", "category": "AI Chat",
        "priority": "P1 - Critical", "requirement": "FR-CHAT-001",
        "prereqs": "Logged in as NGO (fatima@amani.org) or donor (sarah@globalhealth.org).",
        "steps": (
            "1. From the sidebar, click 'Chat with Kuja'\n"
            "2. Verify /chat page loads with the AIChatPanel rendered\n"
            "3. Verify the composer textarea is visible and NOT disabled\n"
            "4. Verify three example prompts are shown\n"
            "5. Type 'What should I prioritise today?' into the composer\n"
            "6. Press Enter to send"
        ),
        "data": "Message: 'What should I prioritise today?'",
        "expected": (
            "1. Composer becomes enabled within 5 seconds (after thread open)\n"
            "2. User bubble appears immediately with the typed text\n"
            "3. Assistant 'thinking…' placeholder shows next\n"
            "4. Within ~10 seconds, assistant reply replaces the placeholder\n"
            "5. Reply references the user's actual portfolio (no invented numbers)"
        ),
        "criteria": "Pass: Composer is enabled and reply lands. Fail: Composer stays disabled or no response."
    },
    {
        "id": "TC-321", "name": "Chat Remembers Prior Turns", "category": "AI Chat",
        "priority": "P1 - Critical", "requirement": "FR-CHAT-002",
        "prereqs": "TC-320 completed successfully. Existing chat thread with at least one exchange.",
        "steps": (
            "1. Send a first message: 'Summarise my open risks.'\n"
            "2. Wait for the reply.\n"
            "3. Send a follow-up: 'Now rewrite that in a less formal tone.'\n"
            "4. Wait for the reply."
        ),
        "data": "Two-turn conversation.",
        "expected": (
            "1. Second reply explicitly references the content of the first reply\n"
            "2. Tone is noticeably less formal than the first reply\n"
            "3. No mention of having 'lost context' or 'starting fresh'\n"
            "4. Reset button is visible at the top of the panel"
        ),
        "criteria": "Pass: Second reply builds on the first. Fail: Reply ignores the prior turn."
    },
    {
        "id": "TC-322", "name": "Reset Chat Thread Wipes Messages", "category": "AI Chat",
        "priority": "P2 - High", "requirement": "FR-CHAT-003",
        "prereqs": "TC-321 completed. Chat has at least 2 user messages.",
        "steps": (
            "1. Click 'Reset' in the chat header.\n"
            "2. Confirm the prompt: 'Clear this conversation? The thread starts fresh.'\n"
            "3. Verify the message list clears.\n"
            "4. Verify the empty-state nudge with 3 example prompts reappears.\n"
            "5. Send a fresh message and verify no prior context is referenced."
        ),
        "data": "Reset action.",
        "expected": (
            "1. All bubbles disappear immediately\n"
            "2. Thread title returns to 'New conversation'\n"
            "3. Next reply has no memory of pre-reset turns"
        ),
        "criteria": "Pass: Clean reset. Fail: Bubbles remain or prior context leaks."
    },
    {
        "id": "TC-323", "name": "Per-Scope Chat On Grant Detail", "category": "AI Chat",
        "priority": "P1 - Critical", "requirement": "FR-CHAT-004",
        "prereqs": "Logged in as donor. At least one published grant exists.",
        "steps": (
            "1. Open any grant detail page /grants/<id>\n"
            "2. Scroll to the AIChatPanel at the bottom\n"
            "3. Verify thread shows 'scope: grant'\n"
            "4. Send: 'What are the top 3 risks reviewers will flag?'\n"
            "5. Wait for reply"
        ),
        "data": "Real grant id.",
        "expected": (
            "1. Reply references actual grant title, criteria, or sectors\n"
            "2. Reply does NOT mention other grants the donor owns\n"
            "3. Thread is distinct from /chat global thread (separate history)"
        ),
        "criteria": "Pass: Reply uses scope context. Fail: Generic answer or cross-scope leak."
    },
    {
        "id": "TC-324", "name": "Per-Scope Chat On Application Detail", "category": "AI Chat",
        "priority": "P1 - Critical", "requirement": "FR-CHAT-004",
        "prereqs": "Logged in as NGO. At least one application exists.",
        "steps": (
            "1. Open /applications/<id>\n"
            "2. Scroll to the AIChatPanel\n"
            "3. Send: 'Where is this application weakest vs the criteria?'\n"
            "4. Wait for reply"
        ),
        "data": "Real application id.",
        "expected": (
            "1. Reply references this application's actual criteria responses\n"
            "2. Reply does NOT invent scores or amounts not in the app data"
        ),
        "criteria": "Pass: Scope-aware reply. Fail: Hallucinated content."
    },
    {
        "id": "TC-325", "name": "Per-Scope Chat On Report Detail", "category": "AI Chat",
        "priority": "P2 - High", "requirement": "FR-CHAT-004",
        "prereqs": "Logged in as NGO. At least one report exists.",
        "steps": (
            "1. Open /reports/<id>\n"
            "2. Verify the page renders title, status, dates, attachments count\n"
            "3. Scroll to the scoped AIChatPanel\n"
            "4. Send: 'What evidence is missing from this report?'"
        ),
        "data": "Real report id.",
        "expected": (
            "1. Page renders cleanly (no 'Could not load report')\n"
            "2. Chat reply references actual report content/requirements"
        ),
        "criteria": "Pass: Page + chat both load. Fail: 'Could not load' or generic reply."
    },
    {
        "id": "TC-326", "name": "Chat Thread Isolation Across Users", "category": "AI Chat",
        "priority": "P1 - Critical", "requirement": "FR-CHAT-005",
        "prereqs": "Two test users available: fatima and sarah.",
        "steps": (
            "1. Log in as fatima@amani.org\n"
            "2. Visit /chat, send: 'My secret project is X'\n"
            "3. Log out\n"
            "4. Log in as sarah@globalhealth.org\n"
            "5. Visit /chat\n"
            "6. Verify no trace of fatima's message"
        ),
        "data": "Two test sessions.",
        "expected": (
            "1. sarah's /chat shows empty-state or her own prior history only\n"
            "2. No bleed of fatima's content into sarah's view"
        ),
        "criteria": "Pass: Strict per-user isolation. Fail: Any cross-user leak."
    },
]
test_cases.extend(chat_cases)


# ============================================================================
# CATEGORY 17: REVIEWER AUTO-ASSIGNMENT  (TC-330 to TC-334)  — Phase 24A, 25A, 26B
# ============================================================================

reviewer_auto_cases = [
    {
        "id": "TC-330", "name": "Auto-Assign Fires On Application Submit", "category": "Reviewer Assignment",
        "priority": "P1 - Critical", "requirement": "FR-REV-AUTO-001",
        "prereqs": "Logged in as NGO with a draft application ready to submit.",
        "steps": (
            "1. Open the draft application detail page\n"
            "2. Click 'Submit Application'\n"
            "3. After submit succeeds, navigate to the same application as donor\n"
            "4. Inspect the reviewer panel section"
        ),
        "data": "Draft application.",
        "expected": (
            "1. Panel shows ~3 reviewer rows immediately\n"
            "2. Each row shows reviewer name + match score + 1-2 reasons\n"
            "3. No 'Assign reviewers' empty state"
        ),
        "criteria": "Pass: Panel auto-populated. Fail: Empty panel or zero reviewers."
    },
    {
        "id": "TC-331", "name": "Manual Auto-Assign Is Idempotent", "category": "Reviewer Assignment",
        "priority": "P2 - High", "requirement": "FR-REV-AUTO-002",
        "prereqs": "Application already has 3 auto-assigned reviewers (from TC-330).",
        "steps": (
            "1. Open application detail as donor\n"
            "2. Click the 'Auto-assign reviewers' manual button\n"
            "3. Click it again\n"
            "4. Re-inspect the panel"
        ),
        "data": "Already-assigned application.",
        "expected": (
            "1. No duplicate reviewer rows created\n"
            "2. Original 3 assignments unchanged\n"
            "3. Response includes count of existing assignments skipped"
        ),
        "criteria": "Pass: No dupes. Fail: Duplicate reviewers appear."
    },
    {
        "id": "TC-332", "name": "Auto-Assign Sweep Cron Backfills Unassigned Apps", "category": "Reviewer Assignment",
        "priority": "P2 - High", "requirement": "FR-REV-AUTO-003",
        "prereqs": "Admin login. Optional: at least one application with zero reviewers.",
        "steps": (
            "1. Log in as admin@kuja.org\n"
            "2. POST /api/cron/reviewer-auto-assign-sweep (via curl or admin tool)\n"
            "3. Parse response JSON"
        ),
        "data": "Cron POST {}",
        "expected": (
            "1. Response success=true\n"
            "2. result.scanned, result.apps_assigned, result.reviewers_assigned all present\n"
            "3. If sweep ran today, apps_assigned >= 0"
        ),
        "criteria": "Pass: Drift report shape returned. Fail: 500 or missing fields."
    },
    {
        "id": "TC-333", "name": "Reviewer Briefing Card Visible On Assignment", "category": "Reviewer Assignment",
        "priority": "P2 - High", "requirement": "FR-REV-BRIEF-001",
        "prereqs": "Logged in as reviewer (james@reviewer.org) with at least one assigned review.",
        "steps": (
            "1. Open any assigned application detail page\n"
            "2. Locate the reviewer briefing card"
        ),
        "data": "Assigned review.",
        "expected": (
            "1. One-paragraph AI brief is rendered\n"
            "2. Brief lists applicant context, key strengths, red flags to probe\n"
            "3. No 'Briefing unavailable' message"
        ),
        "criteria": "Pass: Briefing renders. Fail: Card missing or error string."
    },
    {
        "id": "TC-334", "name": "Side-By-Side Reviewer Score Comparison", "category": "Reviewer Assignment",
        "priority": "P2 - High", "requirement": "FR-REV-SBS-001",
        "prereqs": "Application has 2+ completed reviews from different reviewers.",
        "steps": (
            "1. Open application detail as donor\n"
            "2. Locate the score-breakdown section\n"
            "3. Switch to the side-by-side view"
        ),
        "data": "Multi-reviewer application.",
        "expected": (
            "1. Each reviewer's per-criterion scores appear in adjacent columns\n"
            "2. Divergence on any criterion is visually highlighted\n"
            "3. If divergence >20pts, panel calibration card surfaces with action"
        ),
        "criteria": "Pass: All reviewers visible side-by-side. Fail: Only one or aggregated only."
    },
]
test_cases.extend(reviewer_auto_cases)


# ============================================================================
# CATEGORY 18: DONOR COHORT + RISK HEATMAP + COMMAND CENTER  (TC-340 to TC-345)
# ============================================================================

donor_intel_cases = [
    {
        "id": "TC-340", "name": "Donor Dashboard Cohort Analytics Card", "category": "Donor Intelligence",
        "priority": "P2 - High", "requirement": "FR-COHORT-001",
        "prereqs": "Logged in as donor (sarah@globalhealth.org).",
        "steps": (
            "1. Navigate to /dashboard\n"
            "2. Scroll to the 'Portfolio quality vs cohort' card\n"
            "3. Inspect metric rows + verdict pills"
        ),
        "data": "Donor session.",
        "expected": (
            "1. Card renders without 'unavailable' message\n"
            "2. Shows either metric rows OR a 'sparse honesty' fallback message\n"
            "3. With <3 other donors: shows 'Only N other donors on the platform'\n"
            "4. With ≥3: shows your value, cohort median, percentile, verdict pill per metric"
        ),
        "criteria": "Pass: Card honest about sample. Fail: Fake numbers or 'unavailable'."
    },
    {
        "id": "TC-341", "name": "Admin Can Inspect Any Donor's Cohort", "category": "Donor Intelligence",
        "priority": "P2 - High", "requirement": "FR-COHORT-002",
        "prereqs": "Logged in as admin.",
        "steps": (
            "1. Navigate to /donors/14 (Sarah's donor org)\n"
            "2. Scroll to the bottom of the page\n"
            "3. Verify cohort card appears (admin-only path)"
        ),
        "data": "Admin viewing donor profile.",
        "expected": (
            "1. Cohort card visible at end of /donors/[id]\n"
            "2. Same content as donor sees on their own dashboard\n"
            "3. Other donor profiles also expose this card to admin"
        ),
        "criteria": "Pass: Admin sees cohort card. Fail: Card hidden for admin."
    },
    {
        "id": "TC-342", "name": "Donor Cohort NGOs Stay Anonymous", "category": "Donor Intelligence",
        "priority": "P1 - Critical", "requirement": "FR-COHORT-PRIV-001",
        "prereqs": "Donor on /dashboard cohort card.",
        "steps": (
            "1. Inspect cohort card content carefully\n"
            "2. Check for any NGO name appearing in metric breakdowns"
        ),
        "data": "Page content scan.",
        "expected": (
            "1. NO NGO names appear\n"
            "2. NO specific dollar amounts of other donors\n"
            "3. Only counts + medians + percentiles + verdicts"
        ),
        "criteria": "Pass: Anonymous medians only. Fail: ANY identifying leak."
    },
    {
        "id": "TC-343", "name": "Portfolio Risk Heatmap Renders Grid", "category": "Donor Intelligence",
        "priority": "P2 - High", "requirement": "FR-RISK-001",
        "prereqs": "Logged in as donor.",
        "steps": (
            "1. Navigate to /dashboard\n"
            "2. Locate the 'Portfolio risk' heatmap section\n"
            "3. Inspect sector × country grid cells"
        ),
        "data": "Donor dashboard.",
        "expected": (
            "1. Grid renders with up to 10×10 cells\n"
            "2. Each cell shows grants/risks/overdue-reports counts\n"
            "3. Cell color reflects aggregate risk score\n"
            "4. Empty cells render as muted"
        ),
        "criteria": "Pass: Grid + colors render. Fail: Empty or error."
    },
    {
        "id": "TC-344", "name": "Donor Verdict Card Action Buttons Route Correctly", "category": "Donor Intelligence",
        "priority": "P2 - High", "requirement": "FR-VERDICT-001",
        "prereqs": "Logged in as donor.",
        "steps": (
            "1. Navigate to /dashboard\n"
            "2. Locate top hero 'Today's portfolio decisions' card\n"
            "3. Verify card shows AI synthesis text + 1-3 action buttons\n"
            "4. Click an action button"
        ),
        "data": "Donor dashboard.",
        "expected": (
            "1. Hero loads within ~8s (AI synthesis is async)\n"
            "2. No 'Briefing unavailable' message\n"
            "3. Clicking an action navigates to the appropriate page (review/grant/etc.)"
        ),
        "criteria": "Pass: Verdict + working actions. Fail: 'Unavailable' or broken nav."
    },
    {
        "id": "TC-345", "name": "Donor Broadcast To All Applicants", "category": "Donor Intelligence",
        "priority": "P2 - High", "requirement": "FR-BROADCAST-001",
        "prereqs": "Donor with at least one grant + 1+ applications.",
        "steps": (
            "1. Open a grant detail page\n"
            "2. Click 'Broadcast' button in header\n"
            "3. Fill subject + body\n"
            "4. Select audience ('All applicants')\n"
            "5. Send"
        ),
        "data": "Subject + body + audience.",
        "expected": (
            "1. Response success=true with recipient count\n"
            "2. Donor.broadcast_sent UserEvent recorded\n"
            "3. Applicants receive in-app notification + email per their prefs"
        ),
        "criteria": "Pass: Broadcast delivered. Fail: 500 or no notifications."
    },
]
test_cases.extend(donor_intel_cases)


# ============================================================================
# CATEGORY 19: PWA INSTALL + NATIVE SHARE + WEBAUTHN  (TC-350 to TC-356)
# ============================================================================

pwa_security_cases = [
    {
        "id": "TC-350", "name": "PWA Install Banner Appears On Chrome", "category": "PWA + Security",
        "priority": "P3 - Medium", "requirement": "FR-PWA-001",
        "prereqs": "Chrome desktop or Android, first-time visit (or after clearing localStorage).",
        "steps": (
            "1. Visit the app in Chrome\n"
            "2. Wait up to 10 seconds\n"
            "3. Observe bottom-right corner"
        ),
        "data": "Fresh browser session.",
        "expected": (
            "1. Install banner appears: 'Install Kuja for faster access'\n"
            "2. Two buttons: 'Install' + 'Not now'\n"
            "3. 'X' dismiss control in corner"
        ),
        "criteria": "Pass: Banner shows. Fail: Banner missing or appears in unsupported browsers."
    },
    {
        "id": "TC-351", "name": "Install Banner Dismissal Persists", "category": "PWA + Security",
        "priority": "P3 - Medium", "requirement": "FR-PWA-002",
        "prereqs": "TC-350 completed and banner is showing.",
        "steps": (
            "1. Click 'Not now' (or the X)\n"
            "2. Reload the page\n"
            "3. Wait 10 seconds"
        ),
        "data": "localStorage flag.",
        "expected": (
            "1. Banner does NOT reappear after dismissal\n"
            "2. localStorage shows 'kuja_pwa_install_dismissed_v1' set\n"
            "3. Other sessions on the same browser also respect dismissal"
        ),
        "criteria": "Pass: Dismissal sticky. Fail: Banner reappears."
    },
    {
        "id": "TC-352", "name": "Native Share Button On Donor Profile", "category": "PWA + Security",
        "priority": "P3 - Medium", "requirement": "FR-SHARE-001",
        "prereqs": "Logged in user. Navigate to /donors/<id>.",
        "steps": (
            "1. Click 'Share profile' button in the hero\n"
            "2. On mobile: verify system share sheet opens\n"
            "3. On desktop: verify clipboard receives the URL + toast appears"
        ),
        "data": "Profile URL.",
        "expected": (
            "1. Mobile: system share sheet appears with the donor URL + title\n"
            "2. Desktop: 'Copied' toast appears, clipboard has the URL\n"
            "3. Button label switches to 'Copied' for ~1.5s"
        ),
        "criteria": "Pass: Share works on both. Fail: Button missing or no action."
    },
    {
        "id": "TC-353", "name": "Native Share Button On NGO Profile", "category": "PWA + Security",
        "priority": "P3 - Medium", "requirement": "FR-SHARE-001",
        "prereqs": "Logged in user. Navigate to /ngo/<id>.",
        "steps": (
            "1. Click 'Share profile' button in the hero\n"
            "2. Verify share behavior matches TC-352"
        ),
        "data": "NGO URL.",
        "expected": "Same as TC-352 but with NGO context.",
        "criteria": "Pass: Share works. Fail: Missing or broken."
    },
    {
        "id": "TC-354", "name": "WebAuthn Device Enrolment", "category": "PWA + Security",
        "priority": "P2 - High", "requirement": "FR-AUTH-WEBAUTHN-001",
        "prereqs": "Logged in user. HTTPS production URL. Browser supports WebAuthn.",
        "steps": (
            "1. Navigate to /settings/security\n"
            "2. Click 'Enrol this device'\n"
            "3. Approve the browser's biometric/security-key prompt\n"
            "4. Verify the device appears in the trusted-devices list"
        ),
        "data": "Touch ID / Face ID / Windows Hello / security key.",
        "expected": (
            "1. Browser shows native biometric prompt\n"
            "2. After approval, device appears with label + 'Added today'\n"
            "3. Server stores credential (verify via /api/auth/webauthn/credentials)"
        ),
        "criteria": "Pass: Device enrolled. Fail: Prompt blocked or device not listed."
    },
    {
        "id": "TC-355", "name": "WebAuthn Credential List Returns Shape", "category": "PWA + Security",
        "priority": "P2 - High", "requirement": "FR-AUTH-WEBAUTHN-002",
        "prereqs": "Logged in user (any role).",
        "steps": (
            "1. GET /api/auth/webauthn/credentials\n"
            "2. Parse JSON response"
        ),
        "data": "Authenticated GET.",
        "expected": (
            "1. Response success=true\n"
            "2. 'credentials' array present (may be empty for new user)\n"
            "3. Each row has: id, label, created_at, last_used_at"
        ),
        "criteria": "Pass: Shape correct. Fail: 500 or missing fields."
    },
    {
        "id": "TC-356", "name": "WebAuthn Unsupported-Browser Fallback", "category": "PWA + Security",
        "priority": "P3 - Medium", "requirement": "FR-AUTH-WEBAUTHN-003",
        "prereqs": "Visit /settings/security in a browser that lacks WebAuthn support.",
        "steps": (
            "1. Navigate to /settings/security\n"
            "2. Observe the WebAuthn panel"
        ),
        "data": "Old Chrome/Firefox or no WebAuthn API.",
        "expected": (
            "1. Panel renders a yellow warning card\n"
            "2. Message: 'Your browser doesn't support WebAuthn'\n"
            "3. 'Enrol' button is NOT rendered\n"
            "4. No crash or blank screen"
        ),
        "criteria": "Pass: Graceful fallback. Fail: Broken button or crash."
    },
]
test_cases.extend(pwa_security_cases)


# ============================================================================
# CATEGORY 20: REAL-USER METRICS + NPS MICRO-SURVEY  (TC-360 to TC-366)
# ============================================================================

metrics_cases = [
    {
        "id": "TC-360", "name": "Admin Metrics Dashboard Loads", "category": "Metrics + Feedback",
        "priority": "P1 - Critical", "requirement": "FR-METRICS-001",
        "prereqs": "Logged in as admin.",
        "steps": (
            "1. Navigate to /admin/metrics\n"
            "2. Wait for the page to settle"
        ),
        "data": "Admin session.",
        "expected": (
            "1. Page renders with stat tiles: DAU, WAU, MAU\n"
            "2. WAU breakdowns by role + language render as chips\n"
            "3. Six funnels visible: chat, application, report, review,\n"
            "   readiness→submit, preflight→submit\n"
            "4. Language parity card with chat + search adoption\n"
            "5. Event volume table\n"
            "6. NPS feedback card (sparse-honest if no responses)"
        ),
        "criteria": "Pass: Full page renders. Fail: 'Could not load metrics' or missing sections."
    },
    {
        "id": "TC-361", "name": "Login Records session.start Event", "category": "Metrics + Feedback",
        "priority": "P1 - Critical", "requirement": "FR-METRICS-002",
        "prereqs": "Admin login + access to /admin/metrics.",
        "steps": (
            "1. Note WAU total before login\n"
            "2. Log out\n"
            "3. Log in as a different user\n"
            "4. Refresh /admin/metrics\n"
            "5. Verify WAU total increased OR session.start count incremented"
        ),
        "data": "Two sequential logins.",
        "expected": (
            "1. WAU total reflects the new session\n"
            "2. event_counts_30d shows session.start with count > 0\n"
            "3. by_role and by_language breakdowns also updated"
        ),
        "criteria": "Pass: Event recorded. Fail: Counts unchanged after login."
    },
    {
        "id": "TC-362", "name": "Non-Admin Cannot Access Metrics", "category": "Metrics + Feedback",
        "priority": "P1 - Critical", "requirement": "FR-METRICS-AUTH-001",
        "prereqs": "Logged in as NGO or donor.",
        "steps": (
            "1. GET /api/admin/metrics\n"
            "2. OR navigate to /admin/metrics in browser"
        ),
        "data": "Non-admin role.",
        "expected": (
            "1. API returns 403 with 'Admin access required'\n"
            "2. Page shows 'Admin access required' warning card"
        ),
        "criteria": "Pass: Properly gated. Fail: Metrics exposed to non-admin."
    },
    {
        "id": "TC-363", "name": "Micro-Survey Appears After Application Submit", "category": "Metrics + Feedback",
        "priority": "P2 - High", "requirement": "FR-FEEDBACK-001",
        "prereqs": "Logged in as NGO. Application submitted (not draft). First time on this app detail.",
        "steps": (
            "1. Submit an application\n"
            "2. Navigate to the application detail page\n"
            "3. Wait ~1 second after page load"
        ),
        "data": "Submitted application, own NGO.",
        "expected": (
            "1. Micro-survey card appears bottom-right\n"
            "2. Question: 'How helpful was Kuja in preparing this application?'\n"
            "3. 0-10 score row + optional comment\n"
            "4. 'Not now' and 'X' both dismiss"
        ),
        "criteria": "Pass: Survey shows once. Fail: Doesn't appear or re-appears."
    },
    {
        "id": "TC-364", "name": "Submitting Feedback Persists + Affects NPS", "category": "Metrics + Feedback",
        "priority": "P2 - High", "requirement": "FR-FEEDBACK-002",
        "prereqs": "TC-363 micro-survey visible.",
        "steps": (
            "1. Click score 9\n"
            "2. Optional: type a short comment\n"
            "3. Click Send\n"
            "4. Wait for 'Thanks — feedback saved' confirmation\n"
            "5. Log in as admin and refresh /admin/metrics"
        ),
        "data": "Score 9 with or without comment.",
        "expected": (
            "1. Survey shows 'Thanks' confirmation then dismisses\n"
            "2. NPS card on /admin/metrics shows total_responses incremented\n"
            "3. By-surface table shows application_submit row\n"
            "4. Comment (if any) appears in recent-comments stream"
        ),
        "criteria": "Pass: NPS updates. Fail: Submit fails or NPS unchanged."
    },
    {
        "id": "TC-365", "name": "Survey Won't Re-Prompt Same Surface", "category": "Metrics + Feedback",
        "priority": "P3 - Medium", "requirement": "FR-FEEDBACK-003",
        "prereqs": "TC-364 completed (response saved).",
        "steps": (
            "1. Reload the application detail page\n"
            "2. Wait ~3 seconds"
        ),
        "data": "Same application as TC-363.",
        "expected": (
            "1. Survey does NOT reappear (localStorage flag set)\n"
            "2. The chat panel, score breakdown, etc. all still render normally"
        ),
        "criteria": "Pass: One-and-done. Fail: Survey re-appears."
    },
    {
        "id": "TC-366", "name": "Generic Event Ingest Whitelist", "category": "Metrics + Feedback",
        "priority": "P3 - Medium", "requirement": "FR-METRICS-INGEST-001",
        "prereqs": "Any logged-in user.",
        "steps": (
            "1. POST /api/ai/events/track {event_name: 'feature.tap', props: {feature: 'test'}}\n"
            "2. POST same endpoint with {event_name: 'fake.unknown'}"
        ),
        "data": "Whitelisted + non-whitelisted event names.",
        "expected": (
            "1. First call: 200 with success=true\n"
            "2. Second call: 400 with 'event_name not allowed' + allowed list"
        ),
        "criteria": "Pass: Whitelist enforced. Fail: Arbitrary names accepted."
    },
]
test_cases.extend(metrics_cases)


# ============================================================================
# CATEGORY 21: API ALIASES + WIRING FIXES  (TC-370 to TC-374) — Phase 27, 28
# ============================================================================

wiring_cases = [
    {
        "id": "TC-370", "name": "/api/api/ Double-Prefix Returns 404", "category": "API Wiring",
        "priority": "P1 - Critical", "requirement": "FR-WIRING-001",
        "prereqs": "Logged in as admin.",
        "steps": (
            "1. GET /api/api/reports/1\n"
            "2. GET /api/api/trust-profile/9\n"
            "3. GET /api/api/audit-chain/recent\n"
            "4. GET /api/api/organizations/14/donor-profile"
        ),
        "data": "All URLs with intentional /api/api/ prefix.",
        "expected": (
            "1. All four return 404 (these paths never existed)\n"
            "2. No 200 responses — proves the team's 2026-05-16 bug pattern\n"
            "   would still 404 if it ever recurred via copy-paste"
        ),
        "criteria": "Pass: 404 on all four. Fail: Any 200 means double-prefix accepted."
    },
    {
        "id": "TC-371", "name": "Correct Path Equivalents Return 200", "category": "API Wiring",
        "priority": "P1 - Critical", "requirement": "FR-WIRING-002",
        "prereqs": "Admin login.",
        "steps": (
            "1. GET /api/trust-profile/9\n"
            "2. GET /api/audit-chain/recent?limit=10\n"
            "3. GET /api/organizations/14/donor-profile"
        ),
        "data": "Properly-pathed URLs.",
        "expected": "All three return 200 with shaped JSON.",
        "criteria": "Pass: All 200. Fail: Any 404 or 5xx."
    },
    {
        "id": "TC-372", "name": "/api/search Alias Returns Global Results", "category": "API Wiring",
        "priority": "P2 - High", "requirement": "FR-SEARCH-ALIAS-001",
        "prereqs": "Logged in user.",
        "steps": (
            "1. GET /api/search?q=kenya\n"
            "2. Parse response"
        ),
        "data": "Query: kenya",
        "expected": (
            "1. Response 200\n"
            "2. 'results' array present with matching grants/apps/reports\n"
            "3. Each hit has kind, id, snippet, href"
        ),
        "criteria": "Pass: Alias works. Fail: 404 (regression to pre-Phase-28A)."
    },
    {
        "id": "TC-373", "name": "/api/notifications/preferences Alias", "category": "API Wiring",
        "priority": "P2 - High", "requirement": "FR-NOTIF-ALIAS-001",
        "prereqs": "Logged in user.",
        "steps": (
            "1. GET /api/notifications/preferences\n"
            "2. Parse response"
        ),
        "data": "Alias path.",
        "expected": (
            "1. Response 200, success=true\n"
            "2. 'categories' array + 'catalog' object present\n"
            "3. /api/notification-preferences (original) also still works"
        ),
        "criteria": "Pass: Alias parity. Fail: 404 alias or original broken."
    },
    {
        "id": "TC-374", "name": "Bulk Lockout Clear Resets All Accounts", "category": "API Wiring",
        "priority": "P2 - High", "requirement": "FR-LOCKOUT-CLEAR-001",
        "prereqs": "Logged in as admin. At least one account with failed_login_count > 0 (optional).",
        "steps": (
            "1. POST /api/admin/clear-all-lockouts {}\n"
            "2. Parse response\n"
            "3. Verify NGO can no longer return 403 from login due to lockout"
        ),
        "data": "Admin POST.",
        "expected": (
            "1. success=true\n"
            "2. users_reset + attempts_deleted counts returned\n"
            "3. NGO cannot trigger this endpoint (403)"
        ),
        "criteria": "Pass: Bulk reset works + role-gated. Fail: 500 or open to NGO."
    },
]
test_cases.extend(wiring_cases)


# ============================================================================
# CATEGORY 22: AI CHAT — EDGE CASES + SECURITY  (TC-400 to TC-414)
# ============================================================================

chat_edge_cases = [
    {
        "id": "TC-400", "name": "Empty Chat Message Rejected", "category": "AI Chat Edge",
        "priority": "P2 - High", "requirement": "FR-CHAT-VAL-001",
        "prereqs": "Open /chat, thread is active.",
        "steps": (
            "1. Click in the composer\n"
            "2. Type only whitespace (' ' x10)\n"
            "3. Press Send"
        ),
        "data": "Whitespace-only message.",
        "expected": (
            "1. Send button is disabled when only whitespace present\n"
            "2. If somehow submitted, server returns success=False reason=empty_message\n"
            "3. No assistant call is made, no token spend"
        ),
        "criteria": "Pass: Empty rejected. Fail: Empty message triggers AI call."
    },
    {
        "id": "TC-401", "name": "Chat Message Exceeds Max Length", "category": "AI Chat Edge",
        "priority": "P2 - High", "requirement": "FR-CHAT-VAL-002",
        "prereqs": "Open /chat.",
        "steps": (
            "1. Paste 5000 characters into the composer\n"
            "2. Verify the textarea visually accepts only up to 4000\n"
            "3. Click Send"
        ),
        "data": "Long-form text >4000 chars.",
        "expected": (
            "1. textarea maxLength enforces 4000 hard cap\n"
            "2. Server truncates anything beyond MAX_USER_MESSAGE_CHARS\n"
            "3. AI receives the truncated form, no 413 error"
        ),
        "criteria": "Pass: Bounded to 4000. Fail: Send fails or unbounded."
    },
    {
        "id": "TC-402", "name": "Chat Rate Limit Returns 429", "category": "AI Chat Edge",
        "priority": "P2 - High", "requirement": "FR-CHAT-RATE-001",
        "prereqs": "Logged in. Open /chat.",
        "steps": (
            "1. Send a message and wait for reply\n"
            "2. Send a second message immediately\n"
            "3. Continue sending without pause until rate limit hits"
        ),
        "data": "Repeated rapid sends.",
        "expected": (
            "1. After N messages within the window, server returns 429\n"
            "2. UI shows 'ai.rate_limited' error or graceful banner\n"
            "3. Subsequent send succeeds after window resets"
        ),
        "criteria": "Pass: 429 fired + recovery works. Fail: No limit or stuck after."
    },
    {
        "id": "TC-403", "name": "Composer Disabled While Awaiting Reply", "category": "AI Chat Edge",
        "priority": "P2 - High", "requirement": "FR-CHAT-UX-001",
        "prereqs": "Open /chat.",
        "steps": (
            "1. Send a message\n"
            "2. While 'thinking…' placeholder is showing, try to type\n"
            "3. Try to click Send"
        ),
        "data": "In-flight request.",
        "expected": (
            "1. Composer textarea is disabled during the assistant turn\n"
            "2. Send button is disabled\n"
            "3. Re-enables exactly when the assistant reply lands"
        ),
        "criteria": "Pass: Properly disabled. Fail: Concurrent sends fire."
    },
    {
        "id": "TC-404", "name": "Reset Confirmation Cancel Preserves Thread", "category": "AI Chat Edge",
        "priority": "P3 - Medium", "requirement": "FR-CHAT-RESET-001",
        "prereqs": "Open /chat with at least one exchange.",
        "steps": (
            "1. Click Reset\n"
            "2. When the confirm dialog appears, click Cancel"
        ),
        "data": "Reset cancellation.",
        "expected": (
            "1. Dialog closes\n"
            "2. All messages remain visible\n"
            "3. Thread state unchanged on the server"
        ),
        "criteria": "Pass: Cancel cancels. Fail: Thread wiped on cancel."
    },
    {
        "id": "TC-405", "name": "Thread Title Auto-Set From First Message", "category": "AI Chat Edge",
        "priority": "P3 - Medium", "requirement": "FR-CHAT-TITLE-001",
        "prereqs": "Open a fresh chat thread (after reset or new user).",
        "steps": (
            "1. Send first message 'Summarise my open risks please.'\n"
            "2. Refresh the page"
        ),
        "data": "First user message.",
        "expected": (
            "1. Header now shows the truncated message as the thread title\n"
            "2. Title is limited to ~90 chars\n"
            "3. Refresh preserves the title"
        ),
        "criteria": "Pass: Title set + persisted. Fail: Title empty or wrong."
    },
    {
        "id": "TC-406", "name": "Chat Reply In User's Language (Arabic)", "category": "AI Chat Edge",
        "priority": "P2 - High", "requirement": "FR-CHAT-I18N-001",
        "prereqs": "Logged in as user with language='ar'.",
        "steps": (
            "1. Open /chat\n"
            "2. Send message in Arabic: 'ما هي مخاطري الحالية؟'\n"
            "3. Wait for reply"
        ),
        "data": "Arabic prompt.",
        "expected": (
            "1. Reply lands in Arabic\n"
            "2. RTL layout in the assistant bubble\n"
            "3. No mid-reply language switching"
        ),
        "criteria": "Pass: Arabic reply. Fail: English reply or mixed languages."
    },
    {
        "id": "TC-407", "name": "Chat Reply In Swahili", "category": "AI Chat Edge",
        "priority": "P3 - Medium", "requirement": "FR-CHAT-I18N-002",
        "prereqs": "Logged in as user with language='sw'.",
        "steps": (
            "1. Open /chat\n"
            "2. Send message in Swahili: 'Nipe muhtasari wa hatari zangu.'\n"
            "3. Wait for reply"
        ),
        "data": "Swahili prompt.",
        "expected": "Reply lands in Swahili.",
        "criteria": "Pass: Swahili. Fail: Other language."
    },
    {
        "id": "TC-408", "name": "AI Service Unavailable Graceful Fallback", "category": "AI Chat Edge",
        "priority": "P2 - High", "requirement": "FR-CHAT-FALLBACK-001",
        "prereqs": "Block /api/ai/threads/<id>/messages OR set ANTHROPIC_API_KEY invalid in test env.",
        "steps": (
            "1. Send a message\n"
            "2. Wait for response"
        ),
        "data": "AI call fails.",
        "expected": (
            "1. Server returns success=False with reason='ai_call_failed'\n"
            "2. UI shows '[AI unavailable — please try again in a moment.]'\n"
            "3. Composer re-enables\n"
            "4. No crash, no infinite spinner"
        ),
        "criteria": "Pass: Friendly fallback. Fail: Crash or spinner forever."
    },
    {
        "id": "TC-409", "name": "Anti-Hallucination — Asking About Unloaded Scope", "category": "AI Chat Edge",
        "priority": "P1 - Critical", "requirement": "FR-CHAT-HONESTY-001",
        "prereqs": "Open /chat (global scope).",
        "steps": (
            "1. Ask: 'What was the AI score on application #99999?'\n"
            "2. Wait for reply"
        ),
        "data": "Question referencing data not in scope.",
        "expected": (
            "1. Assistant does NOT invent a score\n"
            "2. Reply says something like 'I don't have that loaded — try the application page'\n"
            "3. No false numbers, no apologies-then-confabulation"
        ),
        "criteria": "Pass: Honest refusal. Fail: Made-up score."
    },
    {
        "id": "TC-410", "name": "XSS Sanitization in User Message", "category": "AI Chat Edge",
        "priority": "P1 - Critical", "requirement": "FR-CHAT-SEC-001",
        "prereqs": "Open /chat.",
        "steps": (
            "1. Send: <script>alert('xss')</script>\n"
            "2. Send: <img src=x onerror=alert(1)>\n"
            "3. Send: <iframe src='evil.com'>"
        ),
        "data": "XSS payloads.",
        "expected": (
            "1. No alert dialogs appear\n"
            "2. Bubbles render the literal escaped text\n"
            "3. View-source confirms HTML entities, not raw tags\n"
            "4. No iframe loads"
        ),
        "criteria": "Pass: All escaped. Fail: ANY dialog or unescaped HTML."
    },
    {
        "id": "TC-411", "name": "Open Chat When User Has No Org", "category": "AI Chat Edge",
        "priority": "P3 - Medium", "requirement": "FR-CHAT-EDGE-001",
        "prereqs": "User account exists but org_id is NULL (rare for production but possible during onboarding).",
        "steps": (
            "1. Open /chat\n"
            "2. Try to send a message"
        ),
        "data": "Orphan user.",
        "expected": (
            "1. Thread opens (user_id is the keying field)\n"
            "2. Message sends and gets reply\n"
            "3. Reply does NOT reference any org-specific data"
        ),
        "criteria": "Pass: Works gracefully. Fail: 500 or blocked."
    },
    {
        "id": "TC-412", "name": "Long Chat History Beyond 12-Message Cap", "category": "AI Chat Edge",
        "priority": "P2 - High", "requirement": "FR-CHAT-COST-001",
        "prereqs": "Open chat with at least 15 prior messages.",
        "steps": (
            "1. Send a new message: 'What did I ask you in the very first message?'\n"
            "2. Wait for reply"
        ),
        "data": "Long thread.",
        "expected": (
            "1. UI shows all 15+ bubbles\n"
            "2. Server only sends the last 12 to Claude\n"
            "3. Reply correctly says it doesn't recall the earliest message\n"
            "4. Cost-tagged endpoint='ai.chat' for budget guard"
        ),
        "criteria": "Pass: History capped at 12. Fail: All 15 sent + cost spike."
    },
    {
        "id": "TC-413", "name": "Two Browser Tabs Same User", "category": "AI Chat Edge",
        "priority": "P3 - Medium", "requirement": "FR-CHAT-CONCURRENT-001",
        "prereqs": "Logged in as fatima. Open /chat in two tabs.",
        "steps": (
            "1. Send a message in tab A\n"
            "2. Switch to tab B and send a different message\n"
            "3. Wait for both replies\n"
            "4. Refresh tab A"
        ),
        "data": "Two-tab session.",
        "expected": (
            "1. Both messages land in the same underlying thread\n"
            "2. After refresh, tab A sees both user messages + both replies\n"
            "3. Server thread is the canonical source of truth"
        ),
        "criteria": "Pass: Convergence. Fail: Two threads created."
    },
    {
        "id": "TC-414", "name": "Reset Thread Without Any Messages", "category": "AI Chat Edge",
        "priority": "P3 - Medium", "requirement": "FR-CHAT-RESET-002",
        "prereqs": "Fresh user. Open /chat. Thread row exists, no messages.",
        "steps": (
            "1. Click Reset\n"
            "2. Confirm prompt"
        ),
        "data": "Empty thread.",
        "expected": (
            "1. Reset returns success=true\n"
            "2. No errors\n"
            "3. Empty-state still rendered"
        ),
        "criteria": "Pass: No-op safe. Fail: 500 or odd state."
    },
]
test_cases.extend(chat_edge_cases)


# ============================================================================
# CATEGORY 23: REVIEWER AUTO-ASSIGN — EDGE CASES  (TC-420 to TC-431)
# ============================================================================

reviewer_edge_cases = [
    {
        "id": "TC-420", "name": "Auto-Assign With Empty Reviewer Pool", "category": "Reviewer Auto Edge",
        "priority": "P2 - High", "requirement": "FR-REV-AUTO-EDGE-001",
        "prereqs": "Test env where no reviewers exist for the application's sector/country.",
        "steps": (
            "1. Submit an application\n"
            "2. POST /api/applications/<id>/auto-assign-reviewers\n"
            "3. Inspect response"
        ),
        "data": "Sector with no matching reviewers.",
        "expected": (
            "1. Response ok=true, assigned=0, reason='no_reviewers_in_pool'\n"
            "2. No reviewer rows created\n"
            "3. No 500 error\n"
            "4. Application status unchanged"
        ),
        "criteria": "Pass: Graceful zero. Fail: 500 or fake assignments."
    },
    {
        "id": "TC-421", "name": "Auto-Assign All-Slipping Pool Fallback", "category": "Reviewer Auto Edge",
        "priority": "P2 - High", "requirement": "FR-REV-AUTO-EDGE-002",
        "prereqs": "Test env where the only matching reviewers are flagged 'busy queue'.",
        "steps": (
            "1. POST auto-assign\n"
            "2. Verify the panel still populates"
        ),
        "data": "Slipping pool only.",
        "expected": (
            "1. Service falls back to slipping reviewers (rather than leaving empty)\n"
            "2. Assigned count > 0\n"
            "3. Reasoning shows 'busy queue — assign carefully' string"
        ),
        "criteria": "Pass: Fallback works. Fail: Empty panel."
    },
    {
        "id": "TC-422", "name": "Auto-Assign Skips Already-Assigned Reviewer", "category": "Reviewer Auto Edge",
        "priority": "P1 - Critical", "requirement": "FR-REV-AUTO-IDEMP-001",
        "prereqs": "Application has 1 manually-assigned reviewer who is also top-ranked.",
        "steps": (
            "1. POST auto-assign with panel_size=3"
        ),
        "data": "Top-ranked reviewer is already assigned.",
        "expected": (
            "1. Manually-assigned reviewer is NOT duplicated\n"
            "2. Service adds 2 NEW reviewers from the remaining ranked list\n"
            "3. Response shows already_assigned=1, assigned=2"
        ),
        "criteria": "Pass: No dupes. Fail: Duplicate or only 2 assigned."
    },
    {
        "id": "TC-423", "name": "Auto-Assign On Non-Existent Application", "category": "Reviewer Auto Edge",
        "priority": "P2 - High", "requirement": "FR-REV-AUTO-VAL-001",
        "prereqs": "Donor login.",
        "steps": (
            "1. POST /api/applications/999999999/auto-assign-reviewers"
        ),
        "data": "Bogus application id.",
        "expected": (
            "1. Response 400 or 404 with reason='application_not_found'\n"
            "2. No reviewer rows created\n"
            "3. No 500"
        ),
        "criteria": "Pass: Clean validation error. Fail: 500."
    },
    {
        "id": "TC-424", "name": "Auto-Assign Forbidden For NGO Role", "category": "Reviewer Auto Edge",
        "priority": "P1 - Critical", "requirement": "FR-REV-AUTO-AUTH-001",
        "prereqs": "NGO login (fatima@amani.org).",
        "steps": (
            "1. POST /api/applications/<id>/auto-assign-reviewers"
        ),
        "data": "NGO trying to trigger.",
        "expected": "403 or 404 (depending on app ownership)",
        "criteria": "Pass: NGO blocked. Fail: NGO can trigger."
    },
    {
        "id": "TC-425", "name": "Auto-Assign Panel Size Cap", "category": "Reviewer Auto Edge",
        "priority": "P2 - High", "requirement": "FR-REV-AUTO-CAP-001",
        "prereqs": "Donor login. Application exists.",
        "steps": (
            "1. POST auto-assign with panel_size=10"
        ),
        "data": "panel_size=10 exceeds MAX_PANEL_SIZE.",
        "expected": (
            "1. Service caps at MAX_PANEL_SIZE (5)\n"
            "2. assigned count <= 5\n"
            "3. No 5xx error"
        ),
        "criteria": "Pass: Capped. Fail: Honors absurd size."
    },
    {
        "id": "TC-426", "name": "Cron Sweep With No Candidates", "category": "Reviewer Auto Edge",
        "priority": "P2 - High", "requirement": "FR-REV-CRON-EDGE-001",
        "prereqs": "Test env where every application already has a reviewer.",
        "steps": (
            "1. POST /api/cron/reviewer-auto-assign-sweep as admin"
        ),
        "data": "Already-assigned ecosystem.",
        "expected": (
            "1. Response success=true\n"
            "2. result.scanned=0\n"
            "3. result.apps_assigned=0\n"
            "4. result.skipped.already_assigned >= number of submitted apps in scan budget"
        ),
        "criteria": "Pass: Clean empty sweep. Fail: 500 or wrong counts."
    },
    {
        "id": "TC-427", "name": "Cron Sweep Caps At MAX_PER_RUN", "category": "Reviewer Auto Edge",
        "priority": "P3 - Medium", "requirement": "FR-REV-CRON-CAP-001",
        "prereqs": "Test env with 100+ submitted apps with zero reviewers (unrealistic in prod).",
        "steps": (
            "1. POST /api/cron/reviewer-auto-assign-sweep"
        ),
        "data": "Large backlog.",
        "expected": (
            "1. result.scanned <= 50\n"
            "2. result.cap=50 in response\n"
            "3. Backlog persists; next sweep picks up more"
        ),
        "criteria": "Pass: Capped. Fail: Unbounded sweep."
    },
    {
        "id": "TC-428", "name": "Cron Sweep With Wrong CRON_SECRET", "category": "Reviewer Auto Edge",
        "priority": "P1 - Critical", "requirement": "FR-REV-CRON-AUTH-001",
        "prereqs": "Not logged in. Use a bad Bearer token.",
        "steps": (
            "1. POST with Authorization: Bearer wrong-secret\n"
            "2. POST with no header at all"
        ),
        "data": "Wrong / missing secret.",
        "expected": "Both return 403 forbidden.",
        "criteria": "Pass: 403 in both cases. Fail: Any 200."
    },
    {
        "id": "TC-429", "name": "Audit Chain Anchor For Auto-Assign", "category": "Reviewer Auto Edge",
        "priority": "P2 - High", "requirement": "FR-REV-AUDIT-001",
        "prereqs": "Submit an application that triggers auto-assign.",
        "steps": (
            "1. Submit application as NGO\n"
            "2. As admin, GET /api/audit-chain/recent?limit=10\n"
            "3. Look for action='reviewer.auto_assigned'"
        ),
        "data": "Audit chain query.",
        "expected": (
            "1. Entry exists with action='reviewer.auto_assigned'\n"
            "2. subject_kind='application', subject_id matches\n"
            "3. details contains panel_size_requested, assigned, reviewer_ids"
        ),
        "criteria": "Pass: Entry anchored. Fail: Missing audit row."
    },
    {
        "id": "TC-430", "name": "Match Score Reasoning Surfaces On Panel", "category": "Reviewer Auto Edge",
        "priority": "P3 - Medium", "requirement": "FR-REV-UI-001",
        "prereqs": "Auto-assigned application detail page as donor.",
        "steps": (
            "1. Open application detail\n"
            "2. Inspect each reviewer row in the panel"
        ),
        "data": "Auto-assigned panel.",
        "expected": (
            "1. Each row shows numeric match score\n"
            "2. Reasons array shown (e.g. 'sector fit: education', 'country fit: Kenya')\n"
            "3. If reviewer was a 'slipping' pick: warning indicator"
        ),
        "criteria": "Pass: Reasoning visible. Fail: Score only or no info."
    },
    {
        "id": "TC-431", "name": "Auto-Assign After Reviewer Removed", "category": "Reviewer Auto Edge",
        "priority": "P3 - Medium", "requirement": "FR-REV-AUTO-REASSIGN-001",
        "prereqs": "Application has 3 auto-assigned reviewers. Donor removes one.",
        "steps": (
            "1. Donor removes 1 reviewer (DELETE Review row)\n"
            "2. POST auto-assign with panel_size=3"
        ),
        "data": "Reduced panel.",
        "expected": (
            "1. Service tops up: 1 new reviewer added\n"
            "2. Existing 2 are not duplicated\n"
            "3. Net 3 reviewers"
        ),
        "criteria": "Pass: Top-up correct. Fail: Wrong total."
    },
]
test_cases.extend(reviewer_edge_cases)


# ============================================================================
# CATEGORY 24: DONOR INTEL EDGE CASES + PRIVACY  (TC-440 to TC-451)
# ============================================================================

donor_edge_cases = [
    {
        "id": "TC-440", "name": "Cohort Card With Zero Awarded Apps", "category": "Donor Intel Edge",
        "priority": "P2 - High", "requirement": "FR-COHORT-EDGE-001",
        "prereqs": "Brand-new donor org with no awarded applications yet.",
        "steps": (
            "1. Log in as fresh donor\n"
            "2. Visit /dashboard\n"
            "3. Scroll to cohort card"
        ),
        "data": "Donor with empty portfolio.",
        "expected": (
            "1. Card shows 'Once you've awarded a couple of grants…' empty state\n"
            "2. No fake metric rows\n"
            "3. No 'sparse' message (cohort has data; donor doesn't)"
        ),
        "criteria": "Pass: Honest empty state. Fail: Fake numbers."
    },
    {
        "id": "TC-441", "name": "Cohort Per-Metric Sparseness", "category": "Donor Intel Edge",
        "priority": "P2 - High", "requirement": "FR-COHORT-EDGE-002",
        "prereqs": "Donor with 1 awarded application (below MIN_SAMPLE=2).",
        "steps": (
            "1. Open cohort card\n"
            "2. Inspect metric rows"
        ),
        "data": "Donor with N=1.",
        "expected": (
            "1. Metrics with self_sample_size < MIN_SAMPLE are OMITTED\n"
            "2. Card may show fewer than 6 metric rows\n"
            "3. Footer accurately shows portfolio_size=1"
        ),
        "criteria": "Pass: Per-metric honesty. Fail: Lies with N=1."
    },
    {
        "id": "TC-442", "name": "Cohort Anonymity — Every Metric Audited", "category": "Donor Intel Edge",
        "priority": "P1 - Critical", "requirement": "FR-COHORT-PRIV-002",
        "prereqs": "Donor with active portfolio + cohort data.",
        "steps": (
            "1. Inspect every visible metric: capacity, AI score, country diversity, sector diversity, small-org share, reports on time\n"
            "2. Open dev tools, inspect raw JSON of /api/dashboard/donor-cohort-analytics"
        ),
        "data": "Full page + raw API.",
        "expected": (
            "1. NO NGO names anywhere in metric rows or raw JSON\n"
            "2. NO specific donor names\n"
            "3. NO grant titles or IDs\n"
            "4. Only: codes, medians, percentiles, counts, verdicts"
        ),
        "criteria": "Pass: Zero identifying data. Fail: ANY leak."
    },
    {
        "id": "TC-443", "name": "Cross-Donor Cohort Isolation", "category": "Donor Intel Edge",
        "priority": "P1 - Critical", "requirement": "FR-COHORT-AUTH-001",
        "prereqs": "Two donors: sarah (org 14) and david (org 7).",
        "steps": (
            "1. Log in as sarah\n"
            "2. GET /api/dashboard/donor-cohort-analytics\n"
            "3. Log out\n"
            "4. Log in as david\n"
            "5. GET same endpoint\n"
            "6. Compare results"
        ),
        "data": "Two donor sessions.",
        "expected": (
            "1. Each call scopes to caller's own donor_org_id\n"
            "2. Results may differ in self_value, never expose other donor's identifiers\n"
            "3. Cohort_size for each excludes the caller (other_donors filter)"
        ),
        "criteria": "Pass: Strict isolation. Fail: Cross-donor leak."
    },
    {
        "id": "TC-444", "name": "Admin Cohort Inspect — Bad donor_org_id", "category": "Donor Intel Edge",
        "priority": "P2 - High", "requirement": "FR-COHORT-VAL-001",
        "prereqs": "Admin login.",
        "steps": (
            "1. GET /api/dashboard/donor-cohort-analytics (no donor_org_id)\n"
            "2. GET with donor_org_id=abc (non-int)\n"
            "3. GET with donor_org_id=99999 (non-existent)\n"
            "4. GET with donor_org_id=14 (real)"
        ),
        "data": "Validation cases.",
        "expected": (
            "1. No param → 400 'admin must pass donor_org_id'\n"
            "2. Non-int → 400 'donor_org_id must be int'\n"
            "3. Non-existent → 200 with reason='not_donor' or success=False\n"
            "4. Real → 200 with full shape"
        ),
        "criteria": "Pass: All validations clean. Fail: 500 on any."
    },
    {
        "id": "TC-445", "name": "Heatmap With Single Grant", "category": "Donor Intel Edge",
        "priority": "P3 - Medium", "requirement": "FR-RISK-EDGE-001",
        "prereqs": "Donor with exactly 1 grant.",
        "steps": (
            "1. Visit /dashboard, see heatmap"
        ),
        "data": "Single-grant donor.",
        "expected": (
            "1. Heatmap renders with 1 cell\n"
            "2. No 'sparse' message\n"
            "3. No 5xx error"
        ),
        "criteria": "Pass: Renders. Fail: Blank or error."
    },
    {
        "id": "TC-446", "name": "Heatmap With No Grants", "category": "Donor Intel Edge",
        "priority": "P3 - Medium", "requirement": "FR-RISK-EDGE-002",
        "prereqs": "Fresh donor with no grants.",
        "steps": (
            "1. Visit /dashboard"
        ),
        "data": "Empty portfolio.",
        "expected": (
            "1. Heatmap shows 0 cells\n"
            "2. Empty state explanation visible\n"
            "3. total_grants=0 in response"
        ),
        "criteria": "Pass: Honest empty. Fail: Fake cells."
    },
    {
        "id": "TC-447", "name": "Heatmap Forbidden For NGO", "category": "Donor Intel Edge",
        "priority": "P1 - Critical", "requirement": "FR-RISK-AUTH-001",
        "prereqs": "NGO login.",
        "steps": (
            "1. GET /api/dashboard/portfolio-risk-heatmap"
        ),
        "data": "NGO request.",
        "expected": "403 'Role not supported'",
        "criteria": "Pass: 403. Fail: Data leak."
    },
    {
        "id": "TC-448", "name": "Cohort Excludes Caller From 'cohort_size'", "category": "Donor Intel Edge",
        "priority": "P2 - High", "requirement": "FR-COHORT-CALC-001",
        "prereqs": "Total donors on platform = N.",
        "steps": (
            "1. Log in as any donor\n"
            "2. Inspect cohort_size in response"
        ),
        "data": "Cohort calculation.",
        "expected": (
            "1. cohort_size = N - 1 (caller excluded)\n"
            "2. Caller's own portfolio not double-counted in any metric"
        ),
        "criteria": "Pass: Excluded. Fail: Includes self."
    },
    {
        "id": "TC-449", "name": "Risk Heatmap Sectors Use get_sectors() Helper", "category": "Donor Intel Edge",
        "priority": "P3 - Medium", "requirement": "FR-RISK-CALC-001",
        "prereqs": "Donor with grants where sectors JSON has mixed types (strings, dicts).",
        "steps": (
            "1. Fetch heatmap"
        ),
        "data": "Messy JSON.",
        "expected": (
            "1. Service uses get_sectors() helper\n"
            "2. Filters out non-string entries silently\n"
            "3. No 500 error"
        ),
        "criteria": "Pass: Robust parsing. Fail: 500 on dict in list."
    },
    {
        "id": "TC-450", "name": "Donor Broadcast Empty Audience", "category": "Donor Intel Edge",
        "priority": "P2 - High", "requirement": "FR-BROADCAST-EDGE-001",
        "prereqs": "Donor with grant that has zero applications.",
        "steps": (
            "1. Send broadcast"
        ),
        "data": "Empty audience.",
        "expected": (
            "1. Response ok=true, recipient_count=0\n"
            "2. No notifications dispatched\n"
            "3. No 500 error"
        ),
        "criteria": "Pass: Clean zero. Fail: 500."
    },
    {
        "id": "TC-451", "name": "Donor Broadcast Cross-Donor Forbidden", "category": "Donor Intel Edge",
        "priority": "P1 - Critical", "requirement": "FR-BROADCAST-AUTH-001",
        "prereqs": "Two donors; donor A tries to broadcast on donor B's grant.",
        "steps": (
            "1. Log in as donor A\n"
            "2. POST broadcast on donor B's grant"
        ),
        "data": "Cross-donor attempt.",
        "expected": "403 or 404 (grant not visible).",
        "criteria": "Pass: Blocked. Fail: Broadcast sent."
    },
]
test_cases.extend(donor_edge_cases)


# ============================================================================
# CATEGORY 25: PWA + WEBAUTHN NEGATIVES  (TC-460 to TC-474)
# ============================================================================

webauthn_pwa_neg_cases = [
    {
        "id": "TC-460", "name": "Install Banner Hidden On iOS Safari", "category": "PWA Negatives",
        "priority": "P2 - High", "requirement": "FR-PWA-IOS-001",
        "prereqs": "iPhone running Safari (NOT Chrome on iOS).",
        "steps": (
            "1. Open the app in mobile Safari\n"
            "2. Wait 10 seconds"
        ),
        "data": "iOS Safari session.",
        "expected": (
            "1. No install banner appears (beforeinstallprompt not fired)\n"
            "2. App is still fully usable\n"
            "3. Manual 'Add to Home Screen' via Safari share sheet works"
        ),
        "criteria": "Pass: No banner, normal usage. Fail: Banner shows or app breaks."
    },
    {
        "id": "TC-461", "name": "Install Banner Hidden When PWA Already Installed", "category": "PWA Negatives",
        "priority": "P3 - Medium", "requirement": "FR-PWA-INSTALLED-001",
        "prereqs": "PWA installed (display-mode: standalone).",
        "steps": (
            "1. Launch the installed PWA\n"
            "2. Wait 10 seconds"
        ),
        "data": "Standalone session.",
        "expected": (
            "1. Banner does NOT appear\n"
            "2. Check via window.matchMedia('(display-mode: standalone)').matches → true\n"
            "3. localStorage may already have dismiss flag"
        ),
        "criteria": "Pass: No banner. Fail: Banner shown to installed users."
    },
    {
        "id": "TC-462", "name": "appinstalled Event Sets Dismiss Flag", "category": "PWA Negatives",
        "priority": "P3 - Medium", "requirement": "FR-PWA-APPINSTALLED-001",
        "prereqs": "Banner is showing.",
        "steps": (
            "1. Click 'Install'\n"
            "2. Approve browser's install dialog\n"
            "3. PWA installs\n"
            "4. Check localStorage"
        ),
        "data": "Install flow.",
        "expected": (
            "1. appinstalled event fires\n"
            "2. localStorage 'kuja_pwa_install_dismissed_v1' = '1'\n"
            "3. Banner state hidden"
        ),
        "criteria": "Pass: Flag set. Fail: Banner re-shows post-install."
    },
    {
        "id": "TC-463", "name": "localStorage Blocked (Incognito)", "category": "PWA Negatives",
        "priority": "P3 - Medium", "requirement": "FR-PWA-LS-001",
        "prereqs": "Browser in incognito with localStorage disabled.",
        "steps": (
            "1. Open the app\n"
            "2. See install banner\n"
            "3. Click 'Not now'\n"
            "4. Reload"
        ),
        "data": "No localStorage.",
        "expected": (
            "1. No crash on read attempt\n"
            "2. Banner may re-appear (no persistence) — acceptable in incognito\n"
            "3. No console errors thrown to user"
        ),
        "criteria": "Pass: Graceful degradation. Fail: Crash."
    },
    {
        "id": "TC-464", "name": "Native Share When navigator.share Unavailable", "category": "PWA Negatives",
        "priority": "P3 - Medium", "requirement": "FR-SHARE-FALLBACK-001",
        "prereqs": "Desktop browser without navigator.share API (older Chrome).",
        "steps": (
            "1. Click 'Share profile' button"
        ),
        "data": "Desktop browser.",
        "expected": (
            "1. Code detects missing navigator.share\n"
            "2. Falls back to navigator.clipboard.writeText(url)\n"
            "3. Toast 'Link copied' appears\n"
            "4. Button label switches to 'Copied' for ~1.5s"
        ),
        "criteria": "Pass: Clipboard fallback works. Fail: Silent failure."
    },
    {
        "id": "TC-465", "name": "Native Share AbortError When User Cancels", "category": "PWA Negatives",
        "priority": "P3 - Medium", "requirement": "FR-SHARE-CANCEL-001",
        "prereqs": "Mobile browser with native share.",
        "steps": (
            "1. Click 'Share profile'\n"
            "2. When share sheet opens, dismiss it (back gesture)"
        ),
        "data": "Aborted share.",
        "expected": (
            "1. No clipboard fallback (AbortError detected)\n"
            "2. No error toast\n"
            "3. Button returns to idle state"
        ),
        "criteria": "Pass: Clean cancellation. Fail: Error or unwanted fallback."
    },
    {
        "id": "TC-466", "name": "WebAuthn Registration Cancelled By User", "category": "WebAuthn Negatives",
        "priority": "P2 - High", "requirement": "FR-WEBAUTHN-CANCEL-001",
        "prereqs": "User on /settings/security.",
        "steps": (
            "1. Click 'Enrol this device'\n"
            "2. When biometric prompt appears, cancel it"
        ),
        "data": "Cancelled enrolment.",
        "expected": (
            "1. No credential created\n"
            "2. No error toast (AbortError suppressed)\n"
            "3. Trusted devices list unchanged\n"
            "4. Button returns to enabled state"
        ),
        "criteria": "Pass: Silent cancel. Fail: Error or partial enrolment."
    },
    {
        "id": "TC-467", "name": "WebAuthn Register With Same Device Twice", "category": "WebAuthn Negatives",
        "priority": "P2 - High", "requirement": "FR-WEBAUTHN-DUPE-001",
        "prereqs": "User has enrolled device A.",
        "steps": (
            "1. Click 'Enrol this device' again on the same device\n"
            "2. Approve the biometric prompt"
        ),
        "data": "Re-enrolment attempt.",
        "expected": (
            "1. Browser may prevent (exclude_credentials list passed)\n"
            "2. OR new credential id allocated, treated as separate\n"
            "3. Either way, no crash"
        ),
        "criteria": "Pass: Handled. Fail: Crash."
    },
    {
        "id": "TC-468", "name": "WebAuthn Auth Begin With No Credentials", "category": "WebAuthn Negatives",
        "priority": "P2 - High", "requirement": "FR-WEBAUTHN-AUTH-EDGE-001",
        "prereqs": "User without any enrolled credentials.",
        "steps": (
            "1. POST /api/auth/webauthn/authenticate/begin"
        ),
        "data": "No credentials.",
        "expected": (
            "1. Response 400 with reason='no_credentials'\n"
            "2. Client UI hides the 'Use biometric' option"
        ),
        "criteria": "Pass: Honest empty. Fail: 500 or fake challenge."
    },
    {
        "id": "TC-469", "name": "WebAuthn Sign-Count Regression (Clone)", "category": "WebAuthn Negatives",
        "priority": "P1 - Critical", "requirement": "FR-WEBAUTHN-CLONE-001",
        "prereqs": "Manually set the stored sign_count to a higher value than what the device will return.",
        "steps": (
            "1. Manipulate DB: UPDATE webauthn_credentials SET sign_count=999 WHERE id=X\n"
            "2. User authenticates (device returns sign_count=2)\n"
            "3. Server verifies"
        ),
        "data": "Sign-count regression.",
        "expected": (
            "1. Response success=False, reason='sign_count_regression'\n"
            "2. Server logs WARNING with user_id, cred_id, both counts\n"
            "3. User cannot complete biometric auth until credential rotated"
        ),
        "criteria": "Pass: Hard fail. Fail: Authentication proceeds."
    },
    {
        "id": "TC-470", "name": "Re-Auth Token Single-Use", "category": "WebAuthn Negatives",
        "priority": "P1 - Critical", "requirement": "FR-WEBAUTHN-TOKEN-001",
        "prereqs": "User has just completed biometric auth and has a valid reauth_token.",
        "steps": (
            "1. Use the token in X-Reauth-Token header on a gated request\n"
            "2. Use the same token again on a second request"
        ),
        "data": "Token reuse.",
        "expected": (
            "1. First request: gate passes\n"
            "2. Second request: 403 reauth_invalid_or_expired"
        ),
        "criteria": "Pass: Token consumed once. Fail: Reusable token."
    },
    {
        "id": "TC-471", "name": "Re-Auth Token Expiration (5 min)", "category": "WebAuthn Negatives",
        "priority": "P2 - High", "requirement": "FR-WEBAUTHN-TOKEN-002",
        "prereqs": "Fresh reauth_token issued.",
        "steps": (
            "1. Wait 6 minutes\n"
            "2. Use the token"
        ),
        "data": "Expired token.",
        "expected": "403 reauth_invalid_or_expired.",
        "criteria": "Pass: Expires. Fail: Token never expires."
    },
    {
        "id": "TC-472", "name": "Re-Auth Token Cross-User Attempt", "category": "WebAuthn Negatives",
        "priority": "P1 - Critical", "requirement": "FR-WEBAUTHN-TOKEN-003",
        "prereqs": "User A has a valid token; user B is logged in.",
        "steps": (
            "1. User B sends a gated request with user A's X-Reauth-Token"
        ),
        "data": "Cross-user token reuse.",
        "expected": "403 — token bound to user_id at issue.",
        "criteria": "Pass: Blocked. Fail: Cross-user bypass."
    },
    {
        "id": "TC-473", "name": "require_reauth() Bypass For Unenrolled User", "category": "WebAuthn Negatives",
        "priority": "P2 - High", "requirement": "FR-WEBAUTHN-OPTIN-001",
        "prereqs": "User logged in with NO WebAuthn credentials.",
        "steps": (
            "1. Send a gated request WITHOUT X-Reauth-Token header"
        ),
        "data": "Unenrolled user.",
        "expected": (
            "1. Helper returns None (treated as pass)\n"
            "2. Request proceeds normally\n"
            "3. Logged in unenrolled users are not blocked from app features"
        ),
        "criteria": "Pass: Opt-in only. Fail: Locks out unenrolled."
    },
    {
        "id": "TC-474", "name": "Revoke Non-Existent Credential", "category": "WebAuthn Negatives",
        "priority": "P3 - Medium", "requirement": "FR-WEBAUTHN-REVOKE-001",
        "prereqs": "User logged in.",
        "steps": (
            "1. DELETE /api/auth/webauthn/credentials/9999999"
        ),
        "data": "Bad credential id.",
        "expected": "404 not_found",
        "criteria": "Pass: 404. Fail: 500 or 200."
    },
]
test_cases.extend(webauthn_pwa_neg_cases)


# ============================================================================
# CATEGORY 26: METRICS EVENT RECORDING + INGEST  (TC-480 to TC-499)
# ============================================================================

metrics_event_cases = [
    {
        "id": "TC-480", "name": "Event Recording Is Non-Blocking", "category": "Metrics Events",
        "priority": "P1 - Critical", "requirement": "FR-METRICS-NONBLOCK-001",
        "prereqs": "Test env where the user_events table is artificially broken (e.g. revoke INSERT permission).",
        "steps": (
            "1. Log in (would normally record session.start)\n"
            "2. Verify login still succeeds"
        ),
        "data": "Broken event storage.",
        "expected": (
            "1. Login completes with 200\n"
            "2. Event recording exception is logged but does NOT raise\n"
            "3. User session is established"
        ),
        "criteria": "Pass: Best-effort. Fail: Event failure breaks login."
    },
    {
        "id": "TC-481", "name": "session.start Event On Every Login", "category": "Metrics Events",
        "priority": "P1 - Critical", "requirement": "FR-METRICS-EVT-001",
        "prereqs": "Admin can view metrics.",
        "steps": (
            "1. Note session.start count\n"
            "2. Log in 3 times (with logouts in between)\n"
            "3. Refresh metrics"
        ),
        "data": "Multiple logins.",
        "expected": "session.start count incremented by 3.",
        "criteria": "Pass: 3 increments. Fail: Missing events."
    },
    {
        "id": "TC-482", "name": "application.start_draft On POST /applications/", "category": "Metrics Events",
        "priority": "P2 - High", "requirement": "FR-METRICS-EVT-002",
        "prereqs": "NGO login.",
        "steps": (
            "1. Create a new application via POST /api/applications/\n"
            "2. Admin checks event_counts_30d"
        ),
        "data": "Draft creation.",
        "expected": "application.start_draft count > 0 with grant_id in props.",
        "criteria": "Pass: Captured. Fail: Missing."
    },
    {
        "id": "TC-483", "name": "application.submit On Submit", "category": "Metrics Events",
        "priority": "P1 - Critical", "requirement": "FR-METRICS-EVT-003",
        "prereqs": "NGO with draft application.",
        "steps": (
            "1. POST /api/applications/<id>/submit\n"
            "2. Check event_counts"
        ),
        "data": "Submit.",
        "expected": "application.submit count incremented with application_id, grant_id, ai_score props.",
        "criteria": "Pass: Captured. Fail: Missing."
    },
    {
        "id": "TC-484", "name": "report.preflight_used On Precheck", "category": "Metrics Events",
        "priority": "P2 - High", "requirement": "FR-METRICS-EVT-004",
        "prereqs": "NGO with draft report.",
        "steps": (
            "1. POST /api/reports/<id>/precheck\n"
            "2. Check event_counts"
        ),
        "data": "Precheck call.",
        "expected": "report.preflight_used count incremented with compliance_score in props.",
        "criteria": "Pass: Captured. Fail: Missing."
    },
    {
        "id": "TC-485", "name": "readiness_check.used On Submission Readiness", "category": "Metrics Events",
        "priority": "P2 - High", "requirement": "FR-METRICS-EVT-005",
        "prereqs": "NGO with application.",
        "steps": (
            "1. POST /api/ai/submission-readiness with application_id\n"
            "2. Check event_counts"
        ),
        "data": "Readiness check.",
        "expected": "readiness_check.used recorded with application_id BEFORE the async wrap (so it persists even if AI fails).",
        "criteria": "Pass: Captured pre-async. Fail: Lost on AI failure."
    },
    {
        "id": "TC-486", "name": "reviewer.assignment_opened Only For Reviewer Role", "category": "Metrics Events",
        "priority": "P2 - High", "requirement": "FR-METRICS-EVT-006",
        "prereqs": "Two users: a reviewer + a donor who can inspect reviews.",
        "steps": (
            "1. Donor opens GET /api/reviews/<id>\n"
            "2. Reviewer opens GET /api/reviews/<id>\n"
            "3. Compare events"
        ),
        "data": "Two role views.",
        "expected": (
            "1. Donor view does NOT record reviewer.assignment_opened\n"
            "2. Reviewer view DOES record the event"
        ),
        "criteria": "Pass: Role-scoped. Fail: Recorded for wrong role."
    },
    {
        "id": "TC-487", "name": "reviewer.review_submitted On Complete", "category": "Metrics Events",
        "priority": "P2 - High", "requirement": "FR-METRICS-EVT-007",
        "prereqs": "Reviewer with assigned review + scores entered.",
        "steps": (
            "1. POST /api/reviews/<id>/complete"
        ),
        "data": "Complete action.",
        "expected": "reviewer.review_submitted event with review_id, application_id, score.",
        "criteria": "Pass: Captured. Fail: Missing."
    },
    {
        "id": "TC-488", "name": "trust_profile.viewed Scopes Correctly", "category": "Metrics Events",
        "priority": "P3 - Medium", "requirement": "FR-METRICS-EVT-008",
        "prereqs": "NGO + donor + reviewer + admin sessions.",
        "steps": (
            "1. NGO views own org (should record own_org=true)\n"
            "2. NGO views another org (should NOT record)\n"
            "3. Donor/reviewer/admin view any org (should record own_org=false)"
        ),
        "data": "Cross-role views.",
        "expected": "Event count matches the access-scope rule.",
        "criteria": "Pass: Matches rule. Fail: Recorded for cross-org NGO."
    },
    {
        "id": "TC-489", "name": "donor.broadcast_sent Only On Success", "category": "Metrics Events",
        "priority": "P3 - Medium", "requirement": "FR-METRICS-EVT-009",
        "prereqs": "Donor login.",
        "steps": (
            "1. POST broadcast with invalid body (no subject)\n"
            "2. POST broadcast with valid body"
        ),
        "data": "Failure + success.",
        "expected": (
            "1. Invalid: no event recorded\n"
            "2. Valid: event recorded with grant_id, audience, recipient count"
        ),
        "criteria": "Pass: Only success counted. Fail: Failures counted too."
    },
    {
        "id": "TC-490", "name": "search.query Records With Length + Count", "category": "Metrics Events",
        "priority": "P3 - Medium", "requirement": "FR-METRICS-EVT-010",
        "prereqs": "Any logged-in user.",
        "steps": (
            "1. GET /api/search?q=kenya"
        ),
        "data": "Search call.",
        "expected": (
            "1. Event search.query recorded\n"
            "2. props.query_length=5\n"
            "3. props.result_count = actual hits"
        ),
        "criteria": "Pass: Captured. Fail: Missing or wrong props."
    },
    {
        "id": "TC-491", "name": "chat.thread_open Idempotent Per Visit", "category": "Metrics Events",
        "priority": "P3 - Medium", "requirement": "FR-METRICS-EVT-011",
        "prereqs": "Logged in user.",
        "steps": (
            "1. POST /api/ai/threads/open three times in a row"
        ),
        "data": "Repeat opens.",
        "expected": (
            "1. Each call records chat.thread_open\n"
            "2. (Not deduplicated server-side; the 'open' event represents intent each time)"
        ),
        "criteria": "Pass: 3 events. Fail: Only 1."
    },
    {
        "id": "TC-492", "name": "chat.message_sent Records With Length", "category": "Metrics Events",
        "priority": "P3 - Medium", "requirement": "FR-METRICS-EVT-012",
        "prereqs": "Chat thread open.",
        "steps": (
            "1. Send a 50-character message\n"
            "2. Check event"
        ),
        "data": "Message.",
        "expected": "chat.message_sent recorded with message_length=50.",
        "criteria": "Pass: Captured. Fail: Missing length."
    },
    {
        "id": "TC-493", "name": "ab_arm() Deterministic Across Calls", "category": "Metrics Events",
        "priority": "P2 - High", "requirement": "FR-AB-DETERMIN-001",
        "prereqs": "Python shell or test.",
        "steps": (
            "1. Call ab_arm('exp_x', org_id=42) ten times\n"
            "2. Call with different org_id"
        ),
        "data": "Bucketing.",
        "expected": (
            "1. Same arm returned every call for org=42\n"
            "2. Different org may land in different arm (over many trials ~50/50)"
        ),
        "criteria": "Pass: Stable. Fail: Random per call."
    },
    {
        "id": "TC-494", "name": "ab_arm() Returns None Without Subject", "category": "Metrics Events",
        "priority": "P3 - Medium", "requirement": "FR-AB-EDGE-001",
        "prereqs": "Caller without org_id and user_id.",
        "steps": (
            "1. ab_arm('exp_x')"
        ),
        "data": "No subject.",
        "expected": "Returns None.",
        "criteria": "Pass: None. Fail: Random arm."
    },
    {
        "id": "TC-495", "name": "ab_arm() Custom Arms Tuple", "category": "Metrics Events",
        "priority": "P3 - Medium", "requirement": "FR-AB-CUSTOM-001",
        "prereqs": "Python shell.",
        "steps": (
            "1. ab_arm('exp', org_id=1, arms=('control','variant','holdout'))"
        ),
        "data": "3-way split.",
        "expected": "Returns one of the three; stable per org_id.",
        "criteria": "Pass: 3-way works. Fail: Errors or wrong arm."
    },
    {
        "id": "TC-496", "name": "Generic Event Ingest Empty event_name", "category": "Metrics Events",
        "priority": "P3 - Medium", "requirement": "FR-INGEST-VAL-001",
        "prereqs": "Logged in user.",
        "steps": (
            "1. POST /api/ai/events/track with empty event_name"
        ),
        "data": "Empty.",
        "expected": "400 'event_name required'",
        "criteria": "Pass: 400. Fail: 500 or silent accept."
    },
    {
        "id": "TC-497", "name": "Generic Event Ingest Non-Whitelisted Name", "category": "Metrics Events",
        "priority": "P2 - High", "requirement": "FR-INGEST-VAL-002",
        "prereqs": "Logged in user.",
        "steps": (
            "1. POST event_name='attack.exfiltrate'"
        ),
        "data": "Attack name.",
        "expected": "400 'event_name not allowed' + 'allowed' list in body.",
        "criteria": "Pass: Whitelist enforced. Fail: Arbitrary names accepted."
    },
    {
        "id": "TC-498", "name": "Generic Event Ingest Non-Dict Props", "category": "Metrics Events",
        "priority": "P3 - Medium", "requirement": "FR-INGEST-VAL-003",
        "prereqs": "Logged in user.",
        "steps": (
            "1. POST props=['array','not','dict']"
        ),
        "data": "Wrong type.",
        "expected": "400 'props must be an object'",
        "criteria": "Pass: 400. Fail: Crash."
    },
    {
        "id": "TC-499", "name": "Generic Event Ingest Oversized Props", "category": "Metrics Events",
        "priority": "P3 - Medium", "requirement": "FR-INGEST-CAP-001",
        "prereqs": "Logged in user.",
        "steps": (
            "1. POST with props containing a 10KB string value"
        ),
        "data": "Oversized payload.",
        "expected": (
            "1. Service replaces props with {'truncated': true}\n"
            "2. 200 success\n"
            "3. Underlying event still recorded with truncated marker"
        ),
        "criteria": "Pass: Truncated. Fail: 500 or unbounded storage."
    },
]
test_cases.extend(metrics_event_cases)


# ============================================================================
# CATEGORY 27: NPS FEEDBACK EDGE CASES  (TC-510 to TC-519)
# ============================================================================

feedback_cases = [
    {
        "id": "TC-510", "name": "Feedback Uniqueness Per Target", "category": "NPS Feedback",
        "priority": "P2 - High", "requirement": "FR-FEEDBACK-UNIQ-001",
        "prereqs": "User has submitted feedback for application=123.",
        "steps": (
            "1. POST /api/feedback with same surface + related_kind + related_id\n"
            "2. Check returned 'created' field"
        ),
        "data": "Re-submission.",
        "expected": (
            "1. Service upserts: returns created=false, feedback.id same as before\n"
            "2. Score updated to the new value\n"
            "3. Only one row exists in DB (uniqueness constraint enforced)"
        ),
        "criteria": "Pass: Upserts. Fail: Duplicate row or 500."
    },
    {
        "id": "TC-511", "name": "Feedback Comment Cap At 500 Chars", "category": "NPS Feedback",
        "priority": "P3 - Medium", "requirement": "FR-FEEDBACK-CAP-001",
        "prereqs": "Logged in user.",
        "steps": (
            "1. POST feedback with comment of 1000 chars"
        ),
        "data": "Long comment.",
        "expected": (
            "1. Server stores only the first 500 chars\n"
            "2. Returned feedback.comment is truncated\n"
            "3. No 400 error"
        ),
        "criteria": "Pass: Cap enforced. Fail: Full stored or 500."
    },
    {
        "id": "TC-512", "name": "NPS Calculation Correctness", "category": "NPS Feedback",
        "priority": "P2 - High", "requirement": "FR-NPS-CALC-001",
        "prereqs": "Submit known mix: 4 promoters (9-10), 3 passives (7-8), 2 detractors (0-6).",
        "steps": (
            "1. Hit /api/admin/metrics as admin\n"
            "2. Inspect nps.overall_nps"
        ),
        "data": "9 responses.",
        "expected": (
            "1. overall_nps = ((4-2)/9)*100 ≈ 22.2\n"
            "2. promoters=4, passives=3, detractors=2"
        ),
        "criteria": "Pass: Correct math. Fail: Wrong calc."
    },
    {
        "id": "TC-513", "name": "NPS Empty State", "category": "NPS Feedback",
        "priority": "P3 - Medium", "requirement": "FR-NPS-EMPTY-001",
        "prereqs": "Fresh DB with no UserFeedback rows.",
        "steps": (
            "1. GET /api/admin/metrics"
        ),
        "data": "No responses.",
        "expected": (
            "1. nps.total_responses=0\n"
            "2. nps.overall_nps=null\n"
            "3. by_surface=[], by_language={}, histogram all-zeros"
        ),
        "criteria": "Pass: Empty shape correct. Fail: Crash or fake numbers."
    },
    {
        "id": "TC-514", "name": "NPS By-Language Rollup", "category": "NPS Feedback",
        "priority": "P3 - Medium", "requirement": "FR-NPS-LANG-001",
        "prereqs": "Responses from users in en, ar, sw, so.",
        "steps": (
            "1. GET admin metrics\n"
            "2. Inspect nps.by_language"
        ),
        "data": "Multi-language responses.",
        "expected": (
            "1. by_language has keys for each language with responses\n"
            "2. Each entry: {responses, nps}\n"
            "3. Per-language NPS calculated independently"
        ),
        "criteria": "Pass: Per-language rollup. Fail: Aggregated only."
    },
    {
        "id": "TC-515", "name": "Feedback Recent Comments Limit", "category": "NPS Feedback",
        "priority": "P3 - Medium", "requirement": "FR-NPS-RECENT-001",
        "prereqs": "30+ feedback rows with comments.",
        "steps": (
            "1. GET admin metrics"
        ),
        "data": "Many comments.",
        "expected": "nps_recent_comments has exactly 10 most-recent entries.",
        "criteria": "Pass: Capped at 10. Fail: Unbounded."
    },
    {
        "id": "TC-516", "name": "Feedback Surface Not In Allowed List", "category": "NPS Feedback",
        "priority": "P2 - High", "requirement": "FR-FEEDBACK-VAL-001",
        "prereqs": "Logged in user.",
        "steps": (
            "1. POST /api/feedback {surface: 'made_up_surface', score: 5}"
        ),
        "data": "Bad surface.",
        "expected": "400 with reason='surface_not_allowed' + allowed list.",
        "criteria": "Pass: Whitelist enforced. Fail: Accepted."
    },
    {
        "id": "TC-517", "name": "Feedback Non-Int Score", "category": "NPS Feedback",
        "priority": "P3 - Medium", "requirement": "FR-FEEDBACK-VAL-002",
        "prereqs": "Logged in user.",
        "steps": (
            "1. POST {surface:'application_submit', score:'nine'}"
        ),
        "data": "String score.",
        "expected": "400 with reason='score_must_be_int'",
        "criteria": "Pass: 400. Fail: 500 or silent coerce."
    },
    {
        "id": "TC-518", "name": "Feedback Unauthenticated", "category": "NPS Feedback",
        "priority": "P1 - Critical", "requirement": "FR-FEEDBACK-AUTH-001",
        "prereqs": "Not logged in.",
        "steps": (
            "1. POST /api/feedback {surface:'application_submit', score:5}"
        ),
        "data": "No session.",
        "expected": "401 (login_required redirects to login)",
        "criteria": "Pass: Auth required. Fail: Anonymous submit allowed."
    },
    {
        "id": "TC-519", "name": "/api/feedback/my Returns Only Caller's", "category": "NPS Feedback",
        "priority": "P1 - Critical", "requirement": "FR-FEEDBACK-AUTH-002",
        "prereqs": "Two users with submitted feedback.",
        "steps": (
            "1. Log in as user A\n"
            "2. GET /api/feedback/my\n"
            "3. Log in as user B\n"
            "4. GET /api/feedback/my"
        ),
        "data": "Two sessions.",
        "expected": "Each user only sees their own feedback rows.",
        "criteria": "Pass: Strict isolation. Fail: Cross-user leak."
    },
]
test_cases.extend(feedback_cases)


# ============================================================================
# CATEGORY 28: COMPREHENSIVE END-TO-END WORKFLOWS  (TC-530 to TC-535)
# ============================================================================

e2e_phase_cases = [
    {
        "id": "TC-530", "name": "Full NGO Journey: Register → Submit → Survey", "category": "E2E Workflows",
        "priority": "P1 - Critical", "requirement": "E2E-NGO-001",
        "prereqs": "Fresh NGO account.",
        "steps": (
            "1. Log in as NGO\n"
            "2. Visit /dashboard, complete onboarding checklist\n"
            "3. Visit /assessments, complete a capacity assessment\n"
            "4. Visit /grants, find a matching grant\n"
            "5. Click Apply, create draft (event: application.start_draft)\n"
            "6. Fill responses\n"
            "7. Run readiness check (event: readiness_check.used)\n"
            "8. Apply 1 AI suggestion (event: ai_assist.suggestion_accepted via ingest)\n"
            "9. Submit (events: application.submit + auto-assign fires)\n"
            "10. Open chat on the app, ask a follow-up (event: chat.message_sent)\n"
            "11. Wait for micro-survey, submit score 9"
        ),
        "data": "Full NGO flow.",
        "expected": (
            "1. All events recorded in /admin/metrics\n"
            "2. Reviewer panel auto-populated after submit\n"
            "3. NPS bumped by 1 with score 9\n"
            "4. Audit chain has 'reviewer.auto_assigned' entry"
        ),
        "criteria": "Pass: All events visible + state correct. Fail: Any missing."
    },
    {
        "id": "TC-531", "name": "Full Donor Journey: Publish → Decide → Cohort", "category": "E2E Workflows",
        "priority": "P1 - Critical", "requirement": "E2E-DONOR-001",
        "prereqs": "Fresh donor account.",
        "steps": (
            "1. Log in as donor\n"
            "2. Visit /grants/new, create grant\n"
            "3. Publish\n"
            "4. Wait for applicants\n"
            "5. Send broadcast to all applicants (event: donor.broadcast_sent)\n"
            "6. Review apps with auto-assigned panels\n"
            "7. Record decision with reason code (event: donor.decision_recorded)\n"
            "8. Visit /dashboard\n"
            "9. Inspect cohort card\n"
            "10. Inspect risk heatmap"
        ),
        "data": "Full donor flow.",
        "expected": (
            "1. All events captured\n"
            "2. Cohort sparse-honest if <3 other donors\n"
            "3. Risk heatmap renders with cells"
        ),
        "criteria": "Pass: Full flow + visible metrics. Fail: Any break."
    },
    {
        "id": "TC-532", "name": "Full Reviewer Journey: Open → Score → Submit", "category": "E2E Workflows",
        "priority": "P1 - Critical", "requirement": "E2E-REV-001",
        "prereqs": "Reviewer with auto-assigned reviews.",
        "steps": (
            "1. Log in as reviewer\n"
            "2. Visit /dashboard, see throughput card\n"
            "3. Visit /reviews\n"
            "4. Open an assigned review (event: reviewer.assignment_opened)\n"
            "5. Read briefing card\n"
            "6. Enter per-criterion scores\n"
            "7. POST complete (event: reviewer.review_submitted)\n"
            "8. Return to /dashboard, verify throughput updated"
        ),
        "data": "Full reviewer flow.",
        "expected": (
            "1. assignment_opened fires only for reviewer role\n"
            "2. review_submitted fires on POST complete\n"
            "3. Application's human_score recalculated"
        ),
        "criteria": "Pass: Both events + state. Fail: Either missing."
    },
    {
        "id": "TC-533", "name": "Full Report Cycle: Draft → Preflight → Submit → Decision", "category": "E2E Workflows",
        "priority": "P2 - High", "requirement": "E2E-REPORT-001",
        "prereqs": "NGO with awarded application.",
        "steps": (
            "1. Create report draft (event: report.start_draft)\n"
            "2. Run pre-flight (event: report.preflight_used)\n"
            "3. Fix flagged issues\n"
            "4. Submit (event: report.submit)\n"
            "5. Wait for AI compliance score\n"
            "6. Micro-survey appears, submit score 8\n"
            "7. Donor logs in, reviews report, marks accepted\n"
            "8. NGO checks /reports, sees status updated"
        ),
        "data": "Report lifecycle.",
        "expected": (
            "1. Full funnel report.start_draft → report.submit captured\n"
            "2. Preflight event also captured\n"
            "3. NPS for report_submit incremented\n"
            "4. Report status transitions correctly"
        ),
        "criteria": "Pass: Complete cycle + events. Fail: Any drop."
    },
    {
        "id": "TC-534", "name": "Cross-Role A/B Experiment Walkthrough", "category": "E2E Workflows",
        "priority": "P2 - High", "requirement": "E2E-AB-001",
        "prereqs": "Developer adds ab_arm() call at app.submit site.",
        "steps": (
            "1. Two NGO orgs in different ab arms submit applications\n"
            "2. Wait for event recording\n"
            "3. Admin checks /admin/metrics → ab_application_submit by_arm"
        ),
        "data": "A/B run.",
        "expected": (
            "1. by_arm shows non-empty buckets (A + B)\n"
            "2. Each org's submits land in their assigned arm\n"
            "3. Same org always lands in same arm (deterministic)"
        ),
        "criteria": "Pass: Bucketing visible. Fail: Single bucket or random."
    },
    {
        "id": "TC-535", "name": "Multi-User Concurrent Session Stress", "category": "E2E Workflows",
        "priority": "P3 - Medium", "requirement": "E2E-STRESS-001",
        "prereqs": "Test runner that can fire 10 parallel sessions.",
        "steps": (
            "1. 10 parallel logins (different users)\n"
            "2. Each user submits an application\n"
            "3. Each opens chat + sends 3 messages"
        ),
        "data": "10×5 = 50 events.",
        "expected": (
            "1. All 10 session.start events recorded\n"
            "2. All 10 application.submit events recorded\n"
            "3. ~30 chat.message_sent events (allow some rate-limited if AI quota exhausted)\n"
            "4. No 500 errors\n"
            "5. Rate-limited requests return 429 cleanly"
        ),
        "criteria": "Pass: Clean concurrent execution. Fail: 500s or lost events."
    },
]
test_cases.extend(e2e_phase_cases)


# ============================================================================
# CATEGORY 29: I18N + RTL EDGE CASES FOR NEW COMPONENTS  (TC-550 to TC-557)
# ============================================================================

i18n_new_cases = [
    {
        "id": "TC-550", "name": "Chat Panel Renders Right-To-Left In Arabic", "category": "I18N RTL",
        "priority": "P2 - High", "requirement": "FR-I18N-CHAT-001",
        "prereqs": "User with language='ar'.",
        "steps": (
            "1. Open /chat\n"
            "2. Inspect <html dir> attribute\n"
            "3. Send + receive messages\n"
            "4. Check message bubble alignment"
        ),
        "data": "Arabic session.",
        "expected": (
            "1. html dir='rtl'\n"
            "2. User bubbles right-aligned\n"
            "3. Assistant bubbles right-aligned for Arabic text\n"
            "4. Composer reads right-to-left"
        ),
        "criteria": "Pass: Clean RTL. Fail: Mirrored or LTR-only layout."
    },
    {
        "id": "TC-551", "name": "Micro-Survey In Arabic", "category": "I18N RTL",
        "priority": "P3 - Medium", "requirement": "FR-I18N-SURVEY-001",
        "prereqs": "Arabic-language user submitting an application.",
        "steps": (
            "1. After submit, wait for survey\n"
            "2. Inspect question text and layout"
        ),
        "data": "Arabic survey.",
        "expected": (
            "1. Question text appears in Arabic (or English fallback if not translated)\n"
            "2. Layout reads RTL (button positions mirrored)\n"
            "3. Score buttons 0-10 still in their natural order"
        ),
        "criteria": "Pass: RTL applied. Fail: LTR forced."
    },
    {
        "id": "TC-552", "name": "PWA Install Banner Translated To Swahili", "category": "I18N RTL",
        "priority": "P3 - Medium", "requirement": "FR-I18N-PWA-001",
        "prereqs": "Swahili user, supported browser.",
        "steps": (
            "1. Trigger banner via beforeinstallprompt\n"
            "2. Inspect text"
        ),
        "data": "Swahili banner.",
        "expected": "Title + body + button labels in Swahili (or graceful EN fallback).",
        "criteria": "Pass: Translated. Fail: English only."
    },
    {
        "id": "TC-553", "name": "Cohort Card Verdict Pills Translated", "category": "I18N RTL",
        "priority": "P3 - Medium", "requirement": "FR-I18N-COHORT-001",
        "prereqs": "Donor with language='fr'.",
        "steps": (
            "1. Visit /dashboard, locate cohort card"
        ),
        "data": "French session.",
        "expected": "Verdict pills (above/around/below) shown in French.",
        "criteria": "Pass: Translated. Fail: English only."
    },
    {
        "id": "TC-554", "name": "Mixed Number + Arabic Text Wrapping", "category": "I18N RTL",
        "priority": "P3 - Medium", "requirement": "FR-I18N-MIXED-001",
        "prereqs": "Arabic UI rendering a cohort row with metric value '82%'.",
        "steps": (
            "1. View the row"
        ),
        "data": "Mixed content.",
        "expected": (
            "1. Number renders as '82%' (LTR within RTL row) — correct Unicode bidi\n"
            "2. No '%82' reversal\n"
            "3. No layout shift between numeric + text segments"
        ),
        "criteria": "Pass: Correct bidi. Fail: Reversed."
    },
    {
        "id": "TC-555", "name": "Long Label Truncation In Somali", "category": "I18N RTL",
        "priority": "P3 - Medium", "requirement": "FR-I18N-LABEL-001",
        "prereqs": "Somali user, sidebar visible.",
        "steps": (
            "1. Inspect long nav labels in Somali"
        ),
        "data": "Long words.",
        "expected": "Labels truncate with ellipsis or wrap cleanly; no horizontal scroll.",
        "criteria": "Pass: Clean overflow. Fail: Scroll or overlap."
    },
    {
        "id": "TC-556", "name": "Date Formatting Localized", "category": "I18N RTL",
        "priority": "P3 - Medium", "requirement": "FR-I18N-DATE-001",
        "prereqs": "Donor in en vs ar vs sw.",
        "steps": (
            "1. View deadline shown on a grant card across all three locales"
        ),
        "data": "Same date.",
        "expected": (
            "1. EN: 'Jan 15, 2026'\n"
            "2. AR: '١٥ يناير ٢٠٢٦' (Arabic numerals + Arabic month) — or graceful fallback to en-US for missing locale\n"
            "3. SW: '15 Jan 2026'"
        ),
        "criteria": "Pass: Localized. Fail: Same format everywhere."
    },
    {
        "id": "TC-557", "name": "Switching Language Mid-Session", "category": "I18N RTL",
        "priority": "P2 - High", "requirement": "FR-I18N-SWITCH-001",
        "prereqs": "Logged-in user.",
        "steps": (
            "1. PUT /api/auth/language with new lang code\n"
            "2. Refresh /dashboard"
        ),
        "data": "Mid-session switch.",
        "expected": (
            "1. UI re-renders in the new locale\n"
            "2. html lang + dir attributes update\n"
            "3. Existing chat threads preserved but new replies arrive in new lang"
        ),
        "criteria": "Pass: Clean switch. Fail: Stale strings."
    },
]
test_cases.extend(i18n_new_cases)


# ============================================================================
# CATEGORY 30: PERFORMANCE + CONCURRENCY EDGE CASES  (TC-570 to TC-579)
# ============================================================================

perf_cases = [
    {
        "id": "TC-570", "name": "Rate Limit /api/ai/ Caps At 40/min", "category": "Performance",
        "priority": "P2 - High", "requirement": "FR-PERF-RATE-001",
        "prereqs": "Single IP / single user.",
        "steps": (
            "1. Fire 50 POST requests to /api/ai/insight-narrate?async=true within 60s\n"
            "2. Count 200 vs 429 responses"
        ),
        "data": "Sustained burst.",
        "expected": (
            "1. First ~40 succeed (status 202)\n"
            "2. Remaining return 429 with Retry-After header\n"
            "3. Window resets after 60s; subsequent burst allowed"
        ),
        "criteria": "Pass: 40/min cap honored. Fail: Unbounded or wrong cap."
    },
    {
        "id": "TC-571", "name": "/api/ai/jobs/ Status Polls NOT Rate-Limited", "category": "Performance",
        "priority": "P1 - Critical", "requirement": "FR-PERF-RATE-002",
        "prereqs": "Single IP.",
        "steps": (
            "1. Fire 30 GET requests to /api/ai/jobs/anyid in 30s"
        ),
        "data": "Poll burst.",
        "expected": "0/30 return 429 (cap is 600/min on this specific path).",
        "criteria": "Pass: 0 limited. Fail: Any 429."
    },
    {
        "id": "TC-572", "name": "Login Rate Limit Per IP", "category": "Performance",
        "priority": "P1 - Critical", "requirement": "FR-PERF-LOGIN-001",
        "prereqs": "Single IP.",
        "steps": (
            "1. POST /api/auth/login 31 times within 60s"
        ),
        "data": "Login burst.",
        "expected": "31st request returns 429 (default 30/min/IP).",
        "criteria": "Pass: Capped. Fail: Unbounded."
    },
    {
        "id": "TC-573", "name": "Email Lockout After 5 Failed Logins", "category": "Performance",
        "priority": "P1 - Critical", "requirement": "FR-PERF-LOCKOUT-001",
        "prereqs": "Test email account.",
        "steps": (
            "1. POST 5 logins with wrong password for the same email\n"
            "2. POST a 6th login with the correct password"
        ),
        "data": "Brute force pattern.",
        "expected": (
            "1. 6th attempt returns 429 with 'Email lockout' message\n"
            "2. Lockout persists for LOCKOUT_DURATION_MINUTES (default 15)\n"
            "3. Admin can clear via /api/admin/clear-lockout"
        ),
        "criteria": "Pass: Lockout enforced. Fail: Brute-forceable."
    },
    {
        "id": "TC-574", "name": "Concurrent Application Submits — Same User", "category": "Performance",
        "priority": "P2 - High", "requirement": "FR-PERF-CONCURRENT-001",
        "prereqs": "NGO with 1 draft application.",
        "steps": (
            "1. Fire two simultaneous POST /api/applications/<id>/submit"
        ),
        "data": "Race condition.",
        "expected": (
            "1. One succeeds (status flips to 'submitted')\n"
            "2. Second sees idempotent path: returns success with 'already submitted' message\n"
            "3. No double auto-assign\n"
            "4. No double scoring"
        ),
        "criteria": "Pass: Idempotent. Fail: Double work."
    },
    {
        "id": "TC-575", "name": "Slow AI Response Doesn't Block Other Endpoints", "category": "Performance",
        "priority": "P2 - High", "requirement": "FR-PERF-ASYNC-001",
        "prereqs": "Trigger a slow AI call (e.g. donor-portfolio-insights cold start).",
        "steps": (
            "1. Fire the slow AI request (expect 5-15s)\n"
            "2. Concurrently fire /api/health and /api/version"
        ),
        "data": "Long + short requests.",
        "expected": (
            "1. /api/health returns 200 within 1 second\n"
            "2. /api/version returns 200 within 1 second\n"
            "3. AI call returns when done (5-15s) but doesn't wedge other workers"
        ),
        "criteria": "Pass: Non-blocking. Fail: All endpoints hang."
    },
    {
        "id": "TC-576", "name": "Heavy AI Call Capped At 300s", "category": "Performance",
        "priority": "P2 - High", "requirement": "FR-PERF-AICAP-001",
        "prereqs": "Trigger an AI call (e.g. analyze_report on a large document).",
        "steps": (
            "1. Fire the call\n"
            "2. Time the response"
        ),
        "data": "Large input.",
        "expected": (
            "1. Response within ~300s (max runtime cap)\n"
            "2. Either success or graceful timeout, NOT silent hang\n"
            "3. Job classification (light/medium/heavy) honored"
        ),
        "criteria": "Pass: Bounded. Fail: Unbounded hang."
    },
    {
        "id": "TC-577", "name": "Cron Endpoint Authentication Methods", "category": "Performance",
        "priority": "P2 - High", "requirement": "FR-PERF-CRON-001",
        "prereqs": "CRON_SECRET set.",
        "steps": (
            "1. POST /api/cron/reviewer-auto-assign-sweep with Bearer CRON_SECRET\n"
            "2. POST same URL with admin session cookie\n"
            "3. POST with no auth"
        ),
        "data": "Three auth modes.",
        "expected": "Both #1 and #2 → 200. #3 → 403.",
        "criteria": "Pass: Both auth modes work. Fail: Either blocked."
    },
    {
        "id": "TC-578", "name": "AI Job Polling Backoff", "category": "Performance",
        "priority": "P3 - Medium", "requirement": "FR-PERF-POLL-001",
        "prereqs": "Frontend issues an async AI call.",
        "steps": (
            "1. Inspect network tab during a copilot-rail load\n"
            "2. Time intervals between poll requests"
        ),
        "data": "Poll cadence.",
        "expected": (
            "1. Intervals follow exponential backoff (250ms, 500ms, 1s, 1.5s, 2s, cap)\n"
            "2. Total polls capped at 30 attempts (~50s total)\n"
            "3. After cap, error 'AI job exceeded poll budget' surfaces"
        ),
        "criteria": "Pass: Backoff working. Fail: Tight polling or no cap."
    },
    {
        "id": "TC-579", "name": "Audit Chain Hash Integrity", "category": "Performance",
        "priority": "P1 - Critical", "requirement": "FR-PERF-AUDIT-001",
        "prereqs": "At least 20 audit chain entries in the DB.",
        "steps": (
            "1. GET /api/audit-chain/verify\n"
            "2. Inspect response"
        ),
        "data": "Verify request.",
        "expected": (
            "1. Response success=true\n"
            "2. All entries verified: each row's previous_hash matches the prior row's computed hash\n"
            "3. Any tampering would surface as a specific row's mismatch"
        ),
        "criteria": "Pass: Chain intact. Fail: Hash mismatch or 500."
    },
]
test_cases.extend(perf_cases)


# ============================================================================
# CATEGORY 31: CORE SYSTEM EDGE CASES + SECURITY  (TC-600 to TC-619)
# Fills coverage gaps in older categories: documents, auth, audit chain,
# compliance, risk, search, and notifications.
# ============================================================================

core_edge_cases = [
    {
        "id": "TC-600", "name": "Document Upload Size Limit Enforced", "category": "Core Edge",
        "priority": "P1 - Critical", "requirement": "FR-DOC-LIMIT-001",
        "prereqs": "Logged in user with upload access.",
        "steps": (
            "1. Try uploading a 20MB PDF (above MAX_CONTENT_LENGTH = 16MB)\n"
            "2. Try uploading a 1MB file (well under)\n"
            "3. Try uploading a 17MB file"
        ),
        "data": "Various sizes.",
        "expected": (
            "1. 20MB returns 413 (Payload Too Large) before hitting Gunicorn\n"
            "2. 1MB succeeds\n"
            "3. 17MB returns 413 with clean message"
        ),
        "criteria": "Pass: Cap enforced cleanly. Fail: 502 from Gunicorn or silent truncation."
    },
    {
        "id": "TC-601", "name": "Document Upload Disallowed File Type", "category": "Core Edge",
        "priority": "P2 - High", "requirement": "FR-DOC-TYPE-001",
        "prereqs": "Upload UI.",
        "steps": (
            "1. Try uploading .exe\n"
            "2. Try uploading .zip\n"
            "3. Try uploading .docx (allowed)"
        ),
        "data": "Mixed file types.",
        "expected": (
            "1. .exe rejected with clear error\n"
            "2. .zip rejected\n"
            "3. .docx accepted"
        ),
        "criteria": "Pass: Whitelist enforced. Fail: Any disallowed accepted."
    },
    {
        "id": "TC-602", "name": "Application Submit With No Responses Filled", "category": "Core Edge",
        "priority": "P1 - Critical", "requirement": "FR-APP-SUBMIT-VAL-001",
        "prereqs": "NGO with draft application, zero criterion responses.",
        "steps": (
            "1. POST /api/applications/<id>/submit"
        ),
        "data": "Empty responses.",
        "expected": (
            "1. 400 with 'missing_criteria' list of every criterion label\n"
            "2. Application status remains 'draft'\n"
            "3. No AI scoring, no auto-assign"
        ),
        "criteria": "Pass: Pre-submit validation. Fail: Submitted empty."
    },
    {
        "id": "TC-603", "name": "Application Submit After Deadline", "category": "Core Edge",
        "priority": "P2 - High", "requirement": "FR-APP-DEADLINE-001",
        "prereqs": "Grant with deadline in the past + NGO draft.",
        "steps": (
            "1. POST submit"
        ),
        "data": "Past deadline.",
        "expected": "400 'The application deadline has passed'.",
        "criteria": "Pass: Blocked. Fail: Late submit accepted."
    },
    {
        "id": "TC-604", "name": "Duplicate Application Attempt", "category": "Core Edge",
        "priority": "P1 - Critical", "requirement": "FR-APP-DUPE-001",
        "prereqs": "NGO already has application for grant X.",
        "steps": (
            "1. POST /api/applications/ with same grant_id again"
        ),
        "data": "Duplicate attempt.",
        "expected": (
            "1. 409 Conflict\n"
            "2. existing_application_id returned in body\n"
            "3. Banner on UI shows link to existing"
        ),
        "criteria": "Pass: Properly de-duped. Fail: Duplicate row created."
    },
    {
        "id": "TC-605", "name": "Report Submit On Already-Submitted Report", "category": "Core Edge",
        "priority": "P2 - High", "requirement": "FR-REP-IDEMP-001",
        "prereqs": "Report already in status 'submitted'.",
        "steps": (
            "1. POST /api/reports/<id>/submit again"
        ),
        "data": "Re-submit.",
        "expected": "400 'Report cannot be submitted in current status' OR idempotent 200; never double-records AI analysis.",
        "criteria": "Pass: Safe re-submit. Fail: Double AI run."
    },
    {
        "id": "TC-606", "name": "Cross-Org Application Access Denied", "category": "Core Edge",
        "priority": "P1 - Critical", "requirement": "FR-APP-AUTH-001",
        "prereqs": "Two NGOs. NGO A has application X; NGO B logged in.",
        "steps": (
            "1. NGO B GET /api/applications/<X.id>\n"
            "2. NGO B POST /api/applications/<X.id>/submit"
        ),
        "data": "Cross-org probe.",
        "expected": "Both return 403 or 404.",
        "criteria": "Pass: Isolated. Fail: Cross-org access."
    },
    {
        "id": "TC-607", "name": "Global Search Special Characters", "category": "Core Edge",
        "priority": "P3 - Medium", "requirement": "FR-SEARCH-CHAR-001",
        "prereqs": "Logged in.",
        "steps": (
            "1. GET /api/search?q=' OR 1=1 --\n"
            "2. GET /api/search?q=<script>alert(1)</script>\n"
            "3. GET /api/search?q=日本語\n"
            "4. GET /api/search?q=العربية"
        ),
        "data": "Adversarial + unicode queries.",
        "expected": (
            "1. SQL injection attempts safely parameterized (no SQL error)\n"
            "2. XSS payload returns escaped in JSON\n"
            "3. Unicode queries succeed (Japanese, Arabic both work)"
        ),
        "criteria": "Pass: Safe + unicode-aware. Fail: SQL error or XSS reflection."
    },
    {
        "id": "TC-608", "name": "Global Search Very Short Query Rejected", "category": "Core Edge",
        "priority": "P3 - Medium", "requirement": "FR-SEARCH-MIN-001",
        "prereqs": "Logged in.",
        "steps": (
            "1. GET /api/search?q=a\n"
            "2. GET /api/search?q=ab"
        ),
        "data": "Short queries.",
        "expected": "Returns empty results with 'minimum length' message, not an error.",
        "criteria": "Pass: Graceful. Fail: 500 or unbounded scan."
    },
    {
        "id": "TC-609", "name": "Audit Chain Tamper Detection", "category": "Core Edge",
        "priority": "P1 - Critical", "requirement": "FR-AUDIT-TAMPER-001",
        "prereqs": "DB access. Audit chain has entries.",
        "steps": (
            "1. Manually edit one audit_chain row's 'details' field in the DB\n"
            "2. GET /api/audit-chain/verify"
        ),
        "data": "Tampered row.",
        "expected": (
            "1. Verify returns success=false\n"
            "2. Specific row reported as 'hash_mismatch'\n"
            "3. Subsequent row also flagged (chain broken)"
        ),
        "criteria": "Pass: Tamper detected. Fail: Tampering silently accepted."
    },
    {
        "id": "TC-610", "name": "Audit Chain Append-Only Enforcement", "category": "Core Edge",
        "priority": "P1 - Critical", "requirement": "FR-AUDIT-IMMUT-001",
        "prereqs": "Admin user.",
        "steps": (
            "1. Attempt DELETE /api/audit-chain/<id> (no such endpoint)\n"
            "2. Attempt PATCH /api/audit-chain/<id>"
        ),
        "data": "Mutation attempts.",
        "expected": "Both return 404 — no mutation endpoint exists by design.",
        "criteria": "Pass: 404. Fail: Any 200 (would break immutability)."
    },
    {
        "id": "TC-611", "name": "Compliance Health Band Boundaries", "category": "Core Edge",
        "priority": "P2 - High", "requirement": "FR-COMPLIANCE-BAND-001",
        "prereqs": "NGOs with varied compliance health scores: 79, 80, 89, 90.",
        "steps": (
            "1. Inspect each NGO's compliance_health_band"
        ),
        "data": "Boundary scores.",
        "expected": (
            "1. 79 → 'attention' or appropriate band\n"
            "2. 80 → next band starts here\n"
            "3. Bands consistent with documented thresholds\n"
            "4. No off-by-one"
        ),
        "criteria": "Pass: Boundaries correct. Fail: Off-by-one."
    },
    {
        "id": "TC-612", "name": "Risk Status Lifecycle Transitions", "category": "Core Edge",
        "priority": "P2 - High", "requirement": "FR-RISK-LIFE-001",
        "prereqs": "A Risk row in 'open' status.",
        "steps": (
            "1. Transition open → mitigated\n"
            "2. Transition mitigated → resolved\n"
            "3. Attempt resolved → open (illegal)"
        ),
        "data": "State transitions.",
        "expected": "Legal transitions succeed; illegal return 400.",
        "criteria": "Pass: State machine enforced. Fail: Any path accepted."
    },
    {
        "id": "TC-613", "name": "Org Merge Refuses Same-Org", "category": "Core Edge",
        "priority": "P2 - High", "requirement": "FR-MERGE-SELF-001",
        "prereqs": "Admin login.",
        "steps": (
            "1. POST merge with source_id = target_id"
        ),
        "data": "Self merge.",
        "expected": "400 with clear message.",
        "criteria": "Pass: Refused. Fail: Self-merge succeeds."
    },
    {
        "id": "TC-614", "name": "Org Merge Confirm Name Mismatch", "category": "Core Edge",
        "priority": "P1 - Critical", "requirement": "FR-MERGE-CONFIRM-001",
        "prereqs": "Admin login.",
        "steps": (
            "1. POST merge with confirm_name='Wrong Org Name'"
        ),
        "data": "Wrong confirmation.",
        "expected": "400 with 'confirm_name does not match' error.",
        "criteria": "Pass: Safety check works. Fail: Merge proceeds."
    },
    {
        "id": "TC-615", "name": "Notification Digest Cadence Off", "category": "Core Edge",
        "priority": "P3 - Medium", "requirement": "FR-DIGEST-OFF-001",
        "prereqs": "User has digest_cadence='off'.",
        "steps": (
            "1. Trigger the digest cron"
        ),
        "data": "Off setting.",
        "expected": "User is excluded from the digest batch.",
        "criteria": "Pass: Respected. Fail: User still receives digest."
    },
    {
        "id": "TC-616", "name": "Notification Preferences Invalid Channel", "category": "Core Edge",
        "priority": "P3 - Medium", "requirement": "FR-NOTIF-VAL-001",
        "prereqs": "Logged in.",
        "steps": (
            "1. PUT /api/notification-preferences with channel='telepathy'"
        ),
        "data": "Bogus channel.",
        "expected": "Channel silently filtered out (in_app retained as minimum).",
        "criteria": "Pass: Filtered. Fail: Stored as-is."
    },
    {
        "id": "TC-617", "name": "GDPR Right To Erasure", "category": "Core Edge",
        "priority": "P1 - Critical", "requirement": "FR-GDPR-001",
        "prereqs": "Test user account with applications + reports.",
        "steps": (
            "1. Admin triggers GDPR erasure for the user"
        ),
        "data": "Erasure request.",
        "expected": (
            "1. User row marked deleted/anonymized\n"
            "2. Personal data scrubbed from related records\n"
            "3. Audit chain shows the erasure event (without re-storing the PII)"
        ),
        "criteria": "Pass: Compliant erasure. Fail: PII remains."
    },
    {
        "id": "TC-618", "name": "Session Expiry On Inactivity", "category": "Core Edge",
        "priority": "P2 - High", "requirement": "FR-SESSION-EXP-001",
        "prereqs": "Logged in user.",
        "steps": (
            "1. Note session cookie expiry\n"
            "2. Wait past the expiry\n"
            "3. Make a request"
        ),
        "data": "Expired session.",
        "expected": "401 returned; UI redirects to /login.",
        "criteria": "Pass: Expires. Fail: Indefinite session."
    },
    {
        "id": "TC-619", "name": "Health/Ready Probes Always Fast", "category": "Core Edge",
        "priority": "P1 - Critical", "requirement": "FR-PROBE-FAST-001",
        "prereqs": "Production URL.",
        "steps": (
            "1. GET /api/health (10x in 10s)\n"
            "2. GET /api/ready (10x in 10s)\n"
            "3. Time each"
        ),
        "data": "Probe storm.",
        "expected": (
            "1. /api/health < 200ms each\n"
            "2. /api/ready < 500ms each (includes DB ping)\n"
            "3. No 5xx even if AI calls are saturating workers"
        ),
        "criteria": "Pass: Probes resilient. Fail: Hang or 5xx."
    },
]
test_cases.extend(core_edge_cases)


# ============================================================================
# ── BUILD THE WORD DOCUMENT ──────────────────────────────────────────────────
# ============================================================================

def build_document():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ── Set default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # ── COVER PAGE ──
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KUJA GRANT MANAGEMENT SYSTEM")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(26, 35, 126)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Version 5.0 - Test Cases Document (Phase 24-31 coverage)")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(63, 81, 181)

    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta2.add_run("Prepared for: Adeso / Kuja Link")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    meta3 = doc.add_paragraph()
    meta3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta3.add_run(f"Total Test Cases: {len(test_cases)}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(26, 35, 126)
    run.bold = True

    meta4 = doc.add_paragraph()
    meta4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta4.add_run("Production URL: https://web-production-6f8a.up.railway.app")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ── TABLE OF CONTENTS (manual) ──
    doc.add_heading("Table of Contents", level=1)

    categories = []
    seen = set()
    for tc in test_cases:
        cat = tc["category"]
        if cat not in seen:
            seen.add(cat)
            categories.append(cat)

    for i, cat in enumerate(categories, 1):
        count = sum(1 for tc in test_cases if tc["category"] == cat)
        p = doc.add_paragraph(f"{i}. {cat} ({count} test cases)")
        p.style = doc.styles['Normal']
        p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph(f"{len(categories) + 1}. Test Summary & Traceability Matrix")
    p = doc.add_paragraph(f"{len(categories) + 2}. Test Account Reference")
    p = doc.add_paragraph(f"{len(categories) + 3}. Test Files Reference")

    doc.add_page_break()

    # ── DOCUMENT INFO ──
    doc.add_heading("Document Information", level=1)

    info_table = doc.add_table(rows=8, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ("Document Title", "Kuja Grant Management System v3.0 - Test Cases"),
        ("Version", "3.0"),
        ("Date", datetime.now().strftime("%Y-%m-%d")),
        ("Author", "QA Team / Adeso"),
        ("Application", "Kuja Grant Management System"),
        ("Production URL", "https://web-production-6f8a.up.railway.app"),
        ("Technology Stack", "Python/Flask, Vanilla JS SPA, SQLite/PostgreSQL, Claude AI"),
        ("Total Test Cases", str(len(test_cases))),
    ]
    for i, (label, value) in enumerate(info_data):
        set_cell_shading(info_table.rows[i].cells[0], "E8EAF6")
        set_cell_text(info_table.rows[i].cells[0], label, bold=True, size=10)
        set_cell_text(info_table.rows[i].cells[1], value, size=10)

    doc.add_paragraph("")

    # Priority breakdown
    p1 = sum(1 for tc in test_cases if "P1" in tc["priority"])
    p2 = sum(1 for tc in test_cases if "P2" in tc["priority"])
    p3 = sum(1 for tc in test_cases if "P3" in tc["priority"])

    doc.add_heading("Priority Breakdown", level=2)
    prio_table = doc.add_table(rows=4, cols=3)
    prio_table.style = 'Table Grid'
    prio_data = [
        ("Priority", "Count", "Description"),
        ("P1 - Critical", str(p1), "Must-pass tests for core functionality"),
        ("P2 - High", str(p2), "Important tests for secondary functionality"),
        ("P3 - Low", str(p3), "Nice-to-have tests for edge cases"),
    ]
    for i, (col1, col2, col3) in enumerate(prio_data):
        if i == 0:
            set_cell_shading(prio_table.rows[i].cells[0], "1A237E")
            set_cell_shading(prio_table.rows[i].cells[1], "1A237E")
            set_cell_shading(prio_table.rows[i].cells[2], "1A237E")
            set_cell_text(prio_table.rows[i].cells[0], col1, bold=True, size=10, color=(255, 255, 255))
            set_cell_text(prio_table.rows[i].cells[1], col2, bold=True, size=10, color=(255, 255, 255))
            set_cell_text(prio_table.rows[i].cells[2], col3, bold=True, size=10, color=(255, 255, 255))
        else:
            set_cell_text(prio_table.rows[i].cells[0], col1, bold=True, size=10)
            set_cell_text(prio_table.rows[i].cells[1], col2, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(prio_table.rows[i].cells[2], col3, size=10)

    doc.add_page_break()

    # ── TEST ACCOUNT REFERENCE ──
    doc.add_heading("Test Account Reference", level=1)
    p = doc.add_paragraph("All passwords: pass123")
    p.runs[0].bold = True

    acct_table = doc.add_table(rows=11, cols=4)
    acct_table.style = 'Table Grid'
    acct_headers = ["Role", "Email", "Organization", "Country"]
    acct_data = [
        ("Admin", "admin@kuja.org", "System Admin", "N/A"),
        ("NGO", "fatima@amani.org", "Amani Community Development", "Kenya"),
        ("NGO", "ahmed@salamrelief.org", "Salam Relief Foundation", "Somalia"),
        ("NGO", "thandi@ubuntu.org", "Ubuntu Education Trust", "South Africa"),
        ("NGO", "peter@hopebridges.org", "Hope Bridges Initiative", "Uganda"),
        ("NGO", "aisha@sahelwomen.org", "Sahel Women's Network", "Nigeria"),
        ("Donor", "sarah@globalhealth.org", "Global Health Fund", "Switzerland"),
        ("Donor", "david@eatrust.org", "East Africa Development Trust", "Kenya"),
        ("Reviewer", "james@reviewer.org", "Independent Reviewer", "N/A"),
        ("Reviewer", "maria@reviewer.org", "Independent Reviewer", "N/A"),
    ]
    # Header row
    for j, h in enumerate(acct_headers):
        set_cell_shading(acct_table.rows[0].cells[j], "1A237E")
        set_cell_text(acct_table.rows[0].cells[j], h, bold=True, size=9, color=(255, 255, 255))
    # Data rows
    for i, (role, email, org, country) in enumerate(acct_data, 1):
        set_cell_text(acct_table.rows[i].cells[0], role, size=9)
        set_cell_text(acct_table.rows[i].cells[1], email, size=9)
        set_cell_text(acct_table.rows[i].cells[2], org, size=9)
        set_cell_text(acct_table.rows[i].cells[3], country, size=9)
        if i % 2 == 0:
            for j in range(4):
                set_cell_shading(acct_table.rows[i].cells[j], "F5F5F5")

    doc.add_page_break()

    # ── TEST FILES REFERENCE ──
    doc.add_heading("Test Files Reference", level=1)
    p = doc.add_paragraph("Location: test-files-v3/ directory")

    doc.add_heading("Positive Test Files", level=2)
    pos_files = [
        ("valid_financial_report.pdf", "Clean financial report with revenue, expenses, auditor notes"),
        ("valid_registration_cert.pdf", "Valid registration certificate (Kenya, expires 2027)"),
        ("valid_audit_report.pdf", "Independent audit report with unqualified opinion"),
        ("valid_psea_policy.pdf", "Comprehensive PSEA policy with all required sections"),
        ("valid_project_report.pdf", "Project report with indicators, progress, lessons learned"),
        ("valid_budget_detail.xlsx", "Detailed budget spreadsheet with formulas"),
        ("valid_strategic_plan.pdf", "Strategic plan with vision, mission, 3 goals"),
    ]
    for fname, desc in pos_files:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(fname)
        run.bold = True
        run.font.size = Pt(9)
        p.add_run(f" - {desc}").font.size = Pt(9)

    doc.add_heading("Negative Test Files", level=2)
    neg_files = [
        ("invalid_empty.pdf", "Valid PDF structure but no meaningful content"),
        ("invalid_wrong_extension.txt.pdf", "Text file with .pdf extension (not a real PDF)"),
        ("invalid_financial_incomplete.pdf", "Financial report missing auditor opinion and expense breakdown"),
        ("expired_registration.pdf", "Registration certificate expired 2023-06-30"),
        ("malicious_script.html", "HTML file - should be rejected by file type filter"),
    ]
    for fname, desc in neg_files:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(fname)
        run.bold = True
        run.font.size = Pt(9)
        p.add_run(f" - {desc}").font.size = Pt(9)

    doc.add_page_break()

    # ── TEST CASES BY CATEGORY ──
    current_category = None
    cat_num = 0

    for tc in test_cases:
        if tc["category"] != current_category:
            current_category = tc["category"]
            cat_num += 1
            if cat_num > 1:
                doc.add_page_break()

            # Category count
            cat_count = sum(1 for t in test_cases if t["category"] == current_category)
            cat_p1 = sum(1 for t in test_cases if t["category"] == current_category and "P1" in t["priority"])
            cat_p2 = sum(1 for t in test_cases if t["category"] == current_category and "P2" in t["priority"])
            cat_p3 = sum(1 for t in test_cases if t["category"] == current_category and "P3" in t["priority"])

            doc.add_heading(f"{cat_num}. {current_category}", level=1)

            summary = doc.add_paragraph()
            summary.add_run(f"Total: {cat_count} test cases").bold = True
            summary.add_run(f" | P1: {cat_p1} | P2: {cat_p2} | P3: {cat_p3}")
            doc.add_paragraph("")

        add_test_case_table(doc, tc)

    # ── TRACEABILITY MATRIX ──
    doc.add_page_break()
    doc.add_heading("Traceability Matrix", level=1)
    p = doc.add_paragraph("This matrix maps test cases to functional requirements.")

    # Create traceability table
    trace_table = doc.add_table(rows=1, cols=4)
    trace_table.style = 'Table Grid'
    headers = ["Requirement ID", "Requirement Area", "Test Cases", "Coverage"]
    for j, h in enumerate(headers):
        set_cell_shading(trace_table.rows[0].cells[j], "1A237E")
        set_cell_text(trace_table.rows[0].cells[j], h, bold=True, size=9, color=(255, 255, 255))

    req_map = {
        "FR-AUTH": ("Authentication & Session Management", [tc["id"] for tc in test_cases if tc["category"] == "Authentication"]),
        "FR-DASH": ("Dashboard & Navigation", [tc["id"] for tc in test_cases if tc["category"] == "Dashboard"]),
        "FR-GRNT": ("Grant Management", [tc["id"] for tc in test_cases if tc["category"] == "Grants"]),
        "FR-APPL": ("Application Workflow", [tc["id"] for tc in test_cases if tc["category"] == "Applications"]),
        "FR-ASMT": ("Capacity Assessment", [tc["id"] for tc in test_cases if tc["category"] == "Assessments"]),
        "FR-AI": ("AI Services & Analysis", [tc["id"] for tc in test_cases if tc["category"] == "AI Services"]),
        "FR-COMP": ("Compliance & Sanctions", [tc["id"] for tc in test_cases if tc["category"] == "Compliance & Sanctions"]),
        "FR-REG": ("Registry Verification", [tc["id"] for tc in test_cases if tc["category"] == "Registry Verification"]),
        "FR-RPT": ("Reports & Deadlines", [tc["id"] for tc in test_cases if tc["category"] == "Reports"]),
        "FR-I18N": ("Internationalization & i18n", [tc["id"] for tc in test_cases if tc["category"] == "Internationalization"]),
        "FR-RBAC": ("Role-Based Access Control", [tc["id"] for tc in test_cases if tc["category"] == "Access Control"]),
        "FR-SYS": ("System Health & Infrastructure", [tc["id"] for tc in test_cases if tc["category"] == "System Health"]),
        "FR-DOC": ("Document Management", [tc["id"] for tc in test_cases if tc["category"] == "Documents"]),
        "FR-ORG": ("Organization Profiles", [tc["id"] for tc in test_cases if tc["category"] == "Organization Profile"]),
        "FR-E2E": ("End-to-End Workflows", [tc["id"] for tc in test_cases if tc["category"] == "End-to-End Workflows"]),
    }

    for req_id, (area, tcs) in req_map.items():
        row = trace_table.add_row()
        set_cell_text(row.cells[0], req_id, bold=True, size=8)
        set_cell_text(row.cells[1], area, size=8)
        set_cell_text(row.cells[2], ", ".join(tcs), size=7)
        set_cell_text(row.cells[3], f"{len(tcs)} tests", size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph("")

    # ── SUMMARY TABLE ──
    doc.add_heading("Test Summary", level=1)

    sum_table = doc.add_table(rows=1, cols=5)
    sum_table.style = 'Table Grid'
    sum_headers = ["Category", "Total", "P1", "P2", "P3"]
    for j, h in enumerate(sum_headers):
        set_cell_shading(sum_table.rows[0].cells[j], "1A237E")
        set_cell_text(sum_table.rows[0].cells[j], h, bold=True, size=9, color=(255, 255, 255))

    total_all = 0
    total_p1 = 0
    total_p2 = 0
    total_p3 = 0

    for cat in categories:
        cat_total = sum(1 for tc in test_cases if tc["category"] == cat)
        cat_p1 = sum(1 for tc in test_cases if tc["category"] == cat and "P1" in tc["priority"])
        cat_p2 = sum(1 for tc in test_cases if tc["category"] == cat and "P2" in tc["priority"])
        cat_p3 = sum(1 for tc in test_cases if tc["category"] == cat and "P3" in tc["priority"])

        total_all += cat_total
        total_p1 += cat_p1
        total_p2 += cat_p2
        total_p3 += cat_p3

        row = sum_table.add_row()
        set_cell_text(row.cells[0], cat, bold=True, size=9)
        set_cell_text(row.cells[1], str(cat_total), size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], str(cat_p1), size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[3], str(cat_p2), size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[4], str(cat_p3), size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Total row
    row = sum_table.add_row()
    for j in range(5):
        set_cell_shading(row.cells[j], "E8EAF6")
    set_cell_text(row.cells[0], "TOTAL", bold=True, size=10)
    set_cell_text(row.cells[1], str(total_all), bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[2], str(total_p1), bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[3], str(total_p2), bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[4], str(total_p3), bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ── SAVE ──
    output_path = r"C:\Users\IdirisLoyan\kuja-grant\docs\Kuja_Grant_v5.0_Test_Cases.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print(f"Total test cases: {len(test_cases)}")
    print(f"P1: {total_p1}, P2: {total_p2}, P3: {total_p3}")
    print(f"Categories: {len(categories)}")


if __name__ == "__main__":
    build_document()
