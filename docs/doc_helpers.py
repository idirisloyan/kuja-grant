"""
Shared helper functions and constants for Kuja Grant document generation.
Used by generate_gtm_external.py and generate_product_analysis.py.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Color Constants ─────────────────────────────────────────────────────────
BLUE = RGBColor(0x1E, 0x40, 0xAF)       # #1E40AF
TEAL = RGBColor(0x0D, 0x94, 0x88)       # #0D9488
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
RED = RGBColor(0xDC, 0x26, 0x26)
GREEN = RGBColor(0x05, 0x96, 0x69)
AMBER = RGBColor(0xD9, 0x77, 0x06)

BLUE_HEX = "1E40AF"
TEAL_HEX = "0D9488"
LIGHT_BLUE_HEX = "DBEAFE"
ALT_ROW_HEX = "F0F7FF"
LIGHT_TEAL_HEX = "CCFBF1"
CALLOUT_HEX = "EFF6FF"

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


# ─── Cell Helpers ─────────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, font_size=Pt(9), font_color=BLACK,
                  font_name="Calibri", alignment=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = font_size
    run.font.color.rgb = font_color
    run.font.bold = bold
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


# ─── Table Helpers ────────────────────────────────────────────────────────────

def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a professionally styled table with blue header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set table borders
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders)

    # Header row
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_shading(cell, BLUE_HEX)
        set_cell_text(cell, header, bold=True, font_size=Pt(9), font_color=WHITE)

    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.cell(i + 1, j)
            if i % 2 == 1:
                set_cell_shading(cell, ALT_ROW_HEX)
            set_cell_text(cell, val, font_size=Pt(9), font_color=DARK_GRAY)

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


# ─── Heading Helpers ──────────────────────────────────────────────────────────

def add_heading1(doc, text):
    """Add a Heading 1 with blue color."""
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = BLUE
        run.font.name = "Calibri"
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(8)
    return h


def add_heading2(doc, text):
    """Add a Heading 2 with teal color."""
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = TEAL
        run.font.name = "Calibri"
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    return h


def add_heading3(doc, text):
    """Add a Heading 3."""
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = BLUE
        run.font.name = "Calibri"
        run.font.size = Pt(12)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    return h


# ─── Text Helpers ─────────────────────────────────────────────────────────────

def add_body(doc, text, bold=False, italic=False, space_after=Pt(6)):
    """Add body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = Pt(15)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.font.name = "Calibri"
        run_bold.font.size = Pt(10.5)
        run_bold.font.color.rgb = DARK_GRAY
        run_bold.bold = True
        run_rest = p.add_run(text)
        run_rest.font.name = "Calibri"
        run_rest.font.size = Pt(10.5)
        run_rest.font.color.rgb = DARK_GRAY
    else:
        p.clear()
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK_GRAY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(15)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * level)
    return p


def add_numbered_item(doc, number, title, description):
    """Add a numbered advantage/item with bold title and description."""
    p = doc.add_paragraph()
    run_num = p.add_run(f"{number}. {title}: ")
    run_num.font.name = "Calibri"
    run_num.font.size = Pt(10.5)
    run_num.font.color.rgb = BLUE
    run_num.bold = True
    run_desc = p.add_run(description)
    run_desc.font.name = "Calibri"
    run_desc.font.size = Pt(10.5)
    run_desc.font.color.rgb = DARK_GRAY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(15)
    return p


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


# ─── Document Setup ──────────────────────────────────────────────────────────

def setup_document(doc=None):
    """Create and configure a new document with standard styles and margins.
    Returns the configured Document."""
    if doc is None:
        doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = DARK_GRAY

    # Update heading styles
    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = "Calibri"

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    return doc


# ─── Cover Page ───────────────────────────────────────────────────────────────

def create_cover_page(doc, title, subtitle, confidential_text="CONFIDENTIAL"):
    """Create a professional cover page."""
    # Spacer paragraphs for vertical centering
    for _ in range(6):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)
        sp.paragraph_format.space_before = Pt(0)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(36)
    run.font.color.rgb = BLUE
    run.bold = True
    p_title.paragraph_format.space_after = Pt(8)

    # Decorative line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line.add_run("_" * 60)
    run.font.color.rgb = TEAL
    run.font.size = Pt(10)
    p_line.paragraph_format.space_after = Pt(16)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run(subtitle)
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_GRAY
    p_sub.paragraph_format.space_after = Pt(24)

    # Date
    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_ver.add_run("March 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = LIGHT_GRAY
    p_ver.paragraph_format.space_after = Pt(12)

    # Organization
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_org.add_run("Adeso \u2014 African Development Solutions")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = TEAL
    run.bold = True
    p_org.paragraph_format.space_after = Pt(48)

    # Confidential notice
    p_conf = doc.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_conf.add_run(confidential_text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RED
    run.bold = True

    add_page_break(doc)


# ─── Table of Contents ────────────────────────────────────────────────────────

def create_toc(doc, items):
    """Create a table of contents from a list of (number, title) tuples."""
    add_heading1(doc, "Table of Contents")

    for number, title in items:
        p = doc.add_paragraph()
        run = p.add_run(f"{number}.   {title}")
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.font.color.rgb = DARK_GRAY
        p.paragraph_format.space_after = Pt(4)
        # Add tab leader dots (simplified)
        p.paragraph_format.left_indent = Cm(0.5)

    add_page_break(doc)


# ─── Footer ──────────────────────────────────────────────────────────────────

def create_footer(doc, footer_text="Kuja Grant Platform | Adeso \u2014 African Development Solutions | 2026"):
    """Add footer text to the document."""
    doc.add_paragraph()
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_footer.add_run(footer_text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MEDIUM_GRAY


# ─── Callout Box ──────────────────────────────────────────────────────────────

def add_callout_box(doc, title, text, bg_hex=None):
    """Add a shaded callout box (1-row, 1-col table with background)."""
    if bg_hex is None:
        bg_hex = CALLOUT_HEX

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Borders
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{BLUE_HEX}"/>'
        f'  <w:left w:val="single" w:sz="6" w:space="0" w:color="{BLUE_HEX}"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{BLUE_HEX}"/>'
        f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="{BLUE_HEX}"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders)

    cell = table.cell(0, 0)
    set_cell_shading(cell, bg_hex)

    # Title
    cell.text = ""
    p = cell.paragraphs[0]
    run_t = p.add_run(title)
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = BLUE
    run_t.bold = True
    p.paragraph_format.space_after = Pt(4)

    # Body
    p2 = cell.add_paragraph()
    run_b = p2.add_run(text)
    run_b.font.name = "Calibri"
    run_b.font.size = Pt(10)
    run_b.font.color.rgb = DARK_GRAY
    p2.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()  # spacer after box
    return table
