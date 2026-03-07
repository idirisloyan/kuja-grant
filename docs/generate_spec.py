#!/usr/bin/env python3
"""
Generate Kuja Grant Management System Technical Design Specification v3.3.4
Professional .docx document using python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Constants ──────────────────────────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
DARK_BLUE = RGBColor(0x2C, 0x3E, 0x6B)
MEDIUM_BLUE = RGBColor(0x34, 0x4E, 0x8A)
ACCENT_BLUE = RGBColor(0x4A, 0x90, 0xD9)
LIGHT_BLUE = RGBColor(0xD6, 0xE8, 0xF7)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
TABLE_HEADER_BG = "1B2A4A"
TABLE_ALT_ROW = "EBF3FC"
TABLE_BORDER_COLOR = "B0B0B0"
RED = RGBColor(0xCC, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x7A, 0x33)

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           "Kuja_Grant_v3.3.4_Technical_Design_Specification.docx")


# ── Helpers ────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            border = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="4" '
                f'w:space="0" w:color="{val}"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)


def set_table_borders(table, color="B0B0B0"):
    """Set borders on entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_styled_table(doc, headers, rows, col_widths=None, first_col_bold=False):
    """Create a professionally styled table with header and alternating rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    set_table_borders(table, TABLE_BORDER_COLOR)

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, TABLE_HEADER_BG)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.name = "Calibri"
            run.font.color.rgb = DARK_GRAY
            if first_col_bold and c_idx == 0:
                run.bold = True
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_ROW)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)

    return table


def add_para(doc, text, style='Normal', bold=False, italic=False, size=None,
             color=None, space_after=None, space_before=None, alignment=None):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.name = "Calibri"
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if alignment is not None:
        p.alignment = alignment
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        run2 = p.add_run(text)
        run2.font.size = Pt(10)
        run2.font.name = "Calibri"
        run2.font.color.rgb = DARK_GRAY
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    return p


def add_code_block(doc, text):
    """Add a monospaced code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    # Add shading to paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shading)
    for line in text.split("\n"):
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = DARK_GRAY
    return p


def add_page_break(doc):
    """Insert a page break."""
    doc.add_page_break()


def add_section_heading(doc, text, level=1):
    """Add a heading with consistent formatting."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.color.rgb = NAVY
            run.font.size = Pt(20)
        elif level == 2:
            run.font.color.rgb = DARK_BLUE
            run.font.size = Pt(15)
        elif level == 3:
            run.font.color.rgb = MEDIUM_BLUE
            run.font.size = Pt(12)
    return h


def add_horizontal_rule(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def setup_header_footer(doc):
    """Add headers and footers to all sections."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ""
        run_left = hp.add_run("Kuja Grant Management System")
        run_left.font.size = Pt(8)
        run_left.font.color.rgb = MEDIUM_GRAY
        run_left.font.name = "Calibri"
        run_left.italic = True
        hp.add_run("    ")
        run_right = hp.add_run("Design Specification v3.3.4  |  CONFIDENTIAL")
        run_right.font.size = Pt(8)
        run_right.font.color.rgb = MEDIUM_GRAY
        run_right.font.name = "Calibri"
        run_right.italic = True
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add bottom border to header
        pPr = hp._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = ""
        run_f = fp.add_run("Adeso  |  Kuja Engineering Team  |  March 2026")
        run_f.font.size = Pt(8)
        run_f.font.color.rgb = MEDIUM_GRAY
        run_f.font.name = "Calibri"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page number - add as field
        fp2 = footer.add_paragraph()
        fp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_pn = fp2.add_run("Page ")
        run_pn.font.size = Pt(8)
        run_pn.font.color.rgb = MEDIUM_GRAY
        run_pn.font.name = "Calibri"
        # PAGE field
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_f1 = fp2.add_run()
        run_f1._r.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_f2 = fp2.add_run()
        run_f2._r.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_f3 = fp2.add_run()
        run_f3._r.append(fldChar2)


# ── Document Generation ───────────────────────────────────────────────────

def generate():
    doc = Document()

    # ── Page Setup ─────────────────────────────────────────────────────
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Style Setup ────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK_GRAY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15

    for lvl, (sz, clr) in enumerate([(24, NAVY), (18, DARK_BLUE), (14, MEDIUM_BLUE)], 1):
        hs = doc.styles[f'Heading {lvl}']
        hs.font.name = 'Calibri'
        hs.font.size = Pt(sz)
        hs.font.color.rgb = clr
        hs.font.bold = True
        hs.paragraph_format.space_before = Pt(18 if lvl == 1 else 12)
        hs.paragraph_format.space_after = Pt(8)

    # ══════════════════════════════════════════════════════════════════
    #  COVER PAGE
    # ══════════════════════════════════════════════════════════════════

    # Spacer
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KUJA GRANT\nMANAGEMENT SYSTEM")
    run.font.size = Pt(36)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"
    run.bold = True
    p.paragraph_format.space_after = Pt(6)

    # Divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = ACCENT_BLUE
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(12)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Technical Design Specification")
    run.font.size = Pt(24)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("System Architecture, Database Design & Technical Reference")
    run.font.size = Pt(13)
    run.font.color.rgb = MEDIUM_GRAY
    run.font.name = "Calibri"
    run.italic = True
    p.paragraph_format.space_after = Pt(40)

    # Metadata table (centered, no borders)
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Version:", "3.3.4"),
        ("Date:", "March 2026"),
        ("Classification:", "CONFIDENTIAL"),
        ("Author:", "Kuja Engineering Team"),
        ("Organization:", "Adeso"),
        ("Document ID:", "KUJA-SPEC-2026-003"),
    ]
    for i, (label, value) in enumerate(meta_data):
        c0 = meta_table.rows[i].cells[0]
        c1 = meta_table.rows[i].cells[1]
        c0.text = ""
        c1.text = ""
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r0 = p0.add_run(label)
        r0.font.size = Pt(11)
        r0.font.color.rgb = MEDIUM_GRAY
        r0.font.name = "Calibri"
        r0.bold = True
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(value)
        r1.font.size = Pt(11)
        r1.font.name = "Calibri"
        if label == "Classification:":
            r1.font.color.rgb = RED
            r1.bold = True
        else:
            r1.font.color.rgb = DARK_GRAY
        c0.width = Inches(2.5)
        c1.width = Inches(3.0)

    # Remove borders from meta table
    for row in meta_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)

    # Footer on cover
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("This document is the property of Adeso and contains confidential information.\n"
                     "Unauthorized reproduction or distribution is prohibited.")
    run.font.size = Pt(8)
    run.font.color.rgb = MEDIUM_GRAY
    run.font.name = "Calibri"
    run.italic = True

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    #  TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════════

    add_section_heading(doc, "Table of Contents", level=1)
    add_horizontal_rule(doc)

    toc_entries = [
        ("1.", "Executive Summary", "4"),
        ("2.", "System Architecture", "5"),
        ("  2.1", "Technology Stack", "5"),
        ("  2.2", "Architecture Diagram", "6"),
        ("  2.3", "Blueprint Architecture", "7"),
        ("  2.4", "Data Flow Diagrams", "8"),
        ("3.", "Database Design", "10"),
        ("  3.1", "Entity-Relationship Diagram", "10"),
        ("  3.2", "Database Schema Tables", "12"),
        ("  3.3", "Indexes & Constraints", "20"),
        ("  3.4", "JSON Field Schemas", "21"),
        ("4.", "API Reference", "23"),
        ("  4.1", "Authentication", "23"),
        ("  4.2", "Dashboard", "24"),
        ("  4.3", "Grants", "24"),
        ("  4.4", "Applications", "25"),
        ("  4.5", "Assessments", "26"),
        ("  4.6", "Documents", "27"),
        ("  4.7", "AI Services", "28"),
        ("  4.8", "Compliance", "29"),
        ("  4.9", "Reviews", "30"),
        ("  4.10", "Reports", "31"),
        ("  4.11", "Admin", "32"),
        ("5.", "Security Architecture", "33"),
        ("  5.1", "Authentication & Authorization", "33"),
        ("  5.2", "Security Headers", "34"),
        ("  5.3", "CSRF Protection", "34"),
        ("  5.4", "Rate Limiting", "35"),
        ("  5.5", "File Upload Security", "35"),
        ("  5.6", "CORS Configuration", "36"),
        ("6.", "AI Integration", "37"),
        ("  6.1", "Claude API Integration", "37"),
        ("  6.2", "AI Features", "37"),
        ("  6.3", "Scoring Algorithm", "38"),
        ("  6.4", "Assessment Scoring", "39"),
        ("  6.5", "Fallback System", "40"),
        ("7.", "Internationalization (i18n)", "41"),
        ("  7.1", "Supported Languages", "41"),
        ("  7.2", "Translation Architecture", "41"),
        ("  7.3", "RTL Support", "42"),
        ("8.", "Scalability & Performance", "43"),
        ("  8.1", "Worker Configuration", "43"),
        ("  8.2", "Database Connection Pooling", "43"),
        ("  8.3", "Caching Strategy", "44"),
        ("  8.4", "Thread Safety", "44"),
        ("9.", "Deployment", "45"),
        ("  9.1", "Railway Configuration", "45"),
        ("  9.2", "Environment Variables", "45"),
        ("  9.3", "Migration Strategy", "46"),
        ("10.", "Compliance & External Integrations", "47"),
        ("  10.1", "Sanctions Screening", "47"),
        ("  10.2", "Registry Verification", "48"),
        ("A.", "Appendix A: Glossary of Terms", "49"),
        ("B.", "Appendix B: Configuration Reference", "50"),
    ]

    toc_table = doc.add_table(rows=len(toc_entries), cols=3)
    toc_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (num, title, pg) in enumerate(toc_entries):
        r = toc_table.rows[i]
        is_main = not num.startswith("  ")

        c0 = r.cells[0]
        c0.text = ""
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(num.strip())
        r0.font.size = Pt(10)
        r0.font.name = "Calibri"
        r0.font.color.rgb = NAVY if is_main else MEDIUM_GRAY
        r0.bold = is_main
        c0.width = Inches(0.6)

        c1 = r.cells[1]
        c1.text = ""
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(title)
        r1.font.size = Pt(10)
        r1.font.name = "Calibri"
        r1.font.color.rgb = DARK_GRAY if is_main else MEDIUM_GRAY
        r1.bold = is_main
        c1.width = Inches(5.0)

        c2 = r.cells[2]
        c2.text = ""
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run(pg)
        r2.font.size = Pt(10)
        r2.font.name = "Calibri"
        r2.font.color.rgb = MEDIUM_GRAY
        c2.width = Inches(0.6)

    # Remove borders from TOC table
    for row in toc_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    #  1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════

    add_section_heading(doc, "1. Executive Summary", level=1)
    add_horizontal_rule(doc)

    add_para(doc,
        "Kuja Grant is an enterprise-grade grant management platform purpose-built for the humanitarian sector, "
        "designed to serve over 1,000 non-governmental organizations (NGOs) operating across Africa and the "
        "broader international development landscape. The platform streamlines the entire grant lifecycle, from "
        "opportunity discovery through application, assessment, compliance verification, and post-award reporting.",
        space_after=8)

    add_para(doc,
        "The system is architected as a single-page application (SPA) with a Python/Flask backend serving a "
        "RESTful API consumed by a vanilla JavaScript frontend. PostgreSQL provides the persistent data layer "
        "in production, while SQLite is used for local development. The application is continuously deployed "
        "via Railway from a GitHub repository, ensuring rapid iteration and zero-downtime deployments.",
        space_after=8)

    add_para(doc,
        "A distinguishing feature of the platform is its deep integration with Anthropic's Claude AI. The AI "
        "subsystem powers automated document analysis with per-requirement scoring, intelligent application "
        "evaluation, capacity assessment guidance, and natural-language chat support for users. This AI layer "
        "operates alongside human reviewers, providing preliminary scores and insights that augment rather than "
        "replace human judgment.",
        space_after=8)

    add_para(doc,
        "Compliance is a first-class concern. The platform performs live sanctions screening against the "
        "OpenSanctions database with automatic fallback to direct downloads from the UN Consolidated List, "
        "OFAC SDN List, and EU Sanctions List. Government registry verification is supported for seven African "
        "countries, with live integration for South Africa and Uganda and portal-based checks for five others.",
        space_after=8)

    add_para(doc, "Key platform metrics:", bold=True, space_after=4)

    metrics = [
        ("Codebase:", " ~5,300 lines backend (Python) + ~5,800 lines frontend (JavaScript)"),
        ("Database:", " 10 SQLAlchemy models, 13+ indexes, full JSON field support"),
        ("API Surface:", " 55+ endpoints across 12 Flask blueprints"),
        ("Languages:", " 4 supported (English, Arabic with RTL, French, Spanish), 675 translation keys"),
        ("AI Features:", " 6 distinct AI capabilities powered by Claude claude-sonnet-4-20250514"),
        ("Compliance:", " Live sanctions screening + 7-country registry verification"),
        ("Frameworks:", " 5 capacity assessment frameworks (Kuja, STEP, UN-HACT, CHS, NUPAS)"),
        ("Security:", " Multi-layer brute-force protection, DB-backed rate limiting, WSGI upload guard, enterprise security headers"),
        ("CI/CD:", " GitHub Actions E2E regression gate with automated post-deploy testing"),
    ]
    for label, text in metrics:
        add_bullet(doc, text, bold_prefix=label)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    #  2. SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════

    add_section_heading(doc, "2. System Architecture", level=1)
    add_horizontal_rule(doc)

    # ── 2.1 Technology Stack ───────────────────────────────────────────
    add_section_heading(doc, "2.1 Technology Stack", level=2)

    add_styled_table(doc,
        ["Layer", "Technology", "Version", "Purpose"],
        [
            ["Backend", "Python", "3.11+", "Core language for server-side logic"],
            ["Framework", "Flask", "3.1", "Lightweight WSGI web framework"],
            ["ORM", "SQLAlchemy", "2.0", "Database abstraction and query building"],
            ["Migrations", "Alembic", "1.13+", "Database schema version management"],
            ["WSGI Server", "Gunicorn", "22.0+", "Production-grade HTTP server"],
            ["Frontend", "Vanilla JavaScript", "ES2020+", "Single-page application (SPA)"],
            ["Styling", "CSS3", "N/A", "Custom styles with RTL support (2,109 lines)"],
            ["Database (Prod)", "PostgreSQL", "15+", "Production relational database"],
            ["Database (Dev)", "SQLite", "3.x", "Local development database"],
            ["AI", "Anthropic Claude API", "claude-sonnet-4-20250514", "Document analysis, scoring, guidance"],
            ["Sanctions API", "OpenSanctions", "v3", "Entity sanctions screening"],
            ["Hosting", "Railway", "N/A", "Auto-deploy PaaS with PostgreSQL"],
            ["Version Control", "Git / GitHub", "N/A", "Source code management and CI/CD trigger"],
        ],
        col_widths=[1.2, 1.6, 1.2, 2.5],
        first_col_bold=True)

    doc.add_paragraph()  # spacer

    # ── 2.2 Architecture Diagram ───────────────────────────────────────
    add_section_heading(doc, "2.2 Layered Architecture Diagram", level=2)

    add_para(doc, "The system follows a layered architecture with clear separation of concerns. "
             "Each layer communicates only with its adjacent layers, ensuring modularity and testability.",
             space_after=8)

    arch_layers = [
        ("PRESENTATION LAYER", "F0F7FF",
         "HTML5 / CSS3 / Vanilla JavaScript SPA (5,429 lines)\n"
         "Internationalization: 4 languages (EN, AR, FR, ES) | 627 translation keys\n"
         "RTL Arabic support: 48 dedicated CSS rules | Noto Sans Arabic font\n"
         "Responsive design | Accessible UI | Role-based view rendering"),
        ("API LAYER", "E8F0FE",
         "12 Flask Blueprints | 55 RESTful API endpoints\n"
         "CORS (configurable origins) | CSRF (custom header validation)\n"
         "Rate Limiting (IP:email tracking) | Session management (Flask-Login)\n"
         "Request validation | Error handling | JSON serialization"),
        ("SERVICE LAYER", "DCE8FC",
         "AIService: Claude API integration, scoring, document analysis\n"
         "ComplianceService: Sanctions screening, fallback pipeline\n"
         "RegistryService: 7-country government registry verification\n"
         "ScoringEngine: Weighted multi-factor application scoring"),
        ("DATA LAYER", "D0E0FA",
         "10 SQLAlchemy ORM Models | Alembic migrations (idempotent)\n"
         "Connection pooling (pool_size=10, max_overflow=15)\n"
         "JSON field support for flexible schemas\n"
         "Transaction management | Query optimization"),
        ("EXTERNAL INTEGRATIONS", "C4D8F8",
         "Anthropic Claude API (claude-sonnet-4-20250514) | OpenSanctions API v3\n"
         "UN Consolidated List (XML) | OFAC SDN (CSV) | EU Sanctions (CSV)\n"
         "Government Registries: ZA (CIPC), UG (URSB), NG, KE, TZ, SO, ET"),
        ("INFRASTRUCTURE", "B8D0F6",
         "Railway PaaS (auto-deploy from GitHub) | GitHub Actions E2E gate\n"
         "Gunicorn: 4 workers x 4 threads (gthread) | 180s timeout | --preload\n"
         "PostgreSQL 15+ | 1000 max-requests per worker | ON_FAILURE restart"),
    ]

    arch_table = doc.add_table(rows=len(arch_layers), cols=1)
    arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(arch_table, "8899AA")

    for i, (title, bg, desc) in enumerate(arch_layers):
        cell = arch_table.rows[i].cells[0]
        cell.text = ""
        set_cell_shading(cell, bg)

        p_title = cell.paragraphs[0]
        r_title = p_title.add_run(title)
        r_title.bold = True
        r_title.font.size = Pt(11)
        r_title.font.name = "Calibri"
        r_title.font.color.rgb = NAVY
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_after = Pt(4)
        p_title.paragraph_format.space_before = Pt(6)

        p_desc = cell.add_paragraph()
        r_desc = p_desc.add_run(desc)
        r_desc.font.size = Pt(8.5)
        r_desc.font.name = "Calibri"
        r_desc.font.color.rgb = DARK_GRAY
        p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_desc.paragraph_format.space_after = Pt(6)

        # Arrow between layers (except last)
        if i < len(arch_layers) - 1:
            p_arrow = cell.add_paragraph()
            r_arrow = p_arrow.add_run("\u2193  \u2193  \u2193")
            r_arrow.font.size = Pt(10)
            r_arrow.font.color.rgb = ACCENT_BLUE
            p_arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_arrow.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # ── 2.3 Blueprint Architecture ─────────────────────────────────────
    add_section_heading(doc, "2.3 Blueprint Architecture", level=2)

    add_para(doc,
        "The Flask application is organized into 12 blueprints, each encapsulating a distinct domain "
        "of functionality. This modular architecture enables independent development, testing, and "
        "maintenance of each subsystem.",
        space_after=8)

    add_styled_table(doc,
        ["Blueprint", "URL Prefix", "Endpoints", "Auth", "Description"],
        [
            ["auth", "/api/auth", "4", "Mixed", "Login, logout, session, language preference"],
            ["dashboard", "/api/dashboard", "1", "Required", "Aggregated statistics per role"],
            ["organizations", "/api/organizations", "5", "Required", "Organization CRUD and profile management"],
            ["grants", "/api/grants", "7", "Mixed", "Grant lifecycle: create, publish, browse, apply"],
            ["applications", "/api/applications", "6", "Required", "Application submission and tracking"],
            ["assessments", "/api/assessments", "5", "Required", "Capacity assessment across 5 frameworks"],
            ["documents", "/api/documents", "4", "Required", "File upload, download, AI analysis"],
            ["ai", "/api/ai", "5", "Required", "AI chat, scoring, analysis, guidance"],
            ["compliance", "/api/compliance", "4", "Admin/Donor", "Sanctions screening and registry checks"],
            ["reviews", "/api/reviews", "4", "Reviewer", "Application review and scoring"],
            ["reports", "/api/reports", "6", "Required", "Post-award reporting and donor review"],
            ["admin", "/api/admin", "4", "Admin", "System administration and user management"],
        ],
        col_widths=[1.0, 1.2, 0.7, 0.7, 2.9],
        first_col_bold=True)

    p = doc.add_paragraph()
    run = p.add_run("Total: 55 API endpoints across 12 blueprints")
    run.bold = True
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"

    add_page_break(doc)

    # ── 2.4 Data Flow Diagrams ─────────────────────────────────────────
    add_section_heading(doc, "2.4 Data Flow Diagrams", level=2)

    # DFD 1
    add_section_heading(doc, "DFD 1: Grant Application Flow", level=3)
    dfd1_steps = [
        ("1", "NGO Login", "NGO user authenticates via /api/auth/login"),
        ("2", "Browse Grants", "Query published grants with filters (sector, country, search)"),
        ("3", "Select Grant", "View grant details including eligibility criteria and document requirements"),
        ("4", "Eligibility Check", "NGO responds to eligibility questionnaire; system validates prerequisites"),
        ("5", "Proposal Response", "NGO provides written responses to each grant criterion"),
        ("6", "Document Upload", "Required documents uploaded (registration, financials, proposals, etc.)"),
        ("7", "AI Document Analysis", "Claude analyzes each document against donor-specific requirements"),
        ("8", "Submit Application", "Application status changes from 'draft' to 'submitted'"),
        ("9", "AI Scoring", "System generates preliminary AI score using weighted algorithm"),
        ("10", "Reviewer Assignment", "Admin assigns one or more human reviewers"),
        ("11", "Human Scoring", "Reviewers provide scores and comments per criterion"),
        ("12", "Final Score", "Final = (Criteria x 0.60) + (Documents x 0.20) + (Eligibility x 0.20)"),
        ("13", "Decision", "Application awarded or rejected based on final score and reviewer recommendation"),
    ]
    add_styled_table(doc,
        ["Step", "Stage", "Description"],
        dfd1_steps,
        col_widths=[0.5, 1.5, 4.5],
        first_col_bold=True)

    doc.add_paragraph()

    # DFD 2
    add_section_heading(doc, "DFD 2: Compliance Screening Flow", level=3)
    dfd2_steps = [
        ("1", "Trigger", "Admin or donor initiates screening for an organization"),
        ("2", "Primary Check", "Query OpenSanctions API with organization name and aliases"),
        ("3", "API Failure?", "If API returns error or timeout, activate fallback pipeline"),
        ("4", "Fallback: UN XML", "Download and parse UN Consolidated Sanctions List (XML format)"),
        ("5", "Fallback: OFAC CSV", "Download and parse US OFAC SDN List (CSV format)"),
        ("6", "Fallback: EU CSV", "Download and parse EU Consolidated Sanctions List (CSV format)"),
        ("7", "Fuzzy Matching", "Apply fuzzy name matching with 0.75 similarity threshold"),
        ("8", "Store Results", "Persist ComplianceCheck record with match details"),
        ("9", "Status Update", "Organization flagged (match found) or cleared (no match)"),
    ]
    add_styled_table(doc,
        ["Step", "Stage", "Description"],
        dfd2_steps,
        col_widths=[0.5, 1.5, 4.5],
        first_col_bold=True)

    doc.add_paragraph()

    # DFD 3
    add_section_heading(doc, "DFD 3: Report Submission Flow", level=3)
    dfd3_steps = [
        ("1", "Create Report", "NGO creates report for a specific grant and reporting period"),
        ("2", "Fill Template", "NGO completes structured template sections per grant requirements"),
        ("3", "Attach Documents", "Supporting documents uploaded (financial reports, deliverables, etc.)"),
        ("4", "Submit", "Report status changes to 'submitted'; timestamp recorded"),
        ("5", "AI Analysis", "Claude analyzes report content against grant reporting requirements"),
        ("6", "Per-Requirement Scoring", "AI generates individual scores for each requirement"),
        ("7", "Risk Flags", "System identifies compliance gaps and potential risk areas"),
        ("8", "Donor Review", "Donor reviews report with AI insights and per-requirement breakdown"),
        ("9", "Decision", "Donor accepts report or requests revision with specific feedback"),
    ]
    add_styled_table(doc,
        ["Step", "Stage", "Description"],
        dfd3_steps,
        col_widths=[0.5, 1.5, 4.5],
        first_col_bold=True)

    doc.add_paragraph()

    # DFD 4
    add_section_heading(doc, "DFD 4: Capacity Assessment Flow", level=3)
    dfd4_steps = [
        ("1", "Select Framework", "NGO chooses assessment framework (Kuja, STEP, UN-HACT, CHS, or NUPAS)"),
        ("2", "Complete Checklist", "NGO responds to framework-specific checklist items"),
        ("3", "Upload Evidence", "Supporting documents uploaded for verification"),
        ("4", "Category Scoring", "System calculates scores per assessment category"),
        ("5", "Weighted Aggregation", "Overall score computed using framework-specific category weights"),
        ("6", "Gap Identification", "System identifies areas where the organization falls below thresholds"),
        ("7", "Recommendations", "AI generates targeted improvement recommendations"),
    ]
    add_styled_table(doc,
        ["Step", "Stage", "Description"],
        dfd4_steps,
        col_widths=[0.5, 1.5, 4.5],
        first_col_bold=True)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    #  3. DATABASE DESIGN
    # ══════════════════════════════════════════════════════════════════

    add_section_heading(doc, "3. Database Design", level=1)
    add_horizontal_rule(doc)

    # ── 3.1 ER Diagram ────────────────────────────────────────────────
    add_section_heading(doc, "3.1 Entity-Relationship Overview", level=2)

    add_para(doc,
        "The Kuja Grant database consists of 10 core entities that model the complete grant management "
        "lifecycle. The schema leverages PostgreSQL JSON fields extensively to support flexible, "
        "domain-specific data structures while maintaining relational integrity for core entities.",
        space_after=8)

    er_rows = [
        ["User", "Organization", "Many-to-One", "user.org_id \u2192 organization.id", "Users belong to an organization"],
        ["Grant", "Organization", "Many-to-One", "grant.donor_org_id \u2192 organization.id", "Grants created by donor organizations"],
        ["Application", "Grant", "Many-to-One", "application.grant_id \u2192 grant.id", "Applications target a specific grant"],
        ["Application", "Organization", "Many-to-One", "application.ngo_org_id \u2192 organization.id", "NGO submitting the application"],
        ["Assessment", "Organization", "Many-to-One", "assessment.org_id \u2192 organization.id", "Assessments belong to an organization"],
        ["Document", "Application", "Many-to-One", "document.application_id \u2192 application.id", "Documents attached to applications"],
        ["Document", "Assessment", "Many-to-One", "document.assessment_id \u2192 assessment.id", "Documents supporting assessments"],
        ["Review", "Application", "Many-to-One", "review.application_id \u2192 application.id", "Reviews score an application"],
        ["Review", "User", "Many-to-One", "review.reviewer_user_id \u2192 user.id", "Reviewer who created the review"],
        ["Report", "Grant", "Many-to-One", "report.grant_id \u2192 grant.id", "Reports for a specific grant"],
        ["Report", "Application", "Many-to-One", "report.application_id \u2192 application.id", "Reports linked to awarded application"],
        ["Report", "Organization", "Many-to-One", "report.submitted_by_org_id \u2192 organization.id", "NGO submitting the report"],
        ["ComplianceCheck", "Organization", "Many-to-One", "compliance.org_id \u2192 organization.id", "Compliance checks for an organization"],
        ["RegistrationVerification", "Organization", "Many-to-One", "reg.org_id \u2192 organization.id", "Registry verification for an organization"],
        ["RegistrationVerification", "Document", "Many-to-One", "reg.document_id \u2192 document.id", "Supporting registration document"],
        ["RegistrationVerification", "User", "Many-to-One", "reg.verified_by_user_id \u2192 user.id", "User who verified registration"],
    ]

    add_styled_table(doc,
        ["Entity A", "Entity B", "Cardinality", "Foreign Key", "Description"],
        er_rows,
        col_widths=[1.2, 1.2, 0.9, 1.8, 1.4],
        first_col_bold=True)

    add_page_break(doc)

    # ── 3.2 Database Schema Tables ─────────────────────────────────────
    add_section_heading(doc, "3.2 Database Schema Tables", level=2)

    # Define all 10 models
    models = {
        "User": [
            ["id", "Integer", "No", "auto", "PK", "Unique user identifier"],
            ["email", "String(120)", "No", "-", "Unique", "User email address (login)"],
            ["password_hash", "String(256)", "No", "-", "-", "Bcrypt-compatible password hash"],
            ["name", "String(100)", "No", "-", "-", "Full display name"],
            ["role", "String(20)", "No", "-", "Index", "One of: admin, ngo, donor, reviewer"],
            ["org_id", "Integer", "Yes", "NULL", "FK \u2192 Organization", "Parent organization"],
            ["language", "String(5)", "No", "'en'", "-", "Preferred UI language code"],
            ["avatar_url", "String(500)", "Yes", "NULL", "-", "Profile image URL"],
            ["created_at", "DateTime", "No", "utcnow", "-", "Account creation timestamp"],
            ["updated_at", "DateTime", "No", "utcnow", "-", "Last modification timestamp"],
            ["is_active", "Boolean", "No", "True", "-", "Account active status flag"],
        ],
        "Organization": [
            ["id", "Integer", "No", "auto", "PK", "Unique organization identifier"],
            ["name", "String(200)", "No", "-", "-", "Organization legal name"],
            ["org_type", "String(20)", "No", "-", "Index", "Type: ngo, donor, reviewer, government"],
            ["country", "String(100)", "Yes", "NULL", "-", "Country of registration"],
            ["city", "String(100)", "Yes", "NULL", "-", "City of primary operations"],
            ["year_established", "Integer", "Yes", "NULL", "-", "Year organization was founded"],
            ["annual_budget", "Numeric(12,2)", "Yes", "NULL", "-", "Annual operating budget (USD)"],
            ["staff_count", "Integer", "Yes", "NULL", "-", "Total number of staff"],
            ["sectors", "JSON", "Yes", "NULL", "-", "Array of operational sectors"],
            ["description", "Text", "Yes", "NULL", "-", "Organization description"],
            ["mission", "Text", "Yes", "NULL", "-", "Mission statement"],
            ["registration_status", "String(20)", "Yes", "NULL", "-", "Gov registration status"],
            ["registration_number", "String(100)", "Yes", "NULL", "-", "Official registration number"],
            ["verified", "Boolean", "No", "False", "Index", "Platform verification status"],
            ["website", "String(500)", "Yes", "NULL", "-", "Organization website URL"],
            ["logo_url", "String(500)", "Yes", "NULL", "-", "Logo image URL"],
            ["assess_score", "Float", "Yes", "NULL", "-", "Latest assessment overall score"],
            ["assess_date", "DateTime", "Yes", "NULL", "-", "Date of latest assessment"],
            ["geographic_areas", "JSON", "Yes", "NULL", "-", "Array of operational regions"],
            ["focus_areas", "JSON", "Yes", "NULL", "-", "Array of thematic focus areas"],
            ["sdg_ids", "JSON", "Yes", "NULL", "-", "Array of aligned SDG numbers"],
            ["created_at", "DateTime", "No", "utcnow", "-", "Record creation timestamp"],
            ["updated_at", "DateTime", "No", "utcnow", "-", "Last modification timestamp"],
        ],
        "Grant": [
            ["id", "Integer", "No", "auto", "PK", "Unique grant identifier"],
            ["donor_org_id", "Integer", "No", "-", "FK \u2192 Organization, Index", "Donor organization that created the grant"],
            ["title", "String(300)", "No", "-", "-", "Grant opportunity title"],
            ["description", "Text", "Yes", "NULL", "-", "Full grant description"],
            ["total_funding", "Numeric(12,2)", "Yes", "NULL", "-", "Total available funding amount"],
            ["currency", "String(3)", "No", "'USD'", "-", "ISO 4217 currency code"],
            ["deadline", "DateTime", "Yes", "NULL", "-", "Application submission deadline"],
            ["status", "String(20)", "No", "'draft'", "Index", "draft, published, closed, awarded"],
            ["sectors", "JSON", "Yes", "NULL", "-", "Target sectors for the grant"],
            ["countries", "JSON", "Yes", "NULL", "-", "Eligible countries"],
            ["eligibility", "JSON", "Yes", "NULL", "-", "Eligibility criteria definitions"],
            ["criteria", "JSON", "Yes", "NULL", "-", "Evaluation criteria with weights"],
            ["doc_requirements", "JSON", "Yes", "NULL", "-", "Required document types and specs"],
            ["reporting_requirements", "JSON", "Yes", "NULL", "-", "Post-award reporting requirements"],
            ["grant_document", "String(500)", "Yes", "NULL", "-", "Uploaded grant agreement filename"],
            ["report_template", "JSON", "Yes", "NULL", "-", "Template for NGO report submission"],
            ["reporting_frequency", "String(20)", "Yes", "NULL", "-", "monthly, quarterly, semi-annual, annual"],
            ["created_at", "DateTime", "No", "utcnow", "-", "Record creation timestamp"],
            ["updated_at", "DateTime", "No", "utcnow", "-", "Last modification timestamp"],
            ["published_at", "DateTime", "Yes", "NULL", "-", "Date grant was published"],
        ],
        "Application": [
            ["id", "Integer", "No", "auto", "PK", "Unique application identifier"],
            ["grant_id", "Integer", "No", "-", "FK \u2192 Grant, Index", "Target grant"],
            ["ngo_org_id", "Integer", "No", "-", "FK \u2192 Organization, Index", "Applying NGO organization"],
            ["status", "String(20)", "No", "'draft'", "Index", "draft, submitted, under_review, awarded, rejected"],
            ["responses", "JSON", "Yes", "NULL", "-", "Criterion responses {criterion_id: text}"],
            ["eligibility_responses", "JSON", "Yes", "NULL", "-", "Eligibility questionnaire answers"],
            ["ai_score", "Float", "Yes", "NULL", "-", "AI-generated preliminary score (0-100)"],
            ["human_score", "Float", "Yes", "NULL", "-", "Aggregated human reviewer score (0-100)"],
            ["final_score", "Float", "Yes", "NULL", "-", "Weighted final score (0-100)"],
            ["submitted_at", "DateTime", "Yes", "NULL", "-", "Submission timestamp"],
            ["created_at", "DateTime", "No", "utcnow", "-", "Record creation timestamp"],
            ["updated_at", "DateTime", "No", "utcnow", "-", "Last modification timestamp"],
        ],
        "Assessment": [
            ["id", "Integer", "No", "auto", "PK", "Unique assessment identifier"],
            ["org_id", "Integer", "No", "-", "FK \u2192 Organization", "Assessed organization"],
            ["assess_type", "String(20)", "Yes", "NULL", "-", "Assessment classification type"],
            ["framework", "String(20)", "No", "-", "-", "kuja, step, unhact, chs, nupas"],
            ["status", "String(20)", "No", "'draft'", "-", "draft, in_progress, completed"],
            ["overall_score", "Float", "Yes", "NULL", "-", "Weighted overall score (0-100)"],
            ["category_scores", "JSON", "Yes", "NULL", "-", "Per-category score breakdown"],
            ["checklist_responses", "JSON", "Yes", "NULL", "-", "Checklist item responses {key: bool}"],
            ["gaps", "JSON", "Yes", "NULL", "-", "Identified capacity gaps"],
            ["completed_at", "DateTime", "Yes", "NULL", "-", "Assessment completion timestamp"],
            ["created_at", "DateTime", "No", "utcnow", "-", "Record creation timestamp"],
            ["updated_at", "DateTime", "No", "utcnow", "-", "Last modification timestamp"],
        ],
        "Document": [
            ["id", "Integer", "No", "auto", "PK", "Unique document identifier"],
            ["application_id", "Integer", "Yes", "NULL", "FK \u2192 Application", "Parent application (if applicable)"],
            ["assessment_id", "Integer", "Yes", "NULL", "FK \u2192 Assessment", "Parent assessment (if applicable)"],
            ["doc_type", "String(50)", "No", "-", "-", "Document classification type"],
            ["original_filename", "String(500)", "No", "-", "-", "User's original filename"],
            ["stored_filename", "String(500)", "No", "-", "Index, Unique", "Obfuscated storage filename"],
            ["file_size", "Integer", "Yes", "NULL", "-", "File size in bytes"],
            ["mime_type", "String(100)", "Yes", "NULL", "-", "MIME content type"],
            ["ai_analysis", "JSON", "Yes", "NULL", "-", "AI analysis results and scores"],
            ["score", "Float", "Yes", "NULL", "-", "Document quality score (0-100)"],
            ["uploaded_at", "DateTime", "No", "utcnow", "-", "Upload timestamp"],
            ["updated_at", "DateTime", "No", "utcnow", "-", "Last modification timestamp"],
        ],
        "Review": [
            ["id", "Integer", "No", "auto", "PK", "Unique review identifier"],
            ["application_id", "Integer", "No", "-", "FK \u2192 Application", "Reviewed application"],
            ["reviewer_user_id", "Integer", "No", "-", "FK \u2192 User, Index", "Assigned reviewer"],
            ["scores", "JSON", "Yes", "NULL", "-", "Per-criterion scores {criterion_id: score}"],
            ["comments", "JSON", "Yes", "NULL", "-", "Per-criterion comments {criterion_id: text}"],
            ["overall_score", "Float", "Yes", "NULL", "-", "Reviewer's overall score (0-100)"],
            ["status", "String(20)", "No", "'assigned'", "Index", "assigned, in_progress, completed"],
            ["completed_at", "DateTime", "Yes", "NULL", "-", "Review completion timestamp"],
            ["created_at", "DateTime", "No", "utcnow", "-", "Record creation timestamp"],
            ["updated_at", "DateTime", "No", "utcnow", "-", "Last modification timestamp"],
        ],
        "Report": [
            ["id", "Integer", "No", "auto", "PK", "Unique report identifier"],
            ["grant_id", "Integer", "No", "-", "FK \u2192 Grant, Index", "Parent grant"],
            ["application_id", "Integer", "No", "-", "FK \u2192 Application", "Associated awarded application"],
            ["submitted_by_org_id", "Integer", "No", "-", "FK \u2192 Organization, Index", "Submitting NGO organization"],
            ["report_type", "String(50)", "Yes", "NULL", "-", "narrative, financial, progress, final"],
            ["reporting_period", "String(50)", "Yes", "NULL", "-", "Period covered (e.g., Q1 2026)"],
            ["title", "String(300)", "Yes", "NULL", "-", "Report title"],
            ["content", "JSON", "Yes", "NULL", "-", "Structured report sections"],
            ["attachments", "JSON", "Yes", "NULL", "-", "Array of attached file references"],
            ["status", "String(20)", "No", "'draft'", "Index", "draft, submitted, reviewed, accepted, revision_requested"],
            ["due_date", "DateTime", "Yes", "NULL", "-", "Report submission deadline"],
            ["submitted_at", "DateTime", "Yes", "NULL", "-", "Actual submission timestamp"],
            ["reviewed_at", "DateTime", "Yes", "NULL", "-", "Donor review timestamp"],
            ["reviewer_notes", "Text", "Yes", "NULL", "-", "Donor feedback and notes"],
            ["ai_analysis", "JSON", "Yes", "NULL", "-", "AI analysis of report content"],
            ["created_at", "DateTime", "No", "utcnow", "-", "Record creation timestamp"],
            ["updated_at", "DateTime", "No", "utcnow", "-", "Last modification timestamp"],
        ],
        "ComplianceCheck": [
            ["id", "Integer", "No", "auto", "PK", "Unique compliance check identifier"],
            ["org_id", "Integer", "No", "-", "FK \u2192 Organization, Index", "Screened organization"],
            ["check_type", "String(50)", "No", "-", "-", "sanctions, registry, aml"],
            ["status", "String(20)", "No", "-", "-", "clear, flagged, error, pending"],
            ["result", "JSON", "Yes", "NULL", "-", "Screening results and match details"],
            ["checked_at", "DateTime", "No", "utcnow", "Index", "Screening execution timestamp"],
            ["updated_at", "DateTime", "No", "utcnow", "-", "Last modification timestamp"],
        ],
        "RegistrationVerification": [
            ["id", "Integer", "No", "auto", "PK", "Unique verification identifier"],
            ["org_id", "Integer", "No", "-", "FK \u2192 Organization", "Verified organization"],
            ["status", "String(20)", "No", "'pending'", "-", "pending, verified, failed, expired"],
            ["registration_number", "String(100)", "Yes", "NULL", "-", "Registration number verified"],
            ["registration_authority", "String(200)", "Yes", "NULL", "-", "Issuing government authority"],
            ["registry_check_result", "JSON", "Yes", "NULL", "-", "Raw registry check response"],
            ["registration_date", "DateTime", "Yes", "NULL", "-", "Original registration date"],
            ["expiry_date", "DateTime", "Yes", "NULL", "-", "Registration expiry date"],
            ["country", "String(100)", "Yes", "NULL", "-", "Country of registration"],
            ["ai_analysis", "JSON", "Yes", "NULL", "-", "AI analysis of registration documents"],
            ["ai_confidence", "Float", "Yes", "NULL", "-", "AI confidence score (0.0-1.0)"],
            ["document_id", "Integer", "Yes", "NULL", "FK \u2192 Document", "Supporting registration document"],
            ["verified_by_user_id", "Integer", "Yes", "NULL", "FK \u2192 User", "Admin who verified"],
            ["verified_at", "DateTime", "Yes", "NULL", "-", "Verification timestamp"],
            ["notes", "Text", "Yes", "NULL", "-", "Verification notes"],
            ["registry_url", "String(500)", "Yes", "NULL", "-", "Government registry URL used"],
            ["created_at", "DateTime", "No", "utcnow", "-", "Record creation timestamp"],
            ["updated_at", "DateTime", "No", "utcnow", "-", "Last modification timestamp"],
        ],
    }

    for model_name, fields in models.items():
        add_section_heading(doc, f"Table: {model_name}", level=3)
        add_styled_table(doc,
            ["Field Name", "Data Type", "Nullable", "Default", "Constraints", "Description"],
            fields,
            col_widths=[1.2, 1.0, 0.6, 0.6, 1.2, 1.9],
            first_col_bold=True)
        doc.add_paragraph()  # spacer

    add_page_break(doc)

    # ── 3.3 Indexes & Constraints ──────────────────────────────────────
    add_section_heading(doc, "3.3 Indexes & Constraints", level=2)

    add_para(doc,
        "The following indexes are defined to optimize query performance for the most common access patterns. "
        "All indexes use B-tree type unless otherwise specified.",
        space_after=8)

    index_rows = [
        ["ix_users_role", "User", "role", "Accelerates role-based user queries (dashboard, admin)"],
        ["ix_users_org_id", "User", "org_id", "Fast lookup of users belonging to an organization"],
        ["ix_orgs_type", "Organization", "org_type", "Filter organizations by type (ngo, donor, etc.)"],
        ["ix_orgs_verified", "Organization", "verified", "Quick identification of verified organizations"],
        ["ix_grants_donor_status", "Grant", "donor_org_id, status", "Composite: donor's grants filtered by status"],
        ["ix_applications_ngo_status", "Application", "ngo_org_id, status", "Composite: NGO's applications by status"],
        ["ix_applications_grant_status", "Application", "grant_id, status", "Composite: applications per grant by status"],
        ["ix_documents_stored_filename", "Document", "stored_filename", "Unique file retrieval by stored name"],
        ["ix_reviews_user_status", "Review", "reviewer_user_id, status", "Composite: reviewer's reviews by status"],
        ["ix_reports_org_status", "Report", "submitted_by_org_id, status", "Composite: org's reports by status"],
        ["ix_reports_grant_status", "Report", "grant_id, status", "Composite: grant's reports by status"],
        ["ix_reports_submitted_by_org", "Report", "submitted_by_org_id", "Fast lookup of reports by submitting org"],
        ["ix_compliance_org_date", "ComplianceCheck", "org_id, checked_at", "Composite: org's checks ordered by date"],
    ]

    add_styled_table(doc,
        ["Index Name", "Table", "Column(s)", "Rationale"],
        index_rows,
        col_widths=[1.8, 1.2, 1.5, 2.0],
        first_col_bold=True)

    doc.add_paragraph()

    add_para(doc, "Unique Constraints:", bold=True, space_after=4)
    add_styled_table(doc,
        ["Constraint Name", "Table", "Column(s)", "Purpose"],
        [
            ["uq_application_grant_ngo", "Application", "grant_id, ngo_org_id",
             "Prevents duplicate applications: one NGO can submit at most one application per grant"],
        ],
        col_widths=[2.0, 1.2, 1.5, 1.8],
        first_col_bold=True)

    doc.add_paragraph()

    # ── 3.4 JSON Field Schemas ─────────────────────────────────────────
    add_section_heading(doc, "3.4 JSON Field Schemas", level=2)

    add_para(doc,
        "Several fields use PostgreSQL JSON columns to store flexible, domain-specific data structures. "
        "The following documents the expected schema for each major JSON field.",
        space_after=8)

    json_schemas = [
        ("Grant.criteria", "Array of evaluation criterion definitions",
         '[\n  {\n    "id": "string (UUID)",\n    "label": "string",\n'
         '    "description": "string",\n    "instructions": "string",\n'
         '    "example": "string",\n    "maxWords": "integer",\n'
         '    "weight": "float (0.0-1.0)"\n  }\n]'),
        ("Grant.eligibility", "Array of eligibility requirements",
         '[\n  {\n    "id": "string (UUID)",\n    "category": "string",\n'
         '    "label": "string",\n    "parameters": "object",\n'
         '    "weight": "float (0.0-1.0)",\n    "required": "boolean",\n'
         '    "helpText": "string"\n  }\n]'),
        ("Grant.reporting_requirements", "Post-award reporting specifications",
         '[\n  {\n    "id": "string (UUID)",\n    "type": "string",\n'
         '    "title": "string",\n    "description": "string",\n'
         '    "frequency": "string (monthly|quarterly|annual)",\n'
         '    "due_days": "integer"\n  }\n]'),
        ("Application.responses", "Criterion responses keyed by criterion ID",
         '{\n  "criterion_id_1": "Response text for criterion 1",\n'
         '  "criterion_id_2": "Response text for criterion 2"\n}'),
        ("Assessment.checklist_responses", "Checklist item completion status",
         '{\n  "governance_01": true,\n  "governance_02": false,\n'
         '  "finance_01": true\n}'),
        ("Document.ai_analysis", "AI document analysis results",
         '{\n  "score": 85.5,\n  "findings": ["Finding 1", "Finding 2"],\n'
         '  "recommendations": ["Rec 1", "Rec 2"],\n'
         '  "per_requirement_scores": {\n    "req_id_1": 90,\n'
         '    "req_id_2": 75\n  }\n}'),
        ("ComplianceCheck.result", "Sanctions screening result details",
         '{\n  "list": "UN Consolidated List",\n  "match_score": 0.82,\n'
         '  "matched_entity": "Entity Name",\n  "reason": "Name similarity",\n'
         '  "datasets": ["un_sc_sanctions"],\n  "source": "opensanctions",\n'
         '  "records_searched": 12450\n}'),
    ]

    for field_name, desc, schema in json_schemas:
        p = doc.add_paragraph()
        run = p.add_run(field_name)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"
        run2 = p.add_run(f"  \u2014  {desc}")
        run2.font.size = Pt(9)
        run2.font.color.rgb = MEDIUM_GRAY
        run2.font.name = "Calibri"
        run2.italic = True
        p.paragraph_format.space_after = Pt(2)

        add_code_block(doc, schema)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    #  4. API REFERENCE
    # ══════════════════════════════════════════════════════════════════

    add_section_heading(doc, "4. API Reference", level=1)
    add_horizontal_rule(doc)

    add_para(doc,
        "All API endpoints return JSON responses and use standard HTTP status codes. "
        "Authentication is managed via Flask-Login sessions. Mutating requests (POST, PUT, DELETE) "
        "require a valid CSRF token in the X-Requested-With header. Rate limiting is applied "
        "to sensitive endpoints as documented below.",
        space_after=8)

    add_para(doc, "Common Response Codes:", bold=True, space_after=4)
    add_styled_table(doc,
        ["Code", "Status", "Description"],
        [
            ["200", "OK", "Request succeeded"],
            ["201", "Created", "Resource created successfully"],
            ["400", "Bad Request", "Invalid request body or parameters"],
            ["401", "Unauthorized", "Not authenticated"],
            ["403", "Forbidden", "Insufficient permissions for the requested operation"],
            ["404", "Not Found", "Resource does not exist"],
            ["429", "Too Many Requests", "Rate limit exceeded"],
            ["500", "Internal Server Error", "Unexpected server error"],
        ],
        col_widths=[0.7, 1.5, 4.3],
        first_col_bold=True)

    doc.add_paragraph()

    # Define all API blueprints
    api_sections = [
        ("4.1 Authentication", "auth", "/api/auth", [
            ["POST", "/api/auth/login", "Public", "{email, password}", "{success, user}", "400, 401, 403, 429", "Authenticate user; creates session. Rate limited: 5/IP:email, 30-min lockout"],
            ["POST", "/api/auth/logout", "login_required", "-", "{success, message}", "-", "Destroy current session"],
            ["GET", "/api/auth/me", "login_required", "-", "{user}", "401", "Return current authenticated user profile"],
            ["PUT", "/api/auth/language", "login_required", "{language}", "{success, language}", "400", "Update preferred language (en, ar, fr, es)"],
        ]),
        ("4.2 Dashboard", "dashboard", "/api/dashboard", [
            ["GET", "/api/dashboard/stats", "login_required", "?role", "{stats, role}", "-", "Aggregated statistics; response varies by role. Cached 30s."],
        ]),
        ("4.3 Grants", "grants", "/api/grants", [
            ["GET", "/api/grants/", "Public", "?status, sector, country, search, page, per_page", "{grants[], total, page, pages}", "-", "List grants with pagination and filters"],
            ["GET", "/api/grants/<id>", "Public", "-", "{grant, application_count, user_application}", "404", "Get grant details including applicant context"],
            ["POST", "/api/grants/", "donor, admin", "Full grant object (JSON)", "{success, grant}", "400, 403", "Create a new grant opportunity"],
            ["PUT", "/api/grants/<id>", "donor, admin", "Partial grant fields", "{success, grant}", "400, 403, 404", "Update an existing grant"],
            ["POST", "/api/grants/<id>/publish", "donor, admin", "-", "{success, grant}", "403, 404", "Publish a draft grant (sets published_at)"],
            ["POST", "/api/grants/<id>/upload-grant-doc", "donor, admin", "multipart/form-data (file)", "{success, filename}", "400, 403, 413", "Upload grant agreement document (16MB max)"],
            ["POST", "/api/grants/<id>/extract-requirements", "donor, admin", "-", "{success, requirements[]}", "403, 404", "AI extracts reporting requirements from grant document"],
        ]),
        ("4.4 Applications", "applications", "/api/applications", [
            ["GET", "/api/applications/", "login_required", "?grant_id, status, page", "{applications[], total}", "-", "List applications (filtered by role context)"],
            ["GET", "/api/applications/<id>", "login_required", "-", "{application, grant, documents, reviews}", "403, 404", "Get full application details with related data"],
            ["POST", "/api/applications/", "ngo", "{grant_id}", "{success, application}", "400, 403, 409", "Create a draft application for a grant (unique constraint)"],
            ["PUT", "/api/applications/<id>", "ngo", "{responses, eligibility_responses}", "{success, application}", "400, 403, 404", "Update draft application responses"],
            ["POST", "/api/applications/<id>/submit", "ngo", "-", "{success, application}", "400, 403", "Submit application; triggers AI scoring"],
            ["POST", "/api/applications/<id>/score", "admin", "-", "{success, scores}", "403, 404", "Trigger AI scoring for a submitted application"],
        ]),
        ("4.5 Assessments", "assessments", "/api/assessments", [
            ["GET", "/api/assessments/", "login_required", "?org_id, framework", "{assessments[]}", "-", "List assessments for an organization"],
            ["GET", "/api/assessments/<id>", "login_required", "-", "{assessment, documents}", "403, 404", "Get assessment details with supporting documents"],
            ["POST", "/api/assessments/", "ngo, admin", "{org_id, framework}", "{success, assessment}", "400, 403", "Start a new capacity assessment"],
            ["PUT", "/api/assessments/<id>", "ngo, admin", "{checklist_responses, category_scores}", "{success, assessment}", "400, 403, 404", "Update assessment responses and scores"],
            ["POST", "/api/assessments/<id>/complete", "ngo, admin", "-", "{success, assessment}", "400, 403", "Mark assessment as completed; calculates final scores"],
        ]),
        ("4.6 Documents", "documents", "/api/documents", [
            ["POST", "/api/documents/upload", "login_required", "multipart/form-data (file, application_id|assessment_id, doc_type)", "{success, document}", "400, 403, 413", "Upload a document (16MB max, validated extension & magic bytes)"],
            ["GET", "/api/documents/<id>", "login_required", "-", "Binary file stream", "403, 404", "Download a document by ID"],
            ["POST", "/api/documents/<id>/analyze", "login_required", "-", "{success, analysis}", "403, 404", "Trigger AI analysis of an uploaded document"],
            ["DELETE", "/api/documents/<id>", "login_required", "-", "{success, message}", "403, 404", "Delete a document (owner or admin only)"],
        ]),
        ("4.7 AI Services", "ai", "/api/ai", [
            ["POST", "/api/ai/chat", "login_required", "{message, context}", "{success, response}", "400, 429", "AI chat for grant guidance. Rate limited: 20/min"],
            ["POST", "/api/ai/analyze-document", "login_required", "{document_id, grant_id}", "{success, analysis}", "400, 403, 429", "Analyze document against grant-specific requirements"],
            ["POST", "/api/ai/score-application", "admin, donor", "{application_id}", "{success, scores}", "403, 404, 429", "Generate AI scores for an application"],
            ["POST", "/api/ai/evaluate-report", "donor, admin", "{report_id}", "{success, evaluation}", "403, 404, 429", "AI evaluation of submitted report"],
            ["POST", "/api/ai/assess-capacity", "login_required", "{assessment_id}", "{success, analysis}", "403, 404, 429", "AI capacity assessment recommendations"],
        ]),
        ("4.8 Compliance", "compliance", "/api/compliance", [
            ["POST", "/api/compliance/screen", "admin, donor", "{org_id}", "{success, result}", "400, 403", "Run sanctions screening for an organization"],
            ["GET", "/api/compliance/history/<org_id>", "admin, donor", "?check_type", "{checks[]}", "403, 404", "Get compliance check history for an organization"],
            ["POST", "/api/compliance/verify-registration", "admin, donor", "{org_id, country, registration_number}", "{success, verification}", "400, 403", "Verify organization against government registry"],
            ["GET", "/api/compliance/status/<org_id>", "login_required", "-", "{sanctions_status, registration_status, last_checked}", "404", "Get current compliance status summary"],
        ]),
        ("4.9 Reviews", "reviews", "/api/reviews", [
            ["GET", "/api/reviews/", "reviewer, admin", "?application_id, status", "{reviews[]}", "-", "List reviews assigned to the current reviewer"],
            ["GET", "/api/reviews/<id>", "reviewer, admin", "-", "{review, application, grant}", "403, 404", "Get review details with application context"],
            ["PUT", "/api/reviews/<id>", "reviewer", "{scores, comments, overall_score}", "{success, review}", "400, 403, 404", "Submit scores and comments for a review"],
            ["POST", "/api/reviews/<id>/complete", "reviewer", "-", "{success, review}", "400, 403", "Mark review as completed; updates application scores"],
        ]),
        ("4.10 Reports", "reports", "/api/reports", [
            ["GET", "/api/reports/", "login_required", "?grant_id, status, org_id", "{reports[], total}", "-", "List reports (filtered by role)"],
            ["GET", "/api/reports/<id>", "login_required", "-", "{report, grant, application}", "403, 404", "Get report details with grant context"],
            ["POST", "/api/reports/", "ngo", "{grant_id, application_id, report_type, reporting_period}", "{success, report}", "400, 403", "Create a new report for an awarded grant"],
            ["PUT", "/api/reports/<id>", "ngo", "{content, title}", "{success, report}", "400, 403, 404", "Update draft report content"],
            ["POST", "/api/reports/<id>/submit", "ngo", "-", "{success, report}", "400, 403", "Submit report for donor review"],
            ["PUT", "/api/reports/<id>/review", "donor, admin", "{status, reviewer_notes}", "{success, report}", "400, 403, 404", "Donor reviews report: accept or request revision"],
        ]),
        ("4.11 Admin", "admin", "/api/admin", [
            ["GET", "/api/admin/users", "admin", "?role, search, page", "{users[], total, page, pages}", "403", "List all users with filters and pagination"],
            ["PUT", "/api/admin/users/<id>", "admin", "{role, is_active, ...}", "{success, user}", "400, 403, 404", "Update user profile and permissions"],
            ["GET", "/api/admin/stats", "admin", "-", "{total_users, total_orgs, total_grants, ...}", "403", "System-wide statistics for admin dashboard"],
            ["POST", "/api/admin/assign-reviewer", "admin", "{application_id, reviewer_id}", "{success, review}", "400, 403, 404", "Assign a reviewer to an application"],
        ]),
    ]

    for section_title, bp_name, prefix, endpoints in api_sections:
        add_section_heading(doc, section_title, level=2)

        add_styled_table(doc,
            ["Method", "Endpoint", "Auth", "Request", "Response", "Errors", "Notes"],
            endpoints,
            col_widths=[0.5, 1.5, 0.6, 1.0, 0.9, 0.5, 1.5],
            first_col_bold=True)
        doc.add_paragraph()

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    #  5. SECURITY ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════

    add_section_heading(doc, "5. Security Architecture", level=1)
    add_horizontal_rule(doc)

    # 5.1
    add_section_heading(doc, "5.1 Authentication & Authorization", level=2)

    add_para(doc,
        "The platform uses Flask-Login for session-based authentication with Werkzeug's "
        "bcrypt-compatible password hashing. Authorization is enforced through a custom "
        "@role_required decorator that checks user roles before granting access to protected endpoints.",
        space_after=8)

    add_styled_table(doc,
        ["Property", "Configuration", "Description"],
        [
            ["Session Backend", "Server-side (Flask-Login)", "Secure session stored server-side; cookie contains session ID only"],
            ["Password Hashing", "Werkzeug PBKDF2-SHA256", "Bcrypt-compatible hashing with automatic salt generation"],
            ["Cookie: HttpOnly", "True", "Prevents JavaScript access to session cookie"],
            ["Cookie: SameSite", "Lax", "Protects against CSRF from external domains"],
            ["Cookie: Secure", "True (production)", "Cookie only sent over HTTPS in production"],
            ["Session Lifetime", "8 hours", "Sessions expire after 8 hours of inactivity"],
            ["Role Enforcement", "@role_required decorator", "Checks user.role against allowed roles per endpoint"],
            ["Roles", "admin, ngo, donor, reviewer", "Four distinct roles with graduated permissions"],
        ],
        col_widths=[1.5, 2.0, 3.0],
        first_col_bold=True)

    doc.add_paragraph()

    # 5.2
    add_section_heading(doc, "5.2 Security Headers", level=2)

    add_styled_table(doc,
        ["Header", "Value", "Purpose"],
        [
            ["X-Content-Type-Options", "nosniff", "Prevents MIME type sniffing"],
            ["X-Frame-Options", "DENY", "Prevents clickjacking via iframe embedding"],
            ["X-XSS-Protection", "1; mode=block", "Legacy XSS protection for older browsers"],
            ["Strict-Transport-Security", "max-age=31536000; includeSubDomains", "Enforces HTTPS for 1 year"],
            ["Content-Security-Policy", "default-src 'self'; ...", "Restricts resource loading to trusted sources"],
            ["Referrer-Policy", "strict-origin-when-cross-origin", "Limits referrer information in cross-origin requests"],
            ["Permissions-Policy", "camera=(), microphone=(), geolocation=()", "Disables sensitive browser APIs"],
            ["Cache-Control", "no-store, no-cache, must-revalidate", "Prevents caching of sensitive responses"],
        ],
        col_widths=[2.0, 2.5, 2.0],
        first_col_bold=True)

    doc.add_paragraph()

    # 5.3
    add_section_heading(doc, "5.3 CSRF Protection", level=2)
    add_para(doc,
        "Cross-Site Request Forgery protection is implemented through a custom header validation scheme. "
        "All mutating requests (POST, PUT, DELETE) must include the X-Requested-With: XMLHttpRequest header. "
        "This approach leverages the browser's same-origin policy for custom headers, which prevents "
        "cross-origin requests from injecting the header without a CORS preflight that the server would reject.",
        space_after=8)

    add_para(doc,
        "The SPA frontend automatically includes this header in all fetch() calls via a centralized "
        "API client utility, ensuring seamless protection without developer intervention per endpoint.",
        space_after=8)

    # 5.4
    add_section_heading(doc, "5.4 Rate Limiting", level=2)

    add_styled_table(doc,
        ["Endpoint", "Limit", "Window", "Lockout", "Key", "Store"],
        [
            ["POST /api/auth/login (IP)", "20 attempts", "5 minutes", "Reject until window expires", "Client IP", "PostgreSQL (login_attempts table)"],
            ["POST /api/auth/login (email)", "5 attempts", "15 minutes", "15-min email lockout", "Email address", "PostgreSQL (login_attempts table)"],
            ["POST /api/auth/login (account)", "5 attempts", "5-min window", "15-min account lock", "User ID", "PostgreSQL (users table, atomic SQL)"],
            ["POST /api/ai/*", "20 requests", "1 minute", "60-sec cooldown", "User ID", "PostgreSQL (DbRateLimiter)"],
        ],
        col_widths=[1.5, 0.8, 0.7, 0.9, 0.7, 1.9],
        first_col_bold=True)

    add_para(doc,
        "Login rate limiting uses a three-layer defense, all backed by PostgreSQL for consistency "
        "across Gunicorn's 4 workers x 4 threads. Layer 1: IP-based limiting prevents any single IP "
        "from exceeding 20 login attempts in 5 minutes. Layer 2: per-email limiting works for both "
        "existing and non-existing accounts, preventing email enumeration attacks. Layer 3: account-level "
        "lockout uses atomic SQL (UPDATE ... SET failed_login_count = failed_login_count + 1) to prevent "
        "race conditions under concurrency.",
        space_after=4)

    add_para(doc,
        "The AI rate limiter uses DbRateLimiter, a database-backed implementation that stores attempts "
        "in the login_attempts table with a limiter-specific tag. This shares state across all Gunicorn "
        "workers. If the database is unavailable (e.g. in tests), it falls back to an in-memory "
        "RateLimiter automatically.",
        space_after=8)

    # 5.5
    add_section_heading(doc, "5.5 File Upload Security", level=2)

    add_styled_table(doc,
        ["Control", "Implementation", "Details"],
        [
            ["WSGI Upload Guard", "OversizedUploadGuard middleware", "WSGI-level defense: drains oversized request body (up to 50 MB) before returning clean JSON 413, preventing client connection resets"],
            ["Size Limit", "16 MB maximum", "Enforced at WSGI middleware and Flask MAX_CONTENT_LENGTH levels"],
            ["Minimum Size", "100 bytes", "Rejects empty or corrupt files below 100 bytes"],
            ["Extension Whitelist", ".pdf, .doc, .docx, .xls, .xlsx, .csv, .jpg, .jpeg, .png, .txt", "Only permitted file extensions accepted"],
            ["Magic Byte Validation", "File header inspection", "Verifies first bytes match declared type (e.g., %PDF for .pdf, PK for .docx)"],
            ["Filename Obfuscation", "UUID-based renaming", "Original filenames stored in DB; files stored with uuid4().hex names"],
            ["Storage Isolation", "Dedicated upload directory", "Uploaded files stored outside web-accessible directory with access-controlled serving"],
            ["Access Control", "Role-based file access", "Admin: all files; NGO/Donor: own org files only; Reviewer: assigned grant files"],
        ],
        col_widths=[1.5, 1.8, 3.2],
        first_col_bold=True)

    doc.add_paragraph()

    # 5.6
    add_section_heading(doc, "5.6 CORS Configuration", level=2)
    add_para(doc,
        "CORS is configured via the CORS_ORIGINS environment variable (comma-separated). The default "
        "origin list is environment-aware: in production only the Railway deployment URL is allowed; "
        "in development, localhost:5000 and 127.0.0.1:5000 are additionally included. This prevents "
        "development origins from leaking into production response headers.",
        space_after=4)

    add_styled_table(doc,
        ["Environment", "Default Origins", "Credentials"],
        [
            ["Production", "https://web-production-6f8a.up.railway.app", "True (cookies supported)"],
            ["Development", "Railway URL + http://localhost:5000 + http://127.0.0.1:5000", "True (cookies supported)"],
            ["Override", "CORS_ORIGINS env var (comma-separated)", "True (always enabled)"],
        ],
        col_widths=[1.3, 3.5, 1.7],
        first_col_bold=True)

    doc.add_paragraph()

    # 5.7
    add_section_heading(doc, "5.7 Audit Logging & Error Handling", level=2)

    add_para(doc,
        "All security-relevant events are logged via Python's structured logging framework (logger 'kuja'). "
        "In production, logs are emitted as JSON to stdout, captured by Railway's log aggregation. "
        "In development, human-readable format with timestamps is used.",
        space_after=4)

    add_styled_table(doc,
        ["Event", "Log Level", "Data Captured"],
        [
            ["Successful login", "INFO", "Email, role, client IP (from X-Forwarded-For)"],
            ["Failed login", "WARNING", "Email, client IP, remaining attempts before lockout"],
            ["Account lockout", "WARNING", "Email, client IP, failure count, lockout duration"],
            ["IP rate limit exceeded", "WARNING", "Client IP, threshold, window"],
            ["API mutation", "INFO", "Method, path, user email, client IP (before_request hook)"],
            ["File upload", "INFO", "Filename, size, user, outcome (success/rejected)"],
            ["Database error", "ERROR", "Exception details, SQL context, rollback status"],
            ["Unhandled exception", "ERROR", "Full traceback (Sentry integration if configured)"],
        ],
        col_widths=[1.5, 0.8, 4.2],
        first_col_bold=True)

    doc.add_paragraph()

    add_para(doc, "Error Handler Responses:", bold=True, space_after=4)
    add_styled_table(doc,
        ["HTTP Status", "Trigger", "Response"],
        [
            ["400", "Malformed request", "JSON error message"],
            ["401", "Unauthenticated access", "JSON error + Flask-Login redirect"],
            ["403", "Unauthorized role", "JSON error (role_required decorator)"],
            ["404", "Unknown endpoint", "JSON error"],
            ["405", "Wrong HTTP method", "JSON error"],
            ["413", "Oversized upload (>16 MB)", "Clean JSON response after WSGI body drain"],
            ["429", "Rate limit exceeded", "JSON error with retry guidance"],
            ["500", "Internal error", "JSON error + DB session rollback"],
            ["503", "Service unavailable", "JSON error (proxy kill on oversized upload)"],
        ],
        col_widths=[0.8, 2.0, 3.7],
        first_col_bold=True)

    doc.add_paragraph()

    add_para(doc, "Sentry Integration (Optional):", bold=True, space_after=4)
    add_para(doc,
        "When the SENTRY_DSN environment variable is set, the application initializes Sentry SDK 2.0.0 "
        "with Flask and SQLAlchemy integrations. Traces are sampled at 10% (traces_sample_rate=0.1) and "
        "tagged with the environment name. All unhandled exceptions are automatically reported.",
        space_after=8)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    #  6. AI INTEGRATION
    # ══════════════════════════════════════════════════════════════════

    add_section_heading(doc, "6. AI Integration", level=1)
    add_horizontal_rule(doc)

    # 6.1
    add_section_heading(doc, "6.1 Claude API Integration", level=2)
    add_para(doc,
        "The platform integrates with Anthropic's Claude API using the claude-sonnet-4-20250514 model. "
        "The API key is configured via the ANTHROPIC_API_KEY environment variable. All AI calls are "
        "made through a centralized AIService class that handles request formatting, error handling, "
        "response parsing, and fallback logic. The service implements structured JSON output parsing "
        "to extract scores, findings, and recommendations from Claude's responses.",
        space_after=8)

    # 6.2
    add_section_heading(doc, "6.2 AI Features", level=2)

    add_styled_table(doc,
        ["Feature", "Trigger", "Input", "Output", "Use Case"],
        [
            ["Document Analysis", "Document upload + analyze", "Document text + grant requirements", "Score, findings, per-requirement scores, recommendations",
             "Evaluate document quality against donor specifications"],
            ["Application Scoring", "Application submission", "Proposal responses + criteria + documents", "Per-criterion scores, overall AI score, strengths, weaknesses",
             "Preliminary scoring to assist human reviewers"],
            ["Report Analysis", "Report submission", "Report content + grant reporting requirements", "Per-requirement compliance scores, risk flags, gaps",
             "Automated compliance check before donor review"],
            ["Capacity Evaluation", "Assessment completion", "Checklist responses + framework definition", "Gap analysis, targeted recommendations, score validation",
             "Guide NGO organizational strengthening"],
            ["Registration Verification", "Registry check trigger", "Registration documents + registry data", "Confidence score, extracted fields, verification notes",
             "AI-assisted verification of government registration documents"],
            ["Chat & Guidance", "User message", "User question + role context", "Natural language response with actionable guidance",
             "On-demand grant writing and compliance guidance"],
        ],
        col_widths=[1.1, 0.9, 1.2, 1.5, 1.8],
        first_col_bold=True)

    doc.add_paragraph()

    # 6.3
    add_section_heading(doc, "6.3 Application Scoring Algorithm", level=2)

    add_para(doc, "The final application score is computed using a weighted formula that combines three dimensions:",
             space_after=4)

    add_code_block(doc,
        "Final Score = (Criteria Average x 0.60) + (Documents Average x 0.20) + (Eligibility x 0.20)")

    add_para(doc, "Component Breakdown:", bold=True, space_after=4)

    add_styled_table(doc,
        ["Component", "Weight", "Source", "Scoring Method"],
        [
            ["Criteria Average", "60%", "AI + Human reviewer scores", "Average of per-criterion scores (0-100). AI evaluates completeness, relevance, depth, and specificity of each response."],
            ["Documents Average", "20%", "AI document analysis", "Average of per-document scores. Each document scored on compliance with requirements, completeness, and quality."],
            ["Eligibility", "20%", "System calculation", "Percentage of mandatory eligibility criteria met. Binary pass/fail per criterion, aggregated as percentage."],
        ],
        col_widths=[1.2, 0.7, 1.5, 3.1],
        first_col_bold=True)

    doc.add_paragraph()

    add_para(doc, "AI Scoring Sub-Dimensions (per criterion):", bold=True, space_after=4)
    add_styled_table(doc,
        ["Dimension", "Weight", "Description"],
        [
            ["Completeness", "30%", "Does the response address all aspects of the criterion?"],
            ["Relevance", "30%", "Is the response directly relevant to the grant's focus and requirements?"],
            ["Depth", "25%", "Does the response demonstrate substantive knowledge and specific examples?"],
            ["Specificity", "15%", "Are concrete data points, timelines, and measurable outcomes included?"],
        ],
        col_widths=[1.5, 0.8, 4.2],
        first_col_bold=True)

    doc.add_paragraph()

    # 6.4
    add_section_heading(doc, "6.4 Assessment Scoring (5 Frameworks)", level=2)

    add_para(doc,
        "Capacity assessments are calculated using the formula: Overall = Sum(category_score x weight). "
        "Each framework defines its own categories and weight distribution.",
        space_after=8)

    frameworks = [
        ("Kuja Framework", [
            ["Governance & Leadership", "25%", "Board structure, policies, strategic planning, oversight"],
            ["Financial Management", "25%", "Accounting systems, controls, audit history, budgeting"],
            ["Program Management", "20%", "Project cycle management, M&E, reporting"],
            ["Human Resources", "15%", "Staff capacity, retention, development, policies"],
            ["External Relations", "15%", "Partnerships, coordination, community engagement"],
        ]),
        ("STEP Framework", [
            ["Systems", "25%", "IT, data management, knowledge management"],
            ["Technical Capacity", "25%", "Sector expertise, methodologies, innovation"],
            ["Effectiveness", "25%", "Impact measurement, outcomes, efficiency"],
            ["Partnerships", "25%", "Collaboration, networking, co-implementation"],
        ]),
        ("UN-HACT Framework", [
            ["Programme Management", "20%", "Planning, implementation, monitoring"],
            ["Organizational Structure", "20%", "Governance, management, staffing"],
            ["Financial Management", "20%", "Accounting, reporting, controls"],
            ["Procurement", "20%", "Policies, processes, value for money"],
            ["Human Resources", "20%", "Staff management, capacity building"],
        ]),
        ("CHS (Core Humanitarian Standard)", [
            ["Appropriateness", "15%", "Needs-based response, community participation"],
            ["Effectiveness", "15%", "Timely, quality, accountable delivery"],
            ["Strengthening Capacity", "15%", "Local capacity building, sustainability"],
            ["Communication & Participation", "15%", "Information sharing, feedback mechanisms"],
            ["Complaints & Response", "15%", "Safe, accessible complaints handling"],
            ["Coordination & Complementarity", "10%", "Inter-agency coordination, avoiding duplication"],
            ["Learning & Improvement", "15%", "Monitoring, evaluation, adaptive management"],
        ]),
        ("NUPAS Framework", [
            ["Institutional Viability", "20%", "Governance, sustainability, strategy"],
            ["Management Systems", "20%", "Operations, processes, decision-making"],
            ["Human Resource Capacity", "20%", "Staffing, skills, development"],
            ["Financial Management", "20%", "Controls, reporting, resource mobilization"],
            ["Program/Service Delivery", "20%", "Quality, reach, impact measurement"],
        ]),
    ]

    for fw_name, categories in frameworks:
        add_section_heading(doc, fw_name, level=3)
        add_styled_table(doc,
            ["Category", "Weight", "Assessment Areas"],
            categories,
            col_widths=[2.0, 0.8, 3.7],
            first_col_bold=True)
        doc.add_paragraph()

    # 6.5
    add_section_heading(doc, "6.5 AI Fallback System", level=2)
    add_para(doc,
        "When the Claude API is unavailable (network error, rate limit, or API key issue), the system "
        "activates a template-based fallback that provides basic functionality without AI enhancement:",
        space_after=4)

    add_bullet(doc, "Document analysis returns a generic template acknowledging the document was received but could not be AI-analyzed")
    add_bullet(doc, "Application scoring falls back to manual-only mode, notifying reviewers that AI scores are unavailable")
    add_bullet(doc, "Chat responses return a polite message directing users to try again later or contact support")
    add_bullet(doc, "Report analysis provides a checklist-based assessment without AI-powered natural language evaluation")
    add_bullet(doc, "All fallback responses include a visible indicator that AI analysis was not available")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    #  7. INTERNATIONALIZATION
    # ══════════════════════════════════════════════════════════════════

    add_section_heading(doc, "7. Internationalization (i18n)", level=1)
    add_horizontal_rule(doc)

    # 7.1
    add_section_heading(doc, "7.1 Supported Languages", level=2)

    add_styled_table(doc,
        ["Code", "Language", "Direction", "Font Family", "Coverage"],
        [
            ["en", "English", "LTR", "Calibri, system default", "Full (primary)"],
            ["ar", "Arabic", "RTL", "Noto Sans Arabic, Arial", "Full (627 keys)"],
            ["fr", "French", "LTR", "Calibri, system default", "Full (627 keys)"],
            ["es", "Spanish", "LTR", "Calibri, system default", "Full (627 keys)"],
        ],
        col_widths=[0.6, 1.0, 0.8, 2.0, 2.1],
        first_col_bold=True)

    doc.add_paragraph()

    # 7.2
    add_section_heading(doc, "7.2 Translation Architecture", level=2)

    add_para(doc,
        "The internationalization system uses a flat-file JSON approach with dot-notation namespaced keys. "
        "Each language file contains 627 translation keys organized by functional area. Translations are "
        "loaded at application startup and served to the frontend via the initial page load, enabling "
        "client-side language switching without additional API calls.",
        space_after=8)

    add_para(doc, "Key Namespaces:", bold=True, space_after=4)
    add_styled_table(doc,
        ["Namespace", "Keys", "Description"],
        [
            ["common.*", "~80", "Shared labels: buttons, status badges, navigation"],
            ["auth.*", "~25", "Login, logout, session messages"],
            ["grants.*", "~85", "Grant listing, details, creation, publishing"],
            ["applications.*", "~70", "Application form, submission, tracking"],
            ["assessments.*", "~65", "Assessment frameworks, checklists, scoring"],
            ["compliance.*", "~45", "Sanctions screening, registry verification"],
            ["reports.*", "~55", "Reporting forms, deadlines, review"],
            ["dashboard.*", "~40", "Dashboard statistics, charts, summaries"],
            ["admin.*", "~35", "User management, system settings"],
            ["errors.*", "~30", "Error messages, validation feedback"],
            ["ai.*", "~50", "AI chat, analysis, scoring labels"],
            ["documents.*", "~45", "Upload, download, analysis results"],
        ],
        col_widths=[1.5, 0.7, 4.3],
        first_col_bold=True)

    doc.add_paragraph()

    # 7.3
    add_section_heading(doc, "7.3 RTL Support", level=2)
    add_para(doc,
        "Arabic language support includes comprehensive right-to-left (RTL) layout mirroring implemented "
        "through 48 dedicated CSS rules. When the Arabic language is active, the [dir='rtl'] attribute "
        "is set on the document root, triggering the RTL stylesheet overrides.",
        space_after=4)

    add_para(doc, "RTL CSS categories:", bold=True, space_after=4)
    add_styled_table(doc,
        ["Category", "Rules", "Description"],
        [
            ["Layout Direction", "8", "Flex direction reversal, text alignment, float direction"],
            ["Margin & Padding", "12", "Mirror left/right margins and padding"],
            ["Positioning", "8", "Swap left/right absolute and fixed positioning"],
            ["Borders & Shadows", "6", "Mirror border-radius and box-shadow direction"],
            ["Icons & Images", "4", "Flip directional icons (arrows, chevrons)"],
            ["Tables", "4", "Reverse cell padding and text alignment"],
            ["Navigation", "6", "Mirror sidebar, breadcrumbs, and menu layouts"],
        ],
        col_widths=[1.5, 0.7, 4.3],
        first_col_bold=True)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    #  8. SCALABILITY & PERFORMANCE
    # ══════════════════════════════════════════════════════════════════

    add_section_heading(doc, "8. Scalability & Performance", level=1)
    add_horizontal_rule(doc)

    # 8.1
    add_section_heading(doc, "8.1 Worker Configuration", level=2)

    add_styled_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["Worker Class", "gthread", "Threaded worker for mixed I/O and CPU workloads"],
            ["Workers", "4", "Number of Gunicorn worker processes"],
            ["Threads", "4 per worker", "Threads per worker (16 total concurrent handlers)"],
            ["Timeout", "180 seconds", "Worker timeout for long-running AI API calls"],
            ["Max Requests", "1000", "Worker recycled after 1000 requests (memory leak prevention)"],
            ["Max Requests Jitter", "50", "Random jitter to prevent synchronized restarts"],
            ["Graceful Timeout", "30 seconds", "Time for workers to finish current requests on shutdown"],
            ["Keep-Alive", "5 seconds", "HTTP keep-alive connection timeout"],
        ],
        col_widths=[1.5, 1.5, 3.5],
        first_col_bold=True)

    doc.add_paragraph()

    # 8.2
    add_section_heading(doc, "8.2 Database Connection Pooling", level=2)

    add_styled_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["pool_size", "10", "Base number of persistent connections maintained"],
            ["max_overflow", "15", "Maximum additional connections above pool_size (25 total max)"],
            ["pool_recycle", "300 seconds", "Connections recycled after 5 minutes to prevent stale connections"],
            ["pool_timeout", "30 seconds", "Maximum wait time for a connection from the pool"],
            ["pool_pre_ping", "True", "Validate connection health before use"],
        ],
        col_widths=[1.5, 1.5, 3.5],
        first_col_bold=True)

    doc.add_paragraph()

    # 8.3
    add_section_heading(doc, "8.3 Caching Strategy", level=2)

    add_styled_table(doc,
        ["Cache Instance", "TTL", "Contents", "Eviction"],
        [
            ["Sanctions Cache", "1 hour (3600s)", "OpenSanctions API responses, fallback list data", "Time-based expiration"],
            ["Lists Cache", "24 hours (86400s)", "Static reference data: countries, sectors, SDGs, frameworks", "Time-based expiration"],
            ["Dashboard Cache", "30 seconds", "Aggregated dashboard statistics per role", "Time-based expiration"],
        ],
        col_widths=[1.3, 1.2, 2.5, 1.5],
        first_col_bold=True)

    add_para(doc,
        "All caches use Werkzeug's SimpleCache implementation with thread-safe access patterns. "
        "Cache invalidation occurs automatically on TTL expiration; manual invalidation is triggered "
        "when underlying data is modified (e.g., new sanctions check clears the sanctions cache).",
        space_after=8)

    # 8.4
    add_section_heading(doc, "8.4 Thread Safety", level=2)
    add_para(doc,
        "The application ensures thread safety through several mechanisms:",
        space_after=4)

    add_bullet(doc, "All cache instances use threading.Lock for atomic read/write operations", bold_prefix="Lock-based caching: ")
    add_bullet(doc, "Login rate limiting uses atomic SQL (UPDATE ... SET count = count + 1) shared across all workers", bold_prefix="DB-backed rate limits: ")
    add_bullet(doc, "AI rate limiter (DbRateLimiter) writes to PostgreSQL login_attempts table, consistent across 4 workers", bold_prefix="Cross-worker AI limits: ")
    add_bullet(doc, "SQLAlchemy's scoped_session provides per-thread database sessions", bold_prefix="Scoped sessions: ")
    add_bullet(doc, "Flask's application context ensures request-local state isolation", bold_prefix="Request isolation: ")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    #  9. DEPLOYMENT
    # ══════════════════════════════════════════════════════════════════

    add_section_heading(doc, "9. Deployment", level=1)
    add_horizontal_rule(doc)

    # 9.1
    add_section_heading(doc, "9.1 Railway Configuration", level=2)

    add_para(doc,
        "The application is deployed on Railway with automatic deployments triggered by pushes to the "
        "main branch on GitHub. The deployment pipeline is configured through two files:",
        space_after=4)

    add_para(doc, "railway.json:", bold=True, space_after=2)
    add_code_block(doc,
        '{\n'
        '  "$schema": "https://railway.app/railway.schema.json",\n'
        '  "build": {},\n'
        '  "deploy": {\n'
        '    "startCommand": "python pre_deploy.py && flask db upgrade &&\n'
        '      gunicorn --bind 0.0.0.0:${PORT:-5000} --workers 4 --threads 4\n'
        '      --timeout 180 --worker-class gthread --max-requests 1000\n'
        '      --max-requests-jitter 50 --preload -c gunicorn.conf.py\n'
        '      server:app",\n'
        '    "restartPolicyType": "ON_FAILURE",\n'
        '    "restartPolicyMaxRetries": 5\n'
        '  }\n'
        '}')

    add_para(doc, "gunicorn.conf.py:", bold=True, space_after=2)
    add_code_block(doc,
        '# limit_request_body = 0 disables Gunicorn\'s built-in body limit.\n'
        '# The OversizedUploadGuard WSGI middleware (app/middleware.py)\n'
        '# handles rejection with proper body draining so the 413\n'
        '# response reaches the client cleanly (no TCP reset).\n'
        'limit_request_body = 0')

    add_para(doc, "Deployment Pipeline:", bold=True, space_after=4)
    add_styled_table(doc,
        ["Step", "Command", "Description"],
        [
            ["1. Pre-deploy", "python pre_deploy.py", "Resets stale Alembic revisions; validates migration chain "
             "(v300_combined, v301_lockout)"],
            ["2. Migrations", "flask db upgrade", "Applies pending Alembic database migrations (idempotent)"],
            ["3. Server Start", "gunicorn ... server:app", "Starts Gunicorn with --preload and gunicorn.conf.py; "
             "ON_FAILURE restart policy (max 5 retries)"],
        ],
        col_widths=[1.0, 2.5, 3.0],
        first_col_bold=True)

    doc.add_paragraph()

    # 9.2
    add_section_heading(doc, "9.2 Environment Variables", level=2)

    add_styled_table(doc,
        ["Variable", "Required", "Default", "Description"],
        [
            ["ANTHROPIC_API_KEY", "Yes", "-", "Anthropic Claude API key for AI features"],
            ["OPENSANCTIONS_API_KEY", "Yes", "-", "OpenSanctions API key for sanctions screening"],
            ["DATABASE_URL", "Yes", "-", "PostgreSQL connection URL (auto-provisioned by Railway)"],
            ["SECRET_KEY", "Yes", "auto-generated", "Flask session encryption key (32+ characters)"],
            ["CORS_ORIGINS", "No", "env-aware", "Production: Railway domain only; Development: "
             "adds localhost. Comma-separated override."],
            ["SENTRY_DSN", "No", "-", "Sentry error-tracking DSN (enables production error reporting)"],
            ["UPLOAD_FOLDER", "No", "uploads/", "Directory for uploaded files"],
            ["FLASK_ENV", "No", "production", "Flask environment (development or production)"],
            ["PORT", "No", "5000", "HTTP port (set automatically by Railway)"],
        ],
        col_widths=[1.8, 0.7, 1.2, 2.8],
        first_col_bold=True)

    doc.add_paragraph()

    # 9.3
    add_section_heading(doc, "9.3 Migration Strategy", level=2)
    add_para(doc,
        "Database migrations use Alembic with an idempotent approach. The migration script uses "
        "SQLAlchemy's inspect() function to check whether tables and columns already exist before "
        "attempting to create them. This prevents errors when re-running migrations on an existing "
        "database and allows the same migration to be safely executed multiple times.",
        space_after=4)

    add_para(doc, "Key migration safeguards:", bold=True, space_after=4)
    add_bullet(doc, "Table existence check via inspect(engine).get_table_names() before CREATE TABLE")
    add_bullet(doc, "Column existence check via inspect(engine).get_columns() before ALTER TABLE ADD COLUMN")
    add_bullet(doc, "Index existence check before CREATE INDEX to avoid duplicate index errors")
    add_bullet(doc, "Transactional DDL where supported (PostgreSQL) for atomic migration application")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    #  10. COMPLIANCE & EXTERNAL INTEGRATIONS
    # ══════════════════════════════════════════════════════════════════

    add_section_heading(doc, "10. Compliance & External Integrations", level=1)
    add_horizontal_rule(doc)

    # 10.1
    add_section_heading(doc, "10.1 Sanctions Screening", level=2)

    add_para(doc,
        "The sanctions screening subsystem provides comprehensive entity screening against multiple "
        "international sanctions lists. The architecture implements a primary-with-fallback pattern "
        "to ensure screening availability even when external APIs are unreachable.",
        space_after=8)

    add_para(doc, "Screening Pipeline:", bold=True, space_after=4)

    add_styled_table(doc,
        ["Priority", "Source", "Format", "Method", "Coverage"],
        [
            ["Primary", "OpenSanctions API v3", "REST JSON", "API call with entity name + type", "Consolidated: UN, OFAC, EU, UK, AU + 40 more lists"],
            ["Fallback 1", "UN Consolidated List", "XML", "Direct download + XML parse", "UN Security Council sanctions"],
            ["Fallback 2", "OFAC SDN List", "CSV", "Direct download + CSV parse", "US Treasury Specially Designated Nationals"],
            ["Fallback 3", "EU Consolidated List", "CSV", "Direct download + CSV parse", "European Union financial sanctions"],
        ],
        col_widths=[0.8, 1.5, 0.8, 1.5, 1.9],
        first_col_bold=True)

    doc.add_paragraph()

    add_para(doc, "Fuzzy Matching Configuration:", bold=True, space_after=4)
    add_styled_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["Algorithm", "Token Set Ratio", "Handles word reordering and partial matches"],
            ["Threshold", "0.75 (75%)", "Minimum similarity score to flag as a potential match"],
            ["Normalization", "Lowercase, strip diacritics", "Name normalization before comparison"],
            ["Multi-name", "Enabled", "Matches against entity aliases and alternate names"],
            ["Cache TTL", "3600 seconds (1 hour)", "Screening results cached to avoid redundant API calls"],
        ],
        col_widths=[1.5, 1.5, 3.5],
        first_col_bold=True)

    doc.add_paragraph()

    # 10.2
    add_section_heading(doc, "10.2 Government Registry Verification", level=2)

    add_para(doc,
        "The platform supports verification of NGO registration status against government registries "
        "in seven African countries. Integration depth varies by country based on registry digitization "
        "and API availability.",
        space_after=8)

    add_styled_table(doc,
        ["Country", "Registry", "Integration", "Method", "Status"],
        [
            ["South Africa", "CIPC (Companies & IP Commission)", "Live", "Web scraping of NPO search portal", "Automated real-time verification"],
            ["Uganda", "URSB (Registration Services Bureau)", "Live", "Web scraping of register portal", "Automated real-time verification"],
            ["Nigeria", "CAC (Corporate Affairs Commission)", "Portal", "Portal URL provided; manual check", "Semi-automated with link generation"],
            ["Kenya", "NGO Coordination Board", "Portal", "Portal URL provided; manual check", "Semi-automated with link generation"],
            ["Tanzania", "RITA / NGO Registrar", "Portal", "Portal URL provided; manual check", "Semi-automated with link generation"],
            ["Somalia", "Ministry of Interior", "Manual", "No digital registry available", "Manual verification with document upload"],
            ["Ethiopia", "ACSO (Charities & Societies Agency)", "Manual", "Limited digital access", "Manual verification with document upload"],
        ],
        col_widths=[0.9, 1.6, 0.7, 1.6, 1.7],
        first_col_bold=True)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    #  11. CI/CD & RELEASE PROCESS
    # ══════════════════════════════════════════════════════════════════

    add_section_heading(doc, "11. CI/CD & Release Process", level=1)
    add_horizontal_rule(doc)

    # 11.1
    add_section_heading(doc, "11.1 GitHub Actions E2E Regression Gate", level=2)

    add_para(doc,
        "Every push to main and every pull request triggers an automated end-to-end regression suite "
        "via GitHub Actions. The workflow can also be dispatched manually with a custom target URL.",
        space_after=4)

    add_para(doc, "Workflow: e2e-regression.yml", bold=True, space_after=4)

    add_styled_table(doc,
        ["Stage", "Action", "Details"],
        [
            ["Checkout", "actions/checkout@v4", "Clone repository at the triggering commit"],
            ["Python Setup", "actions/setup-python@v5", "Python 3.11 with pip requests library"],
            ["Deploy Wait", "sleep 60 + health poll", "On push to main only: waits 60 s for Railway deploy, "
             "then polls /api/health up to 10 times (10 s apart)"],
            ["Build Verify", "curl /api/version", "Logs deployed build hash vs. expected GITHUB_SHA"],
            ["E2E Tests", "python test_e2e_final.py", "Canonical release gate (136 test cases across "
             "auth, RBAC, CRUD, AI, compliance, exports)"],
            ["Artifacts", "actions/upload-artifact@v4", "Uploads test script as artifact; 30-day retention"],
        ],
        col_widths=[1.0, 2.0, 3.5],
        first_col_bold=True)

    doc.add_paragraph()

    # 11.2
    add_section_heading(doc, "11.2 Test Scripts Inventory", level=2)

    add_styled_table(doc,
        ["Script", "Role", "Trigger"],
        [
            ["test_e2e_final.py", "Canonical release gate (CI)", "GitHub Actions on push/PR to main"],
            ["e2e_test.sh", "Quick smoke test (bash)", "Manual; parameterized via KUJA_URL env var"],
            ["test_e2e.py", "Legacy E2E (Python)", "Manual; parameterized via KUJA_URL env var"],
        ],
        col_widths=[2.0, 2.0, 2.5],
        first_col_bold=True)

    doc.add_paragraph()

    # 11.3
    add_section_heading(doc, "11.3 Deployment Flow", level=2)

    add_para(doc,
        "The release pipeline is fully automated from commit to production verification:",
        space_after=4)

    add_bullet(doc, "Developer pushes to main branch on GitHub", bold_prefix="1. Code push: ")
    add_bullet(doc, "Railway detects the push and starts a Nixpacks build", bold_prefix="2. Railway build: ")
    add_bullet(doc, "pre_deploy.py validates Alembic state, flask db upgrade applies migrations", bold_prefix="3. Pre-deploy: ")
    add_bullet(doc, "Gunicorn starts with --preload and gunicorn.conf.py (4 workers, gthread)", bold_prefix="4. Server start: ")
    add_bullet(doc, "GitHub Actions waits 60 s, verifies /api/health, then runs test_e2e_final.py", bold_prefix="5. E2E gate: ")
    add_bullet(doc, "Test results uploaded as artifacts; failures block further releases", bold_prefix="6. Verification: ")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    #  12. TECH DEBT & FUTURE HARDENING
    # ══════════════════════════════════════════════════════════════════

    add_section_heading(doc, "12. Tech Debt & Future Hardening", level=1)
    add_horizontal_rule(doc)

    add_para(doc,
        "The following items are recognized technical debt and hardening opportunities. None are "
        "blocking issues — the system is production-ready — but they should be addressed as the "
        "platform matures.",
        space_after=8)

    add_styled_table(doc,
        ["#", "Item", "Priority", "Current State", "Recommended Action"],
        [
            ["TD-1", "CSP unsafe-inline", "Medium",
             "Content-Security-Policy uses 'unsafe-inline' for both script-src "
             "and style-src (middleware.py). Required by inline event handlers "
             "and dynamically injected styles in the SPA.",
             "Migrate to nonce-based CSP: generate a per-request nonce in the WSGI "
             "middleware, inject it into script/style tags via the template, and "
             "replace 'unsafe-inline' with 'nonce-<value>'. Requires auditing all "
             "inline scripts and styles in app.js and index.html."],
            ["TD-2", "Object storage for uploads", "Medium",
             "File uploads (grant agreements, documents) are stored on the local "
             "filesystem via UPLOAD_FOLDER (config.py, documents.py). Railway's "
             "ephemeral filesystem means files are lost on each redeploy.",
             "Introduce an abstract StorageService with S3-compatible backend "
             "(AWS S3, Railway Volumes, or Cloudflare R2). Keep local filesystem "
             "as the development adapter. Migrate existing upload/download routes "
             "in documents.py to use the service interface."],
            ["TD-3", "In-memory caches (SimpleCache)", "Low",
             "Werkzeug SimpleCache instances (sanctions, lists, dashboard) are "
             "per-worker — 4 workers maintain independent caches with no shared "
             "invalidation.",
             "Evaluate Redis or Railway Redis add-on for shared cache state. "
             "Benefits: cross-worker consistency, survives worker recycling, "
             "enables cache warming on deploy."],
            ["TD-4", "Frontend framework migration", "Low",
             "The SPA is a single 5,810-line vanilla JS file (app.js) with manual "
             "DOM manipulation, string-template rendering, and hand-rolled routing.",
             "Consider a lightweight framework (Preact, Alpine.js, or htmx) to "
             "improve maintainability. Prioritize if the frontend grows beyond "
             "~7,000 lines or if multiple developers contribute simultaneously."],
            ["TD-5", "Automated unit & integration tests", "Medium",
             "Testing relies on E2E regression (test_e2e_final.py, 136 cases) "
             "running against the live deployment. No unit or integration tests "
             "exist for individual Flask routes or service functions.",
             "Add pytest-based unit tests for critical paths: auth, AI scoring, "
             "sanctions screening, file upload validation. Target 70%+ coverage "
             "on business logic. Run in CI before the E2E gate."],
            ["TD-6", "Database backup strategy", "Medium",
             "Railway PostgreSQL auto-provisions the database but backups depend "
             "on Railway's built-in snapshots. No application-level backup or "
             "point-in-time recovery is configured.",
             "Enable Railway's scheduled backups (or pg_dump cron via GitHub "
             "Actions). Store backups in S3. Document recovery procedures and "
             "test restore at least quarterly."],
        ],
        col_widths=[0.5, 1.3, 0.7, 2.0, 2.0],
        first_col_bold=True)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    #  APPENDIX A: GLOSSARY
    # ══════════════════════════════════════════════════════════════════

    add_section_heading(doc, "Appendix A: Glossary of Terms", level=1)
    add_horizontal_rule(doc)

    glossary = [
        ["API", "Application Programming Interface; the HTTP-based interface for client-server communication"],
        ["ACSO", "Agency for Civil Society Organizations (Ethiopia)"],
        ["CAC", "Corporate Affairs Commission (Nigeria)"],
        ["CHS", "Core Humanitarian Standard; a quality and accountability framework for humanitarian response"],
        ["CIPC", "Companies and Intellectual Property Commission (South Africa)"],
        ["CORS", "Cross-Origin Resource Sharing; security mechanism for cross-domain HTTP requests"],
        ["CSRF", "Cross-Site Request Forgery; an attack prevented by custom header validation"],
        ["DFD", "Data Flow Diagram; visual representation of data movement through the system"],
        ["ER", "Entity-Relationship; a data modeling approach for database design"],
        ["FK", "Foreign Key; a database constraint linking records across tables"],
        ["i18n", "Internationalization; the design of software for multi-language support"],
        ["JSON", "JavaScript Object Notation; lightweight data interchange format"],
        ["M&E", "Monitoring and Evaluation; systematic assessment of program performance"],
        ["NGO", "Non-Governmental Organization; the primary user type on the platform"],
        ["NUPAS", "NGO and CSO USAID Pre-Award Survey; USAID's organizational assessment tool"],
        ["OFAC", "Office of Foreign Assets Control; US Treasury sanctions program"],
        ["ORM", "Object-Relational Mapping; database abstraction via SQLAlchemy"],
        ["PK", "Primary Key; unique identifier for database records"],
        ["RTL", "Right-to-Left; text direction for Arabic language support"],
        ["SDG", "Sustainable Development Goals; UN's 17 global development objectives"],
        ["SDN", "Specially Designated Nationals; OFAC's sanctions list"],
        ["SPA", "Single-Page Application; frontend architecture loading one HTML page"],
        ["STEP", "A capacity assessment framework (Systems, Technical, Effectiveness, Partnerships)"],
        ["TTL", "Time To Live; cache expiration duration"],
        ["UN-HACT", "UN Harmonized Approach to Cash Transfers; UN's partner assessment framework"],
        ["URSB", "Uganda Registration Services Bureau"],
        ["WSGI", "Web Server Gateway Interface; Python web server standard"],
    ]

    add_styled_table(doc,
        ["Term", "Definition"],
        glossary,
        col_widths=[1.2, 5.3],
        first_col_bold=True)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    #  APPENDIX B: CONFIGURATION REFERENCE
    # ══════════════════════════════════════════════════════════════════

    add_section_heading(doc, "Appendix B: Configuration Reference", level=1)
    add_horizontal_rule(doc)

    add_para(doc, "Flask Application Configuration:", bold=True, space_after=4)

    add_styled_table(doc,
        ["Setting", "Development", "Production", "Description"],
        [
            ["FLASK_ENV", "development", "production", "Flask environment mode"],
            ["DEBUG", "True", "False", "Debug mode (auto-reload, verbose errors)"],
            ["SECRET_KEY", "dev-key-xxx", "Environment variable", "Session encryption secret"],
            ["SQLALCHEMY_DATABASE_URI", "sqlite:///kuja.db", "DATABASE_URL (PostgreSQL)", "Database connection string"],
            ["SQLALCHEMY_TRACK_MODIFICATIONS", "False", "False", "Disable modification tracking (performance)"],
            ["MAX_CONTENT_LENGTH", "16 MB", "16 MB", "Maximum upload file size"],
            ["SESSION_COOKIE_HTTPONLY", "True", "True", "Prevent JS access to session cookie"],
            ["SESSION_COOKIE_SAMESITE", "Lax", "Lax", "CSRF protection for cookies"],
            ["SESSION_COOKIE_SECURE", "False", "True", "HTTPS-only cookie in production"],
            ["PERMANENT_SESSION_LIFETIME", "8 hours", "8 hours", "Session expiration time"],
        ],
        col_widths=[2.0, 1.2, 1.3, 2.0],
        first_col_bold=True)

    doc.add_paragraph()

    add_para(doc, "Gunicorn Production Configuration:", bold=True, space_after=4)

    add_styled_table(doc,
        ["Parameter", "Value", "Flag"],
        [
            ["Bind", "0.0.0.0:${PORT:-5000}", "--bind"],
            ["Workers", "4", "--workers"],
            ["Threads", "4", "--threads"],
            ["Worker Class", "gthread", "--worker-class"],
            ["Timeout", "180", "--timeout"],
            ["Max Requests", "1000", "--max-requests"],
            ["Max Requests Jitter", "50", "--max-requests-jitter"],
            ["Preload App", "True", "--preload"],
            ["Config File", "gunicorn.conf.py", "-c"],
            ["Request Body Limit", "0 (disabled)", "gunicorn.conf.py"],
            ["Access Log", "- (stdout)", "--access-logfile"],
            ["Error Log", "- (stderr)", "--error-logfile"],
        ],
        col_widths=[2.0, 2.0, 2.5],
        first_col_bold=True)

    # ── Final spacer and end ───────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("--- END OF DOCUMENT ---")
    run.font.size = Pt(10)
    run.font.color.rgb = MEDIUM_GRAY
    run.font.name = "Calibri"
    run.italic = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Kuja Grant Management System  |  Technical Design Specification v3.3.4\n"
                     "Copyright 2026 Adeso. All rights reserved.\n"
                     "This document is confidential and intended for authorized personnel only.")
    run.font.size = Pt(8)
    run.font.color.rgb = MEDIUM_GRAY
    run.font.name = "Calibri"
    run.italic = True

    # ── Apply headers/footers ──────────────────────────────────────────
    setup_header_footer(doc)

    # ── Save ───────────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")
    print(f"File size: {os.path.getsize(OUTPUT_PATH):,} bytes")


if __name__ == "__main__":
    generate()
